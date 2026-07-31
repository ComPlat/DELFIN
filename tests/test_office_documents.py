"""Spreadsheets, PDFs and Word files: reading, editing, filling forms.

The failure modes pinned here are the silent ones — the cases where the
naive implementation reports success and leaves the user with a wrong
file or a wrong number:

* ``read_file`` decoding a container format as text and returning noise
  that reads like data;
* a formula cell with no cached value rendering as an empty cell;
* a filled form whose values a viewer will not display;
* a check box "set" to a state the field does not have;
* a field name typo swallowed instead of reported;
* a document edit that cannot be undone because the journal stored a
  lossy text pre-image.
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

from delfin.agent import change_journal as cj
from delfin.agent import office
from delfin.agent.api_client import (
    _DocToolExecutor,
    _OFFICE_TOOL_NAMES,
    KitToolPermissions,
    ToolSurfaceContext,
    _binary_read_hint,
    advertisable_tools,
    _DOC_TOOLS_OPENAI,
    tool_unavailable_reason,
)

openpyxl = pytest.importorskip("openpyxl")
pypdf = pytest.importorskip("pypdf")


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def ws(tmp_path):
    d = tmp_path / "ws"
    d.mkdir()
    return d


@pytest.fixture
def home(monkeypatch, tmp_path):
    h = tmp_path / "home"
    h.mkdir()
    monkeypatch.setattr(Path, "home", lambda: h)
    return h


@pytest.fixture
def book(ws):
    """A workbook with data, a formula and a chart."""
    from openpyxl.chart import BarChart, Reference

    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "Kosten"
    for row in (["Monat", "Betrag"], ["Jan", 100], ["Feb", 140], ["Mrz", 90]):
        sheet.append(row)
    sheet["D1"] = "=SUM(B2:B4)"
    chart = BarChart()
    chart.add_data(Reference(sheet, min_col=2, min_row=1, max_row=4),
                   titles_from_data=True)
    sheet.add_chart(chart, "F2")
    path = ws / "kosten.xlsx"
    wb.save(path)
    wb.close()
    return path


@pytest.fixture
def form(ws):
    """A PDF with a text field and a check box."""
    reportlab = pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    path = ws / "antrag.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    c.drawString(60, 780, "Antrag")
    acro = c.acroForm
    acro.textfield(name="name", x=180, y=735, width=300, height=18)
    acro.textfield(name="betrag", x=180, y=705, width=300, height=18)
    acro.checkbox(name="genehmigt", x=180, y=675, size=16)
    c.save()
    return path


def _perms(ws: Path, mode: str = "acceptEdits") -> KitToolPermissions:
    perms = KitToolPermissions(workspace=str(ws))
    perms.mode = mode
    perms.task_session_id = "office-test"
    return perms


def _read(ex, perms, path, **kw) -> str:
    return ex._dispatch("read_document", {"path": str(path), **kw}, perms)


# ---------------------------------------------------------------------------
# read_file must refuse container formats
# ---------------------------------------------------------------------------

def test_read_file_refuses_a_spreadsheet_and_names_the_reader(ws, book):
    ex = _DocToolExecutor()
    out = json.loads(ex._execute_read_file({"path": str(book)}, _perms(ws)))
    assert "read_document" in out["error"]
    # The whole point: no binary noise came back as if it were content.
    assert "PK" not in out["error"]


def test_read_file_refuses_pdf_and_word(ws, form):
    ex = _DocToolExecutor()
    for name in ("antrag.pdf", "bericht.docx"):
        p = ws / name
        if not p.exists():
            p.write_bytes(b"PK\x03\x04binary")
        out = json.loads(ex._execute_read_file({"path": str(p)}, _perms(ws)))
        assert "read_document" in out["error"], name


def test_formats_without_a_reader_say_convert_first(ws):
    p = ws / "alt.xls"
    p.write_bytes(b"\xd0\xcf\x11\xe0binary")
    hint = _binary_read_hint(p)
    assert hint is not None
    assert "convert" in hint.lower()
    # It must not send the model back through another tool for the same read.
    assert "read_document" not in hint


def test_unnamed_binary_is_caught_by_content(ws):
    p = ws / "blob.dat"
    p.write_bytes(b"header\x00\x00payload")
    assert _binary_read_hint(p) is not None


def test_text_files_are_untouched_by_the_guard(ws):
    for name, body in (("a.txt", "hello"), ("b.csv", "x,y\n1,2\n"),
                       ("c.md", "# title"), ("d.py", "print(1)\n")):
        p = ws / name
        p.write_text(body, encoding="utf-8")
        assert _binary_read_hint(p) is None, name


def test_csv_is_still_readable_through_read_file(ws):
    p = ws / "liste.csv"
    p.write_text("Name,Betrag\nMeier,100\n", encoding="utf-8")
    ex = _DocToolExecutor()
    out = ex._execute_read_file({"path": str(p)}, _perms(ws))
    assert "Meier" in out


# ---------------------------------------------------------------------------
# Reading spreadsheets
# ---------------------------------------------------------------------------

def test_grid_is_addressable(book):
    result = office.read_sheet(book)
    grid = result["grid"]
    assert "| A" in grid and "| B" in grid       # column letters
    assert "\n  2 |" in grid                     # row numbers
    assert "Jan" in grid
    assert result["sheet"] == "Kosten"


def test_uncached_formula_is_reported_not_shown_as_empty(book):
    result = office.read_sheet(book)
    assert "=SUM(B2:B4)" in result["grid"]
    note = " ".join(result["notes"])
    assert "no cached value" in note
    assert "do not report them as computed" in note


def test_cached_formula_values_are_used_when_present(ws):
    wb = openpyxl.Workbook()
    wb.active["A1"] = "=1+1"
    wb.save(ws / "cached.xlsx")
    wb.close()
    # Simulate what a spreadsheet program leaves behind: a cached result.
    import zipfile
    src = ws / "cached.xlsx"
    with zipfile.ZipFile(src) as zf:
        items = {n: zf.read(n) for n in zf.namelist()}
    items["xl/worksheets/sheet1.xml"] = items["xl/worksheets/sheet1.xml"].replace(
        b"<f>1+1</f>", b"<f>1+1</f><v>2</v>")
    with zipfile.ZipFile(src, "w") as zf:
        for name, data in items.items():
            zf.writestr(name, data)

    result = office.read_sheet(src)
    assert "\n  1 | 2" in result["grid"]
    assert not any("cached" in n for n in result["notes"])


def test_fragile_content_is_named_before_a_rewrite(book):
    assert any("chart" in n for n in office.read_sheet(book)["notes"])


def test_csv_reading_detects_the_delimiter(ws):
    p = ws / "semi.csv"
    p.write_text("Name;Betrag\nMeier;100\nSchmidt;250\n", encoding="utf-8")
    result = office.read_sheet(p)
    assert result["rows"] == 3
    assert "Meier" in result["grid"] and "100" in result["grid"]


def test_paging_reports_the_remainder(ws):
    wb = openpyxl.Workbook()
    for n in range(500):
        wb.active.append([n])
    wb.save(ws / "long.xlsx")
    wb.close()
    result = office.read_sheet(ws / "long.xlsx", max_rows=10)
    assert result["rows"] == 500
    assert any("start_row" in n for n in result["notes"])


def test_missing_sheet_lists_the_real_ones(book):
    with pytest.raises(office.OfficeError) as exc:
        office.read_sheet(book, sheet="Nope")
    assert "Kosten" in str(exc.value)


# ---------------------------------------------------------------------------
# Editing spreadsheets
# ---------------------------------------------------------------------------

def test_edit_sets_cells_and_appends_rows(book):
    result = office.edit_sheet(
        book, edits=[{"cell": "B2", "value": 250}],
        append_rows=[["Apr", 175]])
    assert result["applied"] == [{"cell": "B2", "old": "100", "new": "250"}]
    assert result["first_appended_row"] == 5

    grid = office.read_sheet(book)["grid"]
    assert "250" in grid and "Apr" in grid


def test_editing_preserves_other_formulas(book):
    office.edit_sheet(book, edits=[{"cell": "B2", "value": 250}])
    wb = openpyxl.load_workbook(book)
    assert wb["Kosten"]["D1"].value == "=SUM(B2:B4)"
    wb.close()


def test_writing_a_formula_warns_it_has_no_result_yet(book):
    result = office.edit_sheet(book, edits=[{"cell": "B6", "value": "=B2*2"}])
    assert any("no computed result" in n for n in result["notes"])


def test_edit_names_the_fragile_content_it_rewrote(book):
    result = office.edit_sheet(book, edits=[{"cell": "B2", "value": 1}])
    assert any("chart" in n for n in result["notes"])


def test_bad_cell_reference_is_refused(book):
    with pytest.raises(office.OfficeError) as exc:
        office.edit_sheet(book, edits=[{"cell": "not-a-cell", "value": 1}])
    assert "cell reference" in str(exc.value)


def test_plan_reports_the_change_without_writing(book):
    before = book.read_bytes()
    plan = office.plan_sheet_edits(book, edits=[{"cell": "B2", "value": 999}])
    assert plan["changes"] == [{"cell": "B2", "old": "100", "new": "999"}]
    assert book.read_bytes() == before


def test_create_refuses_to_clobber(ws, book):
    with pytest.raises(office.OfficeError):
        office.create_sheet(book, [["a"]])


def test_create_writes_xlsx_and_csv(ws):
    x = office.create_sheet(ws / "neu.xlsx", [["A", "B"], [1, 2]])
    assert x["rows"] == 2 and (ws / "neu.xlsx").exists()
    c = office.create_sheet(ws / "neu.csv", [["A", "B"], [1, 2]])
    assert (ws / "neu.csv").read_text().startswith("A,B")
    assert c["columns"] == 2


# ---------------------------------------------------------------------------
# PDF forms
# ---------------------------------------------------------------------------

def test_fields_are_listed_with_their_states(form):
    result = office.pdf_form_fields(form)
    by_name = {f["name"]: f for f in result["fields"]}
    assert set(by_name) == {"name", "betrag", "genehmigt"}
    assert by_name["genehmigt"]["type"] == "button/checkbox"
    assert "/Yes" in by_name["genehmigt"]["states"]


def test_filling_writes_values_that_read_back(form, ws):
    out = ws / "antrag_ausgefuellt.pdf"
    result = office.fill_pdf_form(
        form, {"name": "M. Meier", "betrag": "1234,50"}, output=out)
    assert result["verified"] is True
    values = {f["name"]: f["value"]
              for f in office.pdf_form_fields(out)["fields"]}
    assert values["name"] == "M. Meier"
    assert values["betrag"] == "1234,50"


def test_checkbox_true_maps_to_the_fields_own_state(form, ws):
    out = ws / "ok.pdf"
    result = office.fill_pdf_form(form, {"genehmigt": True}, output=out)
    values = {f["name"]: f["value"]
              for f in office.pdf_form_fields(out)["fields"]}
    assert values["genehmigt"] == "/Yes"
    assert any("mapped" in n for n in result["notes"])


def test_checkbox_false_maps_to_off(form, ws):
    out = ws / "off.pdf"
    office.fill_pdf_form(form, {"genehmigt": False}, output=out)
    values = {f["name"]: f["value"]
              for f in office.pdf_form_fields(out)["fields"]}
    assert values["genehmigt"] == "/Off"


def test_need_appearances_is_set_so_viewers_show_the_values(form, ws):
    out = ws / "appear.pdf"
    office.fill_pdf_form(form, {"name": "X"}, output=out)
    reader = pypdf.PdfReader(str(out))
    acro = reader.trailer["/Root"]["/AcroForm"]
    assert bool(acro.get("/NeedAppearances")) is True


def test_unknown_field_name_is_an_error_not_a_silent_skip(form, ws):
    with pytest.raises(office.OfficeError) as exc:
        office.fill_pdf_form(form, {"nmae": "typo"}, output=ws / "x.pdf")
    assert "nmae" in str(exc.value)
    assert "name" in str(exc.value)          # the real names are offered
    assert not (ws / "x.pdf").exists()


def test_filling_in_place_is_refused(form):
    with pytest.raises(office.OfficeError) as exc:
        office.fill_pdf_form(form, {"name": "X"}, output=form)
    assert "blank original" in str(exc.value)


def test_a_pdf_without_fields_is_reported_as_such(ws):
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    plain = ws / "plain.pdf"
    c = canvas.Canvas(str(plain))
    c.drawString(60, 780, "Nur Text")
    c.save()

    result = office.pdf_form_fields(plain)
    assert result["field_count"] == 0
    assert any("no form fields" in n for n in result["notes"])
    with pytest.raises(office.OfficeError):
        office.fill_pdf_form(plain, {"x": "1"}, output=ws / "y.pdf")


def test_flatten_makes_fields_read_only(form, ws):
    out = ws / "flat.pdf"
    result = office.fill_pdf_form(
        form, {"name": "X"}, output=out, flatten=True)
    note = " ".join(result["notes"])
    assert "read-only" in note
    # The note must not claim more than the operation does.
    assert "not merged into the page content" in note
    reader = pypdf.PdfReader(str(out))
    flags = [
        int(a.get_object().get("/Ff", 0) or 0)
        for a in (reader.pages[0].get("/Annots") or [])
    ]
    assert flags and all(f & 1 for f in flags)


def test_pdf_text_extraction_and_scan_detection(ws):
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    doc = ws / "text.pdf"
    c = canvas.Canvas(str(doc))
    c.drawString(60, 780, "Rechnung 4711")
    c.showPage()
    c.showPage()          # deliberately blank: stands in for a scan
    c.save()

    result = office.read_pdf(doc)
    assert "Rechnung 4711" in result["text"]
    assert result["pages"] == 2
    assert any("held no text" in n for n in result["notes"])

    only_blank = office.read_pdf(doc, pages="2")
    assert any("OCR" in n for n in only_blank["notes"])


def test_page_selection(ws):
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    doc = ws / "multi.pdf"
    c = canvas.Canvas(str(doc))
    for n in range(1, 5):
        c.drawString(60, 780, f"Seite {n}")
        c.showPage()
    c.save()
    result = office.read_pdf(doc, pages="2-3")
    assert result["pages_read"] == [2, 3]
    assert "Seite 2" in result["text"] and "Seite 4" not in result["text"]


def test_xfa_forms_are_flagged(form, ws):
    """AcroForm writes do not surface in an XFA form — say so."""
    from pypdf.generic import ArrayObject, NameObject, TextStringObject
    xfa = ws / "xfa.pdf"
    writer = pypdf.PdfWriter(clone_from=str(form))
    acro = writer._root_object["/AcroForm"]
    acro[NameObject("/XFA")] = ArrayObject([TextStringObject("<xdp/>")])
    with xfa.open("wb") as fh:
        writer.write(fh)

    assert office.pdf_form_fields(xfa)["xfa"] is True
    result = office.fill_pdf_form(xfa, {"name": "X"}, output=ws / "out.pdf")
    assert any("XFA" in n for n in result["notes"])


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def test_docx_paragraphs_and_tables(ws):
    docx = pytest.importorskip("docx")
    p = ws / "bericht.docx"
    doc = docx.Document()
    doc.add_paragraph("Zusammenfassung")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Posten"
    table.cell(0, 1).text = "Betrag"
    table.cell(1, 0).text = "Reise"
    table.cell(1, 1).text = "240"
    doc.save(p)

    result = office.read_docx(p)
    assert "Zusammenfassung" in result["text"]
    assert "Reise" in result["text"] and "240" in result["text"]
    assert result["tables"] == 1


def test_legacy_doc_is_refused_with_a_reason(ws):
    p = ws / "alt.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0")
    with pytest.raises(office.OfficeError) as exc:
        office.read_docx(p)
    assert "convert" in str(exc.value)


# ---------------------------------------------------------------------------
# Tool layer
# ---------------------------------------------------------------------------

def test_read_document_returns_a_grid_through_the_tool(ws, book):
    ex = _DocToolExecutor()
    out = _read(ex, _perms(ws), book)
    assert "Jan" in out and "| B" in out
    assert "NOTE:" in out


def test_read_document_registers_the_read(ws, book):
    ex = _DocToolExecutor()
    perms = _perms(ws)
    ex.execute("read_document", {"path": str(book)}, perms)
    assert str(book.resolve()) in perms.read_tracker


def test_edit_sheet_refuses_without_a_prior_read(ws, book):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch(
        "edit_sheet",
        {"path": str(book), "edits": [{"cell": "B2", "value": 1}]},
        _perms(ws)))
    assert "without reading it first" in out["error"]


def test_edit_sheet_applies_after_a_read(ws, book, home):
    ex = _DocToolExecutor()
    perms = _perms(ws)
    ex.execute("read_document", {"path": str(book)}, perms)
    out = json.loads(ex._dispatch(
        "edit_sheet",
        {"path": str(book), "edits": [{"cell": "B2", "value": 250}]},
        perms))
    assert out["status"] == "ok"
    assert "B2: 100 -> 250" in out["summary"]


def test_a_document_edit_can_be_undone(ws, book, home):
    ex = _DocToolExecutor()
    perms = _perms(ws)
    before = book.read_bytes()
    ex.execute("read_document", {"path": str(book)}, perms)
    ex._dispatch("edit_sheet",
                 {"path": str(book), "edits": [{"cell": "B2", "value": 250}]},
                 perms)
    assert book.read_bytes() != before

    result = cj.revert("office-test", scope="last")
    assert result["reverted"] == [str(book)]
    assert book.read_bytes() == before


def test_edit_sheet_can_create_a_new_file(ws, home):
    ex = _DocToolExecutor()
    target = ws / "neu.xlsx"
    out = json.loads(ex._dispatch("edit_sheet", {
        "path": str(target), "create": True,
        "append_rows": [["Posten", "Betrag"], ["Reise", 240]],
    }, _perms(ws)))
    assert out["status"] == "ok"
    assert target.exists()
    assert "Reise" in office.read_sheet(target)["grid"]


def test_create_on_an_existing_file_is_refused(ws, book):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch("edit_sheet", {
        "path": str(book), "create": True, "append_rows": [["x"]],
    }, _perms(ws)))
    assert "already exists" in out["error"]


def test_plan_mode_refuses_document_writes(ws, book):
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="plan")
    for name, args in (
        ("edit_sheet", {"path": str(book),
                        "edits": [{"cell": "A1", "value": 1}]}),
        ("fill_pdf_form", {"path": str(book), "output": str(ws / "o.pdf"),
                           "values": {"a": "b"}}),
    ):
        out = json.loads(ex._dispatch(name, args, perms))
        assert "plan mode" in out["error"], name


def test_diff_approval_asks_with_a_cell_level_preview(ws, book, home):
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="diff_approval")
    seen = {}

    def confirm(name, args, preview):
        seen["preview"] = preview
        return True

    perms.confirm_callback = confirm
    ex.execute("read_document", {"path": str(book)}, perms)
    out = json.loads(ex._dispatch(
        "edit_sheet",
        {"path": str(book), "edits": [{"cell": "B2", "value": 250}]},
        perms))
    assert out["status"] == "ok"
    assert "B2: 100 -> 250" in seen["preview"]
    assert "chart" in seen["preview"]        # what the rewrite touches


def test_a_declined_document_edit_changes_nothing(ws, book, home):
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="diff_approval")
    perms.confirm_callback = lambda *a: False
    before = book.read_bytes()
    ex.execute("read_document", {"path": str(book)}, perms)
    out = json.loads(ex._dispatch(
        "edit_sheet",
        {"path": str(book), "edits": [{"cell": "B2", "value": 250}]},
        perms))
    assert "declined" in out["error"]
    assert book.read_bytes() == before


def test_diff_approval_without_a_dialog_refuses(ws, book):
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="diff_approval")
    ex.execute("read_document", {"path": str(book)}, perms)
    out = json.loads(ex._dispatch(
        "edit_sheet",
        {"path": str(book), "edits": [{"cell": "B2", "value": 1}]},
        perms))
    assert "no approval dialog" in out["error"]


def test_fill_pdf_form_through_the_tool(ws, form, home):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch("fill_pdf_form", {
        "path": str(form), "output": str(ws / "fertig.pdf"),
        "values": {"name": "M. Meier", "genehmigt": True},
    }, _perms(ws)))
    assert out["status"] == "ok"
    assert out["verified"] is True
    assert sorted(out["filled"]) == ["genehmigt", "name"]


def test_fill_pdf_form_reports_a_typo_as_an_error(ws, form):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch("fill_pdf_form", {
        "path": str(form), "output": str(ws / "x.pdf"),
        "values": {"nmae": "typo"},
    }, _perms(ws)))
    assert "nmae" in out["error"]


def test_document_tools_need_permissions(ws, book):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch(
        "edit_sheet", {"path": str(book)}, None))
    assert "permissions" in out["error"]


# ---------------------------------------------------------------------------
# Tool surface
# ---------------------------------------------------------------------------

def test_office_tools_are_hidden_without_a_backend():
    ctx = ToolSurfaceContext(has_office_libs=False)
    advertised = {t["function"]["name"]
                  for t in advertisable_tools(_DOC_TOOLS_OPENAI, ctx)}
    assert not (_OFFICE_TOOL_NAMES & advertised)
    for name in _OFFICE_TOOL_NAMES:
        assert tool_unavailable_reason(name, ctx) == (
            "document backend not installed")


def test_office_tools_are_advertised_when_the_backend_is_there():
    ctx = ToolSurfaceContext(has_office_libs=True)
    advertised = {t["function"]["name"]
                  for t in advertisable_tools(_DOC_TOOLS_OPENAI, ctx)}
    assert _OFFICE_TOOL_NAMES <= advertised


def test_the_guide_role_gets_no_document_tools():
    """The dashboard guide has no file surface at all — office is no gap."""
    for name in _OFFICE_TOOL_NAMES:
        assert tool_unavailable_reason(
            name, ToolSurfaceContext(role="dashboard_agent")) is not None


def test_document_writes_reach_the_audit_trail(ws, book, home):
    from delfin.agent import audit_log as al

    ex = _DocToolExecutor()
    perms = _perms(ws)
    ex.execute("read_document", {"path": str(book)}, perms)
    ex.execute("edit_sheet",
               {"path": str(book), "edits": [{"cell": "B2", "value": 7}]},
               perms)

    report = al.build_changes_report("office-test")
    rendered = al.format_changes_report(report)
    assert "edit_sheet" in rendered
    assert "kosten.xlsx" in rendered


def test_an_unreadable_file_keeps_its_real_error(ws):
    """The sniff must not relabel a permission problem as 'binary'."""
    import os
    p = ws / "locked.log"
    p.write_text("plain text\n", encoding="utf-8")
    os.chmod(p, 0o000)
    try:
        if os.access(p, os.R_OK):        # running as root — nothing to test
            pytest.skip("cannot make a file unreadable here")
        assert _binary_read_hint(p) is None
    finally:
        os.chmod(p, 0o644)


def test_a_denied_read_also_blocks_read_document(tmp_path, ws):
    """The refusal ledger covers every reading tool, this one included."""
    outside = tmp_path / "other" / "geheim.xlsx"
    outside.parent.mkdir(parents=True)
    wb = openpyxl.Workbook()
    wb.active["A1"] = "vertraulich"
    wb.save(outside)
    wb.close()

    perms = _perms(ws)
    perms.mode = "default"
    perms.confirm_callback = lambda *a: False       # the user declines
    ex = _DocToolExecutor()

    first = json.loads(ex.execute(
        "read_document", {"path": str(outside)}, perms))
    assert "declined" in first["error"]

    # Second attempt: refused from the ledger, without asking again.
    perms.confirm_callback = lambda *a: True
    second = json.loads(ex.execute(
        "read_document", {"path": str(outside)}, perms))
    assert "error" in second


def test_a_denied_read_also_blocks_read_document(tmp_path, ws):
    """The refusal ledger covers every reading tool, this one included."""
    outside = tmp_path / "other" / "geheim.xlsx"
    outside.parent.mkdir(parents=True)
    wb = openpyxl.Workbook()
    wb.active["A1"] = "vertraulich"
    wb.save(outside)
    wb.close()

    perms = _perms(ws)
    perms.mode = "default"
    perms.confirm_callback = lambda *a: False       # the user declines
    ex = _DocToolExecutor()

    first = json.loads(ex.execute(
        "read_document", {"path": str(outside)}, perms))
    assert "declined" in first["error"]

    # Second attempt: refused from the ledger, without asking again.
    perms.confirm_callback = lambda *a: True
    second = json.loads(ex.execute(
        "read_document", {"path": str(outside)}, perms))
    assert "already declined" in second["error"]


def test_numeric_arguments_survive_weak_model_formatting(ws, book):
    """A model that sends "2" or 2.0 for a row count must still page."""
    ex = _DocToolExecutor()
    perms = _perms(ws)
    for value in ("2", 2.0, 2):
        out = ex._dispatch(
            "read_document",
            {"path": str(book), "max_rows": value}, perms)
        assert "Jan" in out, repr(value)
        assert "Mrz" not in out, repr(value)      # the window really applied


def test_creating_a_sheet_also_goes_up_for_approval(ws, home):
    """write_file stages a creation; creating a workbook must not slip past."""
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="diff_approval")
    seen = {}

    def confirm(name, args, preview):
        seen["preview"] = preview
        return False

    perms.confirm_callback = confirm
    target = ws / "neu.xlsx"
    out = json.loads(ex._dispatch("edit_sheet", {
        "path": str(target), "create": True,
        "append_rows": [["Posten", "Betrag"], ["Reise", 240]],
    }, perms))

    assert "declined" in out["error"]
    assert not target.exists()
    assert "Create" in seen["preview"] and "Reise" in seen["preview"]


# ---------------------------------------------------------------------------
# Word templates and new documents
# ---------------------------------------------------------------------------

@pytest.fixture
def template(ws):
    """A letter template with a placeholder split across runs.

    Word splits a placeholder into several runs whenever it feels like
    it — a spell-check mark, an edit, a formatting change. Building the
    fixture that way is the point: per-run replacement finds nothing.
    """
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    para = doc.add_paragraph()
    for piece in ("Sehr geehrte(r) {{", "an", "rede", "}}, ", "hiermit ..."):
        para.add_run(piece)
    doc.add_paragraph("{{ort}}, den {{datum}}")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Betrag"
    table.cell(0, 1).text = "{{betrag}}"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Amt für {{abteilung}}"
    section.footer.paragraphs[0].text = "Az. {{aktenzeichen}}"
    path = ws / "vorlage.docx"
    doc.save(path)
    return path


def _docx_text(path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    parts += [c.text for t in doc.tables for r in t.rows for c in r.cells]
    section = doc.sections[0]
    parts += [p.text for p in section.header.paragraphs]
    parts += [p.text for p in section.footer.paragraphs]
    return "\n".join(parts)


def test_placeholders_are_listed_from_body_table_header_and_footer(template):
    names = {p["name"]
             for p in office.docx_placeholders(template)["placeholders"]}
    assert names == {"anrede", "ort", "datum", "betrag",
                     "abteilung", "aktenzeichen"}


def test_a_placeholder_split_across_runs_is_still_replaced(template, ws):
    out = ws / "brief.docx"
    result = office.fill_docx_template(
        template,
        {"anrede": "Frau Meier", "ort": "Karlsruhe", "datum": "31.07.2026",
         "betrag": "1.234,50 EUR", "abteilung": "Beschaffung",
         "aktenzeichen": "IV-7"},
        output=out)
    assert result["complete"] is True
    text = _docx_text(out)
    assert "Sehr geehrte(r) Frau Meier, hiermit" in text
    assert "Beschaffung" in text and "IV-7" in text   # header + footer
    assert "1.234,50 EUR" in text                     # table cell
    assert "{{" not in text


def test_substitution_keeps_the_surrounding_formatting(ws):
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    para = doc.add_paragraph()
    para.add_run("Betrag: ")
    for piece in ("{{be", "trag}}"):
        run = para.add_run(piece)
        run.bold = True
    tail = para.add_run(" (brutto)")
    tail.italic = True
    src = ws / "fmt.docx"
    doc.save(src)

    out = ws / "fmt_out.docx"
    office.fill_docx_template(src, {"betrag": "240,00"}, output=out)
    runs = docx.Document(str(out)).paragraphs[0].runs
    by_text = {r.text: r for r in runs if r.text}
    assert by_text["240,00"].bold is True        # the placeholder's format
    assert by_text["Betrag: "].bold is not True  # untouched neighbours
    assert by_text[" (brutto)"].italic is True


def test_unfilled_placeholders_stay_visible_and_are_reported(template, ws):
    out = ws / "teil.docx"
    result = office.fill_docx_template(
        template, {"anrede": "Herr Yilmaz"}, output=out, strict=False)
    assert result["complete"] is False
    assert "betrag" in result["unfilled"]
    assert any("not ready to hand over" in n for n in result["notes"])
    # Left in the text on purpose: a visible gap is recoverable.
    assert "{{betrag}}" in _docx_text(out)


def test_a_placeholder_typo_is_refused(template, ws):
    with pytest.raises(office.OfficeError) as exc:
        office.fill_docx_template(
            template, {"anrde": "x"}, output=ws / "x.docx")
    assert "anrde" in str(exc.value) and "anrede" in str(exc.value)
    assert not (ws / "x.docx").exists()


def test_filling_a_template_in_place_is_refused(template):
    with pytest.raises(office.OfficeError) as exc:
        office.fill_docx_template(
            template, {"anrede": "x"}, output=template)
    assert "blank original" in str(exc.value)


def test_a_document_without_placeholders_says_so(ws):
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    doc.add_paragraph("Nur Text")
    p = ws / "plain.docx"
    doc.save(p)
    result = office.docx_placeholders(p)
    assert result["placeholders"] == []
    assert any("MERGEFIELD" in n for n in result["notes"])


def test_create_docx_writes_headings_paragraphs_and_tables(ws):
    out = ws / "bericht.docx"
    result = office.create_docx(out, [
        {"heading": "Quartalsbericht", "level": 1},
        {"paragraph": "Zusammenfassung der Ausgaben."},
        {"table": [["Posten", "Betrag"], ["Reise", "240"]],
         "header_row": True},
        {"page_break": True},
        {"heading": "Anhang", "level": 2},
    ])
    assert result["headings"] == 2 and result["tables"] == 1
    text = _docx_text(out)
    assert "Quartalsbericht" in text and "Reise" in text and "Anhang" in text


def test_create_docx_refuses_a_bad_block(ws):
    with pytest.raises(office.OfficeError) as exc:
        office.create_docx(ws / "x.docx", [{"nonsense": 1}])
    assert "block 1" in str(exc.value)


def test_create_docx_refuses_to_clobber(ws):
    out = ws / "einmal.docx"
    office.create_docx(out, [{"paragraph": "eins"}])
    with pytest.raises(office.OfficeError):
        office.create_docx(out, [{"paragraph": "zwei"}])


def test_read_document_lists_placeholders_like_pdf_fields(ws, template):
    ex = _DocToolExecutor()
    out = _read(ex, _perms(ws), template, fields=True)
    assert "anrede" in out and "aktenzeichen" in out


def test_fill_docx_template_through_the_tool(ws, template, home):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch("fill_docx_template", {
        "path": str(template), "output": str(ws / "fertig.docx"),
        "values": {"anrede": "Frau Meier", "ort": "KA", "datum": "1.8.",
                   "betrag": "10", "abteilung": "IT", "aktenzeichen": "A-1"},
    }, _perms(ws)))
    assert out["status"] == "ok" and out["complete"] is True


def test_the_tool_reports_an_incomplete_letter(ws, template, home):
    ex = _DocToolExecutor()
    out = json.loads(ex._dispatch("fill_docx_template", {
        "path": str(template), "output": str(ws / "halb.docx"),
        "values": {"anrede": "Frau Meier"}, "strict": False,
    }, _perms(ws)))
    assert out["complete"] is False
    assert "betrag" in out["unfilled"]


def test_create_docx_through_the_tool_is_undoable(ws, home):
    ex = _DocToolExecutor()
    target = ws / "neu.docx"
    out = json.loads(ex._dispatch("create_docx", {
        "path": str(target),
        "blocks": [{"heading": "Titel", "level": 1},
                   {"paragraph": "Inhalt"}],
    }, _perms(ws)))
    assert out["status"] == "ok" and target.exists()

    assert cj.revert("office-test", scope="last")["reverted"] == [str(target)]
    assert not target.exists()


def test_word_writes_are_refused_in_plan_mode(ws, template):
    ex = _DocToolExecutor()
    perms = _perms(ws, mode="plan")
    for name, args in (
        ("fill_docx_template", {"path": str(template),
                                "output": str(ws / "o.docx"),
                                "values": {"anrede": "x"}}),
        ("create_docx", {"path": str(ws / "n.docx"),
                         "blocks": [{"paragraph": "x"}]}),
    ):
        out = json.loads(ex._dispatch(name, args, perms))
        assert "plan mode" in out["error"], name
