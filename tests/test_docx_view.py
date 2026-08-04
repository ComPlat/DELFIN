"""Word documents: one renderer for showing, searching and editing.

The view used to be produced by mammoth. Good HTML, but nothing in it says
which paragraph of the document a line came from, so the search could count
matches and never jump to one, and there was nothing an edit could be
written back to.

Every block now carries its address, and an edit is spliced into the runs
that are already there rather than rebuilding the paragraph -- which is the
difference between editing a sentence and flattening every bold word in it.
"""

from __future__ import annotations

import pytest

docx = pytest.importorskip('docx')

from delfin.dashboard import docx_view as dv


@pytest.fixture
def letter(tmp_path):
    document = docx.Document()
    document.add_heading('Angebot 2026', level=1)
    para = document.add_paragraph('Sehr geehrte Damen, ')
    para.add_run('vielen Dank').bold = True
    para.add_run(' für Ihre Anfrage.')
    document.add_paragraph('Erster Punkt', style='List Bullet')
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Artikel'
    table.cell(0, 1).text = 'Preis'
    table.cell(1, 0).text = 'Schraube'
    table.cell(1, 1).text = '3,50'
    path = tmp_path / 'angebot.docx'
    document.save(path)
    return path


# ---------------------------------------------------------------------------
# Reading
# ---------------------------------------------------------------------------

def test_every_paragraph_carries_its_address(letter):
    read = dv.read_document(letter)
    addresses = [block.address for block in read.blocks]
    assert 'p:0' in addresses
    assert any(a.startswith('t:0:') for a in addresses), 'table cells are missing'
    assert len(set(addresses)) == len(addresses), 'two blocks share an address'


def test_headings_and_lists_are_recognised(letter):
    read = dv.read_document(letter)
    by_address = {block.address: block for block in read.blocks}
    assert by_address['p:0'].level == 1
    assert by_address['p:2'].listed is True
    assert by_address['p:1'].level == 0


def test_table_cells_know_which_cell_they_are(letter):
    read = dv.read_document(letter)
    cells = [b for b in read.blocks if b.in_table]
    assert cells and all(b.table is not None for b in cells)
    assert {b.text for b in cells} >= {'Artikel', 'Preis', 'Schraube'}


def test_a_doc_file_is_refused_with_a_reason(tmp_path):
    old = tmp_path / 'alt.doc'
    old.write_bytes(b'\xd0\xcf\x11\xe0legacy')
    with pytest.raises(dv.DocxError) as excinfo:
        dv.read_document(old)
    assert 'converted' in str(excinfo.value)


def test_an_oversized_document_is_named_not_truncated(letter, monkeypatch):
    monkeypatch.setattr(dv, 'MAX_FILE_BYTES', 10)
    with pytest.raises(dv.DocxError) as excinfo:
        dv.read_document(letter)
    assert 'MB' in str(excinfo.value)


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def test_the_view_and_the_editor_render_the_same_blocks(letter):
    read = dv.read_document(letter)
    view = dv.render_html(read, editable=False)
    edit = dv.render_html(read, editable=True)
    assert view.count('data-a=') == edit.count('data-a=') == len(read.blocks)
    assert 'contenteditable' not in view
    assert edit.count('contenteditable="true"') == len(read.blocks)


def test_cell_text_is_escaped(tmp_path):
    document = docx.Document()
    document.add_paragraph('<script>alert(1)</script>')
    path = tmp_path / 'x.docx'
    document.save(path)
    html = dv.render_html(dv.read_document(path), editable=True)
    assert '<script>' not in html
    assert '&lt;script&gt;' in html


def test_a_truncated_document_is_shown_but_not_edited(letter, monkeypatch):
    monkeypatch.setattr(dv, 'MAX_BLOCKS', 2)
    read = dv.read_document(letter)
    assert len(read.blocks) == 2
    assert not dv.is_editable(read), (
        'editing a truncated view would address paragraphs by an index the '
        'view no longer agrees with')


# ---------------------------------------------------------------------------
# Searching
# ---------------------------------------------------------------------------

def test_search_finds_hits_in_body_and_tables(letter):
    read = dv.read_document(letter)
    assert [h.address for h in dv.search(read, 'preis')] == ['t:0:0:1:0']
    assert dv.search(read, 'dank')[0].address == 'p:1'


def test_search_offsets_point_into_the_block_not_the_document(letter):
    read = dv.read_document(letter)
    hit = dv.search(read, 'Anfrage')[0]
    block = next(b for b in read.blocks if b.address == hit.address)
    assert block.text[hit.start:hit.end] == 'Anfrage'


def test_search_is_case_insensitive_and_finds_repeats(tmp_path):
    document = docx.Document()
    document.add_paragraph('Rechnung und rechnung und RECHNUNG')
    path = tmp_path / 'r.docx'
    document.save(path)
    assert len(dv.search(dv.read_document(path), 'rechnung')) == 3


def test_an_empty_term_finds_nothing(letter):
    assert dv.search(dv.read_document(letter), '') == []


def test_the_jump_wraps_the_match_instead_of_rewriting_the_paragraph():
    """Rewriting innerHTML would flatten the formatting of the paragraph it
    landed in, and in a Word document that is the document."""
    script = dv.focus_js('calc-scope-1', 'p:3', 4, 9)
    assert 'surroundContents' in script
    assert 'innerHTML' not in script
    assert 'scrollIntoView' in script
    assert '"p:3"' in script


# ---------------------------------------------------------------------------
# Editing
# ---------------------------------------------------------------------------

def test_editing_a_word_keeps_the_bold_run_beside_it(letter):
    result = dv.apply_edits(
        letter, {'p:1': 'Sehr geehrte Damen, vielen Dank für Ihre Nachricht.'})
    dv.save(result['document'], letter)

    runs = docx.Document(str(letter)).paragraphs[1].runs
    assert [(r.text, r.bold) for r in runs if r.text] == [
        ('Sehr geehrte Damen, ', None),
        ('vielen Dank', True),
        (' für Ihre Nachricht.', None),
    ]


def test_the_changed_range_is_the_smallest_one():
    assert dv.changed_range('hello world', 'hello there') == (6, 11, 11)
    assert dv.changed_range('abc', 'abc') == (3, 3, 3)
    assert dv.changed_range('', 'new') == (0, 0, 3)
    assert dv.changed_range('gone', '') == (0, 4, 0)


def test_editing_a_table_cell_writes_to_that_cell(letter):
    result = dv.apply_edits(letter, {'t:0:1:1:0': '4,20'})
    dv.save(result['document'], letter)
    table = docx.Document(str(letter)).tables[0]
    assert table.cell(1, 1).text == '4,20'
    assert table.cell(1, 0).text == 'Schraube', 'a neighbouring cell changed'


def test_an_empty_paragraph_can_be_written_into(tmp_path):
    document = docx.Document()
    document.add_paragraph('')
    path = tmp_path / 'leer.docx'
    document.save(path)
    result = dv.apply_edits(path, {'p:0': 'jetzt mit Text'})
    dv.save(result['document'], path)
    assert docx.Document(str(path)).paragraphs[0].text == 'jetzt mit Text'


def test_an_address_that_does_not_resolve_is_an_error_not_a_silent_skip(letter):
    before = letter.read_bytes()
    with pytest.raises(dv.DocxError) as excinfo:
        dv.apply_edits(letter, {'p:999': 'nope'})
    assert 'left unchanged' in str(excinfo.value)
    assert letter.read_bytes() == before


def test_saving_nothing_is_refused(letter):
    with pytest.raises(dv.DocxError):
        dv.apply_edits(letter, {})


def test_a_round_trip_keeps_what_it_was_not_asked_to_change(letter):
    """python-docx writes back the package it read. The heading, the list
    style and the table survive an edit to an unrelated paragraph."""
    result = dv.apply_edits(letter, {'p:1': 'Kurz.'})
    dv.save(result['document'], letter)

    read = dv.read_document(letter)
    by_address = {b.address: b for b in read.blocks}
    assert by_address['p:0'].level == 1
    assert by_address['p:2'].listed is True
    assert read.tables == 1
    assert by_address['p:1'].text == 'Kurz.'


# ---------------------------------------------------------------------------
# The browser side
# ---------------------------------------------------------------------------

def test_a_block_is_reported_when_it_is_left(letter):
    script = dv.edit_js('calc-scope-1')
    assert "addEventListener('focusout'" in script


def test_it_is_also_reported_while_typing_once_typing_pauses():
    """Leaving the block is the precise moment but not a reliable one:
    pressing Save is a click, and whether that moves focus out of the
    paragraph first is up to the browser. When it did not, the edit was
    never sent and saving reported nothing to save."""
    script = dv.edit_js('calc-scope-1')
    assert "addEventListener('input'" in script
    assert 'setTimeout' in script.split("addEventListener('input'")[1][:300], (
        'reporting every keystroke would put the kernel behind the typing')
    for per_character in ("addEventListener('keyup'", "addEventListener('keypress'"):
        assert per_character not in script


def test_enter_does_not_create_a_paragraph_the_view_cannot_address():
    script = dv.edit_js('calc-scope-1')
    assert "e.key !== 'Enter'" in script
    assert 'preventDefault' in script


def test_marking_saved_clears_the_marks_without_rebuilding():
    script = dv.mark_saved_js('calc-scope-1')
    assert 'dw-dirty' in script
    assert 'innerHTML' not in script


def test_bold_and_italic_survive_into_the_view(letter):
    """A Word view that drops the emphasis is a text dump with headings."""
    read = dv.read_document(letter)
    html = dv.render_html(read, editable=True)
    assert 'font-weight:700' in html
    block = [b for b in read.blocks if b.address == 'p:1'][0]
    assert [r[0] for r in block.runs] == [
        'Sehr geehrte Damen, ', 'vielen Dank', ' für Ihre Anfrage.']
    assert block.runs[1][1] is True


def test_the_block_text_still_reads_as_the_paragraph(letter):
    """The spans must not change what innerText reports: the address an edit
    is written back to and the offsets a search jump counts are both taken
    from that text."""
    import html as _html
    import re

    read = dv.read_document(letter)
    markup = dv.render_html(read, editable=True)
    block = next(b for b in read.blocks if b.address == 'p:1')
    rendered = re.search(r'data-a="p:1"[^>]*>(.*?)</div>', markup, re.S).group(1)
    as_text = _html.unescape(re.sub(r'<[^>]+>', '', rendered))
    assert as_text == block.text


# ---------------------------------------------------------------------------
# Insertions: an empty range sits between runs, not inside one
# ---------------------------------------------------------------------------

def test_appending_at_the_end_of_a_paragraph_is_written(letter):
    """The commonest edit there is. A pure insertion is an empty range, and
    an empty range matches no run, so the splice wrote nothing and the save
    reported success over an unchanged file."""
    read = dv.read_document(letter)
    original = next(b for b in read.blocks if b.address == 'p:1').text

    result = dv.apply_edits(letter, {'p:1': original + ' NACHTRAG'})
    dv.save(result['document'], letter)

    assert docx.Document(str(letter)).paragraphs[1].text == original + ' NACHTRAG'


def test_appending_keeps_the_formatting_beside_it(letter):
    read = dv.read_document(letter)
    original = next(b for b in read.blocks if b.address == 'p:1').text
    dv.save(dv.apply_edits(letter, {'p:1': original + ' ANHANG'})['document'], letter)

    runs = [(r.text, r.bold) for r in docx.Document(str(letter)).paragraphs[1].runs
            if r.text]
    assert ('vielen Dank', True) in runs, 'the bold run was flattened'
    assert runs[-1][0].endswith(' ANHANG')


def test_inserting_at_the_very_start_is_written(letter):
    read = dv.read_document(letter)
    original = next(b for b in read.blocks if b.address == 'p:1').text
    dv.save(dv.apply_edits(letter, {'p:1': 'VORAB ' + original})['document'], letter)
    assert docx.Document(str(letter)).paragraphs[1].text == 'VORAB ' + original


def test_inserting_on_a_run_boundary_is_written(letter):
    """Between 'Sehr geehrte Damen, ' and the bold 'vielen Dank'."""
    read = dv.read_document(letter)
    original = next(b for b in read.blocks if b.address == 'p:1').text
    cut = original.index('vielen Dank')
    changed = original[:cut] + 'X' + original[cut:]
    dv.save(dv.apply_edits(letter, {'p:1': changed})['document'], letter)
    assert docx.Document(str(letter)).paragraphs[1].text == changed


def test_deleting_from_the_middle_still_works(letter):
    read = dv.read_document(letter)
    original = next(b for b in read.blocks if b.address == 'p:1').text
    changed = original.replace('vielen Dank', '')
    dv.save(dv.apply_edits(letter, {'p:1': changed})['document'], letter)
    assert docx.Document(str(letter)).paragraphs[1].text == changed


def test_typing_into_a_paragraph_whose_runs_are_empty(tmp_path):
    """Word leaves empty runs behind routinely -- a deleted line, a stray
    formatting mark. There is no character to splice against, and the
    insertion path reached for one: "Saving failed: string index out of
    range" on an ordinary blank line."""
    document = docx.Document()
    para = document.add_paragraph('')
    para.add_run('')          # a run that holds nothing
    para.add_run('')
    path = tmp_path / 'leer.docx'
    document.save(path)

    result = dv.apply_edits(path, {'p:0': 'jetzt steht hier etwas'})
    dv.save(result['document'], path)
    assert docx.Document(str(path)).paragraphs[0].text == 'jetzt steht hier etwas'


def test_clearing_a_paragraph_completely(tmp_path):
    document = docx.Document()
    document.add_paragraph('weg damit')
    path = tmp_path / 'weg.docx'
    document.save(path)

    dv.save(dv.apply_edits(path, {'p:0': ''})['document'], path)
    assert docx.Document(str(path)).paragraphs[0].text == ''


# ---------------------------------------------------------------------------
# Formatting
# ---------------------------------------------------------------------------

def test_the_controls_appear_only_while_editing(letter):
    read = dv.read_document(letter)
    assert 'dw-bar' not in dv.render_html(read, editable=False)
    assert 'dw-bar' in dv.render_html(read, editable=True)


def test_the_styles_offered_are_the_ones_a_letter_is_built_from():
    codes = [code for _label, code in dv.PARAGRAPH_STYLES]
    assert codes[:4] == ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']
    assert 'List Bullet' in codes and 'List Number' in codes


def test_a_style_this_view_does_not_set_is_refused():
    with pytest.raises(dv.DocxError):
        dv.check_style('Intense Quote')


def test_a_heading_is_made_by_naming_the_style(letter):
    """Not by turning the text big and bold: the style is what carries the
    look, and the table of contents reads styles."""
    result = dv.apply_edits(letter, {'p:1': {'style': 'Heading 2'}})
    dv.save(result['document'], letter)

    reopened = docx.Document(str(letter))
    assert reopened.paragraphs[1].style.name == 'Heading 2'
    assert reopened.paragraphs[1].text.startswith('Sehr geehrte'), (
        'the text was rewritten as well')


def test_a_paragraph_becomes_a_list_and_comes_back(letter):
    dv.save(dv.apply_edits(letter, {'p:1': {'style': 'List Bullet'}})['document'],
            letter)
    assert docx.Document(str(letter)).paragraphs[1].style.name == 'List Bullet'
    dv.save(dv.apply_edits(letter, {'p:1': {'style': 'Normal'}})['document'],
            letter)
    assert docx.Document(str(letter)).paragraphs[1].style.name == 'Normal'


def test_a_document_without_the_style_says_so(letter, monkeypatch):
    """A template that does not define one is not a broken document."""
    document = docx.Document(str(letter))

    class Missing:
        def __getitem__(self, name):
            raise KeyError(name)

    monkeypatch.setattr(type(document), 'styles', property(lambda self: Missing()))
    with pytest.raises(dv.DocxError) as excinfo:
        dv.set_paragraph_style(document, 'p:1', 'Heading 1')
    assert 'template' in str(excinfo.value)


def test_a_text_edit_still_keeps_every_run_as_it_was(letter):
    """Only the emphasis changing may redraw the runs; a typing edit must
    not, or every bold word in the paragraph would be flattened."""
    read = dv.read_document(letter)
    block = next(b for b in read.blocks if b.address == 'p:1')
    runs = [{'t': r[0], 'b': r[1], 'i': r[2], 'u': r[3]} for r in block.runs]
    runs[-1]['t'] = ' für Ihre Nachricht.'

    dv.save(dv.apply_edits(letter, {'p:1': {'runs': runs}})['document'], letter)
    after = [(r.text, r.bold) for r in docx.Document(str(letter)).paragraphs[1].runs
             if r.text]
    assert ('vielen Dank', True) in after
    assert after[-1][0] == ' für Ihre Nachricht.'


def test_changing_the_emphasis_redraws_the_runs(letter):
    read = dv.read_document(letter)
    block = next(b for b in read.blocks if b.address == 'p:1')
    runs = [{'t': r[0], 'b': r[1], 'i': r[2], 'u': r[3]} for r in block.runs]
    runs[0]['i'] = True

    dv.save(dv.apply_edits(letter, {'p:1': {'runs': runs}})['document'], letter)
    after = [(r.text, r.bold, r.italic)
             for r in docx.Document(str(letter)).paragraphs[1].runs if r.text]
    assert after[0][2] is True
    assert after[1][1] is True, 'the bold run was lost along the way'


def test_the_look_the_paragraph_was_written_in_is_carried_over(tmp_path):
    """Redrawing the runs cannot keep them, but it can keep the font."""
    document = docx.Document()
    para = document.add_paragraph()
    run = para.add_run('Angebot')
    run.font.name = 'Georgia'
    run.font.size = docx.shared.Pt(14)
    path = tmp_path / 'schrift.docx'
    document.save(path)

    dv.save(dv.apply_edits(path, {'p:0': {'runs': [
        {'t': 'Angebot', 'b': 1}]}})['document'], path)

    after = docx.Document(str(path)).paragraphs[0].runs[0]
    assert after.bold is True
    assert after.font.name == 'Georgia'
    assert after.font.size == docx.shared.Pt(14)


def test_only_an_emphasis_change_counts_as_one():
    before = [('Hallo ', False, False, False), ('Welt', True, False, False)]
    same_text = [{'t': 'Hallo ', 'b': 0}, {'t': 'Welt', 'b': 1}]
    assert dv.runs_differ(before, same_text) is False
    changed = [{'t': 'Hallo ', 'b': 1}, {'t': 'Welt', 'b': 1}]
    assert dv.runs_differ(before, changed) is True
    split = [{'t': 'Hallo', 'b': 0}, {'t': ' ', 'b': 0}, {'t': 'Welt', 'b': 1}]
    assert dv.runs_differ(before, split) is True


def test_the_browser_sends_the_runs_not_the_text():
    """An emphasis change has to arrive as one, or the paragraph would be
    written back with its old runs and the new words."""
    script = dv.edit_js('calc-scope-1')
    assert 'function runsOf(block)' in script
    assert 'runs: runsOf(block)' in script


def test_the_buttons_act_before_the_selection_is_gone():
    """By the time click fires, the selection in the paragraph is lost."""
    script = dv.edit_js('calc-scope-1')
    body = script[script.index("var press = {"):][:600]
    assert "addEventListener('mousedown'" in body
    assert 'preventDefault' in body


def test_the_shortcuts_are_the_ones_people_already_use():
    script = dv.edit_js('calc-scope-1')
    assert "key === 'b' ? 'bold'" in script
    assert "key === 'i' ? 'italic'" in script
    assert "key === 'u' ? 'underline'" in script


# ---------------------------------------------------------------------------
# Size and alignment
# ---------------------------------------------------------------------------

def test_a_run_carries_its_size(tmp_path):
    document = docx.Document()
    para = document.add_paragraph()
    para.add_run('klein').font.size = docx.shared.Pt(9)
    para.add_run(' normal')
    path = tmp_path / 'groessen.docx'
    document.save(path)

    block = dv.read_document(path).blocks[0]
    assert block.runs[0][4] == 9.0
    assert block.runs[1][4] is None


def test_a_size_shows_in_the_view(tmp_path):
    document = docx.Document()
    document.add_paragraph().add_run('gross').font.size = docx.shared.Pt(18)
    path = tmp_path / 'gross.docx'
    document.save(path)
    assert 'font-size:18pt' in dv.render_html(dv.read_document(path))


def test_a_size_is_written_back(letter):
    read = dv.read_document(letter)
    block = next(b for b in read.blocks if b.address == 'p:1')
    runs = [{'t': r[0], 'b': r[1], 'i': r[2], 'u': r[3], 's': r[4]}
            for r in block.runs]
    runs[0]['s'] = 14

    dv.save(dv.apply_edits(letter, {'p:1': {'runs': runs}})['document'], letter)
    after = docx.Document(str(letter)).paragraphs[1].runs
    assert after[0].font.size == docx.shared.Pt(14)


def test_a_size_change_counts_as_a_formatting_change():
    """Otherwise it would go through the splice, which writes text only,
    and the size would be dropped without a word."""
    before = [('Hallo', False, False, False, None)]
    assert dv.runs_differ(before, [{'t': 'Hallo', 's': 14}]) is True
    assert dv.runs_differ(before, [{'t': 'Hallo'}]) is False


def test_a_size_that_is_not_one_is_refused():
    assert dv.check_size('') is None
    assert dv.check_size(12) == 12.0
    for bad in ('riesig', 0, 500):
        with pytest.raises(dv.DocxError):
            dv.check_size(bad)


def test_a_paragraph_can_be_centred(letter):
    dv.save(dv.apply_edits(letter, {'p:1': {'align': 'center'}})['document'],
            letter)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    assert docx.Document(str(letter)).paragraphs[1].alignment == (
        WD_ALIGN_PARAGRAPH.CENTER)
    assert dv.read_document(letter).blocks[1].align == 'center'


def test_an_alignment_this_view_does_not_set_is_refused():
    with pytest.raises(dv.DocxError):
        dv.check_alignment('middle-ish')


def test_the_alignment_shows_in_the_view(letter):
    dv.save(dv.apply_edits(letter, {'p:1': {'align': 'right'}})['document'],
            letter)
    assert 'text-align:right' in dv.render_html(dv.read_document(letter))


def test_the_controls_are_all_on_the_bar(letter):
    markup = dv.render_html(dv.read_document(letter), editable=True)
    assert 'dw-size' in markup
    assert markup.count('dw-align') == len(dv.ALIGNMENTS)
    assert 'dw-style' in markup


def test_the_size_is_written_in_points_not_in_html_sizes():
    """execCommand only offers the seven HTML sizes, so the size that is
    asked for would not be the size that is written."""
    script = dv.edit_js('calc-scope-1')
    body = script[script.index('sizeBox.addEventListener'):][:800]
    assert "fontSize = points + 'pt'" in body
    assert 'execCommand' not in body


def test_the_letters_on_the_buttons_carry_no_markup_of_their_own():
    """A <b>, an <i> and a <u> inside three buttons have three different
    line boxes, and the three buttons stopped lining up."""
    markup = dv.toolbar_html()
    assert '<b>B</b>' not in markup and '<i>I</i>' not in markup
    assert '>B</button>' in markup and '>I</button>' in markup
