"""Word documents in the dashboard file browser: show, search, edit.

One renderer serves all three. That is the point of the module: the view
and the editor produce the same DOM, so what the search jumps to is what
the user sees, and what they type into is a block whose address in the
document is written on it.

The previous view converted the file with mammoth. Good HTML, but nothing
in it says which paragraph of the document a line came from, so it could
carry a view and nothing else. Here every block is rendered with its
address -- ``p:12`` for the twelfth body paragraph, ``t:1:3:2:0`` for the
first paragraph of a table cell -- and an edit is written straight back to
that paragraph.

Writing goes through :func:`delfin.agent.office._splice_runs`, the same
splice the template filler uses: the new text is placed into the runs that
are already there instead of rebuilding the paragraph, so bold, spacing and
font changes inside the line survive an edit that did not touch them.

What a round-trip preserves: python-docx writes back the package it read
and leaves untouched everything it was not asked to change -- headers and
footers, numbering, charts, text boxes, comments, tracked changes. That is
why editing text here is safe in a way that editing a workbook is not, and
why the warning a spreadsheet needs has no counterpart in this module.

Emphasis and paragraph style can be set. A heading is made by naming the
document's own style, never by turning text big and bold: the style is
what carries the look and what the table of contents reads.

Enter makes a paragraph, as it does in Word. That shifts the addresses every
later block is written back through, which is why it was withheld: the answer
is that a new paragraph is not written straight away but carried as an insert
against the block it follows, and the addresses are resolved once, in order, at
save time. The paragraph it starts takes the style Word would give it -- the
style's own "style for the following paragraph", so Enter after a heading
returns to body text.

What is still not offered: editing anything outside the body and the tables.
"""

from __future__ import annotations

import base64
import html as _html
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Mapping, Optional, Sequence, Tuple

WORD_SUFFIXES = ('.docx',)

# A document this far past the comm channel's comfort would arrive as a
# blank page after a long wait. Reported rather than truncated silently.
MAX_FILE_BYTES = 40 * 1024 * 1024
MAX_BLOCKS = 4000
MAX_IMAGE_BYTES = 4 * 1024 * 1024

# Paragraph styles that carry meaning worth keeping in the view. Anything
# else renders as body text -- the point is legibility, not a style engine.
_HEADING_RE = re.compile(r'^(?:heading|überschrift)\s*([1-6])$', re.IGNORECASE)
_LIST_RE = re.compile(r'(list|liste|aufz)', re.IGNORECASE)


class DocxError(RuntimeError):
    """A document could not be read or written, with a reason to show."""


@dataclass
class Block:
    """One editable paragraph, and where it lives in the document."""

    address: str
    text: str
    level: int = 0            # heading level, 0 for body text
    listed: bool = False
    table: Optional[Tuple[int, int, int]] = None   # table, row, cell
    # (text, bold, italic, underline) per run. What makes the view look like
    # the document rather than like a text dump of it; an edit still works on
    # the paragraph's plain text, and the splice on write puts the changed
    # span back into the run it belongs to.
    runs: List[Tuple[str, bool, bool, bool, Optional[float]]] = field(
        default_factory=list)
    align: str = ''

    @property
    def in_table(self) -> bool:
        return self.table is not None


@dataclass
class DocxDocument:
    """What the browser needs to show and search a document."""

    path: Path
    blocks: List[Block] = field(default_factory=list)
    images: Dict[str, str] = field(default_factory=dict)
    tables: int = 0
    notes: List[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        """The searchable buffer, matching the DOM's textContent exactly."""
        return '\n'.join(block.text for block in self.blocks)


def _docx():
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - depends on the install
        raise DocxError(
            'Word support is missing. Install it with: pip install python-docx'
        ) from exc
    return docx


# ---------------------------------------------------------------------------
# Reading
# ---------------------------------------------------------------------------

def _runs_of(paragraph):
    """The paragraph's runs with the emphasis, size and font each one carries."""
    out = []
    for run in getattr(paragraph, 'runs', []) or []:
        text = run.text or ''
        if not text:
            continue
        size = None
        try:
            if run.font.size is not None:
                size = float(run.font.size.pt)
        except Exception:
            size = None
        name = ''
        try:
            name = str(run.font.name or '')
        except Exception:
            name = ''
        out.append((text, bool(run.bold), bool(run.italic),
                    bool(run.underline), size, name))
    return out


def _alignment_of(paragraph) -> str:
    try:
        value = paragraph.alignment
    except Exception:
        return ''
    if value is None:
        return ''
    return {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}.get(
        int(value), '')


def _style_of(paragraph) -> Tuple[int, bool]:
    """Heading level and whether the paragraph is a list item."""
    name = ''
    try:
        name = str(paragraph.style.name or '')
    except Exception:
        pass
    match = _HEADING_RE.match(name.strip())
    if match:
        return int(match.group(1)), False
    if name.strip().lower() in ('title', 'titel'):
        return 1, False
    return 0, bool(_LIST_RE.search(name))


def _cell_paragraphs(document) -> List[Tuple[str, Any, Tuple[int, int, int]]]:
    out = []
    for t_index, table in enumerate(getattr(document, 'tables', []) or []):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for p_index, para in enumerate(cell.paragraphs):
                    address = f't:{t_index}:{r_index}:{c_index}:{p_index}'
                    out.append((address, para, (t_index, r_index, c_index)))
    return out


def read_document(path) -> DocxDocument:
    """Read a .docx into addressed blocks.

    Body paragraphs first, then table cells, which is the order the view
    renders them in. Every block carries the address an edit is written
    back to, so the two can never drift apart.
    """
    path = Path(path)
    if path.suffix.lower() not in WORD_SUFFIXES:
        raise DocxError(
            f'{path.name}: only .docx is supported. The older .doc format '
            'has to be converted first.')
    try:
        size = path.stat().st_size
    except OSError as exc:
        raise DocxError(f'File cannot be read: {exc}') from exc
    if size > MAX_FILE_BYTES:
        raise DocxError(
            f'Document is {size / (1024 * 1024):.0f} MB; the viewer is '
            f'limited to {MAX_FILE_BYTES // (1024 * 1024)} MB.')

    docx = _docx()
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocxError(f'Document could not be opened: {exc}') from exc

    result = DocxDocument(path=path)
    blocks: List[Block] = []
    for index, para in enumerate(document.paragraphs):
        level, listed = _style_of(para)
        blocks.append(Block(address=f'p:{index}', text=para.text,
                            level=level, listed=listed,
                            runs=_runs_of(para), align=_alignment_of(para)))
    for address, para, where in _cell_paragraphs(document):
        level, listed = _style_of(para)
        blocks.append(Block(address=address, text=para.text,
                            level=level, listed=listed, table=where,
                            runs=_runs_of(para), align=_alignment_of(para)))

    if len(blocks) > MAX_BLOCKS:
        result.notes.append(
            f'Document has {len(blocks)} paragraphs; the first {MAX_BLOCKS} '
            'are shown. Editing is therefore switched off here.')
        blocks = blocks[:MAX_BLOCKS]

    result.blocks = blocks
    result.tables = len(getattr(document, 'tables', []) or [])
    result.images = _inline_images(document)
    if not any(block.text.strip() for block in blocks):
        result.notes.append('The document holds no readable text.')
    return result


def _inline_images(document) -> Dict[str, str]:
    """Embedded images as data URIs, keyed by relationship id.

    Skipped rather than shrunk when oversized: the browser gets the page
    either way, and a hundred megabytes of base64 over the comm channel
    would cost the whole document.
    """
    images: Dict[str, str] = {}
    try:
        parts = document.part.related_parts
    except Exception:
        return images
    for rel_id, part in list(parts.items()):
        content_type = str(getattr(part, 'content_type', '') or '')
        if not content_type.startswith('image/'):
            continue
        try:
            blob = part.blob
        except Exception:
            continue
        if not blob or len(blob) > MAX_IMAGE_BYTES:
            continue
        encoded = base64.b64encode(blob).decode('ascii')
        images[str(rel_id)] = f'data:{content_type};base64,{encoded}'
    return images


def is_editable(document: DocxDocument) -> bool:
    """A truncated document is shown but not edited: an edit would address
    a paragraph by an index the view no longer agrees with."""
    return not any('Editing is therefore switched off' in note
                   for note in document.notes)


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _run_html(runs, fallback: str) -> str:
    """The paragraph's text, with the emphasis it carries.

    Wrapped in spans rather than <b>/<i>, so that what the block's
    innerText reports back is exactly the paragraph text -- the addresses
    and the offsets an edit and a search jump use are both counted in that
    text.
    """
    if not runs:
        return _html.escape(fallback) or '&nbsp;'
    out = []
    for run in runs:
        text, bold, italic, underline = run[:4]
        size = run[4] if len(run) > 4 else None
        styles = []
        if bold:
            styles.append('font-weight:700')
        if italic:
            styles.append('font-style:italic')
        if underline:
            styles.append('text-decoration:underline')
        if size:
            styles.append(f'font-size:{size:g}pt')
        escaped = _html.escape(text)
        out.append(f'<span style="{";".join(styles)}">{escaped}</span>'
                   if styles else escaped)
    return ''.join(out) or '&nbsp;'


def _block_html(block: Block, editable: bool) -> str:
    text = _run_html(block.runs, block.text)
    classes = ['dw-b']
    if block.level:
        classes.append(f'dw-h{block.level}')
    if block.listed:
        classes.append('dw-li')
    if block.in_table:
        classes.append('dw-tc')
    attrs = (
        f' class="{" ".join(classes)}"'
        f' data-a="{_html.escape(block.address, quote=True)}"'
    )
    if block.align:
        attrs += f' style="text-align:{block.align}"'
    if editable:
        attrs += ' contenteditable="true" spellcheck="false"'
    return f'<div{attrs}>{text}</div>'


DOC_CSS = (
    '.dw-page { font-family: Calibri, Carlito, Arial, sans-serif;'
    ' font-size:15px; line-height:1.55; color:#1a1a1a; background:#fff;'
    ' padding:30px 34px; margin:0 auto; max-width:900px;'
    ' box-shadow:0 1px 4px rgba(0,0,0,0.14); box-sizing:border-box; }'
    '.dw-b { margin:0 0 9px 0; min-height:1em; white-space:pre-wrap;'
    ' overflow-wrap:anywhere; outline:none; border-radius:3px; }'
    '.dw-b[contenteditable="true"]:hover { background:#f4f8ff; }'
    '.dw-b[contenteditable="true"]:focus { background:#eef4ff;'
    ' box-shadow:0 0 0 2px #90caf9; }'
    '.dw-b.dw-dirty { background:#fff6cc; }'
    '.dw-h1 { font-size:1.7em; font-weight:600; color:#12447a;'
    ' margin:18px 0 9px 0; }'
    '.dw-h2 { font-size:1.4em; font-weight:600; color:#12447a;'
    ' margin:16px 0 8px 0; }'
    '.dw-h3 { font-size:1.18em; font-weight:600; margin:14px 0 6px 0; }'
    '.dw-h4, .dw-h5, .dw-h6 { font-size:1.05em; font-weight:600;'
    ' margin:12px 0 5px 0; }'
    '.dw-li { padding-left:22px; position:relative; }'
    '.dw-li:before { content:"•"; position:absolute; left:8px; color:#666; }'
    '.dw-table { border-collapse:collapse; margin:12px 0; width:100%; }'
    '.dw-table td { border:1px solid #d6d6d6; padding:6px 8px;'
    ' vertical-align:top; }'
    '.dw-img { max-width:100%; height:auto; margin:10px 0;'
    ' border:1px solid #e0e0e0; border-radius:3px; }'
    '.dw-match { background:#fff59d; }'
    '.dw-match.dw-current { background:#ffb74d; }'
    '.dw-bar { position:sticky; top:0; z-index:6; display:flex; gap:6px;'
    ' align-items:center; padding:7px 12px; margin:0 0 14px 0;'
    ' background:linear-gradient(#fdfdfe,#f1f4f7);'
    ' border:1px solid #d6dbe0; border-radius:5px;'
    ' box-shadow:0 1px 3px rgba(0,0,0,0.07); box-sizing:border-box; }'
    '.dw-btn { font-size:14px; min-width:30px; height:28px; cursor:pointer;'
    ' border:1px solid #c8ced4; border-radius:4px; background:#fff;'
    ' color:#1f2937; line-height:1; padding:0;'
    ' display:inline-flex; align-items:center; justify-content:center; }'
    # 'dw-b' is the editable paragraph block, which carries a 9px bottom
    # margin -- so the Bold button wearing the same class sat 9px above its
    # neighbours.  The button is 'dw-bold'.
    '.dw-bold { font-weight:700; }'
    '.dw-zen { margin-left:auto; }'
    '.dw-i { font-style:italic; font-family:Georgia,serif; }'
    '.dw-u { text-decoration:underline; text-underline-offset:2px; }'
    '.dw-btn:hover { border-color:#1565c0; color:#12447a; }'
    '.dw-btn:active { background:#e8f0fe; }'
    '.dw-bar-sep { width:1px; height:20px; background:#d6dbe0; margin:0 4px; }'
    '.dw-style { font-size:13px; height:28px; padding:0 6px;'
    ' border:1px solid #c8ced4; border-radius:4px; background:#fff;'
    ' color:#1f2937; min-width:150px; }'
    '.dw-style:hover { border-color:#1565c0; }'
    '.dw-size { font-size:13px; height:28px; padding:0 4px;'
    ' border:1px solid #c8ced4; border-radius:4px; background:#fff;'
    ' color:#1f2937; min-width:62px; }'
    '.dw-size:hover { border-color:#1565c0; }'
    '.dw-align { font-size:15px; }'
    '.dw-al-right { transform:scaleX(-1); }'
)


# Drawn rather than written: four words in a row would be a sentence, and
# these are the marks Word uses for the same four things.
_ALIGN_MARK = {'left': '&#8801;', 'center': '&#8803;',
               'right': '&#8802;', 'justify': '&#9776;'}


def toolbar_html(current: str = 'Normal') -> str:
    """The formatting controls, shown while a document is being edited."""
    out = ['<div class="dw-bar">']
    # The letters are plain and the button carries the styling. A <b>, an
    # <i> and a <u> inside three buttons have three different line boxes,
    # so the three buttons stopped lining up with each other.
    out.append('<button class="dw-btn dw-bold" title="Bold (Ctrl+B)">B</button>')
    out.append('<button class="dw-btn dw-i" title="Italic (Ctrl+I)">I</button>')
    out.append('<button class="dw-btn dw-u" title="Underline (Ctrl+U)">U</button>')
    out.append('<select class="dw-font" title="Font">')
    out.append('<option value="">Font</option>')
    for name in FONT_NAMES:
        out.append(f'<option value="{_html.escape(name, quote=True)}">{_html.escape(name)}</option>')
    out.append('</select>')
    out.append('<select class="dw-size" title="Font size">')
    out.append('<option value="">Size</option>')
    for size in FONT_SIZES:
        out.append(f'<option value="{size}">{size}</option>')
    out.append('</select>')
    out.append('<span class="dw-bar-sep"></span>')
    for label, code in ALIGNMENTS:
        out.append(f'<button class="dw-btn dw-align dw-al-{code}"'
                   f' data-align="{code}" title="{_html.escape(label, quote=True)}">'
                   f'{_ALIGN_MARK[code]}</button>')
    out.append('<span class="dw-bar-sep"></span>')
    out.append('<select class="dw-style" title="Paragraph style">')
    for label, code in PARAGRAPH_STYLES:
        chosen = ' selected' if code == current else ''
        out.append(f'<option value="{_html.escape(code, quote=True)}"{chosen}>'
                   f'{_html.escape(label)}</option>')
    out.append('</select>')
    # Last of everything and pushed into the corner: a control that changes
    # the whole frame rather than the text belongs away from the ones that
    # act on what is selected.
    out.append('<button class="dw-btn dw-zen"'
               ' title="Fullscreen: the document alone (Esc leaves)">'
               '&#9974;</button>')
    out.append('</div>')
    return ''.join(out)


def render_html(document: DocxDocument, *, editable: bool = False) -> str:
    """The document as one block per paragraph, each carrying its address.

    Table cells are laid out as a table so the document still reads like
    one; the editable unit inside a cell is the same addressed block as
    everywhere else.
    """
    out: List[str] = []
    if editable:
        out.append(toolbar_html())
    out.append('<div class="dw-page">')
    for block in document.blocks:
        if block.in_table:
            continue
        out.append(_block_html(block, editable))

    by_table: Dict[int, Dict[Tuple[int, int], List[Block]]] = {}
    for block in document.blocks:
        if not block.in_table:
            continue
        t_index, r_index, c_index = block.table
        by_table.setdefault(t_index, {}).setdefault(
            (r_index, c_index), []).append(block)

    for t_index in sorted(by_table):
        cells = by_table[t_index]
        rows = sorted({r for r, _c in cells})
        out.append('<table class="dw-table"><tbody>')
        for r_index in rows:
            out.append('<tr>')
            for _r, c_index in sorted(k for k in cells if k[0] == r_index):
                out.append('<td>')
                for block in cells[(r_index, c_index)]:
                    out.append(_block_html(block, editable))
                out.append('</td>')
            out.append('</tr>')
        out.append('</tbody></table>')

    for uri in document.images.values():
        out.append(f'<img class="dw-img" src="{uri}" alt="">')

    out.append('</div>')
    return ''.join(out)


# ---------------------------------------------------------------------------
# Writing
# ---------------------------------------------------------------------------

# The paragraph styles offered, by the name shown in the menu. Small on
# purpose: these are the ones a letter or a report is actually built from,
# and a document's own styles are what carry its look and its table of
# contents -- so a heading is made by naming the style, never by making the
# text big and bold.
# Sizes in points, as Word lists them.
FONT_SIZES: Tuple[int, ...] = (8, 9, 10, 11, 12, 14, 16, 18, 20, 24, 28, 32)

# The typefaces Word offers by default, which are the ones a reader is likely
# to have.  A font a machine does not have is substituted by the reader's
# machine, so offering a long list would promise more than a document can keep.
FONT_NAMES: Tuple[str, ...] = (
    'Calibri', 'Cambria', 'Arial', 'Times New Roman', 'Helvetica',
    'Georgia', 'Verdana', 'Courier New', 'Consolas',
)

# Where the text sits in the line. Paragraph-level, like the style.
ALIGNMENTS: Tuple[Tuple[str, str], ...] = (
    ('Left', 'left'),
    ('Centre', 'center'),
    ('Right', 'right'),
    ('Justified', 'justify'),
)
_KNOWN_ALIGNMENTS = {code for _label, code in ALIGNMENTS}


PARAGRAPH_STYLES: Tuple[Tuple[str, str], ...] = (
    ('Body text', 'Normal'),
    ('Heading 1', 'Heading 1'),
    ('Heading 2', 'Heading 2'),
    ('Heading 3', 'Heading 3'),
    ('Bullet list', 'List Bullet'),
    ('Numbered list', 'List Number'),
)
_KNOWN_STYLES = {code for _label, code in PARAGRAPH_STYLES}


def check_style(name: Any) -> str:
    """A paragraph style this view offers."""
    text = str(name or '').strip()
    if text not in _KNOWN_STYLES:
        raise DocxError(f'{text!r} is not a paragraph style this view sets.')
    return text


def check_font(name: Any) -> str:
    """A typeface this view offers, or '' for "leave it as it was"."""
    text = str(name or '').strip()
    if not text:
        return ''
    for offered in FONT_NAMES:
        if offered.lower() == text.lower():
            return offered
    raise DocxError(f'{text!r} is not one of the fonts this view offers.')


def check_alignment(name: Any) -> str:
    text = str(name or '').strip().lower()
    if text not in _KNOWN_ALIGNMENTS:
        raise DocxError(f'{name!r} is not an alignment this view sets.')
    return text


def check_size(value: Any) -> Optional[float]:
    """A font size in points, or None for "as the style says"."""
    if value in (None, '', 'style'):
        return None
    try:
        size = float(value)
    except (TypeError, ValueError):
        raise DocxError(f'{value!r} is not a font size.') from None
    if not 1 <= size <= 409:        # what the format itself allows
        raise DocxError('A font size has to be between 1 and 409 points.')
    return size


def set_paragraph_alignment(document, address: str, alignment: str) -> None:
    """Put a paragraph left, centred, right or justified."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = _paragraph_at(document, address)
    if paragraph is None:
        raise DocxError(f'There is no paragraph {address!r} in the document.')
    paragraph.alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[check_alignment(alignment)]


def set_paragraph_style(document, address: str, style: str) -> None:
    """Give a paragraph one of the document's own styles.

    Named rather than imitated: a heading made by turning text big and bold
    is not a heading, and the table of contents will not have it.
    """
    paragraph = _paragraph_at(document, address)
    if paragraph is None:
        raise DocxError(f'There is no paragraph {address!r} in the document.')
    style = check_style(style)
    try:
        paragraph.style = document.styles[style]
    except KeyError as exc:
        raise DocxError(
            f'This document has no {style!r} style. It was made from a '
            'template that does not define one.') from exc


def _set_paragraph_runs(paragraph, runs: Sequence[Mapping[str, Any]]) -> None:
    """Rebuild a paragraph from runs that carry their own emphasis.

    Used only when the emphasis changed. A plain text edit goes through the
    splice instead, which keeps every run exactly as it was -- this path
    cannot, because the runs are being redrawn. What it can keep is the
    look the paragraph was written in, so the font, size and colour of the
    first run are carried onto all of them.
    """
    existing = list(paragraph.runs)
    base = existing[0] if existing else None
    for run in existing:
        run._element.getparent().remove(run._element)
    for spec in runs:
        text = str(spec.get('t') or '')
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bool(spec.get('b'))
        run.italic = bool(spec.get('i'))
        run.underline = bool(spec.get('u'))
        size = check_size(spec.get('s'))
        if base is not None:
            run.font.name = base.font.name
            run.font.size = base.font.size
            if base.font.color is not None and base.font.color.rgb is not None:
                run.font.color.rgb = base.font.color.rgb
        # After the base, not before: a size that was asked for has to win
        # over the one the paragraph was written in, or setting it would
        # look like it did nothing.
        if size is not None:
            from docx.shared import Pt

            run.font.size = Pt(size)
        font_name = check_font(spec.get('f'))
        if font_name:
            # Word stores the typeface three times over -- Latin, complex
            # script and East Asian -- and a reader that only finds one of
            # them falls back to the theme font, which is how a font change
            # ends up looking as though it did not happen.
            from docx.oxml.ns import qn

            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            fonts = rpr.get_or_add_rFonts()
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                fonts.set(qn(attr), font_name)
    if not paragraph.runs:
        paragraph.add_run('')


def runs_differ(before: Sequence[Sequence[Any]],
                after: Sequence[Mapping[str, Any]]) -> bool:
    """Whether the emphasis changed, as opposed to only the text.

    Deciding this is what lets an ordinary edit keep the paragraph's runs
    untouched and only a formatting change redraw them.
    """
    def _size(value):
        try:
            return round(float(value), 1) if value else None
        except (TypeError, ValueError):
            return None

    was = [(bool(run[1]), bool(run[2]), bool(run[3]),
            _size(run[4] if len(run) > 4 else None),
            str(run[5] if len(run) > 5 else '') or '') for run in before]
    now = [(bool(r.get('b')), bool(r.get('i')), bool(r.get('u')),
            _size(r.get('s')), str(r.get('f') or '')) for r in after
           if str(r.get('t') or '')]
    if len(was) != len(now):
        return True
    return was != now


def _paragraph_at(document, address: str):
    """Resolve a block address back to the paragraph it came from."""
    parts = str(address).split(':')
    try:
        if parts[0] == 'p' and len(parts) == 2:
            return document.paragraphs[int(parts[1])]
        if parts[0] == 't' and len(parts) == 5:
            _kind, t_index, r_index, c_index, p_index = parts
            table = document.tables[int(t_index)]
            cell = table.rows[int(r_index)].cells[int(c_index)]
            return cell.paragraphs[int(p_index)]
    except (IndexError, ValueError):
        return None
    return None


def changed_range(before: str, after: str) -> Tuple[int, int, int]:
    """The one span that differs, as ``(start, end_before, end_after)``.

    Splicing the whole paragraph would put all of its text into the first
    run and empty the others, which flattens every bold word and font
    change in the line -- including the ones the edit never touched. Only
    the span that actually changed is written, so the rest keeps the runs
    it already had.
    """
    start = 0
    limit = min(len(before), len(after))
    while start < limit and before[start] == after[start]:
        start += 1
    end_before, end_after = len(before), len(after)
    while (end_before > start and end_after > start
           and before[end_before - 1] == after[end_after - 1]):
        end_before -= 1
        end_after -= 1
    return start, end_before, end_after


def _set_paragraph_text(paragraph, text: str) -> None:
    """Put ``text`` into a paragraph without rebuilding it.

    A paragraph with no runs at all (an empty line) has nothing to splice
    into, so one run is added.
    """
    from delfin.agent.office import _splice_runs

    runs = list(paragraph.runs)
    if not runs:
        if text:
            paragraph.add_run(text)
        return
    current = ''.join(run.text for run in runs)
    if current == text:
        return
    if not current:
        # Runs that are all empty: there is no character to splice against,
        # and the insertion branch below would reach for current[0]. Word
        # leaves such runs behind routinely -- a deleted line, a stray
        # formatting mark -- so this is the ordinary case of typing into a
        # blank paragraph, not a broken document.
        runs[0].text = text
        return
    start, end_before, end_after = changed_range(current, text)
    if start == end_before and end_after > start:
        # A pure insertion is an empty range, and an empty range sits between
        # two runs rather than inside one: the splice skips every run and
        # writes nothing. Typing at the end of a paragraph is exactly that
        # case, so it silently saved the paragraph unchanged.
        #
        # Widen it onto one neighbouring character and carry that character
        # through. The text then lands in the run that character belongs to,
        # which is also where Word would put it.
        inserted = text[start:end_after]
        if start > 0:
            _splice_runs(runs, start - 1, start, current[start - 1] + inserted)
        else:
            _splice_runs(runs, 0, 1, inserted + current[0])
        return
    _splice_runs(runs, start, end_before, text[start:end_after])


def apply_edits(path, edits: Dict[str, str]) -> Dict[str, Any]:
    """Write block edits back into the document.

    Returns what was written and what was refused. An address that does
    not resolve is an error rather than a silent skip: it means the view
    and the file have drifted apart, and the user is owed that news
    before they close the tab believing the edit landed.
    """
    path = Path(path)
    if not isinstance(edits, dict) or not edits:
        raise DocxError('There is nothing to save.')

    docx = _docx()
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocxError(f'Document could not be opened: {exc}') from exc

    written = 0
    unknown: List[str] = []
    for address, change in edits.items():
        paragraph = _paragraph_at(document, str(address))
        if paragraph is None:
            unknown.append(str(address))
            continue
        if isinstance(change, Mapping):
            style = change.get('style')
            if style:
                set_paragraph_style(document, str(address), style)
            if change.get('align'):
                set_paragraph_alignment(document, str(address), change['align'])
            runs = change.get('runs')
            if runs is not None:
                before = _runs_of(paragraph)
                if runs_differ(before, runs):
                    # The emphasis changed, so the runs have to be redrawn.
                    _set_paragraph_runs(paragraph, runs)
                else:
                    # Only the text did: the splice keeps every run as it is.
                    _set_paragraph_text(
                        paragraph, ''.join(str(r.get('t') or '') for r in runs))
            elif 'text' in change:
                _set_paragraph_text(paragraph, str(change.get('text') or ''))
        else:
            _set_paragraph_text(paragraph, '' if change is None else str(change))
        written += 1

    if unknown:
        raise DocxError(
            'These paragraphs are not in the document (any more): '
            + ', '.join(unknown[:5])
            + '. The file was left unchanged.')

    return {'document': document, 'written': written}


def save(document, path) -> None:
    """Write a python-docx document to ``path``."""
    try:
        document.save(str(path))
    except Exception as exc:
        raise DocxError(f'Saving failed: {exc}') from exc


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------

@dataclass
class DocxHit:
    """One match, addressed so the view can jump to it."""

    block: int
    address: str
    start: int
    end: int


def search(document: DocxDocument, term: str, *, limit: int = 2000
           ) -> List[DocxHit]:
    """Every occurrence of ``term``, case-insensitively, block by block.

    Offsets are into the block's own text, not into a flattened document,
    so a jump lands on a character rather than near one.
    """
    term = str(term or '')
    if not term:
        return []
    needle = term.lower()
    hits: List[DocxHit] = []
    for index, block in enumerate(document.blocks):
        haystack = block.text.lower()
        start = haystack.find(needle)
        while start >= 0:
            hits.append(DocxHit(block=index, address=block.address,
                                start=start, end=start + len(needle)))
            if len(hits) >= limit:
                return hits
            start = haystack.find(needle, start + 1)
    return hits


def search_matches(
    document: DocxDocument,
    term: str,
    *,
    match_case: bool = False,
    whole_word: bool = False,
    limit: int = 2000,
) -> List[DocxHit]:
    """Every occurrence of *term*, with Word's two search switches.

    ``search`` above is the tab's own always-insensitive find; this is what
    Find and Replace works from, where the two switches change which matches
    exist and therefore which of them get replaced.
    """
    term = str(term or '')
    if not term:
        return []
    flags = 0 if match_case else re.IGNORECASE
    pattern = re.escape(term)
    if whole_word:
        # \b does not fire between two non-word characters, so a term that
        # starts or ends with punctuation would never match with the switch on.
        prefix = r'\b' if term[:1].isalnum() or term[:1] == '_' else ''
        suffix = r'\b' if term[-1:].isalnum() or term[-1:] == '_' else ''
        pattern = prefix + pattern + suffix
    finder = re.compile(pattern, flags)
    hits: List[DocxHit] = []
    for index, block in enumerate(document.blocks):
        for found in finder.finditer(block.text):
            hits.append(DocxHit(block=index, address=block.address,
                                start=found.start(), end=found.end()))
            if len(hits) >= limit:
                return hits
    return hits


def replace_all(
    path,
    term: str,
    replacement: str,
    *,
    match_case: bool = False,
    whole_word: bool = False,
) -> Dict[str, Any]:
    """Replace every occurrence, and say how many and where.

    Word replaces the *characters*, not the paragraph: the formatting on either
    side of a replaced word stays as it was.  That is why this goes through the
    same text splice an ordinary edit uses instead of rewriting the paragraph --
    replacing a word in a bold sentence must not unbolden the sentence.

    Returns ``{'document', 'replaced', 'paragraphs'}``; nothing is written to
    disk here, so the caller decides when to save.
    """
    term = str(term or '')
    if not term:
        raise DocxError('There is nothing to search for.')
    replacement = '' if replacement is None else str(replacement)

    path = Path(path)
    docx = _docx()
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocxError(f'Document could not be opened: {exc}') from exc

    flags = 0 if match_case else re.IGNORECASE
    pattern = re.escape(term)
    if whole_word:
        prefix = r'\b' if term[:1].isalnum() or term[:1] == '_' else ''
        suffix = r'\b' if term[-1:].isalnum() or term[-1:] == '_' else ''
        pattern = prefix + pattern + suffix
    finder = re.compile(pattern, flags)

    replaced = 0
    touched = 0
    for paragraph in _all_paragraphs(document):
        before = paragraph.text
        after, count = finder.subn(replacement, before)
        if not count:
            continue
        _set_paragraph_text(paragraph, after)
        replaced += count
        touched += 1
    return {'document': document, 'replaced': replaced, 'paragraphs': touched}


def _following_style(anchor):
    """What Word makes the next paragraph when Enter is pressed.

    A style carries the style for the paragraph that follows it, which is why
    Enter after a heading returns to body text instead of writing a second
    heading.  Where a template does not say, the paragraph keeps its own style
    -- which is right for body text and for a list, where Enter continues it.
    """
    try:
        following = anchor.style.next_paragraph_style
    except Exception:
        following = None
    return following if following is not None else getattr(anchor, 'style', None)


def open_document(path):
    """The document as python-docx sees it, with our error on failure."""
    docx = _docx()
    try:
        return docx.Document(str(Path(path)))
    except Exception as exc:
        raise DocxError(f'Document could not be opened: {exc}') from exc


def address_order(address: str) -> Tuple[int, int, int, int, int]:
    """Sort key for a block address, so deletions can run bottom-up."""
    parts = str(address).split(':')
    try:
        if parts[0] == 'p' and len(parts) == 2:
            return (0, int(parts[1]), 0, 0, 0)
        if parts[0] == 't' and len(parts) == 5:
            return (1, int(parts[1]), int(parts[2]), int(parts[3]), int(parts[4]))
    except ValueError:
        pass
    return (2, 0, 0, 0, 0)


def insert_paragraph_after(document, anchor, text: str = '', *, style=None):
    """Put a paragraph after *anchor* (a paragraph object), and return it."""
    fresh = anchor.insert_paragraph_before(text or '')
    anchor._p.addnext(fresh._p)
    chosen = style
    if chosen is None:
        chosen = _following_style(anchor)
    if isinstance(chosen, str):
        try:
            chosen = document.styles[check_style(chosen)]
        except KeyError:
            chosen = None
    if chosen is not None:
        try:
            fresh.style = chosen
        except Exception:
            pass                   # a template without it keeps what it has
    return fresh


def insert_paragraph(document, address: str, text: str = '',
                     *, style=None, before: bool = False):
    """Add a paragraph next to the one at *address*, and return it.

    Word puts a new paragraph *after* the one the cursor is in when Enter is
    pressed, which is what this does unless told otherwise.
    """
    anchor = _paragraph_at(document, str(address))
    if anchor is None:
        raise DocxError(f'There is no paragraph {address!r} in the document.')
    if not before:
        return insert_paragraph_after(document, anchor, text, style=style)
    fresh = anchor.insert_paragraph_before(text or '')
    chosen = style if style is not None else getattr(anchor, 'style', None)
    if isinstance(chosen, str):
        try:
            chosen = document.styles[check_style(chosen)]
        except KeyError:
            chosen = None
    if chosen is not None:
        try:
            fresh.style = chosen
        except Exception:
            pass
    return fresh


def delete_paragraph(document, address: str) -> None:
    """Take a paragraph out of the document, as Word does on a whole-line delete."""
    paragraph = _paragraph_at(document, str(address))
    if paragraph is None:
        raise DocxError(f'There is no paragraph {address!r} in the document.')
    element = paragraph._p
    parent = element.getparent()
    if parent is None:
        raise DocxError('That paragraph is not part of the document any more.')
    parent.remove(element)


def add_table(document, rows: int, cols: int, *, style: str = 'Table Grid',
              after: str = ''):
    """Put a table into the document, with lines drawn as Word's default does.

    A table without a style has no borders, which looks like a broken table
    rather than a plain one, so 'Table Grid' stands in where the template has
    it.
    """
    rows, cols = int(rows), int(cols)
    if rows < 1 or cols < 1:
        raise DocxError('A table needs at least one row and one column.')
    if rows * cols > 4000:
        raise DocxError('That table is too large to add in one step.')
    table = document.add_table(rows=rows, cols=cols)
    try:
        table.style = document.styles[style]
    except KeyError:
        pass                       # the template does not define it; leave plain
    if after:
        anchor = _paragraph_at(document, str(after))
        if anchor is not None:
            anchor._p.addnext(table._tbl)
    return table


def add_table_row(document, table_index: int):
    """Append a row, as Tab at the end of a table does in Word."""
    tables = list(document.tables)
    index = int(table_index)
    if not (0 <= index < len(tables)):
        raise DocxError(f'There is no table {index + 1} in the document.')
    return tables[index].add_row()


def drop_table_row(document, table_index: int, row_index: int) -> None:
    tables = list(document.tables)
    index = int(table_index)
    if not (0 <= index < len(tables)):
        raise DocxError(f'There is no table {index + 1} in the document.')
    table = tables[index]
    rows = list(table.rows)
    if not (0 <= int(row_index) < len(rows)):
        raise DocxError('That row is not in the table.')
    if len(rows) == 1:
        raise DocxError('A table cannot lose its last row.')
    element = rows[int(row_index)]._tr
    element.getparent().remove(element)


def add_table_column(document, table_index: int, width_emu: Optional[int] = None):
    tables = list(document.tables)
    index = int(table_index)
    if not (0 <= index < len(tables)):
        raise DocxError(f'There is no table {index + 1} in the document.')
    table = tables[index]
    reference = table.columns[-1].width if table.columns else None
    return table.add_column(width_emu or reference or 914400)


def drop_table_column(document, table_index: int, col_index: int) -> None:
    """Remove a column: in the file that is one cell out of every row."""
    tables = list(document.tables)
    index = int(table_index)
    if not (0 <= index < len(tables)):
        raise DocxError(f'There is no table {index + 1} in the document.')
    table = tables[index]
    column = int(col_index)
    if not (0 <= column < len(table.columns)):
        raise DocxError('That column is not in the table.')
    if len(table.columns) == 1:
        raise DocxError('A table cannot lose its last column.')
    for row in table.rows:
        cells = list(row.cells)
        if column < len(cells):
            element = cells[column]._tc
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
    grid = table._tbl.find(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid')
    if grid is not None:
        columns = list(grid)
        if column < len(columns):
            grid.remove(columns[column])


def _all_paragraphs(document):
    """Every paragraph, the body ones and the ones inside tables."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def focus_js(scope_class: str, address: str, start: int, end: int) -> str:
    """Mark one hit and scroll it into view, without touching the rest.

    The match is wrapped with a Range rather than by rewriting innerHTML:
    a rewrite would flatten the formatting of the paragraph it landed in,
    and in a Word document that is the document.
    """
    return _FOCUS_JS_TEMPLATE.replace(
        '__SCOPE__', json.dumps(str(scope_class))
    ).replace(
        '__ADDR__', json.dumps(str(address))
    ).replace('__START__', str(int(start))).replace('__END__', str(int(end)))


def edit_js(scope_class: str) -> str:
    """Report a block to the kernel when the user leaves it.

    Reported on leaving rather than on every keystroke: a paragraph is a
    unit of text, and sending one message per character would put the
    kernel behind the typing.
    """
    return _EDIT_JS_TEMPLATE.replace('__SCOPE__', json.dumps(str(scope_class)))


def mark_saved_js(scope_class: str) -> str:
    """Clear the changed marks, without rebuilding anything.

    Rebuilding would take the caret out of the paragraph the user is
    still standing in, which is not what saving does.
    """
    return (
        "(function(){var r=document.querySelector('.' + "
        + json.dumps(str(scope_class))
        + ");if(!r)return;Array.prototype.forEach.call("
        "r.querySelectorAll('.dw-b.dw-dirty'),function(b){"
        "b.classList.remove('dw-dirty');});})();"
    )


_EDIT_JS_TEMPLATE = r"""
(function(){
  var SCOPE = __SCOPE__;
  var tries = 0;

  function boot(){
    var root = document.querySelector('.' + SCOPE);
    var page = root && root.querySelector('.dw-page');
    if (!page) { if (++tries < 40) setTimeout(boot, 50); return; }
    if (page.dataset.editBound === '1') return;
    page.dataset.editBound = '1';

    function send(payload){
      var field = root.querySelector('.calc-sheet-payload textarea')
               || root.querySelector('.calc-sheet-payload input');
      if (!field) return;
      var proto = field.tagName === 'TEXTAREA'
        ? HTMLTextAreaElement.prototype : HTMLInputElement.prototype;
      var desc = Object.getOwnPropertyDescriptor(proto, 'value');
      if (desc && desc.set) desc.set.call(field, JSON.stringify(payload));
      else field.value = JSON.stringify(payload);
      field.dispatchEvent(new Event('input', {bubbles: true}));
      field.dispatchEvent(new Event('change', {bubbles: true}));
      var holder = root.querySelector('.calc-sheet-action');
      var btn = holder && (holder.tagName === 'BUTTON'
        ? holder : holder.querySelector('button'));
      if (btn) btn.click();
    }

    /* The runs of a block, read back off the page. What is sent is the
       structure, not the text: an emphasis change has to arrive as one, or
       the paragraph would be written back with its old runs and the new
       words. */
    function runsOf(block){
      var runs = [];
      (function walk(node, bold, italic, under, size, face){
        for (var i = 0; i < node.childNodes.length; i++) {
          var child = node.childNodes[i];
          if (child.nodeType === 3) {
            if (child.nodeValue) {
              runs.push({t: child.nodeValue, b: bold ? 1 : 0,
                         i: italic ? 1 : 0, u: under ? 1 : 0,
                         s: size || 0, f: face || ''});
            }
            continue;
          }
          if (child.nodeType !== 1) continue;
          var name = child.tagName;
          var style = child.style || {};
          var weight = style.fontWeight || '';
          /* A size set on this element wins over the one inherited: it
             is what a font-size control on the selection produces. */
          var own = (style.fontSize || '').trim();
          var pt = size;
          if (own) {
            var px = parseFloat(own);
            if (!isNaN(px)) {
              pt = own.indexOf('pt') >= 0 ? px : Math.round(px * 0.75 * 10) / 10;
            }
          }
          var ownFace = (style.fontFamily || '').split(',')[0]
                          .replace(/["']/g, '').trim();
          walk(child,
               bold || name === 'B' || name === 'STRONG'
                 || weight === 'bold' || weight === '700' || weight === '600',
               italic || name === 'I' || name === 'EM'
                 || style.fontStyle === 'italic',
               under || name === 'U'
                 || (style.textDecoration || '').indexOf('underline') >= 0,
               pt, ownFace || face);
        }
      })(block, false, false, false, 0, '');
      return runs;
    }

    function report(block, force){
      if (!block || block.getAttribute('contenteditable') !== 'true') return;
      if (!force && block.innerText === block.dataset.was) return;
      block.classList.add('dw-dirty');
      block.dataset.was = block.innerText;   /* sent; only resend on change */
      send({kind: 'docx', address: block.getAttribute('data-a'),
            runs: runsOf(block)});
    }

    /* The block being worked on, whether or not it still has focus: a
       toolbar button takes focus away the moment it is pressed. */
    var active = null;
    function currentBlock(){
      var node = document.activeElement;
      var block = node && node.closest ? node.closest('.dw-b') : null;
      return block || active;
    }

    function emphasise(command){
      var block = currentBlock();
      if (!block) return;
      block.focus();
      document.execCommand(command, false, null);
      report(block, true);
    }

    var bar = root.querySelector('.dw-bar');
    if (bar) {
      var press = {'.dw-bold': 'bold', '.dw-i': 'italic', '.dw-u': 'underline'};
      Object.keys(press).forEach(function(sel){
        var button = bar.querySelector(sel);
        if (!button) return;
        /* mousedown, not click: by the time click fires, the selection in
           the paragraph is already gone. */
        button.addEventListener('mousedown', function(e){
          e.preventDefault();
          emphasise(press[sel]);
        });
      });
      /* execCommand has no font-size in points, only the seven HTML sizes.
         The selection is wrapped by hand so the size that is asked for is
         the size that is written. */
      var sizeBox = bar.querySelector('.dw-size');
      if (sizeBox) sizeBox.addEventListener('change', function(){
        var block = currentBlock();
        var points = parseFloat(sizeBox.value);
        if (!block || !points) { sizeBox.value = ''; return; }
        block.focus();
        var range = window.getSelection();
        if (range && range.rangeCount && !range.isCollapsed) {
          var span = document.createElement('span');
          span.style.fontSize = points + 'pt';
          try {
            range.getRangeAt(0).surroundContents(span);
          } catch (_err) {
            /* A selection across several runs cannot be wrapped in one
               element; the whole paragraph is the honest fallback. */
            block.style.fontSize = points + 'pt';
          }
        } else {
          block.style.fontSize = points + 'pt';
        }
        report(block, true);
        sizeBox.value = '';
      });

      /* Same shape as the size control: wrap what is selected, or the whole
         paragraph when nothing is. */
      var fontBox = bar.querySelector('.dw-font');
      if (fontBox) fontBox.addEventListener('change', function(){
        var block = currentBlock();
        var face = fontBox.value;
        if (!block || !face) { fontBox.value = ''; return; }
        block.focus();
        var picked = window.getSelection();
        if (picked && picked.rangeCount && !picked.isCollapsed) {
          var span = document.createElement('span');
          span.style.fontFamily = face;
          try {
            picked.getRangeAt(0).surroundContents(span);
          } catch (_err) {
            block.style.fontFamily = face;
          }
        } else {
          block.style.fontFamily = face;
        }
        report(block, true);
        fontBox.value = '';
      });

      Array.prototype.forEach.call(bar.querySelectorAll('.dw-align'),
        function(button){
          button.addEventListener('mousedown', function(e){
            e.preventDefault();
            var block = currentBlock();
            if (!block) return;
            var where = button.getAttribute('data-align');
            block.style.textAlign = where;
            block.classList.add('dw-dirty');
            send({kind: 'docx', address: block.getAttribute('data-a'),
                  align: where});
          });
        });

      /* Fullscreen: the kernel owns the state, so the tab's own button and
         this one cannot disagree about it. */
      var zenBtn = bar.querySelector('.dw-zen');
      if (zenBtn) zenBtn.addEventListener('click', function(){
        send({kind: 'docx', zen: 1});
      });
      document.addEventListener('keydown', function(e){
        if (e.key !== 'Escape') return;
        var scope = page.closest('.calc-tab');
        if (!scope || !scope.classList.contains('calc-zen-doc')) return;
        e.preventDefault();
        send({kind: 'docx', zen: 1});
      }, true);

      var styleBox = bar.querySelector('.dw-style');
      if (styleBox) styleBox.addEventListener('change', function(){
        var block = currentBlock();
        if (!block) return;
        block.classList.add('dw-dirty');
        send({kind: 'docx', address: block.getAttribute('data-a'),
              style: styleBox.value});
        styleBox.blur();
      });
    }

    page.addEventListener('keydown', function(e){
      if (!(e.ctrlKey || e.metaKey)) return;
      var key = String(e.key || '').toLowerCase();
      var command = key === 'b' ? 'bold'
                  : key === 'i' ? 'italic'
                  : key === 'u' ? 'underline' : '';
      if (!command) return;
      e.preventDefault();
      emphasise(command);
    });

    page.addEventListener('focusin', function(e){
      var block = e.target.closest && e.target.closest('.dw-b');
      if (block) { block.dataset.was = block.innerText; active = block; }
    }, true);

    page.addEventListener('focusout', function(e){
      report(e.target.closest && e.target.closest('.dw-b'));
    }, true);

    /* Also while typing, once it pauses. Leaving the block is the precise
       moment, but it is not a reliable one: pressing Save is a click on a
       button, and whether that moves focus out of the paragraph first is up
       to the browser. When it does not, the edit was never sent and saving
       reported nothing to save while the text sat on screen. */
    var typing = null;
    page.addEventListener('input', function(e){
      var block = e.target.closest && e.target.closest('.dw-b');
      if (!block) return;
      if (typing) clearTimeout(typing);
      typing = setTimeout(function(){ typing = null; report(block); }, 350);
    });

    /* ---------- Enter makes a paragraph ----------
       The block under the caret is split where the caret stands: what is in
       front stays, what is behind moves into a new block after it, and the
       caret goes with it.  The new block carries a temporary address; the
       kernel resolves it against the block it follows, once, at save time. */
    var minted = 0;

    function caretTo(block, atEnd){
      block.focus();
      var range = document.createRange();
      range.selectNodeContents(block);
      range.collapse(!atEnd);
      var sel = window.getSelection();
      sel.removeAllRanges();
      sel.addRange(range);
    }

    page.addEventListener('keydown', function(e){
      if (e.key !== 'Enter' || e.shiftKey) return;
      var block = e.target.closest && e.target.closest('.dw-b');
      if (!block) return;
      e.preventDefault();

      /* Everything from the caret to the end of the paragraph moves out. */
      var tail = '';
      var sel = window.getSelection();
      if (sel && sel.rangeCount) {
        var caret = sel.getRangeAt(0);
        caret.deleteContents();
        var rest = document.createRange();
        rest.selectNodeContents(block);
        rest.setStart(caret.endContainer, caret.endOffset);
        var holder = document.createElement('div');
        holder.appendChild(rest.extractContents());
        tail = holder.innerText || '';
      }
      report(block, true);

      minted += 1;
      var fresh = document.createElement('div');
      fresh.className = block.className.replace(/\bdw-dirty\b/g, '').trim();
      fresh.setAttribute('contenteditable', 'true');
      fresh.setAttribute('spellcheck', 'false');
      fresh.setAttribute('data-a', 'new:' + minted);
      fresh.textContent = tail;
      block.parentNode.insertBefore(fresh, block.nextSibling);
      fresh.dataset.was = fresh.innerText;
      fresh.classList.add('dw-dirty');
      send({kind: 'docx', address: 'new:' + minted,
            after: block.getAttribute('data-a'),
            runs: [{t: tail, b: 0, i: 0, u: 0, s: 0, f: ''}]});
      active = fresh;
      caretTo(fresh, false);
    });

    /* Backspace at the very start folds the paragraph into the one above,
       which is the other half of Enter: without it a paragraph could be made
       but never unmade. */
    page.addEventListener('keydown', function(e){
      if (e.key !== 'Backspace') return;
      var block = e.target.closest && e.target.closest('.dw-b');
      if (!block) return;
      var sel = window.getSelection();
      if (!sel || !sel.rangeCount || !sel.isCollapsed) return;
      var at = sel.getRangeAt(0);
      var before = document.createRange();
      before.selectNodeContents(block);
      before.setEnd(at.startContainer, at.startOffset);
      if (before.toString().length) return;       /* not at the start */
      var previous = block.previousElementSibling;
      if (!previous || !previous.classList.contains('dw-b')) return;
      e.preventDefault();
      var joined = (previous.innerText || '') + (block.innerText || '');
      var wasLength = (previous.innerText || '').length;
      previous.textContent = joined;
      block.parentNode.removeChild(block);
      previous.classList.add('dw-dirty');
      report(previous, true);
      send({kind: 'docx', address: block.getAttribute('data-a'), drop: 1});
      active = previous;
      caretTo(previous, true);
      if (wasLength === 0) caretTo(previous, false);
    });
  }

  boot();
})();
"""


_FOCUS_JS_TEMPLATE = r"""
(function(){
  var root = document.querySelector('.' + __SCOPE__);
  if (!root) return;
  var page = root.querySelector('.dw-page');
  if (!page) return;

  var old = page.querySelector('.dw-match.dw-current');
  if (old && old.parentNode) {
    var parent = old.parentNode;
    parent.replaceChild(document.createTextNode(old.textContent || ''), old);
    parent.normalize();
  }

  var block = page.querySelector('[data-a="' + __ADDR__ + '"]');
  if (!block) return;

  /* Walk to the character offset: a paragraph is several text nodes once
     it has any formatting, so an offset into its text is not an offset
     into its first node. */
  function locate(node, index){
    var walker = document.createTreeWalker(node, NodeFilter.SHOW_TEXT, null);
    var seen = 0, current = walker.nextNode();
    while (current) {
      var length = (current.nodeValue || '').length;
      if (seen + length >= index) return {node: current, offset: index - seen};
      seen += length;
      current = walker.nextNode();
    }
    return null;
  }

  var from = locate(block, __START__);
  var to = locate(block, __END__);
  if (!from || !to) return;

  var range = document.createRange();
  try {
    range.setStart(from.node, from.offset);
    range.setEnd(to.node, to.offset);
    var mark = document.createElement('mark');
    mark.className = 'dw-match dw-current';
    range.surroundContents(mark);
    mark.scrollIntoView({block: 'center'});
  } catch (_err) {
    /* A match spanning several runs cannot be wrapped in one element;
       showing the paragraph is still better than not moving at all. */
    block.scrollIntoView({block: 'center'});
  }
})();
"""
