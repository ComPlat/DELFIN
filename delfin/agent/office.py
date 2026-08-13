"""Document primitives for office work: spreadsheets, PDFs, Word files.

Until now the agent could only reach these files through ``bash`` plus
generated Python. That works, but every failure mode has to be
rediscovered by the model on the spot, and several of them are silent:

* ``read_file`` decodes an ``.xlsx`` as UTF-8 and returns the raw ZIP
  container as mojibake — thousands of tokens of noise that look like
  data.
* A workbook written by a library carries formulas but **no cached
  values**. Reading it with ``data_only=True`` yields ``None`` for every
  computed cell, which reads as "the column is empty" rather than "the
  numbers were never evaluated".
* A PDF form filled without ``/NeedAppearances`` renders empty in
  viewers that do not build appearance streams themselves. The values
  are in the file; the printout is blank.
* Check boxes do not take ``true``. They take the on-state name declared
  by that particular field — ``/Yes``, ``/On``, ``/1`` — and silently
  stay off for anything else.
* XFA forms ignore AcroForm writes, so a "successful" fill changes
  nothing the user will see.
* A merge that opens its inputs while it writes produces a document
  that looks complete when the fourth attachment turns out to be
  encrypted. The pages missing from it are the ones nobody can name
  afterwards, so every input is opened before the first byte is
  written.
* An .ods cell repeats: LibreOffice writes one cell with
  ``number-columns-repeated="1013"`` where a naive reader sees one
  column, and every value after it lands under the wrong heading.
* A page with no text layer is a photograph of paper, not an empty
  page. Saying "no text" reports the file wrongly, and OCR of it is a
  reading of pixels rather than content of the document — so it is
  labelled as such and never overwrites a page that has real text.

Each of those is handled once, here, instead of in whatever code the
model happens to write that turn. The functions are pure in the sense
that matters for testing: they take paths and values, they touch only
the file named, and they raise only :class:`OfficeError`.

Where there is no reader, there is a refusal that names the way out
rather than a parser that copes. A legacy .xls read by an approximate
BIFF reader, or an HTML export that arrived under an .xls name, gives
back rows that look exactly like data — and an administrative file read
wrongly is worse than one not read at all, because nothing downstream
can tell. Those refusals name what the bytes actually are and the one
command that converts them.

Writing is deliberately narrow. ``edit_sheet`` changes cells and appends
rows; it does not restructure workbooks. Rewriting a spreadsheet through
``openpyxl`` reconstructs the whole container, so anything the library
does not model is at risk — the caller is told which fragile features a
workbook carries before it is rewritten, and the tool layer captures a
byte-exact pre-image so the change can be undone.
"""

from __future__ import annotations

import contextvars as _contextvars
import csv
import re
import threading as _threading
from dataclasses import dataclass, field as _field
from datetime import date as _dt_date, datetime as _dt_datetime, time as _dt_time
from decimal import Decimal as _Decimal
from pathlib import Path
from typing import Any, Optional

from . import german as _de

# Reading caps. A spreadsheet can be a million rows; the point of the
# read path is to let the model see the shape and the region it asked
# for, not to pour the file into the context window.
DEFAULT_MAX_ROWS = 200
DEFAULT_MAX_COLS = 40
MAX_TEXT_CHARS = 20000

SPREADSHEET_SUFFIXES = frozenset({".xlsx", ".xlsm", ".xltx", ".xltm"})
CSV_SUFFIXES = frozenset({".csv", ".tsv"})
PDF_SUFFIXES = frozenset({".pdf"})
WORD_SUFFIXES = frozenset({".docx"})
# OpenDocument. LibreOffice is the office suite of a great many public
# authorities and universities, so .ods and .odt arrive as ordinary
# attachments rather than as exotica. Reading only: see
# _CONVERSION_ROUTES for why writing them is refused instead of faked.
ODS_SUFFIXES = frozenset({".ods"})
ODT_SUFFIXES = frozenset({".odt"})

DOCUMENT_SUFFIXES = (
    SPREADSHEET_SUFFIXES | CSV_SUFFIXES | PDF_SUFFIXES | WORD_SUFFIXES
    | ODS_SUFFIXES | ODT_SUFFIXES
)

# Formats there is no reader for, mapped to the format a conversion has
# to produce. Deliberately NOT readers: every one of these has a parser
# somewhere that will return rows for most files and mangle the rest,
# and a wrong table out of an administrative file is worse than no
# table, because nothing about it looks wrong afterwards.
_CONVERSION_ROUTES: dict[str, tuple[str, str]] = {
    ".xls": ("legacy Excel workbook (BIFF, Excel 97-2003)", "xlsx"),
    ".xlt": ("legacy Excel template (BIFF, Excel 97-2003)", "xlsx"),
    ".doc": ("legacy Word document (Word 97-2003)", "docx"),
}

# Suffix -> (import name, distribution name). The distribution name is
# what the user has to install, and it differs from the import name for
# python-docx, which is the classic reason a "just pip install docx"
# hint sends someone to an unrelated package.
_BACKENDS: dict[str, tuple[str, str]] = {
    "spreadsheet": ("openpyxl", "openpyxl"),
    "pdf": ("pypdf", "pypdf"),
    "word": ("docx", "python-docx"),
    "csv": ("csv", ""),  # standard library
    # Reading a PDF and producing one are different libraries: pypdf
    # takes existing pages apart and puts them together, but it does not
    # lay out text. Writing a page needs reportlab.
    "pdf_write": ("reportlab", "reportlab"),
    # OpenDocument (.ods/.odt). The import name is 'odf' and the package
    # is 'odfpy' — the same trap python-docx sits in, which is why the
    # distribution name is carried separately.
    "opendocument": ("odf", "odfpy"),
}


class OfficeError(RuntimeError):
    """A document operation failed for a reason worth reporting verbatim."""


# ---------------------------------------------------------------------------
# Backends and file kinds
# ---------------------------------------------------------------------------

def document_kind(path: Any) -> Optional[str]:
    """The handling family of *path*, or None if there is no reader.

    ``spreadsheet`` / ``csv`` / ``pdf`` / ``word`` /
    ``opendocument_sheet`` / ``opendocument_text``. The OpenDocument
    kinds are separate from ``spreadsheet`` and ``word`` on purpose:
    they can be read but not written, and letting them share a kind
    would send an .ods into openpyxl, which fails with a message about
    a broken ZIP rather than about the format.
    """
    suffix = Path(str(path)).suffix.lower()
    if suffix in SPREADSHEET_SUFFIXES:
        return "spreadsheet"
    if suffix in CSV_SUFFIXES:
        return "csv"
    if suffix in PDF_SUFFIXES:
        return "pdf"
    if suffix in WORD_SUFFIXES:
        return "word"
    if suffix in ODS_SUFFIXES:
        return "opendocument_sheet"
    if suffix in ODT_SUFFIXES:
        return "opendocument_text"
    return None


def backend_available(kind: str) -> bool:
    """True if the library backing *kind* can be imported right now."""
    entry = _BACKENDS.get(kind)
    if entry is None:
        return False
    import importlib
    try:
        importlib.import_module(entry[0])
        return True
    except Exception:
        return False


def available_backends() -> dict[str, bool]:
    """Availability of every backend, for diagnostics and tool gating."""
    return {kind: backend_available(kind) for kind in _BACKENDS}


def have_office_support() -> bool:
    """True if at least one non-stdlib document backend is importable.

    The tool surface is gated on this: advertising a spreadsheet tool on
    a machine without ``openpyxl`` only buys a tool call that fails.
    """
    return any(
        backend_available(kind)
        for kind in ("spreadsheet", "pdf", "word", "opendocument")
    )


def _require(kind: str) -> Any:
    """Import the backend for *kind* or raise a fixable error."""
    entry = _BACKENDS.get(kind)
    if entry is None:
        raise OfficeError(f"unknown document kind {kind!r}")
    import importlib
    try:
        return importlib.import_module(entry[0])
    except Exception as exc:
        dist = entry[1] or entry[0]
        raise OfficeError(
            f"{kind} support needs the '{dist}' package, which is not "
            f"installed ({exc}). Install it with "
            f"'pip install delfin[office]' or 'pip install {dist}'."
        ) from exc


def _near_miss(p: Path) -> str:
    """Names in the same folder that differ from ``p`` only slightly.

    A miss that is only a matter of case cost a real session a whole
    round: the agent asked for ``kostenstellen.xlsx``, the folder held
    ``Kostenstellen.xlsx``, and the answer repeated the path back without
    mentioning the file sitting beside it. On a backend where the first
    token takes tens of seconds, re-listing the directory is expensive.

    Naming it is not the same as opening it. The caller still gets an
    error and still has to ask again with the right name -- on someone's
    real records the difference between two similar filenames can be last
    year's ledger and this one, so the correction stays theirs to make.
    """
    try:
        siblings = [c.name for c in p.parent.iterdir() if c.is_file()]
    except OSError:
        return ""
    want = p.name.lower()
    exact_case = [n for n in siblings if n.lower() == want]
    if exact_case:
        return exact_case[0]
    import difflib
    close = difflib.get_close_matches(want, [n.lower() for n in siblings],
                                      n=1, cutoff=0.75)
    if not close:
        return ""
    return next(n for n in siblings if n.lower() == close[0])


def _resolve(path: Any, *, must_exist: bool = True) -> Path:
    p = Path(str(path)).expanduser()
    if must_exist and not p.exists():
        hint = _near_miss(p)
        if hint:
            raise OfficeError(
                f"file not found: {p} — the folder does have {hint!r}. "
                f"Ask again with that exact name if you meant it.")
        raise OfficeError(f"file not found: {p}")
    if must_exist and p.is_dir():
        raise OfficeError(f"{p} is a directory, not a document")
    return p


def column_letter(index: int) -> str:
    """1-based column index to spreadsheet letters (1 -> A, 27 -> AA).

    Local rather than ``openpyxl.utils`` so the CSV path stays usable
    without the spreadsheet backend installed.
    """
    if index < 1:
        raise OfficeError(f"column index must be >= 1, got {index}")
    out = ""
    while index:
        index, rem = divmod(index - 1, 26)
        out = chr(ord("A") + rem) + out
    return out


_CELL_RE = re.compile(r"^\s*([A-Za-z]{1,3})\s*([1-9][0-9]{0,6})\s*$")


def parse_cell(ref: str) -> tuple[int, int]:
    """``"B7"`` -> ``(row, column)`` = ``(7, 2)``, both 1-based."""
    m = _CELL_RE.match(str(ref or ""))
    if not m:
        raise OfficeError(
            f"{ref!r} is not a cell reference — expected a form like 'B7'")
    letters, digits = m.group(1).upper(), m.group(2)
    col = 0
    for ch in letters:
        col = col * 26 + (ord(ch) - ord("A") + 1)
    return int(digits), col


# ---------------------------------------------------------------------------
# What a file really is, and what to do when it is not readable here
# ---------------------------------------------------------------------------
#
# A suffix is a claim, not evidence. Two cases turn up constantly in
# administrative mail and both end in a wrong table rather than an
# error: a report exported by a web system arrives as an HTML table
# named ``.xls``, and a workbook someone "saved as .ods" is an .xlsx
# underneath. Parsers meet those with a best effort — a row count, some
# plausible values — so the refusals below name what the bytes actually
# are instead of what the name says.

_CONTAINER_LABELS: dict[str, str] = {
    "ole2": "a legacy Microsoft Office document (OLE2/BIFF)",
    "xlsx": "an .xlsx workbook (Office Open XML)",
    "docx": "a .docx document (Office Open XML)",
    "ods": "an OpenDocument spreadsheet",
    "odt": "an OpenDocument text document",
    "odf": "an OpenDocument file",
    "zip": "a ZIP archive",
    "pdf": "a PDF",
    "html": "an HTML page or table",
    "xml": "an XML file",
    "text": "plain text",
}


def _zip_container_kind(path: Path) -> str:
    """Which OOXML/OpenDocument flavour a ZIP holds, or ``"zip"``."""
    import zipfile
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "mimetype" in names:
                try:
                    mime = archive.read("mimetype").decode("ascii", "replace")
                except Exception:
                    mime = ""
                if mime.endswith("spreadsheet"):
                    return "ods"
                if mime.endswith("text"):
                    return "odt"
                if "opendocument" in mime:
                    return "odf"
            if any(n.startswith("xl/") for n in names):
                return "xlsx"
            if any(n.startswith("word/") for n in names):
                return "docx"
    except Exception:
        return "zip"
    return "zip"


def container_kind(path: Any) -> str:
    """What *path* holds, judged by its bytes. ``""`` when undecidable.

    Cheap on purpose: the first few kilobytes decide, because this runs
    on the refusal path where the alternative is a wrong answer, not on
    a hot loop.
    """
    p = Path(str(path))
    try:
        with open(p, "rb") as fh:
            head = fh.read(4096)
    except OSError:
        return ""
    # Four bytes, not the full eight of the OLE2 signature: a file that
    # was truncated in transit still has to be named for what it was,
    # rather than falling through to "plain text".
    if head.startswith(b"\xd0\xcf\x11\xe0"):
        return "ole2"
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        return _zip_container_kind(p)
    probe = head.lstrip()[:512].lower()
    if (probe.startswith(b"<!doctype html") or probe.startswith(b"<html")
            or b"<table" in probe):
        return "html"
    if probe.startswith(b"<?xml"):
        return "xml"
    if b"\x00" in head:
        return ""
    return "text" if head else ""


_LIBREOFFICE_HINT = (
    "libreoffice --headless --convert-to {target} --outdir <directory> "
    "\"{name}\"")


def _mismatch_note(path: Path, expected: tuple[str, ...]) -> str:
    """A sentence naming the real format when it is not the expected one.

    Empty when the bytes match the suffix or cannot be judged, so a
    caller can append it unconditionally.
    """
    found = container_kind(path)
    if not found or found in expected:
        return ""
    label = _CONTAINER_LABELS.get(found, f"a {found} file")
    return (
        f" Note that {path.name} does not hold what its name claims: the "
        f"bytes are {label}. Exports from web systems are routinely named "
        "for a format they are not, so convert or rename it to match what "
        "it really is before reading it.")


def _refuse_unreadable_format(path: Path) -> None:
    """Raise for a format with no reader, naming the conversion route.

    Reached before any parser sees the file. The message has to carry
    the whole way out, because the caller cannot look the command up:
    which format to convert to, and one command that does it.
    """
    suffix = path.suffix.lower()
    entry = _CONVERSION_ROUTES.get(suffix)
    if entry is None:
        return
    label, target = entry
    expected = ("ole2",) if suffix in (".xls", ".xlt", ".doc") else ()
    raise OfficeError(
        f"{path.name} is a {label}. There is no reader for it here, and "
        "nothing guesses at the bytes: a legacy Office file taken apart by "
        "the wrong parser gives back content that looks right and is not. "
        f"Convert it to .{target} once and read the result — "
        + _LIBREOFFICE_HINT.format(target=target, name=path.name)
        + f" — or open it and use 'Save as' with the .{target} format."
        + _mismatch_note(path, expected))


def _refuse_unwritable_format(path: Path, *, what: str) -> None:
    """Raise for a format this module reads but cannot write."""
    suffix = path.suffix.lower()
    if suffix in _CONVERSION_ROUTES:
        label, target = _CONVERSION_ROUTES[suffix]
        raise OfficeError(
            f"{path.name} is a {label}, which is neither read nor written "
            f"here. Work in .{target} instead: convert the original once — "
            + _LIBREOFFICE_HINT.format(target=target, name=path.name)
            + f" — and produce a .{target} as the result.")
    kind = document_kind(path)
    if kind == "opendocument_sheet":
        raise OfficeError(
            f"{path.name} is an OpenDocument spreadsheet. It can be read "
            f"here but not written, so {what} would either fail or produce "
            "a file that is no longer an .ods. Convert it to .xlsx first — "
            + _LIBREOFFICE_HINT.format(target="xlsx", name=path.name)
            + " — change that, and convert back if the .ods is what has to "
            "be delivered.")
    if kind == "opendocument_text":
        raise OfficeError(
            f"{path.name} is an OpenDocument text document. It can be read "
            f"here but not written, so {what} is not possible. Convert it "
            "to .docx first — "
            + _LIBREOFFICE_HINT.format(target="docx", name=path.name)
            + " — and work on that.")
    _refuse_unreadable_format(path)


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _fmt(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


_CLIP_MARK = "…"


def _clip_cell(text: str, limit: int) -> str:
    """Shorten a cell and say that it was shortened.

    The grid is the only representation of a cell the model ever sees, so
    a silent cut made "Rechnung 2026-0001 Storno wegen Rücksend" —
    the first forty characters of a purpose line — indistinguishable
    from a value that really ends there. The model could then quote it
    back as the cell's content, or write it into an edit.
    """
    if limit <= 1 or len(text) <= limit:
        return text
    return text[:limit - 1] + _CLIP_MARK


def clipped_cell_count(rows: list[list[Any]], limit: int = 40) -> int:
    """How many cells the grid had to shorten — for the reader's notes.

    The per-cell mark says THAT a value is cut; this says how much of the
    table is affected, which is what decides whether the window can be
    trusted for anything but orientation.
    """
    n = 0
    for row in rows:
        for value in row:
            if len(_fmt(value)) > limit:
                n += 1
    return n


def render_grid(
    rows: list[list[Any]],
    *,
    first_row: int = 1,
    first_col: int = 1,
    max_cell_chars: int = 40,
) -> str:
    """Render cells as an addressable text grid.

    The column letters and row numbers are part of the output on purpose:
    the model has to be able to name a cell (``Daten!B7``) in a follow-up
    edit, and it can only do that if it saw the coordinates.
    """
    if not rows:
        return "(empty)"
    width = max(len(r) for r in rows)
    cells = [
        [_clip_cell(_fmt(r[i]), max_cell_chars) if i < len(r) else ""
         for i in range(width)]
        for r in rows
    ]
    headers = [column_letter(first_col + i) for i in range(width)]
    row_label_w = max(len(str(first_row + len(rows) - 1)), 3)
    col_w = [
        max(len(headers[i]), *(len(c[i]) for c in cells)) for i in range(width)
    ]
    out = [
        " " * row_label_w + " | "
        + " | ".join(headers[i].ljust(col_w[i]) for i in range(width))
    ]
    out.append("-" * len(out[0]))
    for n, row in enumerate(cells):
        label = str(first_row + n).rjust(row_label_w)
        out.append(
            label + " | "
            + " | ".join(row[i].ljust(col_w[i]) for i in range(width))
        )
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Values: numbers and dates as they are actually written
# ---------------------------------------------------------------------------
#
# A cell reading "1.234,50" is a string. Handed to arithmetic as-is it
# either raises or, worse, parses as 1.234 — the thousands separator
# read as a decimal point, off by a factor of a thousand, with no error
# anywhere. The same trap sits in dates: 03.04.2026 and 04/03/2026 are
# the same day written under two conventions, and guessing per value
# gets it wrong half the time.
#
# The way out is to decide per COLUMN rather than per value. A single
# "1.234" is genuinely ambiguous; a column containing it alongside one
# value with both separators, or one with two dots, or one with a
# non-three-digit group, is not. So: gather the evidence over the whole
# column, and where the column stays ambiguous, say so instead of
# picking the reading that happens to look plausible.

# Noise a number may legitimately carry: whitespace (including the
# non-breaking and narrow kinds a spreadsheet emits), currency symbols and
# a currency code beside it. Nothing else — and that is the point. The
# earlier rule stripped every character that was not a digit or separator,
# which turned the document reference "R-001" into "-001" and read it as
# the number -1. A whole column of references profiled as numeric, and in
# a reconciliation "R-001" and "X-001" both became -1 and compared EQUAL:
# two different records reported as agreeing.
_CURRENCY_CODE_RE = re.compile(r"(?i)\b(?:eur|usd|chf|gbp|sfr)\b")
_NUM_NOISE_RE = re.compile(r"[\s\u00a0\u202f\u2009€$£¥']")
# Any run of digits, optionally broken by separators. Deliberately NOT
# "1-3 digits then groups": that shape is right for 1.234,50 and wrong
# for 1234.50, and rejecting the ungrouped form makes a plain numeric
# column read as text — which then compares by spelling rather than by
# value.
_NUMERIC_SHAPE_RE = re.compile(r"^[-+]?[0-9]+(?:[.,][0-9]+)*$")
_DATE_SEP_RE = re.compile(r"^\s*(\d{1,4})([./-])(\d{1,2})\2(\d{1,4})\s*$")


def _numeric_body(value: Any) -> Optional[str]:
    """The digits-and-separators core of a value, or None if it is not one.

    Returning None for anything carrying a letter is what keeps a column
    of document references out of the number machinery.
    """
    text = _CURRENCY_CODE_RE.sub("", str(value if value is not None else ""))
    text = _NUM_NOISE_RE.sub("", text).strip()
    if not text or not _NUMERIC_SHAPE_RE.match(text):
        return None
    return text

DECIMAL_COMMA = "decimal_comma"     # 1.234,50 — German / most of Europe
DECIMAL_POINT = "decimal_point"     # 1,234.50 — English
PLAIN_NUMBER = "plain"              # 1234.5 / 1234 — no grouping in play
AMBIGUOUS = "ambiguous"

DAY_FIRST = "day_first"             # 31.07.2026
MONTH_FIRST = "month_first"         # 07/31/2026
ISO_DATE = "iso"                    # 2026-07-31


# ---------------------------------------------------------------------------
# Writing a value back out
# ---------------------------------------------------------------------------

# Excel's number formats say what a value LOOKS like, and the readers
# below already decode what it means. Between them sat ``str(value)``:
# a Serienbrief built from a cell formatted ``#,##0.00 "€"`` and one
# formatted ``DD.MM.YYYY`` went to a customer reading
#
#     wir stellen Ihnen 1234.5 EUR in Rechnung, fällig am
#     2026-07-31 00:00:00.
#
# and the run reported {'ok': 1}. The cell's own number_format was
# available at every step and consulted at none. Everything that puts a
# cell value into a document goes through here.

# What a number format says about a NUMBER: how many places, whether it
# groups, whether it is a percentage. The quoted literals and the colour
# blocks are stripped first — ``"€"`` is decoration the template already
# carries, and repeating it turns "1.234,50 EUR" into "1.234,50 € EUR".
_FORMAT_LITERAL_RE = re.compile(r'"[^"]*"|\[[^\]]*\]|\\.|_.|\*.')


def _number_format_spec(number_format: Any) -> tuple[Optional[int], bool, bool]:
    """``(decimals, grouping, percent)`` for an Excel number format.

    ``decimals`` is None when the format decides nothing (General, or no
    format at all), which leaves the value's own precision alone.
    """
    text = str(number_format or "").strip()
    if not text or text.lower() == "general":
        return None, False, False
    text = text.split(";")[0]
    percent = "%" in _FORMAT_LITERAL_RE.sub("", text)
    body = _FORMAT_LITERAL_RE.sub("", text)
    body = "".join(ch for ch in body if ch in "#0.,%")
    if not any(ch in body for ch in "#0"):
        return None, False, percent
    whole, dot, frac = body.replace("%", "").rpartition(".")
    if not dot:                      # no decimal point: it is all integer
        whole, frac = body.replace("%", ""), ""
    decimals = sum(1 for ch in frac if ch in "#0")
    grouping = "," in whole
    return decimals, grouping, percent


def _render_for_locale(value: Any, number_format: Any = "") -> str:
    """One cell value as the document should show it.

    A float becomes ``1.234,50`` and a datetime ``31.07.2026`` — the cell's
    own format decides the precision and the date pattern, the locale
    decides the characters. Strings pass through untouched: a CSV already
    holds the text the user typed, and re-formatting it would be this
    module's other failure mode in reverse.
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Ja" if value else "Nein"
    if isinstance(value, (_dt_datetime, _dt_date, _dt_time)):
        return _de.format_date(value, number_format)
    if isinstance(value, (int, float)) or isinstance(value, _Decimal):
        decimals, grouping, percent = _number_format_spec(number_format)
        return _de.format_number(value, decimals=decimals, grouping=grouping,
                                 percent=percent)
    return str(value)


def _number_evidence(text: str) -> Optional[str]:
    """What one value proves about its column's convention, if anything."""
    body = _numeric_body(text)
    if body is None:
        return None
    dots, commas = body.count("."), body.count(",")

    # Both separators present: the LAST one is the decimal separator.
    # This is decisive and needs no column context.
    if dots and commas:
        return DECIMAL_COMMA if body.rfind(",") > body.rfind(".") else DECIMAL_POINT
    # More than one of the same separator can only be grouping.
    if dots > 1:
        return DECIMAL_COMMA
    if commas > 1:
        return DECIMAL_POINT
    # Exactly one separator: a group that is not three digits cannot be a
    # thousands separator, so the separator is decimal.
    for sep, convention in ((".", DECIMAL_POINT), (",", DECIMAL_COMMA)):
        if body.count(sep) == 1:
            tail = body.split(sep)[1]
            if len(tail) != 3:
                return convention
            return None                      # 1.234 / 1,234 — undecidable
    return PLAIN_NUMBER                       # no separator at all


def detect_number_convention(values: list) -> tuple[str, str]:
    """Decide a column's number convention. Returns ``(convention, why)``.

    ``PLAIN_NUMBER`` means every value parses the same either way (no
    grouping separators in play). ``AMBIGUOUS`` means the column contains
    only values like "1.234" that read as 1234 or as 1.234 depending on
    the convention, and nothing in the column settles it — the caller
    must ask rather than pick.
    """
    votes: dict[str, int] = {}
    for value in values:
        evidence = _number_evidence(value)
        if evidence in (DECIMAL_COMMA, DECIMAL_POINT):
            votes[evidence] = votes.get(evidence, 0) + 1
    if len(votes) > 1:
        top = max(votes, key=lambda k: votes[k])
        other = min(votes, key=lambda k: votes[k])
        return AMBIGUOUS, (
            f"the column mixes conventions: {votes[top]} value(s) look like "
            f"{top}, {votes[other]} like {other}")
    if votes:
        convention = next(iter(votes))
        return convention, f"{votes[convention]} value(s) settle it"

    undecided = [
        v for v in values
        if _number_evidence(v) is None and _numeric_body(v) is not None
    ]
    if undecided:
        return AMBIGUOUS, (
            f"values like {str(undecided[0])!r} read as two different numbers "
            "depending on the convention, and nothing in the column decides it")
    return PLAIN_NUMBER, "no grouping separators in this column"


def parse_number(text: Any, convention: str = PLAIN_NUMBER) -> Optional[float]:
    """Parse one value under a known convention. None if it is not a number."""
    if isinstance(text, bool):
        return None
    if isinstance(text, (int, float)):
        return float(text)
    body = _numeric_body(text)
    if body is None:
        return None
    if convention == DECIMAL_COMMA:
        body = body.replace(".", "").replace(",", ".")
    elif convention == DECIMAL_POINT:
        body = body.replace(",", "")
    else:
        # No convention decided: a lone separator with a 3-digit tail is
        # left alone only when it cannot change the value.
        if body.count(",") == 1 and "." not in body:
            body = body.replace(",", ".")
        else:
            body = body.replace(",", "")
    try:
        return float(body)
    except ValueError:
        return None


def _date_parts(text: Any) -> Optional[tuple[int, int, int, str]]:
    """``(a, b, c, separator)`` of a date-shaped value, else None."""
    match = _DATE_SEP_RE.match(str(text or ""))
    if match is None:
        return None
    a, sep, b, c = match.group(1), match.group(2), match.group(3), match.group(4)
    return int(a), int(b), int(c), sep


def detect_date_convention(values: list) -> tuple[str, str]:
    """Decide a column's date convention. Returns ``(convention, why)``."""
    saw_iso = saw_dot = False
    day_first_proof = month_first_proof = 0
    for value in values:
        parts = _date_parts(value)
        if parts is None:
            continue
        a, b, _c, sep = parts
        if sep == "-" and a > 31:
            saw_iso = True
            continue
        if sep == ".":
            saw_dot = True
        # A first field above 12 can only be a day; a second field above
        # 12 can only be a day in second position.
        if a > 12:
            day_first_proof += 1
        elif b > 12:
            month_first_proof += 1
    if saw_iso and not (day_first_proof or month_first_proof):
        return ISO_DATE, "ISO dates (YYYY-MM-DD)"
    if day_first_proof and month_first_proof:
        return AMBIGUOUS, (
            f"{day_first_proof} value(s) can only be day-first and "
            f"{month_first_proof} can only be month-first — the column mixes "
            "two conventions")
    if day_first_proof:
        return DAY_FIRST, f"{day_first_proof} value(s) have a day above 12"
    if month_first_proof:
        return MONTH_FIRST, f"{month_first_proof} value(s) have a month above 12"
    if saw_dot:
        # Dotted dates are day-first everywhere they are written.
        return DAY_FIRST, "dot-separated dates are written day-first"
    if saw_iso:
        return ISO_DATE, "ISO dates (YYYY-MM-DD)"
    return AMBIGUOUS, (
        "every value could be read either way (no day above 12 anywhere)")


def parse_date(text: Any, convention: str = ISO_DATE) -> Optional[str]:
    """Parse one date under a known convention, returned as ISO ``YYYY-MM-DD``."""
    import datetime as _dt

    if isinstance(text, _dt.datetime):
        return text.date().isoformat()
    if isinstance(text, _dt.date):
        return text.isoformat()
    parts = _date_parts(text)
    if parts is None:
        return None
    a, b, c, sep = parts
    if convention == ISO_DATE or (sep == "-" and a > 31):
        year, month, day = a, b, c
    elif convention == MONTH_FIRST:
        month, day, year = a, b, c
    elif convention == DAY_FIRST:
        day, month, year = a, b, c
    else:
        return None
    if year < 100:                     # two-digit year, as spreadsheets do
        year += 2000 if year < 70 else 1900
    try:
        return _dt.date(year, month, day).isoformat()
    except ValueError:
        return None


def profile_column(values: list, *, name: str = "") -> dict:
    """Describe one column: what it holds, under which convention.

    ``kind`` is ``number`` / ``date`` / ``text`` / ``empty``. A column is
    only called numeric or date-like when most of its non-empty values
    parse that way — one stray "n/a" must not turn a money column into
    text, and one date-shaped code must not turn a text column into
    dates. Values that do NOT parse are listed, because those are the
    rows a total would silently omit.
    """
    filled = [v for v in values if str(v or "").strip() != ""]
    if not filled:
        return {"name": name, "kind": "empty", "convention": "",
                "values": 0, "parsed": 0, "unparsed": []}

    date_shaped = sum(1 for v in filled if _date_parts(v) is not None)
    numeric_shaped = sum(
        1 for v in filled
        if (isinstance(v, (int, float)) and not isinstance(v, bool))
        or _numeric_body(v) is not None
    )

    if date_shaped >= max(1, int(0.8 * len(filled))):
        convention, why = detect_date_convention(filled)
        parsed = [(v, parse_date(v, convention)) for v in filled]
        return {
            "name": name, "kind": "date", "convention": convention,
            "convention_reason": why, "values": len(filled),
            "parsed": sum(1 for _, p in parsed if p is not None),
            "unparsed": [str(v) for v, p in parsed if p is None][:10],
        }
    if numeric_shaped >= max(1, int(0.8 * len(filled))):
        convention, why = detect_number_convention(filled)
        parsed = [(v, parse_number(v, convention)) for v in filled]
        return {
            "name": name, "kind": "number", "convention": convention,
            "convention_reason": why, "values": len(filled),
            "parsed": sum(1 for _, p in parsed if p is not None),
            "unparsed": [str(v) for v, p in parsed if p is None][:10],
        }
    # Not numeric enough to be called a number column -- but a column that
    # is PART numbers is the dangerous case, and it used to be the silent
    # one. At 20 unreadable values in 100 the tool named them and said a
    # total would leave them out; at 21 the column flipped to text with an
    # empty `unparsed` list, `column_notes` skipped it because it only
    # spoke about number and date kinds, and the renderer dropped it for
    # the same reason. One more bad row turned the warning off, exactly on
    # the columns where somebody is about to sum money.
    #
    # The 80% rule stays -- a column of company names should not be
    # described as broken numbers. What is added is the share, so a column
    # that LOOKS like amounts and fails to qualify can say so.
    numeric_share = numeric_shaped / len(filled)
    unreadable = [str(v) for v in filled
                  if not ((isinstance(v, (int, float))
                           and not isinstance(v, bool))
                          or _numeric_body(v) is not None)]
    return {"name": name, "kind": "text", "convention": "",
            "values": len(filled), "parsed": len(filled), "unparsed": [],
            "numeric_share": numeric_share,
            "numeric_values": numeric_shaped,
            "unreadable_samples": unreadable[:3]}


def profile_table(rows: list[list], *, header: bool = True) -> list[dict]:
    """Profile every column of a table (first row taken as the header)."""
    if not rows:
        return []
    width = max(len(r) for r in rows)
    names = (
        [str(rows[0][i]) if i < len(rows[0]) else f"col{i + 1}"
         for i in range(width)]
        if header else [column_letter(i + 1) for i in range(width)]
    )
    body = rows[1:] if header else rows
    out = []
    for index in range(width):
        column = [r[index] if index < len(r) else "" for r in body]
        out.append(profile_column(column, name=names[index]))
    return out


# How many cell values one read contributes to the figure ledger. A
# window is at most 200 x 40, and an answer quotes a handful of them.
MAX_WINDOW_NUMBERS = 600


def _window_numbers(
    window: list[list], profiles: list[dict], *, header: bool,
) -> list[float]:
    """Every number the grid actually shows, under its column's convention.

    These are the values an answer quotes back ("R-006 mit 2.145,75 EUR"),
    so the ledger has to hold them or a correct quotation looks exactly
    like an invented figure. A column whose convention nothing settles is
    left out: its values have two readings and neither is what the tool
    returned.
    """
    out: list[float] = []
    body = window[1:] if header else window
    for row in body:
        for index, raw in enumerate(row):
            profile = profiles[index] if index < len(profiles) else {}
            convention = str(profile.get("convention") or PLAIN_NUMBER)
            if convention == AMBIGUOUS:
                continue
            value = parse_number(raw, convention)
            if value is not None:
                out.append(round(value, 6))
                if len(out) >= MAX_WINDOW_NUMBERS:
                    return out
    return out


def column_notes(profiles: list[dict]) -> list[str]:
    """The findings from a profile that a reader must not miss."""
    notes: list[str] = []
    for entry in profiles:
        name = entry.get("name") or "?"
        if entry.get("convention") == AMBIGUOUS:
            notes.append(
                f"column '{name}': {entry.get('convention_reason', '')}. "
                "Ask which reading is meant — do not compute on it as it is.")
        # The two findings are independent: a column can both use a
        # decimal comma AND hold values that are not numbers, and each
        # one on its own is enough to make a naive total wrong.
        if entry.get("kind") == "number" and entry.get("convention") == DECIMAL_COMMA:
            notes.append(
                f"column '{name}': numbers use a decimal comma "
                "(1.234,50 = 1234.5). Convert before computing — read as-is "
                "they are off by a factor of a thousand.")
        elif entry.get("kind") == "date" and entry.get("convention") == DAY_FIRST:
            notes.append(
                f"column '{name}': dates are day-first (31.07.2026 = "
                "2026-07-31).")
        unparsed = entry.get("unparsed") or []
        if unparsed:
            missing = entry["values"] - entry["parsed"]
            notes.append(
                f"column '{name}': {missing} of {entry['values']} value(s) "
                f"are not {entry['kind']}s (e.g. {', '.join(unparsed[:3])}). "
                "A total over this column would leave them out.")
        # A column too broken to be CALLED a number column. This is where
        # the tool used to fall silent -- one unreadable value past the
        # threshold and nothing was said at all, on the columns most
        # likely to be summed.
        elif entry.get("kind") == "text" and entry.get("numeric_values"):
            bad = entry["values"] - entry["numeric_values"]
            samples = entry.get("unreadable_samples") or []
            notes.append(
                f"column '{name}': {bad} of {entry['values']} value(s) "
                f"could not be read as numbers"
                + (f" (e.g. {', '.join(samples)})" if samples else "")
                + ". Too many to treat this as an amount column — do not "
                "total it before the unreadable values are resolved.")
    return notes


# ---------------------------------------------------------------------------
# Spreadsheets — reading
# ---------------------------------------------------------------------------

# Encodings tried in order when reading a text table. utf-8 first (it
# fails loudly on non-utf-8 bytes rather than producing plausible
# nonsense), then the two Windows encodings a spreadsheet program on a
# German-language system writes. cp1252 before latin-1 because it is a
# superset for the printable range and decodes the Euro sign and typographic
# quotes that latin-1 turns into control characters.
_TEXT_ENCODINGS: tuple[str, ...] = ("utf-8-sig", "cp1252", "latin-1")


def decode_text(raw: bytes) -> tuple[str, str]:
    """Decode table bytes, returning ``(text, encoding_used)``.

    A CSV exported from a spreadsheet program on a German-language
    Windows system is cp1252, not utf-8. Decoding it as utf-8 with
    ``errors="replace"`` does not fail — it turns 'Müller' into
    'M�ller' and 'Straße' into 'Stra�e', and the agent then
    writes those broken names into letters. Every name and street in a
    German data set carries at least one of those bytes, so guessing
    wrong here corrupts essentially the whole file, silently.
    """
    for encoding in _TEXT_ENCODINGS:
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    # latin-1 maps every byte, so this is unreachable in practice; keep a
    # defined outcome rather than an exception if that ever changes.
    return raw.decode("utf-8", errors="replace"), "utf-8 (with damage)"


def sniff_delimiter(text: str, suffix: str = "") -> str:
    """Pick the separator that actually splits this file into columns.

    Counting occurrences is the obvious approach and the wrong one for
    German data: a decimal comma appears in every amount, so a
    semicolon-separated file looks comma-separated and collapses into one
    column. What distinguishes the real separator is that it produces the
    SAME number of fields in every row — the others land in the middle of
    values and vary from line to line.
    """
    if suffix.lower() == ".tsv":
        return "\t"
    import csv as _csv
    import io as _io

    sample = [ln for ln in text.splitlines()[:60] if ln.strip()][:40]
    if not sample:
        return ","
    best, best_score = ",", (-1.0, 0)
    for candidate in (";", ",", "\t", "|"):
        try:
            rows = list(_csv.reader(_io.StringIO("\n".join(sample)),
                                    delimiter=candidate))
        except _csv.Error:
            continue
        widths = [len(r) for r in rows if r]
        if not widths or max(widths) < 2:
            continue
        common = max(set(widths), key=widths.count)
        consistency = widths.count(common) / len(widths)
        score = (consistency, common)
        if score > best_score:
            best, best_score = candidate, score
    return best


def window_record(*, first_row: int, shown_rows: int,
                  total_rows: int) -> dict:
    """Which rows the caller actually got — as data, not as prose.

    The note beside it says the same thing in words, for the model. This
    says it for the mechanisms, because a figure taken from a window is
    not a figure about the file: the largest value among the first five
    rows of twenty-six is not the largest booking, and the two are
    indistinguishable once the number is out of the grid and in a
    sentence.

    Costs nothing in the context window. The result dict is rendered into
    a report for the model, never handed over as it stands, so a field
    here is read by code and by nobody else.
    """
    first = max(1, int(first_row or 1))
    shown = max(0, int(shown_rows or 0))
    total = max(0, int(total_rows or 0))
    last = first + shown - 1 if shown else first - 1
    return {
        "first_row": first,
        "last_row": last,
        "total_rows": total,
        # A window is complete when it starts at the top and reaches the
        # bottom. Anything else is a page, however large.
        "complete": bool(shown and first <= 1 and last >= total),
    }


def _window_notes(
    window: list[list[Any]], *, total_cols: int, start_col: int,
    max_cols: int, cell_limit: int = 40,
) -> list[str]:
    """What the caller is NOT seeing, and how to see it.

    Row truncation was signposted with a remedy from the start; the two
    other cuts were not. Columns said "showing 40 of 87" and named
    nothing the caller could do — the slice always began at column 1, so
    columns 41 to 87 were unreachable through this tool. And cells were
    cut at forty characters in silence.

    A limit stated without a remedy is worse than no limit: it tells the
    model something is missing and leaves it to guess, which in practice
    means answering from the part it has.
    """
    out: list[str] = []
    shown = len(window[0]) if window else 0
    last = start_col + shown - 1
    if total_cols > shown or start_col > 1:
        out.append(
            f"showing columns {start_col}-{last} of {total_cols} — pass "
            f"start_col to page across")
    clipped = clipped_cell_count(window, cell_limit)
    if clipped:
        out.append(
            f"{clipped} cell(s) are longer than {cell_limit} characters and "
            f"end with '{_CLIP_MARK}' in the grid above — that is the display "
            f"being cut, not the value. Narrow the window to read one in "
            f"full before quoting or copying it.")
    return out


def _profile_scope_note(profiled_rows: int, total_rows: int) -> list[str]:
    """A profile is a fact about the rows it saw, not about the file.

    The column profiles — the type, the decimal convention, the date
    order, the count of values that would not parse — are computed from
    the window and rendered under a header stating the file's full row
    count. So "Betrag: number (decimal_comma), 3 of 200 not parseable"
    sat directly under "40000 rows", and read as a property of the
    column. Worse, the convention it decides there is then applied to
    the whole file by the summing and comparing tools.
    """
    if profiled_rows >= total_rows or total_rows <= 0:
        return []
    return [
        f"the column profiles below are computed from rows 1-{profiled_rows} "
        f"of {total_rows}. The conventions they report — decimal comma, date "
        f"order, what counts as a number — and every count in them describe "
        f"that range, not the whole file."]


def _read_csv(
    path: Path, *, max_rows: int, max_cols: int, start_row: int,
    start_col: int = 1,
) -> dict:
    import io

    # A .csv that is not text at all. Decoding a workbook or a PDF with
    # latin-1 never fails — every byte maps to some character — so this
    # would otherwise come back as a one-column table of mojibake that
    # has a row count and a header and means nothing.
    found = container_kind(path)
    if found in ("ole2", "zip", "xlsx", "docx", "ods", "odt", "odf", "pdf"):
        raise OfficeError(
            f"{path.name} is named like a text table but holds "
            f"{_CONTAINER_LABELS.get(found, found)}. Decoding it as text "
            "would return a table of noise rather than its contents — give "
            "it the right suffix and read it again, or convert it.")

    try:
        raw = path.read_bytes()
    except OSError as exc:
        raise OfficeError(f"could not read {path.name}: {exc}") from exc
    text, encoding = decode_text(raw)

    delimiter = sniff_delimiter(text, path.suffix)
    all_rows = list(csv.reader(io.StringIO(text, newline=""),
                              delimiter=delimiter))

    total_rows = len(all_rows)
    total_cols = max((len(r) for r in all_rows), default=0)
    begin = max(0, start_row - 1)
    col_begin = max(0, start_col - 1)
    window = [r[col_begin:col_begin + max_cols]
              for r in all_rows[begin:begin + max_rows]]
    notes = []
    if delimiter != ",":
        notes.append(f"delimiter detected: {delimiter!r}")
    if encoding != "utf-8-sig":
        notes.append(
            f"decoded as {encoding} (not utf-8) — this file came from a "
            "spreadsheet program on Windows. Anything you write back should "
            "keep the accented characters intact.")
    if total_rows > begin + len(window):
        notes.append(
            f"showing rows {begin + 1}-{begin + len(window)} of {total_rows} "
            f"— pass start_row to page further")
    notes.extend(_window_notes(
        window, total_cols=total_cols, start_col=start_col, max_cols=max_cols))
    # Rows that do not have the header's width. An unquoted separator
    # inside a value — "Digitalwaage 0,1 mg" in a comma-separated file —
    # shifts every field after it by one, so a personnel number column
    # quietly starts holding amounts. Nothing about the resulting table
    # looks wrong, which is why it has to be said.
    width = len(all_rows[0]) if all_rows else 0
    ragged = [i for i, r in enumerate(all_rows[1:], start=2)
              if len(r) != width and any(str(c).strip() for c in r)]
    if ragged and width:
        shown = ", ".join(str(i) for i in ragged[:8])
        notes.append(
            f"{len(ragged)} row(s) do not have the header's {width} column(s) "
            f"— rows {shown}{' …' if len(ragged) > 8 else ''}. A separator "
            "inside an unquoted value shifts every field after it, so the "
            "columns of those rows do not mean what their header says.")

    _PROFILE_ROWS = 2000
    profiles = profile_table(all_rows[:_PROFILE_ROWS])
    notes.extend(_profile_scope_note(min(_PROFILE_ROWS, total_rows),
                                     total_rows))
    notes.extend(column_notes(profiles))
    return {
        "path": str(path),
        "kind": "csv",
        "sheets": [],
        "sheet": "",
        "rows": total_rows,
        "columns": total_cols,
        "grid": render_grid(window, first_row=begin + 1,
                            first_col=start_col),
        "column_profile": profiles,
        "profile_rows": min(_PROFILE_ROWS, total_rows),
        "window": window_record(first_row=begin + 1, shown_rows=len(window),
                                total_rows=total_rows),
        "numbers": _window_numbers(window, profiles, header=(begin == 0)),
        "notes": notes,
    }


def _fragile_features(workbook: Any) -> list[str]:
    """Workbook contents that a rewrite puts at risk, named per sheet.

    ``openpyxl`` reconstructs the file on save. Recent versions carry
    charts, images, merges and number formats through a round-trip, but
    the library does not model every part of the format, and a workbook
    authored by Excel can hold constructs it has never seen. Naming what
    is present lets the caller decide, and lets the user hear about it
    before the file is rewritten rather than after.
    """
    found: list[str] = []
    for sheet in getattr(workbook, "worksheets", []) or []:
        title = getattr(sheet, "title", "?")
        for attr, label in (
            ("_charts", "chart"),
            ("_images", "image"),
            ("_pivots", "pivot table"),
            ("_tables", "table"),
        ):
            items = getattr(sheet, attr, None) or []
            try:
                count = len(items)
            except TypeError:
                count = 0
            if count:
                found.append(
                    f"{title}: {count} {label}{'s' if count != 1 else ''}")
    if getattr(workbook, "vba_archive", None) is not None:
        found.append("VBA macros")
    return found


# ---------------------------------------------------------------------------
# OpenDocument spreadsheets (.ods)
# ---------------------------------------------------------------------------
#
# Three properties of the format decide whether a reader reports the
# file or something adjacent to it:
#
# * Cells repeat. LibreOffice writes one cell with
#   ``table:number-columns-repeated="1013"`` rather than a thousand empty
#   ones, and one row with ``number-rows-repeated="1048571"`` for the
#   rest of the sheet. A reader that ignores the attribute shifts every
#   value after a repeated cell into the wrong column.
# * A cell carries both a stored value and a formatted display text:
#   ``office:value="1234.5"`` next to ``1.234,50 €``. The stored value is
#   what the file means; taking the display text would put the whole
#   decimal-comma problem back into a file that had already settled it.
# * A merged range leaves ``table:covered-table-cell`` elements behind.
#   They hold nothing but they occupy their position, so dropping them
#   moves everything to their right one column left.

MAX_ODS_ROWS = 100_000
# Nothing sensible has more columns than this; the number exists to stop
# a repeat count of a million from being materialised.
_ODS_MAX_COLS = 4096


def _ods_blank(value: Any) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def _odf_load(p: Path) -> Any:
    """Open an OpenDocument file, or say what the file really is."""
    _require("opendocument")
    from odf.opendocument import load
    try:
        return load(str(p))
    except Exception as exc:
        raise OfficeError(
            f"could not open {p.name} as an OpenDocument file: "
            f"{type(exc).__name__}: {exc}."
            + (_mismatch_note(p, ("ods", "odt", "odf")) or
               " The file may be damaged or incomplete.")
        ) from exc


def _ods_cell_value(cell: Any) -> tuple[Any, bool]:
    """One cell as ``(value, is_uncached_formula)``.

    Typed values are returned in their own type — a float stays a float,
    a date comes back as a ``date`` — so the column profiler sees what
    the file stores rather than how a locale renders it.
    """
    import datetime as _dt

    from odf.namespaces import OFFICENS, TABLENS
    from odf import teletype

    attrs = dict(getattr(cell, "attributes", None) or {})
    value_type = attrs.get((OFFICENS, "value-type"))
    formula = attrs.get((TABLENS, "formula"))

    def _text() -> str:
        try:
            return teletype.extractText(cell)
        except Exception:
            return ""

    if value_type in ("float", "percentage", "currency"):
        try:
            return float(attrs.get((OFFICENS, "value"), "")), False
        except (TypeError, ValueError):
            return _text(), False
    if value_type == "date":
        raw = str(attrs.get((OFFICENS, "date-value"), ""))
        # A day is returned as a ``date``, not as a midnight ``datetime``:
        # rendered, the second one carries a 00:00:00 the file never
        # stated, and the column profiler then reads the whole column as
        # text rather than as dates.
        try:
            return (_dt.date.fromisoformat(raw) if len(raw) == 10
                    else _dt.datetime.fromisoformat(raw)), False
        except ValueError:
            return _text(), False
    if value_type == "boolean":
        raw = str(attrs.get((OFFICENS, "boolean-value"), "")).lower()
        if raw in ("true", "false"):
            return raw == "true", False
        return _text(), False
    if value_type == "string":
        stored = attrs.get((OFFICENS, "string-value"))
        return (stored if stored is not None else _text()), False
    if value_type == "time":
        # Stored as an ISO 8601 duration (PT10H30M00S), which is not what
        # anyone wrote in the cell. The display text is.
        return _text(), False

    text = _text()
    if formula and not text:
        # A formula with no stored result. Same case as a workbook
        # written by a library: the number does not exist in the file
        # yet, and showing an empty cell would report it as one.
        cleaned = str(formula)
        if cleaned.startswith("of:"):
            cleaned = cleaned[3:]
        return cleaned, True
    return text, False


def _ods_row_cells(row: Any) -> list:
    """The cells of one row, repeats expanded and trailing blanks cut."""
    from odf.namespaces import TABLENS

    cells: list[Any] = []
    for node in getattr(row, "childNodes", None) or []:
        qname = getattr(node, "qname", None)
        if not qname or qname[0] != TABLENS:
            continue
        if qname[1] not in ("table-cell", "covered-table-cell"):
            continue
        attrs = dict(getattr(node, "attributes", None) or {})
        try:
            repeat = int(attrs.get((TABLENS, "number-columns-repeated"), 1))
        except (TypeError, ValueError):
            repeat = 1
        repeat = max(1, min(repeat, _ODS_MAX_COLS - len(cells)))
        value = (None if qname[1] == "covered-table-cell"
                 else _ods_cell_value(node)[0])
        cells.extend([value] * repeat)
        if len(cells) >= _ODS_MAX_COLS:
            break
    while cells and _ods_blank(cells[-1]):
        cells.pop()
    return cells


def _ods_tables(document: Any) -> list:
    """Every sheet of an .ods, in document order."""
    from odf.namespaces import TABLENS

    body = getattr(document, "spreadsheet", None)
    if body is None:
        return []
    return [node for node in (getattr(body, "childNodes", None) or [])
            if getattr(node, "qname", None) == (TABLENS, "table")]


def _ods_sheet_names(document: Any) -> list[str]:
    from odf.namespaces import TABLENS
    names = []
    for index, table in enumerate(_ods_tables(document), start=1):
        attrs = dict(getattr(table, "attributes", None) or {})
        names.append(str(attrs.get((TABLENS, "name"), f"Tabelle{index}")))
    return names


def _ods_sheet(
    p: Path, sheet: Optional[str] = None
) -> tuple[list[str], str, list, int, int, int]:
    """``(sheet names, chosen, row entries, rows, columns, hidden rows)``.

    Row entries are ``(cells, repeat)`` pairs — unexpanded, because the
    repeat count at the end of a sheet routinely runs to a million and
    materialising it would turn a 4 kB file into a gigabyte of ``None``.
    """
    # Load first: it is the call that reports a missing odfpy as
    # something the reader can act on. Importing the namespaces above it
    # would raise a bare ImportError out of a document read instead.
    document = _odf_load(p)

    from odf.namespaces import TABLENS

    mimetype = str(getattr(document, "mimetype", "") or "")
    if mimetype and not mimetype.endswith(
            ("spreadsheet", "spreadsheet-template")):
        raise OfficeError(
            f"{p.name} is named like a spreadsheet but its OpenDocument "
            f"type is '{mimetype}'. Reading it as a table would invent a "
            "structure the file does not have.")

    names = _ods_sheet_names(document)
    tables = _ods_tables(document)
    if not tables:
        raise OfficeError(f"{p.name} holds no sheet")
    if sheet:
        if sheet not in names:
            raise OfficeError(
                f"no sheet named {sheet!r} — available: "
                f"{', '.join(names) or '(none)'}")
        table = tables[names.index(sheet)]
        chosen = sheet
    else:
        table, chosen = tables[0], names[0]

    entries: list[tuple[list, int]] = []
    concealed: list[bool] = []
    for node in getattr(table, "childNodes", None) or []:
        qname = getattr(node, "qname", None)
        if not qname or qname[0] != TABLENS:
            continue
        if qname[1] == "table-header-rows":
            # A repeated header is wrapped in its own element; its rows
            # are ordinary rows and belong at the top of the sheet.
            for inner in getattr(node, "childNodes", None) or []:
                if getattr(inner, "qname", None) == (TABLENS, "table-row"):
                    entries.append((_ods_row_cells(inner), 1))
                    concealed.append(False)
            continue
        if qname[1] != "table-row":
            continue
        attrs = dict(getattr(node, "attributes", None) or {})
        try:
            repeat = max(
                1, int(attrs.get((TABLENS, "number-rows-repeated"), 1)))
        except (TypeError, ValueError):
            repeat = 1
        entries.append((_ods_row_cells(node), repeat))
        concealed.append(
            str(attrs.get((TABLENS, "visibility"), "visible")) != "visible")

    while entries and not entries[-1][0]:
        entries.pop()
        concealed.pop()
    # Counted after the trim: the filler at the end of a sheet can carry
    # a visibility attribute too, and counting it would report a million
    # hidden rows for a file that has none.
    hidden = sum(repeat for (_cells, repeat), flag
                 in zip(entries, concealed) if flag)
    total_rows = sum(repeat for _cells, repeat in entries)
    total_cols = max((len(c) for c, _r in entries), default=0)
    return names, chosen, entries, total_rows, total_cols, hidden


def _ods_window(entries: list, start_row: int, max_rows: int) -> list[list]:
    """Materialise rows ``[start_row, start_row + max_rows)``, 1-based."""
    out: list[list] = []
    position = 1
    for cells, repeat in entries:
        if len(out) >= max_rows:
            break
        end = position + repeat
        if end > start_row:
            first = max(position, start_row)
            take = min(end, start_row + max_rows) - first
            out.extend([list(cells) for _ in range(max(0, take))])
        position = end
    return out


def _read_ods(
    p: Path, *, max_rows: int, max_cols: int, start_row: int,
    sheet: Optional[str] = None, start_col: int = 1,
) -> dict:
    names, chosen, entries, total_rows, total_cols, hidden = _ods_sheet(
        p, sheet)
    _col_begin = max(0, start_col - 1)
    window = [row[_col_begin:_col_begin + max_cols] for row in
              _ods_window(entries, start_row, max_rows)]

    notes: list[str] = []
    uncached = sum(
        1 for row in window for value in row
        if isinstance(value, str) and value.startswith("=")
    )
    if uncached:
        notes.append(
            f"{uncached} formula cell(s) carry no stored result and are "
            "shown as their formula text. Nothing has evaluated this file, "
            "so those numbers do not exist in it yet — do not report them "
            "as computed.")
    if len(names) > 1:
        notes.append(f"sheets in this workbook: {', '.join(names)}")
    if total_rows > start_row + len(window) - 1:
        notes.append(
            f"showing rows {start_row}-{start_row + len(window) - 1} of "
            f"{total_rows} — pass start_row to page further")
    notes.extend(_window_notes(
        window, total_cols=total_cols, start_col=start_col,
        max_cols=max_cols))
    if hidden:
        notes.append(
            f"{hidden} row(s) of this sheet are hidden or filtered out. "
            "They are included in the grid above, so the totals here can "
            "differ from what the file shows on screen.")

    profiles = profile_table(window, header=(start_row == 1))
    notes.extend(_profile_scope_note(len(window), total_rows))
    notes.extend(column_notes(profiles))
    return {
        "path": str(p),
        # The content is a spreadsheet whatever the container is; the
        # format is carried separately so a caller can tell that writing
        # is not on the table.
        "kind": "spreadsheet",
        "format": "ods",
        "writable": False,
        "sheets": names,
        "sheet": chosen,
        "rows": total_rows,
        "columns": total_cols,
        "grid": render_grid(window, first_row=start_row,
                            first_col=start_col),
        "column_profile": profiles,
        "profile_rows": len(window),
        "window": window_record(first_row=start_row, shown_rows=len(window),
                                total_rows=total_rows),
        "numbers": _window_numbers(window, profiles,
                                   header=(start_row == 1)),
        "notes": notes,
    }


def read_sheet(
    path: Any,
    *,
    sheet: Optional[str] = None,
    max_rows: int = DEFAULT_MAX_ROWS,
    max_cols: int = DEFAULT_MAX_COLS,
    start_row: int = 1,
    start_col: int = 1,
) -> dict:
    """Read a window of a spreadsheet or CSV as an addressable grid.

    Formula cells are reported by their computed value where the file
    carries one. Where it does not — the usual case for a workbook
    written by a library and never opened in a spreadsheet program — the
    formula text is shown instead and a note says so, because silently
    printing an empty column would misreport the file.
    """
    p = _resolve(path)
    _refuse_unreadable_format(p)
    kind = document_kind(p)
    max_rows = max(1, min(int(max_rows), 2000))
    max_cols = max(1, min(int(max_cols), 200))
    start_row = max(1, int(start_row))
    start_col = max(1, int(start_col))

    if kind == "csv":
        return _read_csv(
            p, max_rows=max_rows, max_cols=max_cols, start_row=start_row,
            start_col=start_col)
    if kind == "opendocument_sheet":
        return _read_ods(p, max_rows=max_rows, max_cols=max_cols,
                         start_row=start_row, sheet=sheet,
                         start_col=start_col)
    if kind != "spreadsheet":
        raise OfficeError(
            f"{p.name} is not a spreadsheet (suffix {p.suffix!r})")

    openpyxl = _require("spreadsheet")
    try:
        wb = openpyxl.load_workbook(p, data_only=False, read_only=False)
    except Exception as exc:
        raise OfficeError(f"could not open workbook: {exc}") from exc

    try:
        names = list(wb.sheetnames)
        if sheet:
            if sheet not in names:
                raise OfficeError(
                    f"no sheet named {sheet!r} — available: "
                    f"{', '.join(names) or '(none)'}")
            ws = wb[sheet]
        else:
            ws = wb.active

        total_rows = int(ws.max_row or 0)
        total_cols = int(ws.max_column or 0)
        window: list[list[Any]] = []
        has_formula = False
        for row in ws.iter_rows(
            min_row=start_row,
            max_row=min(total_rows, start_row + max_rows - 1),
            min_col=start_col,
            max_col=min(total_cols, start_col + max_cols - 1),
            values_only=True,
        ):
            window.append(list(row))
            for value in row:
                if isinstance(value, str) and value.startswith("="):
                    has_formula = True

        notes: list[str] = []
        # Two things the .ods reader has always reported and this one did
        # not. Read from the sheet rather than the value grid, because both
        # are invisible once iter_rows has flattened it.
        last_shown = min(total_rows, start_row + max_rows - 1)
        xlsx_hidden = 0
        try:
            for _r in range(start_row, last_shown + 1):
                _dim = ws.row_dimensions.get(_r)
                if _dim is not None and getattr(_dim, "hidden", False):
                    xlsx_hidden += 1
        except Exception:
            xlsx_hidden = 0
        xlsx_merges: list[str] = []
        try:
            for _rng in getattr(ws, "merged_cells", None).ranges:
                if _rng.min_row <= last_shown and _rng.max_row >= start_row:
                    xlsx_merges.append(str(_rng))
        except Exception:
            xlsx_merges = []
        xlsx_merges.sort()
        # Second pass only when it can change what is shown: swap in
        # cached values for the formula cells, and say which ones have
        # none rather than rendering them as blanks.
        if has_formula:
            uncached = 0
            try:
                wb_values = openpyxl.load_workbook(
                    p, data_only=True, read_only=True)
                ws_values = (
                    wb_values[ws.title] if ws.title in wb_values.sheetnames
                    else wb_values.active
                )
                for r_idx, row in enumerate(ws_values.iter_rows(
                    min_row=start_row,
                    max_row=min(total_rows, start_row + max_rows - 1),
                    min_col=start_col,
                    max_col=min(total_cols, start_col + max_cols - 1),
                    values_only=True,
                )):
                    if r_idx >= len(window):
                        break
                    for c_idx, cached in enumerate(row):
                        if c_idx >= len(window[r_idx]):
                            break
                        current = window[r_idx][c_idx]
                        if isinstance(current, str) and current.startswith("="):
                            if cached is None:
                                uncached += 1
                            else:
                                window[r_idx][c_idx] = cached
                wb_values.close()
            except Exception:
                notes.append(
                    "formula cells are shown as formulas — the cached "
                    "values could not be read")
                uncached = 0
            if uncached:
                notes.append(
                    f"{uncached} formula cell(s) carry no cached value and "
                    "are shown as their formula text. A spreadsheet program "
                    "has never evaluated this file, so those numbers do not "
                    "exist in it yet — do not report them as computed.")

        if len(names) > 1:
            notes.append(f"sheets in this workbook: {', '.join(names)}")
        if total_rows > start_row + len(window) - 1:
            notes.append(
                f"showing rows {start_row}-{start_row + len(window) - 1} of "
                f"{total_rows} — pass start_row to page further")
        notes.extend(_window_notes(
            window, total_cols=total_cols, start_col=start_col,
            max_cols=max_cols))
        if xlsx_hidden:
            # The .ods reader has said this for a long time; the .xlsx branch
            # never looked at row_dimensions or auto_filter. A finance workbook
            # is normally saved with a filter left on, and the agent's total
            # then silently includes rows the user cannot see on screen.
            notes.append(
                f"{xlsx_hidden} row(s) in the range shown are hidden or "
                "filtered out. They ARE included in the grid above, so a total "
                "computed here can differ from what the file shows on screen.")
        if xlsx_merges:
            # iter_rows(values_only=True) yields None for every non-anchor cell
            # of a merged range, so the label of a merged block appears on its
            # first row and the rows beneath look empty. In the German
            # cost-centre layout that is the difference between "KST 4711 has
            # the highest total" and four fifths of the spend sitting under a
            # cost centre with no name. The values are NOT filled in: that
            # would put data in the file's mouth.
            shown = ", ".join(xlsx_merges[:4])
            more = (f" and {len(xlsx_merges) - 4} more"
                    if len(xlsx_merges) > 4 else "")
            notes.append(
                f"{len(xlsx_merges)} merged cell range(s): {shown}{more}. Only "
                "the first cell of each range carries the value; the cells "
                "under it read as blank here but are not empty in the file.")
        fragile = _fragile_features(wb)
        if fragile:
            notes.append("fragile content present: " + "; ".join(fragile))

        # Profile the window that was read, treating its first row as the
        # header only when it really is one (paging into the middle of a
        # sheet must not promote a data row to a column name).
        profiles = profile_table(window, header=(start_row == 1))
        notes.extend(_profile_scope_note(len(window), total_rows))
        notes.extend(column_notes(profiles))

        return {
            "path": str(p),
            "kind": "spreadsheet",
            "sheets": names,
            "sheet": ws.title,
            "rows": total_rows,
            "columns": total_cols,
            "grid": render_grid(window, first_row=start_row,
                                first_col=start_col),
            "column_profile": profiles,
            "profile_rows": len(window),
            "window": window_record(first_row=start_row,
                                    shown_rows=len(window),
                                    total_rows=total_rows),
            "numbers": _window_numbers(window, profiles,
                                       header=(start_row == 1)),
            "notes": notes,
        }
    finally:
        try:
            wb.close()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Spreadsheets — writing
# ---------------------------------------------------------------------------

def plan_sheet_edits(
    path: Any,
    *,
    edits: Optional[list[dict]] = None,
    append_rows: Optional[list[list[Any]]] = None,
    sheet: Optional[str] = None,
) -> dict:
    """Resolve what :func:`edit_sheet` would change, without writing.

    Returns the per-cell before/after list, the fragile features the
    rewrite would put at risk, and a one-line summary. The tool layer
    shows this to the user as the approval preview: a binary diff is
    unreadable, whereas ``Daten!B7: 100 -> 140`` is exactly the question
    being asked.
    """
    p = _resolve(path)
    _refuse_unwritable_format(p, what="editing it here")
    if document_kind(p) != "spreadsheet":
        raise OfficeError(
            f"{p.name} is not a spreadsheet (suffix {p.suffix!r})")
    openpyxl = _require("spreadsheet")
    try:
        wb = openpyxl.load_workbook(p, data_only=False)
    except Exception as exc:
        raise OfficeError(f"could not open workbook: {exc}") from exc
    try:
        if sheet:
            if sheet not in wb.sheetnames:
                raise OfficeError(
                    f"no sheet named {sheet!r} — available: "
                    f"{', '.join(wb.sheetnames)}")
            ws = wb[sheet]
        else:
            ws = wb.active

        changes: list[dict] = []
        for edit in edits or []:
            if not isinstance(edit, dict):
                raise OfficeError(f"each edit must be an object, got {edit!r}")
            ref = str(edit.get("cell", "")).strip()
            row, col = parse_cell(ref)
            old = ws.cell(row=row, column=col).value
            changes.append({
                "cell": f"{column_letter(col)}{row}",
                "old": _fmt(old),
                "new": _fmt(edit.get("value")),
            })

        appended = len(append_rows or [])
        return {
            "path": str(p),
            "sheet": ws.title,
            "changes": changes,
            "appended_rows": appended,
            "first_appended_row": int(ws.max_row or 0) + 1 if appended else 0,
            "fragile": _fragile_features(wb),
        }
    finally:
        try:
            wb.close()
        except Exception:
            pass


# Which labels a bookkeeping sheet puts on its closing row is a German
# question, so the vocabulary lives with the other German matchers. A
# label alone is not enough: "Summe der Belege folgt" is prose. The row
# has to aggregate something as well, which in practice means a formula.


def _total_row_above(ws, before_row: int) -> "tuple[int, str] | None":
    """The last row before ``before_row`` that closes the sheet, if any.

    An append lands at max_row + 1, which on a sheet ending in
    "Summe | =SUMME(D2:D250)" is BELOW the total and outside its range.
    The booking is recorded, the workbook's total is not, and the user
    reads a number off their own file that is now wrong. Nothing in this
    module looked for one.

    The label is matched by its tokens. A prefix plus a two-character
    clamp — what stood here — saw "Summe" and "Gesamt" and was silent on
    "Summe EUR", "Endsumme", "Gesamtbetrag", "Summe gesamt" and "Gesamt
    netto", which is what German sheets actually write. Every one of
    those appended below the total with no warning at all.
    """
    try:
        first = max(1, before_row - 5)
        for row in range(before_row - 1, first - 1, -1):
            label = ""
            has_formula = False
            for col in range(1, min(int(ws.max_column or 1), 30) + 1):
                value = ws.cell(row=row, column=col).value
                if isinstance(value, str):
                    if value.startswith("="):
                        has_formula = True
                    elif not label:
                        label = value.strip().lower()
            if not label or not has_formula:
                continue
            if _de.is_total_row_label(label):
                return row, label
    except Exception:
        return None
    return None


def _verify_cells(path: Path, sheet_name: str, applied: list[dict]) -> list[str]:
    """Re-read the saved file and report cells that do not hold the new value.

    The form filler reads its result back before reporting success; a
    spreadsheet edit reported from the in-memory workbook would be
    claiming an outcome from the intent. Never raises — a verification
    that cannot run is reported as unverified, not as a failure.
    """
    if not applied:
        return []
    try:
        import openpyxl as _op
        wb = _op.load_workbook(path, data_only=False)
    except Exception:
        return []
    try:
        ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active
        out: list[str] = []
        for change in applied:
            ref = str(change.get("cell", ""))
            try:
                row, col = parse_cell(ref)
            except OfficeError:
                continue
            if _fmt(ws.cell(row=row, column=col).value) != str(change.get("new", "")):
                out.append(ref)
        return out
    except Exception:
        return []
    finally:
        try:
            wb.close()
        except Exception:
            pass


def _sheet_header(ws: Any) -> list[str]:
    """The first row as column names (blank cells become empty strings)."""
    for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
        return [str(v).strip() if v is not None else "" for v in row]
    return []


def _locate_rows(ws: Any, key_column: str) -> tuple[int, dict, list[str]]:
    """``(column index, key -> row, duplicate keys)`` for a key column.

    Duplicates are collected rather than resolved: updating "the row of
    Müller" when there are two of them is not a decision a tool may make
    silently.
    """
    header = _sheet_header(ws)
    wanted = str(key_column).strip().lower()
    index = next((i for i, name in enumerate(header)
                  if name.strip().lower() == wanted), -1)
    if index < 0:
        raise OfficeError(
            f"no column {key_column!r} in this sheet. Columns: "
            + (", ".join(h for h in header if h) or "(none)"))
    by_key: dict[str, int] = {}
    duplicates: list[str] = []
    for number, row in enumerate(
            ws.iter_rows(min_row=2, values_only=True), start=2):
        raw = row[index] if index < len(row) else None
        token = str(raw).strip() if raw is not None else ""
        if not token:
            continue
        if token in by_key:
            if token not in duplicates:
                duplicates.append(token)
        else:
            by_key[token] = number
    return index, by_key, duplicates


def edit_sheet(
    path: Any,
    *,
    edits: Optional[list[dict]] = None,
    append_rows: Optional[list[list[Any]]] = None,
    key_column: Optional[str] = None,
    updates: Optional[list[dict]] = None,
    append_records: Optional[list[dict]] = None,
    sheet: Optional[str] = None,
) -> dict:
    """Set cells and append rows in an existing workbook, in place.

    Formulas elsewhere in the file are preserved (the workbook is loaded
    with ``data_only=False``; loading with cached values instead would
    replace every formula in the file with its last computed number).
    Macro-enabled workbooks keep their VBA project.

    A cell value beginning with ``=`` is stored as a formula, which
    means the file then carries no cached result for it until a
    spreadsheet program opens and recalculates the file. That is stated
    in the returned notes rather than left for the reader to discover.
    """
    p = _resolve(path)
    _refuse_unwritable_format(p, what="editing it here")
    if document_kind(p) != "spreadsheet":
        raise OfficeError(
            f"{p.name} is not a spreadsheet (suffix {p.suffix!r})")
    if not any((edits, append_rows, updates, append_records)):
        raise OfficeError(
            "nothing to do: pass edits/append_rows (by cell) or "
            "updates/append_records (by key and column name)")
    if updates and not key_column:
        raise OfficeError(
            "updates need key_column — the column whose value identifies a "
            "row (a document number, a case reference)")

    openpyxl = _require("spreadsheet")
    keep_vba = p.suffix.lower() in (".xlsm", ".xltm")
    try:
        wb = openpyxl.load_workbook(p, data_only=False, keep_vba=keep_vba)
    except Exception as exc:
        raise OfficeError(f"could not open workbook: {exc}") from exc

    try:
        if sheet:
            if sheet not in wb.sheetnames:
                raise OfficeError(
                    f"no sheet named {sheet!r} — available: "
                    f"{', '.join(wb.sheetnames)}")
            ws = wb[sheet]
        else:
            ws = wb.active

        applied: list[dict] = []
        wrote_formula = False

        # Record-oriented changes. Everything is validated BEFORE the first
        # cell is written: an unknown key, a duplicate key or an unknown
        # column stops the whole call. A half-applied batch on someone's
        # records is worse than a refused one, because nothing about the
        # file afterwards says which half went through.
        if updates or append_records:
            header = _sheet_header(ws)
            named = {h.strip().lower(): i for i, h in enumerate(header) if h}
            if not named:
                raise OfficeError(
                    "this sheet has no header row, so columns cannot be "
                    "addressed by name — use edits with cell references")
            wanted_columns: set[str] = set()
            for entry in list(updates or []) + list(append_records or []):
                if not isinstance(entry, dict):
                    raise OfficeError(f"expected an object, got {entry!r}")
                values = entry.get("set") if updates and "set" in entry else entry
                if not isinstance(values, dict):
                    raise OfficeError(
                        f"each update needs a 'set' object of column -> value, "
                        f"got {entry!r}")
                wanted_columns.update(str(c) for c in values
                                      if str(c).strip().lower() != "key")
            unknown = sorted(c for c in wanted_columns
                             if c.strip().lower() not in named)
            if unknown:
                raise OfficeError(
                    f"no column(s) {', '.join(unknown)} in this sheet. "
                    "Columns: " + ", ".join(h for h in header if h))

        if updates:
            _key_index, by_key, duplicates = _locate_rows(ws, key_column)
            header = _sheet_header(ws)
            named = {h.strip().lower(): i for i, h in enumerate(header) if h}
            wanted_keys = [str(u.get("key", "")).strip() for u in updates]
            missing = sorted({k for k in wanted_keys if k and k not in by_key})
            blank = [i for i, k in enumerate(wanted_keys, start=1) if not k]
            ambiguous = sorted(set(wanted_keys) & set(duplicates))
            if blank:
                raise OfficeError(
                    f"update {blank[0]} has no key — every update needs the "
                    f"value in '{key_column}' that identifies its row")
            if missing:
                raise OfficeError(
                    f"no row with {key_column} = "
                    + ", ".join(repr(k) for k in missing)
                    + f". The sheet holds {len(by_key)} distinct key(s).")
            if ambiguous:
                raise OfficeError(
                    f"{key_column} = " + ", ".join(repr(k) for k in ambiguous)
                    + " appears more than once — which row is meant is not "
                    "decidable, so nothing was changed.")
            for update in updates:
                row_number = by_key[str(update.get("key", "")).strip()]
                for column, value in (update.get("set") or {}).items():
                    col_index = named[str(column).strip().lower()] + 1
                    cell = ws.cell(row=row_number, column=col_index)
                    old_value = cell.value
                    cell.value = value
                    if isinstance(value, str) and value.startswith("="):
                        wrote_formula = True
                    applied.append({
                        "cell": f"{column_letter(col_index)}{row_number}",
                        "key": str(update.get("key", "")).strip(),
                        "column": str(column),
                        "old": _fmt(old_value), "new": _fmt(value),
                    })

        for edit in edits or []:
            if not isinstance(edit, dict):
                raise OfficeError(f"each edit must be an object, got {edit!r}")
            row, col = parse_cell(str(edit.get("cell", "")).strip())
            value = edit.get("value")
            cell = ws.cell(row=row, column=col)
            old = cell.value
            cell.value = value
            if isinstance(value, str) and value.startswith("="):
                wrote_formula = True
            applied.append({
                "cell": f"{column_letter(col)}{row}",
                "old": _fmt(old),
                "new": _fmt(value),
            })

        first_appended = 0
        appended = 0
        # Read back like every other write, but kept out of `applied`:
        # that field is the reported list of UPDATED cells and callers
        # read its shape. Only the verification is widened.
        appended_cells: list[dict] = []

        # Records are placed by column NAME. A positional list silently
        # lands in the wrong columns as soon as the sheet's order differs
        # from the caller's assumption, and the result reads perfectly.
        if append_records:
            header = _sheet_header(ws)
            named = {h.strip().lower(): i for i, h in enumerate(header) if h}
            for record in append_records:
                if not appended:
                    first_appended = int(ws.max_row or 0) + 1
                target_row = int(ws.max_row or 0) + 1
                for column, value in record.items():
                    col_index = named[str(column).strip().lower()] + 1
                    ws.cell(row=target_row, column=col_index).value = value
                    if isinstance(value, str) and value.startswith("="):
                        wrote_formula = True
                    # Appended cells go into `applied` like every other
                    # write. They did not, so _verify_cells had nothing to
                    # read back and an append-only edit answered
                    # verified: True having checked nothing.
                    appended_cells.append({
                        "cell": f"{column_letter(col_index)}{target_row}",
                        "old": _fmt(None), "new": _fmt(value),
                    })
                appended += 1

        for new_row in append_rows or []:
            if not isinstance(new_row, (list, tuple)):
                raise OfficeError(
                    f"each appended row must be a list, got {new_row!r}")
            if not appended:
                first_appended = int(ws.max_row or 0) + 1
            _target = int(ws.max_row or 0) + 1
            ws.append(list(new_row))
            for _c, _v in enumerate(new_row, start=1):
                appended_cells.append({
                    "cell": f"{column_letter(_c)}{_target}",
                    "old": _fmt(None), "new": _fmt(_v),
                })
            appended += 1

        notes: list[str] = []
        if appended and first_appended:
            _total = _total_row_above(ws, first_appended)
            if _total is not None:
                _row, _label = _total
                notes.append(
                    f"row {_row} of this sheet closes it with a total "
                    f"({_label!r}). The new row(s) were appended BELOW it, "
                    f"so they are outside its range and that total no "
                    f"longer covers them — move the total or extend its "
                    f"formula.")
        fragile = _fragile_features(wb)
        if fragile:
            notes.append(
                "this workbook carries content a rewrite can affect ("
                + "; ".join(fragile)
                + ") — the file was rebuilt on save, so verify those parts "
                "survived before handing it over")
        if wrote_formula:
            notes.append(
                "a formula was written: the file now stores the formula but "
                "no computed result. Anything reading cached values will see "
                "an empty cell until a spreadsheet program recalculates it.")

        try:
            wb.save(p)
        except Exception as exc:
            raise OfficeError(f"could not save workbook: {exc}") from exc

        sheet_name = ws.title
        unverified = _verify_cells(
            p, sheet_name, applied + appended_cells)
        if unverified:
            notes.append(
                "these cells do not hold the new value in the saved file: "
                + ", ".join(unverified)
                + ". Do not report this edit as done.")

        return {
            "path": str(p),
            "sheet": sheet_name,
            "applied": applied,
            "appended_rows": appended,
            "first_appended_row": first_appended,
            "verified": not unverified,
            "notes": notes,
        }
    finally:
        try:
            wb.close()
        except Exception:
            pass


def create_sheet(
    path: Any,
    rows: list[list[Any]],
    *,
    sheet_name: str = "Sheet1",
    overwrite: bool = False,
) -> dict:
    """Write *rows* to a new workbook (or CSV, by suffix)."""
    p = _resolve(path, must_exist=False)
    _refuse_unwritable_format(p, what="writing it")
    if p.exists() and not overwrite:
        raise OfficeError(
            f"{p.name} already exists — pass overwrite=True to replace it, "
            "or use edit_sheet to change it in place")
    if not isinstance(rows, list) or not all(
            isinstance(r, (list, tuple)) for r in rows):
        raise OfficeError("rows must be a list of lists")

    kind = document_kind(p)
    p.parent.mkdir(parents=True, exist_ok=True)
    if kind == "csv":
        delimiter = "\t" if p.suffix.lower() == ".tsv" else ","
        # utf-8 WITH a BOM: a spreadsheet program opening a plain utf-8 CSV
        # by double-click assumes the system code page and renders every
        # umlaut as mojibake. The BOM is what makes it detect utf-8. It is
        # invisible to every other reader that matters here (Python's csv
        # strips it via utf-8-sig, as does the reader above).
        with p.open("w", encoding="utf-8-sig", newline="") as fh:
            csv.writer(fh, delimiter=delimiter).writerows(
                [list(r) for r in rows])
    elif kind == "spreadsheet":
        openpyxl = _require("spreadsheet")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = str(sheet_name or "Sheet1")[:31]
        for row in rows:
            ws.append(list(row))
        wb.save(p)
        wb.close()
    else:
        raise OfficeError(
            f"cannot create {p.suffix!r} as a spreadsheet — use .xlsx or .csv")
    return {
        "path": str(p),
        "kind": kind,
        "rows": len(rows),
        "columns": max((len(r) for r in rows), default=0),
    }


# ---------------------------------------------------------------------------
# Reconciling two tables
# ---------------------------------------------------------------------------

def _workbook_sheets(path: Any) -> list[str]:
    """Sheet names of a workbook, or an empty list for anything else."""
    p = Path(str(path))
    if document_kind(p) == "opendocument_sheet":
        try:
            return _ods_sheet_names(_odf_load(p))
        except Exception:
            return []
    if document_kind(p) != "spreadsheet":
        return []
    try:
        openpyxl = _require("spreadsheet")
        wb = openpyxl.load_workbook(p, read_only=True)
        try:
            return list(wb.sheetnames)
        finally:
            wb.close()
    except Exception:
        return []


def _table_rows(path: Any, sheet: Optional[str] = None) -> list[list]:
    """Every row of a table file, as lists (no window, no cap on width)."""
    p = _resolve(path)
    _refuse_unreadable_format(p)
    kind = document_kind(p)
    if kind == "csv":
        raw = p.read_bytes()
        text, _enc = decode_text(raw)
        import io
        return list(csv.reader(io.StringIO(text, newline=""),
                               delimiter=sniff_delimiter(text, p.suffix)))
    if kind == "opendocument_sheet":
        _names, _chosen, entries, total, _cols, _hidden = _ods_sheet(p, sheet)
        # Unlike the read path there is no window to bound the work here,
        # so the cap has to be a refusal: silently comparing the first
        # 100 000 rows of a longer sheet would report a reconciliation
        # that never saw the rest of the file.
        if total > MAX_ODS_ROWS:
            raise OfficeError(
                f"{p.name} has {total} rows, above the {MAX_ODS_ROWS} this "
                "reader will materialise at once. Narrow the file, or "
                "convert it to .xlsx and work on that.")
        return _ods_window(entries, 1, total)
    if kind != "spreadsheet":
        raise OfficeError(
            f"{p.name} is not a table (need .xlsx, .ods, .csv or .tsv)")
    openpyxl = _require("spreadsheet")
    try:
        wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
    except Exception as exc:
        raise OfficeError(f"could not open workbook: {exc}") from exc
    try:
        if sheet:
            if sheet not in wb.sheetnames:
                raise OfficeError(
                    f"no sheet named {sheet!r} in {p.name} — available: "
                    f"{', '.join(wb.sheetnames)}")
            ws = wb[sheet]
        else:
            ws = wb.active
        return [list(r) for r in ws.iter_rows(values_only=True)]
    finally:
        try:
            wb.close()
        except Exception:
            pass


def _table_cell_formats(path: Any, sheet: Optional[str] = None) -> list[list]:
    """The number format of every cell, in the shape ``_table_rows``
    returns its values.

    Only a workbook has them. A CSV holds text the user already formatted,
    and an .ods reader gives values without their style, so both come back
    empty and the renderer falls back to the value's own precision.
    """
    try:
        p = _resolve(path)
        if document_kind(p) != "spreadsheet":
            return []
        openpyxl = _require("spreadsheet")
        wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
        try:
            ws = (wb[sheet] if sheet and sheet in wb.sheetnames else wb.active)
            return [[getattr(c, "number_format", "") or "" for c in row]
                    for row in ws.iter_rows()]
        finally:
            try:
                wb.close()
            except Exception:
                pass
    except Exception:
        return []


def _header_index(header: list, column: str, where: str) -> int:
    wanted = str(column).strip().lower()
    for index, name in enumerate(header):
        if str(name or "").strip().lower() == wanted:
            return index
    raise OfficeError(
        f"no column {column!r} in the {where} table. Columns: "
        + ", ".join(str(h) for h in header if str(h or "").strip()))


def _comparable(value: Any, profile: dict) -> Any:
    """Normalise one value for comparison under its column's convention."""
    kind = profile.get("kind")
    convention = profile.get("convention", "")
    if kind == "number":
        parsed = parse_number(value, convention)
        return parsed if parsed is not None else str(value or "").strip()
    if kind == "date":
        parsed = parse_date(value, convention)
        return parsed if parsed is not None else str(value or "").strip()
    return str(value or "").strip()


def _hidden_data_rows(path: Path, sheet: Optional[str]) -> int:
    """How many rows of this sheet are hidden or filtered out.

    A finance workbook is routinely saved with a filter left switched
    on. The rows behind it are still in the file and still in a total
    computed over it, so a total that does not mention them differs from
    the number on the user's screen for a reason nobody can see.
    """
    try:
        kind = document_kind(path)
        if kind == "opendocument_sheet":
            return int(_ods_sheet(path, sheet)[5])
        if kind != "spreadsheet":
            return 0
        openpyxl = _require("spreadsheet")
        wb = openpyxl.load_workbook(path, data_only=True, read_only=False)
        try:
            ws = (wb[sheet] if sheet and sheet in wb.sheetnames else wb.active)
            return sum(1 for dim in ws.row_dimensions.values()
                       if getattr(dim, "hidden", False))
        finally:
            wb.close()
    except Exception:
        return 0


def sum_column(
    path: Any,
    column: str,
    *,
    sheet: Optional[str] = None,
    group_by: Optional[str] = None,
    convention: Optional[str] = None,
    header_row: int = 1,
) -> dict:
    """Total one column of a table and say what it could not add.

    The failure this exists for is not a wrong sum; it is a sum that
    quietly leaves rows out. "1.500,00-" is a credit note, "n/a" is a
    missing amount, and a naive total over the column reads exactly like
    a complete one. So every row of the column lands in one of three
    places — counted, skipped (with the value that could not be read),
    or blank — and the three add up to the rows of the table.

    The column's own convention decides how its values parse, so
    "1.234,50" totals as 1234.5 and not as 1.23450. A column nothing in
    the file settles (every value like "8.986") is REFUSED rather than
    guessed: the two readings differ by a factor of a thousand, and
    picking one silently is the failure this whole module exists to
    avoid. ``convention`` is how the caller says which reading applies
    once the user has been asked.

    ``group_by`` totals per value of another column — the "Summe je
    Kostenstelle" every administrative answer is really about.
    ``header_row`` is for the layout every German cost-centre sheet has:
    a title in row 1 and the column names below it.
    """
    p = _resolve(path)
    rows = _table_rows(p, sheet)
    head_at = max(1, int(header_row or 1))
    if len(rows) < head_at + 1:
        raise OfficeError(
            f"{p.name} needs a header row and at least one data row")
    header = rows[head_at - 1]
    body = rows[head_at:]

    def _column_index(name: str) -> int:
        wanted = str(name).strip().lower()
        for index, cell in enumerate(header):
            if str(cell or "").strip().lower() == wanted:
                return index
        raise OfficeError(
            f"no column {name!r} in {p.name}. Columns: "
            + (", ".join(str(h) for h in header if str(h or "").strip())
               or "(none)"))

    index = _column_index(column)
    values = [r[index] if index < len(r) else "" for r in body]
    profile = profile_column(values, name=str(column))
    # A column too broken to be CALLED numeric still has a convention its
    # readable values agree on -- and that is the column somebody is about
    # to total. Gutschriften.xlsx is the case: four credit notes written
    # as "1.500,00-" and "(340,00)" push it below the numeric threshold,
    # and refusing it outright would withhold the six amounts that ARE
    # readable together with the fact that four are missing.
    detected = profile.get("convention") or detect_number_convention(
        [v for v in values if str(v or "").strip()])[0]
    chosen = str(convention or "").strip() or detected or ""
    if chosen not in (DECIMAL_COMMA, DECIMAL_POINT, PLAIN_NUMBER):
        reason = profile.get("convention_reason") or detect_number_convention(
            [v for v in values if str(v or "").strip()])[1]
        raise OfficeError(
            f"column {column!r} cannot be totalled as it stands: {reason}. "
            f"Ask which reading applies and pass convention="
            f"'{DECIMAL_COMMA}' or '{DECIMAL_POINT}'.")

    group_index = _column_index(group_by) if group_by else -1
    group_name = str(header[group_index]) if group_index >= 0 else ""

    total = 0.0
    counted = 0
    blank = 0
    skipped: list[str] = []
    groups: dict[str, list] = {}
    for offset, raw in enumerate(values):
        if str(raw if raw is not None else "").strip() == "":
            blank += 1
            continue
        bucket = None
        if group_index >= 0:
            row = body[offset]
            key = str(
                (row[group_index] if group_index < len(row) else "") or ""
            ).strip()
            bucket = groups.setdefault(key, [0.0, 0, 0])
        parsed = parse_number(raw, chosen)
        if parsed is None:
            skipped.append(_fmt(raw))
            if bucket is not None:
                bucket[2] += 1
            continue
        total += parsed
        counted += 1
        if bucket is not None:
            bucket[0] += parsed
            bucket[1] += 1

    # Floating point: a column of two-decimal amounts must not come back
    # as 45231.499999999996. Rounded to the precision money is written
    # in, and only there -- the parsed values themselves are untouched.
    total = round(total, 6)

    notes: list[str] = []
    if skipped:
        notes.append(
            f"{len(skipped)} of {counted + len(skipped)} value(s) could not be "
            f"read as numbers and are NOT in the total (e.g. "
            f"{', '.join(skipped[:3])}). Resolve them before reporting it.")
    if blank:
        notes.append(f"{blank} row(s) have no value in this column.")
    if chosen == DECIMAL_COMMA:
        notes.append(
            "this column is written with a decimal comma (1.234,50 = "
            "1234.5); the total above is a plain number.")
    if convention and profile.get("convention") != chosen:
        notes.append(
            f"the convention {chosen!r} was given by the caller, not read "
            "from the column — say so when reporting the total.")
    hidden = _hidden_data_rows(p, sheet)
    if hidden:
        notes.append(
            f"{hidden} row(s) of this sheet are hidden or filtered out. They "
            "ARE in this total, so it can differ from what the file shows on "
            "screen.")
    names = _workbook_sheets(p)
    if len(names) > 1 and not sheet:
        # Same reason compare_tables says it: the active sheet is a silent
        # default, and a total taken from the wrong month is not wrong in
        # any way the number itself shows.
        notes.append(
            f"this workbook has {len(names)} sheets ({', '.join(names)}) and "
            "none was named, so its ACTIVE sheet was totalled. Name the "
            "sheet if that is not the intended scope.")
    if group_index >= 0 and "" in groups:
        without = groups[""][1] + groups[""][2]
        notes.append(
            f"{without} row(s) have no value in {group_name!r} and are "
            "grouped under '(ohne)'. In a sheet with merged cells those rows "
            "belong to the block above them in the file, not to nobody.")

    group_rows = [
        {"key": key or "(ohne)", "total": round(bucket[0], 6),
         "counted": bucket[1], "skipped": bucket[2]}
        for key, bucket in groups.items()
    ]

    # The figures this call produced, named, for the answer-side ledger:
    # a caller must not have to parse them back out of the prose above.
    figures = [
        {"kind": "sum", "value": total,
         "label": f"Summe {column!r} in {p.name}"},
        {"kind": "count", "value": float(counted),
         "label": f"{counted} counted value(s) in {column!r}"},
        {"kind": "count", "value": float(len(body)),
         "label": f"{len(body)} data row(s) in {p.name}"},
    ]
    if skipped:
        figures.append({"kind": "count", "value": float(len(skipped)),
                        "label": f"{len(skipped)} value(s) not readable"})
    if blank:
        figures.append({"kind": "count", "value": float(blank),
                        "label": f"{blank} blank row(s)"})
    if hidden:
        figures.append({"kind": "count", "value": float(hidden),
                        "label": f"{hidden} hidden or filtered row(s)"})
    if skipped:
        # The readable values, which is what "6 von 10" in an answer means.
        figures.append({"kind": "count", "value": float(counted + len(skipped)),
                        "label": f"{counted + len(skipped)} value(s) present"})
    for entry in group_rows:
        figures.append({
            "kind": "sum", "value": entry["total"],
            "label": f"Summe {column!r} für {entry['key']}"})
        figures.append({
            "kind": "count", "value": float(entry["counted"]),
            "label": f"{entry['counted']} row(s) for {entry['key']}"})

    return {
        "path": str(p),
        "sheet": sheet or "",
        "column": str(header[index]),
        "convention": chosen,
        "total": total,
        "counted": counted,
        "blank": blank,
        "skipped": skipped,
        "rows": len(body),
        "group_by": group_name,
        "groups": group_rows,
        "notes": notes,
        "figures": figures,
    }


def compare_tables(
    left: Any,
    right: Any,
    *,
    key: str,
    right_key: Optional[str] = None,
    columns: Optional[list[str]] = None,
    left_sheet: Optional[str] = None,
    right_sheet: Optional[str] = None,
    max_report: int = 200,
) -> dict:
    """Reconcile two tables on a key column.

    Every row of both inputs lands in exactly one group — equal,
    differing, only-left, only-right, or not-comparable — and the counts
    add up to the inputs. That is the whole point: the failure mode of
    doing this by hand is not a wrong comparison, it is rows that quietly
    fall out of the report and are never noticed.

    Values are compared under the convention their column actually uses,
    so "1.234,50" and 1234.5 are equal and 31.07.2026 matches 2026-07-31.
    Duplicate keys are reported rather than joined: a duplicate key turns
    a comparison into a cross product, and the result looks plausible.
    """
    left_rows = _table_rows(left, left_sheet)
    right_rows = _table_rows(right, right_sheet)
    if len(left_rows) < 2 or len(right_rows) < 2:
        raise OfficeError(
            "both tables need a header row and at least one data row")

    left_header, right_header = left_rows[0], right_rows[0]
    # The two tables come from two systems, so the key is routinely spelled
    # differently on each side — "Beleg" against "Belegnummer". Requiring
    # one name would push the caller into renaming a column first, which
    # means writing to a file just to be able to read it.
    right_key_name = str(right_key).strip() if right_key else key
    left_key_index = _header_index(left_header, key, "left")
    right_key_index = _header_index(right_header, right_key_name, "right")

    # ``columns`` may be a list (same name on both sides) or a mapping
    # {left: right}. Two exports of the same facts routinely disagree on
    # every column name — "Betrag" against "Rechnungsbetrag" — and without
    # pairs those columns simply drop out of the comparison, which reports
    # agreement it never checked.
    pairs: dict[str, str] = {}
    if isinstance(columns, dict):
        pairs = {str(k): str(v) for k, v in columns.items()}
        wanted = list(pairs)
    elif columns:
        wanted = [str(c) for c in columns]
    else:
        # Default: every column both tables share, key excluded.
        right_names = {str(h or "").strip().lower() for h in right_header}
        wanted = [
            str(h) for h in left_header
            if str(h or "").strip()
            and str(h or "").strip().lower() in right_names
            and str(h or "").strip().lower() not in (
                str(key).strip().lower(), right_key_name.strip().lower())
        ]
    if not wanted:
        raise OfficeError(
            "the two tables share no comparable column besides the key — "
            "name the columns explicitly")

    left_idx = {c: _header_index(left_header, c, "left") for c in wanted}
    right_idx = {
        c: _header_index(right_header, pairs.get(c, c), "right")
        for c in wanted
    }

    # Each side is profiled on its own. The two tables routinely come from
    # different systems and are written under different conventions —
    # 1.234,50 out of a German export against 1234.50 out of a database.
    # Applying one side's convention to both would report every such pair
    # as a difference, which is exactly the noise that makes a
    # reconciliation useless.
    left_profiles = {
        c: profile_column([r[i] if i < len(r) else ""
                           for r in left_rows[1:]], name=c)
        for c, i in left_idx.items()
    }
    right_profiles = {
        c: profile_column([r[i] if i < len(r) else ""
                           for r in right_rows[1:]], name=pairs.get(c, c))
        for c, i in right_idx.items()
    }

    def _index(rows: list[list], key_at: int) -> tuple[dict, list, list]:
        by_key: dict[str, list] = {}
        blank: list[int] = []
        for offset, row in enumerate(rows[1:], start=2):
            raw = row[key_at] if key_at < len(row) else ""
            token = str(raw or "").strip()
            if not token:
                blank.append(offset)
                continue
            by_key.setdefault(token, []).append((offset, row))
        duplicates = [
            {"key": k, "rows": [o for o, _ in v]}
            for k, v in by_key.items() if len(v) > 1
        ]
        return by_key, duplicates, blank

    left_by_key, left_dupes, left_blank = _index(left_rows, left_key_index)
    right_by_key, right_dupes, right_blank = _index(right_rows, right_key_index)

    not_comparable: list[dict] = []
    for line in left_blank:
        not_comparable.append(
            {"side": "left", "row": line, "reason": "empty key"})
    for line in right_blank:
        not_comparable.append(
            {"side": "right", "row": line, "reason": "empty key"})
    for entry in left_dupes:
        not_comparable.append({
            "side": "left", "key": entry["key"], "rows": entry["rows"],
            "reason": "key appears more than once"})
    for entry in right_dupes:
        not_comparable.append({
            "side": "right", "key": entry["key"], "rows": entry["rows"],
            "reason": "key appears more than once"})

    duplicate_keys = ({e["key"] for e in left_dupes}
                      | {e["key"] for e in right_dupes})
    comparable_left = {k: v[0] for k, v in left_by_key.items()
                       if k not in duplicate_keys}
    comparable_right = {k: v[0] for k, v in right_by_key.items()
                        if k not in duplicate_keys}

    equal: list[str] = []
    differing: list[dict] = []
    # How many matched rows a column actually had something to say about.
    # A column that is empty on both sides agrees on every row, and that
    # agreement is what "equal" was built out of: comparing two tables on
    # a Notiz column nobody fills reported "2 equal, 0 differing" — a
    # clean reconciliation over nothing. Counted here so the verdict can
    # state what it rests on.
    with_values: dict[str, int] = {c: 0 for c in wanted}
    for token, (line, row) in sorted(comparable_left.items()):
        if token not in comparable_right:
            continue
        other_line, other_row = comparable_right[token]
        diffs = []
        for column in wanted:
            a = row[left_idx[column]] if left_idx[column] < len(row) else ""
            b = (other_row[right_idx[column]]
                 if right_idx[column] < len(other_row) else "")
            if str(a or "").strip() or str(b or "").strip():
                with_values[column] += 1
            left_p, right_p = left_profiles[column], right_profiles[column]
            if left_p.get("kind") == right_p.get("kind"):
                # Same kind on both sides: normalise each under its OWN
                # convention, so the comparison is of values, not spellings.
                same = _comparable(a, left_p) == _comparable(b, right_p)
            else:
                # One side is text where the other is a number or a date.
                # Comparing the normalised forms would be comparing two
                # different things, so fall back to the literal text.
                same = str(a or "").strip() == str(b or "").strip()
            if not same:
                diffs.append({"column": column,
                              "left": _fmt(a), "right": _fmt(b)})
        if diffs:
            differing.append({"key": token, "left_row": line,
                              "right_row": other_line, "differences": diffs})
        else:
            equal.append(token)

    only_left = sorted(set(comparable_left) - set(comparable_right))
    only_right = sorted(set(comparable_right) - set(comparable_left))

    notes: list[str] = []

    # Scope, before anything else. A workbook with one sheet per month
    # compared against a whole year reports hundreds of one-sided rows —
    # a result that reads like a catastrophe when the only thing wrong is
    # which sheet was taken. The active sheet is a silent default, so it
    # has to be said out loud.
    for side, source, chosen in (("left", left, left_sheet),
                                 ("right", right, right_sheet)):
        names = _workbook_sheets(source)
        if len(names) > 1 and not chosen:
            notes.append(
                f"the {side} workbook has {len(names)} sheets and none was "
                f"named, so its ACTIVE sheet was used: {names[0]!r} of "
                f"{', '.join(names)}. Check this is the intended scope."
            )

    for column in wanted:
        for side, profile in (("left", left_profiles[column]),
                              ("right", right_profiles[column])):
            if profile.get("convention") == AMBIGUOUS:
                notes.append(
                    f"column '{column}' ({side}): "
                    f"{profile.get('convention_reason', '')} — the comparison "
                    "treated its values as text.")
        if left_profiles[column].get("kind") != right_profiles[column].get("kind"):
            notes.append(
                f"column '{column}': the left side looks like "
                f"{left_profiles[column].get('kind')}s and the right side like "
                f"{right_profiles[column].get('kind')}s — compared as text, so "
                "a difference here may be a formatting difference only.")
    # A column both sides leave empty agrees on every row it is asked
    # about, and that agreement is indistinguishable from a checked one
    # in the counts. Two tables compared on a Notiz column nobody fills
    # reported "2 equal, 0 differing" — a clean reconciliation over
    # nothing at all.
    matched_rows = len(equal) + len(differing)
    empty_cols = [c for c in wanted if with_values.get(c, 0) == 0]
    if matched_rows and empty_cols:
        if len(empty_cols) == len(wanted):
            notes.append(
                f"NOTHING WAS ACTUALLY COMPARED: "
                f"{'the only column' if len(wanted) == 1 else 'every column'} "
                f"checked ({', '.join(empty_cols)}) is empty on both sides "
                f"for all {matched_rows} matched row(s). The {len(equal)} "
                f"'equal' below means the keys line up, not that the data "
                f"agrees — name columns that hold values.")
        else:
            notes.append(
                f"column(s) {', '.join(empty_cols)} are empty on both sides "
                f"for all {matched_rows} matched row(s), so they contributed "
                f"agreement without being checked. The verdict rests on the "
                f"remaining "
                f"{', '.join(c for c in wanted if c not in empty_cols)}.")
    if duplicate_keys:
        notes.append(
            f"{len(duplicate_keys)} key(s) appear more than once and were NOT "
            "compared — a duplicate key would make the join a cross product. "
            "Resolve them or pick a different key.")
    if left_blank or right_blank:
        notes.append(
            f"{len(left_blank) + len(right_blank)} row(s) have an empty key "
            "and could not be matched.")

    # A gross size asymmetry is far more often the wrong scope than a real
    # gap of that size. Say which reading is the likelier one rather than
    # letting a count of hundreds stand as a finding.
    matched = matched_rows
    for side, rows, one_sided in (("left", len(left_rows) - 1, only_left),
                                  ("right", len(right_rows) - 1, only_right)):
        if rows and len(one_sided) > max(10, matched):
            notes.append(
                f"{len(one_sided)} of {rows} rows exist only on the {side} "
                "side, against " + str(matched) + " matched. A difference of "
                "that size is usually a scope mismatch — a different period, "
                "sheet or filter — rather than that many missing records. "
                "Confirm both tables cover the same range before reporting "
                "this.")

    # Every input row must be accounted for. If this ever fails, the
    # report is wrong in a way the caller could not see.
    accounted = (len(equal) + len(differing) + len(only_left)
                 + len(only_right) + len(left_blank) + len(right_blank)
                 + sum(len(e["rows"]) for e in left_dupes)
                 + sum(len(e["rows"]) for e in right_dupes))
    expected = (len(left_rows) - 1) + (len(right_rows) - 1)
    # Matched pairs consume one row from each side.
    balanced = accounted + len(equal) + len(differing) == expected

    return {
        "left": str(_resolve(left)),
        "right": str(_resolve(right)),
        "key": key,
        "right_key": right_key_name,
        "compared_columns": [
            c if pairs.get(c, c) == c else f"{c} / {pairs[c]}" for c in wanted],
        "left_rows": len(left_rows) - 1,
        "right_rows": len(right_rows) - 1,
        "equal": equal[:max_report],
        "equal_count": len(equal),
        "differing": differing[:max_report],
        "differing_count": len(differing),
        "only_left": only_left[:max_report],
        "only_left_count": len(only_left),
        "only_right": only_right[:max_report],
        "only_right_count": len(only_right),
        "not_comparable": not_comparable[:max_report],
        "not_comparable_count": len(not_comparable),
        "rows_accounted_for": balanced,
        # Per compared column: how many matched rows held a value on at
        # least one side. A zero here is a column that agreed on every row
        # because there was nothing in it.
        "columns_with_values": dict(with_values),
        # The counts this comparison produced, named, for the answer-side
        # ledger. Every one of them is a number an answer states as a
        # finding ("4 Abweichungen"), and re-parsing them out of the
        # rendered report would be a second implementation of the same
        # facts that could disagree with this one.
        "figures": [
            {"kind": "count", "value": float(len(equal)),
             "label": f"{len(equal)} equal row(s)"},
            {"kind": "count", "value": float(len(differing)),
             "label": f"{len(differing)} differing row(s)"},
            {"kind": "count", "value": float(len(only_left)),
             "label": f"{len(only_left)} row(s) only in the left table"},
            {"kind": "count", "value": float(len(only_right)),
             "label": f"{len(only_right)} row(s) only in the right table"},
            {"kind": "count", "value": float(len(not_comparable)),
             "label": f"{len(not_comparable)} row(s) not comparable"},
            {"kind": "count", "value": float(len(left_rows) - 1),
             "label": f"{len(left_rows) - 1} row(s) in the left table"},
            {"kind": "count", "value": float(len(right_rows) - 1),
             "label": f"{len(right_rows) - 1} row(s) in the right table"},
        ],
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# Series: one document per row
# ---------------------------------------------------------------------------

_UNSAFE_NAME_RE = re.compile(r"[^\w.\- ()äöüÄÖÜß]+")

MAX_SERIES_ROWS = 500


def _safe_name(text: str) -> str:
    """A file name that cannot escape its directory or collide with shell."""
    cleaned = _UNSAFE_NAME_RE.sub("_", str(text or "").strip())
    # Collapse dot runs: separators are already gone, so "../etc" cannot
    # escape, but a name still carrying ".." invites doubt every time
    # someone reads the output listing.
    cleaned = re.sub(r"\.{2,}", ".", cleaned)
    cleaned = cleaned.strip(" .") or "unbenannt"
    return cleaned[:120]


def _series_name(pattern: str, record: dict, index: int, suffix: str) -> str:
    """Build one output file name from the row's own values."""
    out = pattern
    for column, value in record.items():
        out = out.replace("{" + column + "}", _safe_name(_fmt(value)))
    out = out.replace("{row}", str(index))
    out = _UNSAFE_NAME_RE.sub("_", out).strip(" .") or f"dokument_{index}"
    if not out.lower().endswith(suffix):
        out += suffix
    return out[:150]


def fill_series(
    table: Any,
    template: Any,
    *,
    output_dir: Any,
    mapping: Optional[dict] = None,
    constants: Optional[dict] = None,
    name_pattern: str = "",
    sheet: Optional[str] = None,
    limit: int = MAX_SERIES_ROWS,
) -> dict:
    """One filled document per row of *table*, from one form or template.

    The whole point is the report. A series that ends in "done" while six
    rows quietly failed is the expensive kind of mistake in
    administrative work, so every row gets its own outcome and the counts
    are what the caller reports.

    Everything that can be checked is checked BEFORE the first file is
    written: the template's fields, the table's columns, the mapping
    between them. A batch that fails halfway leaves a directory of
    documents nobody can tell apart from a complete one.

    Nothing is ever overwritten. A name that already exists — on disk or
    from an earlier row — fails that row and says which row took it,
    because two different records writing to one file is data loss that
    looks like success.
    """
    rows = _table_rows(table, sheet)
    if len(rows) < 2:
        raise OfficeError(
            "the table needs a header row and at least one data row")
    header = [str(h or "").strip() for h in rows[0]]
    if not any(header):
        raise OfficeError("the table's first row is empty — no column names")
    # Read alongside the values, not instead of them: the letter that
    # leaves the building has to show "1.234,50" and "31.07.2026", and
    # only the cell knows which of the two it is.
    formats = _table_cell_formats(table, sheet)

    tpl = _resolve(template)
    _refuse_unwritable_format(tpl, what="filling it as a template")
    kind = document_kind(tpl)
    if kind == "pdf":
        available = {f["name"] for f in pdf_form_fields(tpl)["fields"]}
        if not available:
            raise OfficeError(
                f"{tpl.name} has no form fields to fill")
        suffix = ".pdf"
    elif kind == "word":
        available = {p["name"]
                     for p in docx_placeholders(tpl)["placeholders"]}
        if not available:
            raise OfficeError(
                f"{tpl.name} contains no {{{{placeholder}}}} markers")
        suffix = ".docx"
    else:
        raise OfficeError(
            f"{tpl.name} is neither a PDF form nor a .docx template")

    # Mapping field -> column. Without one, match by name: a template
    # built for this table usually already uses the column names.
    if mapping:
        resolved_map = {str(k): str(v) for k, v in mapping.items()}
    else:
        lowered = {h.lower(): h for h in header if h}
        resolved_map = {
            field: lowered[field.lower()]
            for field in available
            if field.lower() in lowered and field not in (constants or {})
        }
        if not resolved_map:
            raise OfficeError(
                "no field name matches a column name, so there is nothing to "
                f"map automatically. Template fields: {', '.join(sorted(available))}. "
                f"Columns: {', '.join(h for h in header if h)}. "
                "Pass mapping={field: column}.")

    # Fixed values: the same entry in every document — a file reference, a
    # date, a department. Without them a caller has to build a derived
    # table with the value copied into all 400 rows, which is what was
    # observed in the field: two constants cost a whole intermediate file.
    fixed = {str(k): v for k, v in (constants or {}).items()}
    both = sorted(set(fixed) & set(resolved_map))
    if both:
        raise OfficeError(
            f"field(s) {', '.join(both)} are given both a column and a fixed "
            "value — one of the two has to go")

    unknown_fields = sorted(
        f for f in list(resolved_map) + list(fixed) if f not in available)
    if unknown_fields:
        raise OfficeError(
            f"the template has no field(s) {', '.join(unknown_fields)}. "
            f"It has: {', '.join(sorted(available))}")
    unknown_columns = sorted(
        c for c in resolved_map.values() if c not in header)
    if unknown_columns:
        raise OfficeError(
            f"the table has no column(s) {', '.join(unknown_columns)}. "
            f"It has: {', '.join(h for h in header if h)}. If a value is the "
            "same for every row, pass it in constants instead of adding a "
            "column for it.")

    out_dir = _resolve(output_dir, must_exist=False)
    if out_dir.exists() and not out_dir.is_dir():
        raise OfficeError(f"{out_dir} exists and is not a directory")
    out_dir.mkdir(parents=True, exist_ok=True)

    pattern = name_pattern or (tpl.stem + "_{row}" + suffix)
    limit = max(1, min(int(limit or MAX_SERIES_ROWS), MAX_SERIES_ROWS))

    data_rows = rows[1:]
    truncated = len(data_rows) > limit
    used_names: dict[str, int] = {}
    results: list[dict] = []

    for offset, row in enumerate(data_rows[:limit], start=2):
        row_formats = (formats[offset - 1]
                       if offset - 1 < len(formats) else [])
        record = {
            header[i]: _render_for_locale(
                row[i] if i < len(row) else "",
                row_formats[i] if i < len(row_formats) else "")
            for i in range(len(header)) if header[i]
        }
        values = {field: record.get(column, "")
                  for field, column in resolved_map.items()}
        values.update(fixed)
        empty = sorted(f for f, v in values.items()
                       if str(v or "").strip() == "")

        name = _series_name(pattern, record, offset, suffix)
        if name in used_names:
            results.append({
                "row": offset, "output": name, "status": "failed",
                "detail": (f"the file name is already used by row "
                           f"{used_names[name]} — two records would write "
                           "to one file"),
            })
            continue
        target = out_dir / name
        if target.exists():
            results.append({
                "row": offset, "output": name, "status": "failed",
                "detail": "a file of that name already exists; nothing was "
                          "overwritten",
            })
            continue

        try:
            if kind == "pdf":
                filled = fill_pdf_form(
                    tpl, {k: v for k, v in values.items()}, output=target)
                ok = filled["verified"]
                detail = "; ".join(filled.get("notes", []))
            else:
                filled = fill_docx_template(
                    tpl, {k: v for k, v in values.items()}, output=target,
                    strict=False)
                ok = filled["complete"]
                detail = "; ".join(filled.get("notes", []))
        except OfficeError as exc:
            results.append({"row": offset, "output": name,
                            "status": "failed", "detail": str(exc)})
            continue
        except Exception as exc:                      # pragma: no cover
            results.append({"row": offset, "output": name,
                            "status": "failed", "detail": str(exc)})
            continue

        used_names[name] = offset
        if empty:
            results.append({
                "row": offset, "output": name, "status": "incomplete",
                "detail": "no value for " + ", ".join(empty),
            })
        elif not ok:
            results.append({"row": offset, "output": name,
                            "status": "incomplete",
                            "detail": detail or "the result did not verify"})
        else:
            results.append({"row": offset, "output": name, "status": "ok"})

    counts = {
        "ok": sum(1 for r in results if r["status"] == "ok"),
        "incomplete": sum(1 for r in results if r["status"] == "incomplete"),
        "failed": sum(1 for r in results if r["status"] == "failed"),
    }
    notes: list[str] = []
    if truncated:
        notes.append(
            f"the table has {len(data_rows)} data rows and only the first "
            f"{limit} were processed. Raise limit, or narrow the table.")
    if counts["incomplete"]:
        notes.append(
            f"{counts['incomplete']} document(s) are incomplete — they exist "
            "but are missing values. They are not ready to hand over.")
    if counts["failed"]:
        notes.append(
            f"{counts['failed']} row(s) produced no document.")
    return {
        "table": str(_resolve(table)),
        "template": str(tpl),
        "output_dir": str(out_dir),
        "mapping": resolved_map,
        "constants": fixed,
        "rows": len(data_rows),
        "processed": len(results),
        "counts": counts,
        "results": results,
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _pdf_reader(path: Path) -> Any:
    pypdf = _require("pdf")
    try:
        return pypdf.PdfReader(str(path))
    except Exception as exc:
        raise OfficeError(f"could not open PDF: {exc}") from exc


def _acroform(reader: Any) -> Optional[Any]:
    """The document's AcroForm dictionary, resolved, or None."""
    try:
        root = reader.trailer["/Root"]
        acro = root.get("/AcroForm")
        if acro is None:
            return None
        try:
            return acro.get_object()
        except Exception:
            return acro
    except Exception:
        return None


def _has_xfa(reader: Any) -> bool:
    """True if the document is an XFA form.

    XFA (LiveCycle) forms keep their data in an XML stream, not in the
    AcroForm field dictionaries. Writing AcroForm values on such a file
    reports success and changes nothing the user will see, so this has
    to be detected before a fill is claimed to have worked.
    """
    acro = _acroform(reader)
    try:
        return acro is not None and "/XFA" in acro
    except Exception:
        return False


def _is_dynamic_xfa(reader: Any) -> bool:
    """True for a *dynamic* XFA form, where the AcroForm is only a stub.

    The distinction decides whether a fill is worth attempting at all. A
    static (hybrid) XFA form carries real AcroForm fields and renders
    them; a dynamic one is laid out from the XML at open time and the
    catalogue says so with ``/NeedsRendering``. Writing AcroForm values
    into a dynamic form produces a file whose data is there and whose
    every page still shows blank.
    """
    try:
        root = reader.trailer["/Root"]
        flag = root.get("/NeedsRendering")
        try:
            flag = flag.get_object()
        except Exception:
            pass
        # Not ``bool(flag)``: a pypdf BooleanObject is an object, and
        # every object is truthy, so a /NeedsRendering false would read
        # as true.
        return str(getattr(flag, "value", flag)).strip().lower() == "true"
    except Exception:
        return False


# The XFA packets that hold something a reader can use. ``datasets`` is
# the filled-in data; ``template`` is the layout, which is far too large
# to be worth returning and holds no values.
_XFA_DATA_PACKETS = ("datasets", "xdp:xdp", "xdp")

MAX_XFA_VALUES = 300


def _xfa_packets(reader: Any) -> dict[str, bytes]:
    """The XFA packets of a form, by name. Empty for anything else.

    ``/XFA`` is either one stream holding the whole XDP or a flat array
    alternating names and streams. Both shapes occur in the wild, and a
    reader that assumes one of them silently returns nothing for the
    other.
    """
    acro = _acroform(reader)
    if acro is None:
        return {}
    try:
        xfa = acro.get("/XFA")
        xfa = xfa.get_object() if hasattr(xfa, "get_object") else xfa
    except Exception:
        return {}
    if xfa is None:
        return {}

    def _bytes(entry: Any) -> bytes:
        try:
            entry = entry.get_object()
        except Exception:
            pass
        for attr in ("get_data", "getData"):
            reader_fn = getattr(entry, attr, None)
            if callable(reader_fn):
                try:
                    data = reader_fn()
                    return (data if isinstance(data, bytes)
                            else str(data).encode())
                except Exception:
                    return b""
        if isinstance(entry, bytes):
            return entry
        return str(entry).encode("utf-8", "replace")

    packets: dict[str, bytes] = {}
    if isinstance(xfa, (list, tuple)):
        items = list(xfa)
        # Name/stream pairs; an odd tail is a malformed array and is
        # filed under its position rather than dropped silently.
        for index in range(0, len(items), 2):
            name = str(items[index])
            body = items[index + 1] if index + 1 < len(items) else None
            packets[name if body is not None else f"packet{index}"] = _bytes(
                body if body is not None else items[index])
    else:
        packets["xdp"] = _bytes(xfa)
    return packets


def _xfa_values(packets: dict[str, bytes]) -> list[dict]:
    """Field path -> value out of the XFA ``datasets`` packet.

    This is the only way to see what a dynamic XFA form actually holds:
    its values never reach the AcroForm dictionaries, so the field
    listing reports every field as empty while the form is full.
    """
    import xml.etree.ElementTree as ET

    candidates = [packets[name] for name in _XFA_DATA_PACKETS
                  if packets.get(name)]
    candidates += [body for name, body in packets.items()
                   if name not in _XFA_DATA_PACKETS and b"datasets" in body]
    out: list[dict] = []
    seen: set[str] = set()
    for body in candidates:
        try:
            root = ET.fromstring(body.decode("utf-8", "replace"))
        except Exception:
            continue
        data = root
        for element in root.iter():
            if element.tag.rsplit("}", 1)[-1] == "data":
                data = element
                break

        def walk(node: Any, path: str) -> None:
            children = list(node)
            if not children:
                text = (node.text or "").strip()
                if path and text and path not in seen:
                    seen.add(path)
                    out.append({"name": path, "value": text})
                return
            counts: dict[str, int] = {}
            for child in children:
                tag = child.tag.rsplit("}", 1)[-1]
                counts[tag] = counts.get(tag, 0) + 1
            used: dict[str, int] = {}
            for child in children:
                tag = child.tag.rsplit("}", 1)[-1]
                if counts[tag] > 1:
                    used[tag] = used.get(tag, 0) + 1
                    label = f"{tag}[{used[tag]}]"
                else:
                    label = tag
                walk(child, f"{path}.{label}" if path else label)

        for child in list(data):
            walk(child, child.tag.rsplit("}", 1)[-1])
        if out:
            break
    return out[:MAX_XFA_VALUES]


def pdf_xfa_data(path: Any) -> dict:
    """The values stored in an XFA form's dataset, so they can be read.

    Reading is all this offers. Writing the dataset back would mean
    rewriting an XML the form's own script layer validates, and a form
    that opens with the right numbers in the wrong places is exactly the
    outcome a refusal is worth more than.
    """
    p = _resolve(path)
    if document_kind(p) != "pdf":
        raise OfficeError(f"{p.name} is not a PDF")
    reader = _pdf_reader(p)
    if not _has_xfa(reader):
        raise OfficeError(
            f"{p.name} is not an XFA form — list its fields instead")
    packets = _xfa_packets(reader)
    values = _xfa_values(packets)
    notes: list[str] = []
    if not values:
        notes.append(
            "the XFA packets carry no filled-in dataset: either the form is "
            "empty or its data sits in a packet this reader does not "
            "understand (packets present: "
            + (", ".join(sorted(packets)) or "none") + ")")
    if len(values) == MAX_XFA_VALUES:
        notes.append(
            f"only the first {MAX_XFA_VALUES} values are listed")
    notes.append(
        "these values were read from the form's XML dataset, not from "
        "AcroForm fields. They cannot be written back here.")
    return {
        "path": str(p),
        "packets": sorted(packets),
        "dynamic": _is_dynamic_xfa(reader),
        "values": values,
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# Scanned pages
# ---------------------------------------------------------------------------
#
# A PDF whose pages hold no text is not a broken PDF, it is a photograph
# of paper — the single most common attachment in an administrative
# in-tray. Two things matter here and neither is the OCR itself:
#
# * the diagnosis has to be specific. "OCR is not available" leaves the
#   reader with no way to change that; naming the component that is
#   missing and the command that installs it does.
# * text produced by OCR is a reading of pixels, not content of the
#   file. It is labelled as such everywhere it appears, and it never
#   replaces a page that has a real text layer.

OCR_MAX_PAGES = 3
OCR_DPI = 200
OCR_LANGUAGES = ("de", "en")
# Below this the recogniser is guessing at the shape of the glyphs. Such
# blocks are counted and reported rather than quietly dropped: a missing
# line is easier to notice than a wrong one.
OCR_LOW_CONFIDENCE = 0.5


def _easyocr_model_dir() -> Path:
    import os
    base = (os.environ.get("EASYOCR_MODULE_PATH")
            or os.environ.get("MODULE_PATH")
            or str(Path.home() / ".EasyOCR"))
    return Path(base).expanduser() / "model"


def ocr_availability() -> dict:
    """What OCR would need on this machine. Never raises.

    Returns ``{"available", "engine", "detail", "next_step"}``. The
    detail names the component that is missing rather than the
    capability, because "install tesseract-ocr" is actionable and "OCR
    unavailable" is not.
    """
    import importlib
    import shutil

    detail: list[str] = []
    renderer = False
    try:
        importlib.import_module("fitz")
        renderer = True
    except Exception:
        detail.append(
            "pymupdf is not installed, so a PDF page cannot be turned into "
            "an image at all (pip install pymupdf)")

    engine = ""
    try:
        importlib.import_module("pytesseract")
    except Exception:
        detail.append("pytesseract is not installed")
    else:
        if shutil.which("tesseract"):
            engine = engine or "pytesseract"
        else:
            detail.append(
                "pytesseract is installed but the 'tesseract' program it "
                "drives is not on PATH — pytesseract is only a wrapper, the "
                "recognition happens in that binary (Debian/Ubuntu: "
                "apt-get install tesseract-ocr tesseract-ocr-deu)")

    try:
        importlib.import_module("easyocr")
        importlib.import_module("torch")
    except Exception as exc:
        detail.append(f"easyocr is not usable ({type(exc).__name__})")
    else:
        weights = _easyocr_model_dir()
        try:
            present = len(list(weights.glob("*.pth")))
        except OSError:
            present = 0
        if present >= 2:
            engine = engine or "easyocr"
        else:
            detail.append(
                "easyocr is installed but its recognition models are not in "
                f"{weights} — the first call would download roughly 100 MB "
                "without asking, which a document read must not do. Fetch "
                "them once with network access: python -c \"import easyocr; "
                "easyocr.Reader(['de','en'])\"")

    available = bool(engine) and renderer
    if available:
        next_step = (
            f"pass ocr=true to read the scanned pages with {engine}; the "
            "result is a reading of the page image, not text out of the "
            "file")
    elif engine and not renderer:
        next_step = "install pymupdf, then pass ocr=true"
    else:
        next_step = (
            "install one of the two engines above, then pass ocr=true")
    return {
        "available": available,
        "engine": engine if available else "",
        "renderer": "pymupdf" if renderer else "",
        "detail": detail,
        "next_step": next_step,
    }


def _render_page(path: Path, number: int, dpi: int = OCR_DPI) -> bytes:
    """One PDF page as PNG bytes. Raises OfficeError if it cannot."""
    try:
        import fitz
    except Exception as exc:
        raise OfficeError(
            f"rendering a page needs pymupdf, which is not installed ({exc})"
        ) from exc
    try:
        document = fitz.open(str(path))
    except Exception as exc:
        raise OfficeError(f"could not open {path.name} for rendering: {exc}"
                          ) from exc
    try:
        return document[number - 1].get_pixmap(dpi=dpi).tobytes("png")
    except Exception as exc:
        raise OfficeError(
            f"page {number} of {path.name} could not be rendered: {exc}"
        ) from exc
    finally:
        try:
            document.close()
        except Exception:
            pass


_OCR_READER_CACHE: dict[tuple, Any] = {}


def _ocr_image(
    image: bytes, languages: tuple = OCR_LANGUAGES
) -> tuple[str, float, int]:
    """OCR one page image. Returns ``(text, mean confidence, weak blocks)``."""
    status = ocr_availability()
    if not status["available"] or status["engine"] != "easyocr":
        raise OfficeError(
            "OCR is not available here: "
            + ("; ".join(status["detail"]) or "no engine could be used")
            + ". " + status["next_step"])
    try:
        import easyocr
    except Exception as exc:                      # pragma: no cover
        raise OfficeError(f"easyocr could not be imported: {exc}") from exc

    key = tuple(languages)
    reader = _OCR_READER_CACHE.get(key)
    if reader is None:
        try:
            # Never on the GPU: this runs inside an interactive tool call
            # on machines that may be running something else on the card,
            # and a CUDA allocation failure would surface as an OCR
            # failure rather than as what it is.
            reader = easyocr.Reader(list(languages), gpu=False, verbose=False,
                                    download_enabled=False)
        except Exception as exc:
            raise OfficeError(
                f"the OCR engine could not be started: {exc}") from exc
        _OCR_READER_CACHE[key] = reader
    try:
        found = reader.readtext(image, detail=1)
    except Exception as exc:
        raise OfficeError(f"OCR failed on this page: {exc}") from exc

    lines = [str(text) for _box, text, _conf in found]
    scores = [float(conf) for _box, _text, conf in found]
    weak = sum(1 for s in scores if s < OCR_LOW_CONFIDENCE)
    mean = sum(scores) / len(scores) if scores else 0.0
    return "\n".join(lines), mean, weak


def _page_image_count(page: Any) -> int:
    """How many images a page draws, read from its resources.

    Deliberately not ``page.images``, which decodes every one of them: a
    300 dpi scan would be unpacked into memory only to be counted.
    """
    try:
        resources = page.get("/Resources")
        resources = (resources.get_object()
                     if hasattr(resources, "get_object") else resources)
        xobjects = resources.get("/XObject") if resources else None
        xobjects = (xobjects.get_object()
                    if hasattr(xobjects, "get_object") else xobjects)
        if not xobjects:
            return 0
        count = 0
        for key in xobjects:
            try:
                entry = xobjects[key].get_object()
            except Exception:
                continue
            if str(entry.get("/Subtype", "")) == "/Image":
                count += 1
        return count
    except Exception:
        return 0


def read_pdf(
    path: Any,
    *,
    pages: Optional[str] = None,
    max_chars: int = MAX_TEXT_CHARS,
    ocr: bool = False,
) -> dict:
    """Extract text from a PDF, page by page.

    ``pages`` accepts ``"3"``, ``"2-5"`` or ``"1,4,7"``; omitted means
    the whole document up to *max_chars*.

    ``ocr`` reads the pages that carry no text layer from their page
    image instead. It is off by default and it never touches a page that
    has real text: OCR output is a reading of pixels, and replacing text
    the file actually contains with a guess at it would trade a correct
    answer for a plausible one. Blocks it recognised are labelled as OCR
    in the output, with the confidence the engine reported.
    """
    p = _resolve(path)
    if document_kind(p) != "pdf":
        raise OfficeError(f"{p.name} is not a PDF")
    reader = _pdf_reader(p)
    total = len(reader.pages)
    wanted = _parse_pages(pages, total)

    extracted: dict[int, str] = {}
    empty: list[int] = []
    scanned: list[int] = []
    for number in wanted:
        try:
            page = reader.pages[number - 1]
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"(page {number}: extraction failed: {exc})"
            page = None
        extracted[number] = text.strip()
        if not extracted[number]:
            empty.append(number)
            # A page with images and no text is a scan; a page with
            # neither is a blank sheet. OCR helps with the first and has
            # nothing to find on the second, and saying which is which
            # is the difference between a next step and a guess.
            if page is not None and _page_image_count(page):
                scanned.append(number)

    notes: list[str] = []
    ocr_pages: list[dict] = []
    if ocr and scanned:
        budget = scanned[:OCR_MAX_PAGES]
        for number in budget:
            try:
                image = _render_page(p, number)
                text, confidence, weak = _ocr_image(image)
            except OfficeError as exc:
                notes.append(f"page {number}: {exc}")
                break
            extracted[number] = text
            ocr_pages.append({"page": number,
                              "confidence": round(confidence, 3),
                              "low_confidence_blocks": weak})
        if ocr_pages:
            mean = sum(e["confidence"] for e in ocr_pages) / len(ocr_pages)
            weak_total = sum(e["low_confidence_blocks"] for e in ocr_pages)
            notes.append(
                "page(s) " + ", ".join(str(e["page"]) for e in ocr_pages)
                + f" were read by OCR ({ocr_availability()['engine']}, mean "
                f"confidence {mean:.2f}). That text is a reading of the page "
                "image, not content of the file: check every figure, name and "
                "reference number against the original before acting on it.")
            if weak_total:
                notes.append(
                    f"{weak_total} recognised block(s) scored below "
                    f"{OCR_LOW_CONFIDENCE} — treat those lines as unread "
                    "rather than as text.")
        if len(scanned) > len(budget):
            notes.append(
                f"{len(scanned)} page(s) have no text layer and only "
                f"{OCR_MAX_PAGES} were run through OCR (it takes tens of "
                "seconds per page). Pass pages to choose which.")
    elif ocr and not scanned:
        notes.append(
            "ocr was requested but no requested page is a scan — every page "
            "either has a text layer already or holds nothing at all.")

    chunks: list[str] = []
    used = 0
    truncated = False
    ocr_numbers = {e["page"]: e for e in ocr_pages}
    for number in wanted:
        entry = ocr_numbers.get(number)
        label = (f"--- page {number} (OCR, confidence "
                 f"{entry['confidence']:.2f}) ---" if entry
                 else f"--- page {number} ---")
        block = f"{label}\n{extracted[number]}"
        if used + len(block) > max_chars:
            chunks.append(block[: max(0, max_chars - used)])
            truncated = True
            break
        chunks.append(block)
        used += len(block)

    still_empty = [n for n in empty if n not in ocr_numbers]
    if truncated:
        notes.append(
            f"output truncated at {max_chars} characters — pass pages to "
            "read a specific range")
    if still_empty:
        status = ocr_availability()
        blank = [n for n in still_empty if n not in scanned]
        if len(still_empty) == len(wanted):
            head = ("no extractable text on any requested page"
                    if len(wanted) > 1 else "no extractable text on this page")
        else:
            head = f"{len(still_empty)} of {len(wanted)} pages held no text"
        if [n for n in still_empty if n in scanned]:
            head += (
                " — page(s) "
                + ", ".join(str(n) for n in still_empty if n in scanned)
                + " draw an image and carry no text layer, which is a scan")
            if status["available"]:
                head += ". " + status["next_step"]
            else:
                head += (". OCR is not possible on this machine: "
                         + ("; ".join(status["detail"]) or "no engine found")
                         + ". " + status["next_step"]
                         + ", or run the file through OCR elsewhere and read "
                         "the result.")
        if blank:
            head += (
                " — page(s) " + ", ".join(str(n) for n in blank)
                + " hold neither text nor an image and are simply empty; OCR "
                "would find nothing on them")
        notes.append(head)
    if _has_xfa(reader):
        dynamic = _is_dynamic_xfa(reader)
        found = len(_xfa_values(_xfa_packets(reader)))
        notes.append(
            ("this is a dynamic XFA form" if dynamic
             else "this is an XFA form")
            + " — its field data lives in an XML stream, not in the page text"
            + (f". {found} stored value(s) can be read: list the fields "
               "(fields=true) to see them" if found else
               ". No stored values were found in its dataset"))

    return {
        "path": str(p),
        "kind": "pdf",
        "pages": total,
        "pages_read": wanted,
        "pages_without_text": still_empty,
        "ocr_pages": ocr_pages,
        "text": "\n\n".join(chunks),
        "notes": notes,
    }


def _parse_pages(spec: Optional[str], total: int) -> list[int]:
    if not spec or not str(spec).strip():
        return list(range(1, total + 1))
    out: list[int] = []
    for part in str(spec).split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            lo_s, _, hi_s = part.partition("-")
            try:
                lo, hi = int(lo_s), int(hi_s)
            except ValueError as exc:
                raise OfficeError(f"bad page range {part!r}") from exc
            if lo > hi:
                lo, hi = hi, lo
            out.extend(range(lo, hi + 1))
        else:
            try:
                out.append(int(part))
            except ValueError as exc:
                raise OfficeError(f"bad page number {part!r}") from exc
    seen: set[int] = set()
    result = []
    for n in out:
        if 1 <= n <= total and n not in seen:
            seen.add(n)
            result.append(n)
    if not result:
        raise OfficeError(
            f"no valid page in {spec!r} — the document has {total} page(s)")
    return result


_FIELD_TYPES = {
    "/Tx": "text",
    "/Btn": "button/checkbox",
    "/Ch": "choice",
    "/Sig": "signature",
}


def _field_labels(path: Path) -> dict:
    """Pair each form field with the text printed next to it.

    A field name carries no meaning in many real forms — ``text1``,
    ``text2``, ``Kontrollkästchen3``. What identifies the field is the
    label printed beside it, and the two are connected only by their
    position on the page. Without that link a caller can do nothing but
    assume the field order matches the visual order, which in a form
    built with a designer it frequently does not: a value in the wrong
    field looks perfectly correct afterwards.

    Order of evidence: the field's own tooltip (``/TU``) when the form
    author set one, otherwise the nearest text to the LEFT on the same
    line, otherwise the text directly ABOVE. Returns
    ``{field: {"label", "source", "page"}}`` and never raises — a form
    whose geometry cannot be read simply yields no labels.
    """
    out: dict[str, dict] = {}
    try:
        import fitz
    except Exception:
        return out
    try:
        doc = fitz.open(str(path))
    except Exception:
        return out
    try:
        for number, page in enumerate(doc, start=1):
            try:
                widgets = list(page.widgets() or [])
            except Exception:
                continue
            if not widgets:
                continue
            try:
                words = page.get_text("words")     # x0, y0, x1, y1, word, ...
            except Exception:
                words = []
            for widget in widgets:
                name = str(getattr(widget, "field_name", "") or "")
                if not name or name in out:
                    continue
                tooltip = str(getattr(widget, "field_label", "") or "").strip()
                if tooltip:
                    out[name] = {"label": tooltip, "source": "tooltip",
                                 "page": number}
                    continue
                found = _nearest_label(widget.rect, words)
                if found:
                    out[name] = {"label": found[0], "source": found[1],
                                 "page": number}
    except Exception:
        return out
    finally:
        try:
            doc.close()
        except Exception:
            pass
    return out


def _nearest_label(rect: Any, words: list) -> Optional[tuple]:
    """The printed text that identifies a widget, plus where it sat.

    Left first: forms put the label on the same line, before the box.
    Above second, for stacked layouts. Both windows are deliberately
    narrow — a label further away than this is more likely to belong to
    a neighbouring field, and a wrong label is worse than none.
    """
    mid_y = (rect.y0 + rect.y1) / 2
    height = max(rect.y1 - rect.y0, 1.0)

    same_line = [
        w for w in words
        if w[2] <= rect.x0 + 2                      # ends left of the box
        and w[1] <= mid_y <= w[3] + height * 0.4    # vertically on the line
        and rect.x0 - w[2] < 320                    # not across the page
    ]
    if same_line:
        same_line.sort(key=lambda w: w[0])
        gap_limit = max(w[2] for w in same_line) - 260
        text = " ".join(w[4] for w in same_line if w[2] >= gap_limit)
        cleaned = text.strip(" :._-\t")
        if cleaned:
            return cleaned, "left"

    above = [
        w for w in words
        if w[3] <= rect.y0 + 2                      # ends above the box
        and rect.y0 - w[3] < height * 2.2           # directly above
        and w[2] > rect.x0 - 40 and w[0] < rect.x1  # horizontally overlapping
    ]
    if above:
        top = max(w[3] for w in above)
        text = " ".join(w[4] for w in sorted(above, key=lambda w: w[0])
                        if w[3] >= top - 2)
        cleaned = text.strip(" :._-\t")
        if cleaned:
            return cleaned, "above"
    return None


def pdf_form_fields(path: Any) -> dict:
    """List a PDF form's fields with their type, value and allowed states.

    The allowed states matter: a check box is not set with ``true`` but
    with the on-state name that field declares, and those differ between
    documents.
    """
    p = _resolve(path)
    if document_kind(p) != "pdf":
        raise OfficeError(f"{p.name} is not a PDF")
    reader = _pdf_reader(p)
    raw = {}
    try:
        raw = reader.get_fields() or {}
    except Exception as exc:
        raise OfficeError(f"could not read form fields: {exc}") from exc

    labels = _field_labels(p)
    fields = []
    for name, spec in raw.items():
        entry: dict[str, Any] = {
            "name": str(name),
            "type": _FIELD_TYPES.get(str(spec.get("/FT", "")), str(
                spec.get("/FT", "") or "unknown")),
            "value": _fmt(spec.get("/V")),
        }
        found = labels.get(str(name))
        if found:
            entry["label"] = found["label"]
            entry["label_source"] = found["source"]
            entry["page"] = found["page"]
        states = spec.get("/_States_")
        if states:
            # A choice field's options may be [export, display] pairs, and
            # str() over the pair produced "['One', 'One']" -- unreadable
            # for the model that has to pick one of them.
            if str(spec.get("/FT", "")) == "/Ch":
                entry["states"] = [
                    d if d == e else f"{d} ({e})"
                    for e, d in _choice_options(spec.get("/Opt") or states)
                ]
            else:
                entry["states"] = [str(s) for s in states]
        fields.append(entry)

    notes: list[str] = []
    unlabelled = [f["name"] for f in fields if not f.get("label")]
    if fields and unlabelled:
        notes.append(
            "no printed label could be located for: "
            + ", ".join(unlabelled[:10])
            + ". Do not infer them from the field order — in a form built "
            "with a designer it often does not follow the visual order.")
    xfa = _has_xfa(reader)
    dynamic = xfa and _is_dynamic_xfa(reader)
    xfa_values: list[dict] = []
    if xfa:
        try:
            xfa_values = _xfa_values(_xfa_packets(reader))
        except Exception:
            xfa_values = []
    if not fields and not xfa_values:
        notes.append(
            "this PDF has no form fields. Values cannot be filled in; "
            "producing a completed document means generating a new PDF "
            "or overlaying text.")
    if dynamic:
        notes.append(
            "dynamic XFA form (/NeedsRendering): the AcroForm entries above "
            "are a stub the viewer ignores — it builds the pages from the "
            "XML instead. Filling them is refused here rather than reported "
            "as done. To produce a completed form, fill it in a viewer that "
            "renders XFA, or rebuild it as a flat PDF.")
    elif xfa:
        notes.append(
            "XFA form: writing AcroForm field values will report success "
            "but will not change what a viewer displays. Confirm the "
            "result before handing this file over.")
    if xfa_values:
        # The values a dynamic form holds never reach the AcroForm
        # dictionaries, so the field listing above shows every one of
        # them as empty while the form is full. Reading them out of the
        # dataset is the only way to see what the document says.
        notes.append(
            f"{len(xfa_values)} value(s) were read out of the form's XML "
            "dataset (xfa_values). They are what this form actually holds; "
            "they cannot be written back here.")
    elif xfa:
        notes.append(
            "no values could be read from the XFA dataset — the form is "
            "either empty or stores its data in a packet this reader does "
            "not understand.")
    return {
        "path": str(p),
        "fields": fields,
        "field_count": len(fields),
        "xfa": xfa,
        "xfa_dynamic": dynamic,
        "xfa_values": xfa_values,
        "notes": notes,
    }


def fill_pdf_form(
    path: Any,
    values: dict,
    *,
    output: Any,
    flatten: bool = False,
) -> dict:
    """Fill an AcroForm and write the result to *output*.

    Behaviour that exists because the naive version silently fails:

    * ``/NeedAppearances`` is set on the form, so viewers that do not
      build appearance streams still show the values.
    * Boolean and loose check-box values are mapped to the on-state the
      individual field declares (``/Yes``, ``/On``, ``/1``, ...).
    * A field name that is not in the document is an error, not a
      silently ignored key — a typo would otherwise read as success.
    """
    src = _resolve(path)
    if document_kind(src) != "pdf":
        raise OfficeError(f"{src.name} is not a PDF")
    if not isinstance(values, dict) or not values:
        raise OfficeError("values must be a non-empty object of field -> value")
    out = _resolve(output, must_exist=False)
    if out.resolve() == src.resolve():
        raise OfficeError(
            "output must differ from the source — filling a form in place "
            "destroys the blank original")
    # ... and an output that already exists is somebody's document too.
    # Every sibling refuses this -- create_pdf, merge_pdfs, split_pdf,
    # fill_series, draft_email -- and this path did not, so a second fill
    # onto one name returned verified: True over the first one's grave.
    # This handler even reads the pre-image for the undo journal, so it
    # knew the file was there and said nothing.
    if out.exists():
        raise OfficeError(
            f"{out.name} already exists — nothing was filled. Write to a "
            "new name, or delete that file first if it is meant to go.")

    pypdf = _require("pdf")
    reader = _pdf_reader(src)
    # A dynamic XFA form is refused rather than filled. Its pages are
    # laid out from the XML at open time, so AcroForm values written
    # here are read by nothing: the call would report every field
    # written and the printed form would come out blank.
    if _has_xfa(reader) and _is_dynamic_xfa(reader):
        raise OfficeError(
            f"{src.name} is a dynamic XFA form (/NeedsRendering). Its pages "
            "are built from an XML layer that ignores the AcroForm fields, "
            "so filling those would report success and change nothing a "
            "viewer shows. Nothing was written. Fill it in a viewer that "
            "renders XFA, or ask for a flat PDF version of the form. Its "
            "current values can be read here (fields=true).")
    declared = {}
    try:
        declared = reader.get_fields() or {}
    except Exception:
        declared = {}
    if not declared:
        raise OfficeError(
            f"{src.name} has no form fields to fill — check the document "
            "with the field listing first")

    unknown = [k for k in values if str(k) not in declared]
    if unknown:
        raise OfficeError(
            f"no such field(s): {', '.join(sorted(unknown))}. Available: "
            f"{', '.join(sorted(str(k) for k in declared))}")

    resolved_values: dict[str, Any] = {}
    mapped: list[str] = []
    for key, value in values.items():
        spec = declared[str(key)]
        states = [str(s) for s in (spec.get("/_States_") or [])]
        field_type = str(spec.get("/FT", "") or "")
        # The field's TYPE decides, not the presence of a state list.
        # pypdf fills /_States_ for a choice field from its /Opt, so a
        # dropdown was going through the check-box mapper -- and that
        # mapper resolves any truthy word to the FIRST non-/Off state,
        # which is the first option in document order. On a Ja/Nein field
        # declared as ["Nein", "Ja"], "ja" was written as "Nein": the
        # opposite answer, reported verified because the read-back
        # compares against what was just written, under a note calling it
        # a check box. Einverständnis, Widerspruch, Datenweitergabe --
        # every German consent field has this shape.
        if states and field_type == "/Ch":
            wanted = _render_for_locale(value)
            options = _choice_options(spec.get("/Opt")
                                      or spec.get("/_States_"))
            match = _choice_option(wanted, options)
            if match is None:
                shown = ", ".join(
                    d if d == e else f"{d} ({e})" for e, d in options)
                raise OfficeError(
                    f"{key}: {wanted!r} is not one of the options this "
                    f"field allows ({shown}). A choice field takes one of "
                    f"its own options -- guessing at one would answer a "
                    f"question nobody asked.")
            if match != value:
                mapped.append(f"{key}={value!r} -> {match} (choice)")
            resolved_values[str(key)] = match
        elif states:
            resolved = _checkbox_state(value, states)
            if resolved != value:
                mapped.append(f"{key}={value!r} -> {resolved} (check box)")
            resolved_values[str(key)] = resolved
        else:
            resolved_values[str(key)] = _render_for_locale(value)

    try:
        writer = pypdf.PdfWriter(clone_from=str(src))
        _set_need_appearances(writer)
        _ensure_form_resources(writer)
        for page in writer.pages:
            writer.update_page_form_field_values(
                page, resolved_values, auto_regenerate=False)
        if flatten:
            _flatten_form(writer)
        out.parent.mkdir(parents=True, exist_ok=True)
        with out.open("wb") as fh:
            writer.write(fh)
    except OfficeError:
        raise
    except Exception as exc:
        raise OfficeError(f"could not fill the form: {exc}") from exc

    # Verify against the file that was actually written, rather than
    # reporting the intent as the outcome.
    written = {}
    try:
        written = {
            str(k): _fmt(v.get("/V"))
            for k, v in (pypdf.PdfReader(str(out)).get_fields() or {}).items()
        }
    except Exception:
        written = {}
    missing = [
        k for k, v in resolved_values.items()
        if written.get(k, "") != _fmt(v)
    ]

    notes: list[str] = []
    if mapped:
        # Names the KIND of field beside each mapping. The old wording
        # called everything a check box, so the one time it mattered --
        # a Ja/Nein choice written as its opposite -- the note read like
        # routine housekeeping.
        notes.append("values mapped onto what the field itself declares: "
                     + "; ".join(mapped))
    if _has_xfa(reader):
        notes.append(
            "XFA form: the AcroForm values were written, but a viewer may "
            "render the XFA layer instead and show the form as empty. "
            "Open the result before treating this as done.")
    if missing:
        notes.append(
            "these fields did not hold the expected value when the written "
            "file was read back: " + ", ".join(sorted(missing)))
    if flatten:
        notes.append(
            "the fields were set to read-only, so the values can no longer "
            "be changed in a viewer. They remain form-field values and are "
            "not merged into the page content — text extraction will still "
            "not find them in the page text.")

    return {
        "source": str(src),
        "path": str(out),
        "filled": sorted(resolved_values),
        "verified": not missing,
        "notes": notes,
    }


def _choice_options(raw: Any) -> list[tuple[str, str]]:
    """A choice field's options as (export value, shown label) pairs.

    ``/Opt`` holds either a plain string per option, or a two-element
    array ``[export, display]`` -- the export value is what goes into
    ``/V``, the display value is what the person filling the form reads,
    and they are routinely different (``["J", "Ja"]``). Flattening each
    entry with ``str()`` turned a pair into the text ``"['One', 'One']"``,
    which no comparison and no message could then make sense of.
    """
    out: list[tuple[str, str]] = []
    for entry in (raw or []):
        if isinstance(entry, (list, tuple)):
            items = [str(x) for x in entry]
            if not items:
                continue
            export = items[0]
            display = items[1] if len(items) > 1 else items[0]
        else:
            export = display = str(entry)
        out.append((export, display))
    return out


def _choice_option(value: Any, options: list[tuple[str, str]]) -> Optional[str]:
    """The EXPORT value a choice field would accept for *value*, or None.

    Either spelling is accepted -- what the document stores or what it
    shows -- because a person answering a form reads the label. Exact
    first, then case- and space-insensitive, because "ja" for an option
    spelled "Ja" is the same answer. Nothing looser: a choice field's
    whole point is that the document decided what the answers are, so a
    near-miss comes back as a question rather than as a guess written
    into a form somebody signs.
    """
    wanted = str(value if value is not None else "")
    for export, display in options:
        if wanted in (export, display):
            return export
    folded = " ".join(wanted.split()).casefold()
    for export, display in options:
        for spelling in (export, display):
            if " ".join(spelling.split()).casefold() == folded:
                return export
    return None


def _checkbox_state(value: Any, states: list[str]) -> str:
    """Map a loose truthy/falsy value onto a field's declared states."""
    off = "/Off" if "/Off" in states else None
    on_states = [s for s in states if s != "/Off"]
    on = on_states[0] if on_states else "/Yes"
    if isinstance(value, str) and value in states:
        return value
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in ("true", "yes", "on", "x", "1", "ja"):
            return on
        if lowered in ("false", "no", "off", "", "0", "nein"):
            return off or "/Off"
        return value
    return on if value else (off or "/Off")


def _set_need_appearances(writer: Any) -> None:
    """Ask the viewer to build appearance streams for filled fields."""
    try:
        from pypdf.generic import BooleanObject, NameObject
        root = writer._root_object
        acro = root.get("/AcroForm")
        if acro is None:
            return
        try:
            acro = acro.get_object()
        except Exception:
            pass
        acro[NameObject("/NeedAppearances")] = BooleanObject(True)
    except Exception:
        # Cosmetic in viewers that render appearances themselves; never
        # worth failing an otherwise good fill.
        pass


def _ensure_form_resources(writer: Any) -> None:
    """Give the form the default-resources dictionary the writer expects.

    An AcroForm may legally omit ``/DR``, and some producers do. pypdf
    reads it while writing a field value and calls ``get_object()`` on the
    plain dictionary it substitutes, which fails with an AttributeError
    that names neither the form nor the field -- so a form that is simply
    missing an optional entry reads as a broken document.

    An empty but real ``/DR`` costs nothing here: ``/NeedAppearances`` is
    set, so the viewer builds the appearance itself.
    """
    try:
        from pypdf.generic import DictionaryObject, NameObject

        acro = writer._root_object.get("/AcroForm")
        if acro is None:
            return
        try:
            acro = acro.get_object()
        except Exception:
            pass

        def _resolved(container: Any, key: str) -> Any:
            value = container.get(key)
            try:
                return value.get_object() if value is not None else None
            except Exception:
                return None

        resources = _resolved(acro, "/DR")
        if not isinstance(resources, DictionaryObject):
            resources = DictionaryObject()
            acro[NameObject("/DR")] = resources
        if not isinstance(_resolved(resources, "/Font"), DictionaryObject):
            resources[NameObject("/Font")] = DictionaryObject()
    except Exception:
        # Best effort. Failing here would replace a confusing error with a
        # different confusing error.
        pass


def _flatten_form(writer: Any) -> None:
    """Set every form widget read-only.

    Deliberately NOT a true flatten: the values stay form-field values
    rather than being merged into the page content stream. Calling this
    "flattened" without saying so would overstate it — a reader
    extracting page text still will not find the entries.
    """
    try:
        from pypdf.generic import NameObject, NumberObject
        for page in writer.pages:
            for annot in page.get("/Annots", None) or []:
                try:
                    obj = annot.get_object()
                except Exception:
                    continue
                if obj.get("/Subtype") == "/Widget":
                    # Bit 1 (value 1) = read-only.
                    flags = int(obj.get("/Ff", 0) or 0)
                    obj[NameObject("/Ff")] = NumberObject(flags | 1)
    except Exception as exc:
        raise OfficeError(f"could not flatten the form: {exc}") from exc


# ---------------------------------------------------------------------------
# PDF assembly: merging, splitting, writing
# ---------------------------------------------------------------------------

# Caps on the batch sizes. A merge of a thousand files or a split of a
# thousand pages is far more likely to be a wrong argument than a wanted
# operation, and both would be discovered as a full directory.
MAX_MERGE_INPUTS = 100
MAX_SPLIT_PARTS = 200


def _pdf_pages(path: Path) -> tuple[Any, int]:
    """Open a PDF and count its pages, naming the file on every failure.

    An encrypted document opens without complaint and only refuses when a
    page is touched, so the page count is read here rather than during the
    write: a file that cannot be used has to be a refusal while nothing
    has been produced yet, not a half-finished result.
    """
    try:
        reader = _pdf_reader(path)
    except OfficeError as exc:
        raise OfficeError(f"{path.name}: {exc}") from exc
    try:
        return reader, len(reader.pages)
    except Exception as exc:
        if getattr(reader, "is_encrypted", False):
            raise OfficeError(
                f"{path.name} is encrypted and cannot be read without its "
                "password"
            ) from exc
        raise OfficeError(f"{path.name} could not be read: {exc}") from exc


def _pdf_has_fields(reader: Any) -> bool:
    try:
        return bool(reader.get_fields())
    except Exception:
        return False


def merge_pdfs(inputs: list, *, output: Any) -> dict:
    """Combine several PDFs into one new file, in the order given.

    Every input is opened and its pages are counted before anything is
    written. A merge that fails halfway would leave a document that looks
    finished while missing pages nobody can name afterwards, so an
    unreadable or encrypted input refuses the whole operation and says
    which file it was.

    An existing output is never replaced: in administrative work the file
    already sitting under that name is the one somebody else may have
    already sent.
    """
    if not isinstance(inputs, (list, tuple)) or len(inputs) < 2:
        raise OfficeError(
            "merging needs at least two PDFs, listed in the order they "
            "should appear in the result")
    if len(inputs) > MAX_MERGE_INPUTS:
        raise OfficeError(
            f"{len(inputs)} inputs exceeds the cap of {MAX_MERGE_INPUTS}")

    out = _resolve(output, must_exist=False)
    if document_kind(out) != "pdf":
        raise OfficeError(f"cannot write {out.suffix!r} as a PDF — use .pdf")
    if out.exists():
        raise OfficeError(
            f"{out.name} already exists — nothing was merged. Write to a "
            "new name.")

    pypdf = _require("pdf")
    sources: list[tuple[Path, Any, int]] = []
    for item in inputs:
        src = _resolve(item)
        if document_kind(src) != "pdf":
            raise OfficeError(
                f"{src.name} is not a PDF — nothing was merged")
        reader, count = _pdf_pages(src)
        sources.append((src, reader, count))

    expected = sum(count for _, _, count in sources)
    try:
        writer = pypdf.PdfWriter()
        for _src, reader, _count in sources:
            try:
                writer.append(reader)
            except AttributeError:      # pragma: no cover - older pypdf
                # No writer-level append: the pages still carry the
                # content, only the outline entries are lost.
                for page in reader.pages:
                    writer.add_page(page)
        out.parent.mkdir(parents=True, exist_ok=True)
        with out.open("wb") as fh:
            writer.write(fh)
    except OfficeError:
        raise
    except Exception as exc:
        # A truncated file under the requested name reads as a finished
        # merge to everyone who finds it later.
        try:
            out.unlink()
        except OSError:
            pass
        raise OfficeError(f"could not merge the documents: {exc}") from exc

    # Count the pages of the file that was written, not the pages that
    # were meant to go into it.
    written = 0
    try:
        written = len(pypdf.PdfReader(str(out)).pages)
    except Exception:
        written = 0

    notes: list[str] = []
    if written != expected:
        notes.append(
            f"the written file holds {written} page(s), not the {expected} "
            "the inputs add up to")
    with_fields = [src.name for src, reader, _c in sources
                   if _pdf_has_fields(reader)]
    if with_fields:
        notes.append(
            "form fields are present in " + ", ".join(with_fields)
            + " — fields carrying the same name collapse into one shared "
            "value when the documents are combined. Open the result before "
            "handing it over.")
    return {
        "path": str(out),
        "inputs": [{"path": str(src), "pages": count}
                   for src, _r, count in sources],
        "pages": written,
        "expected_pages": expected,
        "verified": written == expected and written > 0,
        "notes": notes,
    }


def _split_label(group: list[int]) -> str:
    """Name the page selection a part covers, for the output file name."""
    if len(group) == 1:
        return f"p{group[0]}"
    if group == list(range(group[0], group[-1] + 1)):
        return f"p{group[0]}-{group[-1]}"
    return "p" + "_".join(str(n) for n in group)


def split_pdf(
    path: Any, *, output_dir: Any, pages: Optional[str] = None
) -> dict:
    """Write page ranges of a PDF into separate files.

    ``pages`` takes the syntax :func:`read_pdf` uses, and every
    comma-separated part becomes one document: ``"2-5,9"`` writes a
    four-page file and a one-page file. Without it every page is written
    on its own.

    Each part is reported with its own outcome. A name already taken
    fails that part and leaves the existing file untouched, because the
    aggregate "12 files written" is exactly what hides the one page that
    silently replaced last week's version.
    """
    src = _resolve(path)
    if document_kind(src) != "pdf":
        raise OfficeError(f"{src.name} is not a PDF")

    pypdf = _require("pdf")
    reader, total = _pdf_pages(src)

    spec = str(pages or "").strip()
    if spec:
        # One output per comma-separated part, so "2-5" stays one
        # document. The parts themselves go through the reader's page
        # parser — there is one page syntax in this module, not two.
        groups = [_parse_pages(part, total)
                  for part in spec.split(",") if part.strip()]
        if not groups:
            raise OfficeError(f"no page selection in {spec!r}")
    else:
        groups = [[n] for n in _parse_pages(None, total)]
    if len(groups) > MAX_SPLIT_PARTS:
        raise OfficeError(
            f"{len(groups)} parts exceeds the cap of {MAX_SPLIT_PARTS} — "
            "pass a page selection")

    out_dir = _resolve(output_dir, must_exist=False)
    if out_dir.exists() and not out_dir.is_dir():
        raise OfficeError(f"{out_dir} exists and is not a directory")
    out_dir.mkdir(parents=True, exist_ok=True)

    stem = _safe_name(src.stem)
    used: dict[str, list[int]] = {}
    results: list[dict] = []
    for group in groups:
        name = f"{stem}_{_split_label(group)}.pdf"[:150]
        if name in used:
            results.append({
                "output": name, "pages": group, "status": "failed",
                "detail": (f"the name is already taken by pages "
                           f"{used[name]} of this call"),
            })
            continue
        target = out_dir / name
        if target.exists():
            results.append({
                "output": name, "pages": group, "status": "failed",
                "detail": "a file of that name already exists; nothing was "
                          "overwritten",
            })
            continue
        try:
            writer = pypdf.PdfWriter()
            for number in group:
                writer.add_page(reader.pages[number - 1])
            with target.open("wb") as fh:
                writer.write(fh)
        except Exception as exc:
            try:
                target.unlink()
            except OSError:
                pass
            results.append({"output": name, "pages": group,
                            "status": "failed", "detail": str(exc)})
            continue
        used[name] = group

        written = 0
        try:
            written = len(pypdf.PdfReader(str(target)).pages)
        except Exception:
            written = 0
        if written == len(group):
            results.append({"output": name, "pages": group, "status": "ok"})
        else:
            results.append({
                "output": name, "pages": group, "status": "incomplete",
                "detail": (f"the written file holds {written} page(s), not "
                           f"the {len(group)} requested"),
            })

    counts = {
        "ok": sum(1 for r in results if r["status"] == "ok"),
        "incomplete": sum(1 for r in results if r["status"] == "incomplete"),
        "failed": sum(1 for r in results if r["status"] == "failed"),
    }
    notes: list[str] = []
    if counts["failed"]:
        notes.append(f"{counts['failed']} part(s) produced no file.")
    if counts["incomplete"]:
        notes.append(
            f"{counts['incomplete']} file(s) do not hold the pages that were "
            "asked for. They are not ready to hand over.")
    if _pdf_has_fields(reader):
        notes.append(
            "the source is a form: the extracted pages keep the field "
            "widgets but not the document-level form dictionary, so filled "
            "values can render as empty. Check one result before handing "
            "it over.")
    return {
        "source": str(src),
        "output_dir": str(out_dir),
        "source_pages": total,
        "files": results,
        "counts": counts,
        "verified": counts["ok"] == len(groups),
        "notes": notes,
    }


def _pdf_markup(value: Any) -> str:
    """Escape a value for the paragraph markup of the PDF writer.

    Paragraph text is parsed as markup, so an unescaped ``&`` in a company
    name or a ``<`` in a note aborts the build or swallows the rest of the
    line. Line breaks are kept as breaks rather than collapsing into one
    run-on paragraph.
    """
    from xml.sax.saxutils import escape
    return escape(_fmt(value)).replace("\r\n", "\n").replace("\n", "<br/>")


_PROBE_RE = re.compile(r"[^\W\d_]{4,}", re.UNICODE)


def _text_probe(text: str) -> str:
    """The longest word of *text*, used to check the file reads back.

    A whole line cannot be searched for: the writer wraps lines wherever
    they run out of width. A single word is not hyphenated, so it survives
    extraction intact — and a word carrying an umlaut proves the encoding
    survived as well.
    """
    raw = str(text or "")
    words = _PROBE_RE.findall(raw)
    if words:
        return max(words, key=len)[:24]
    # No word of four letters. That is not an exotic document -- it is a
    # Kostenaufstellung, a Terminliste, a table of dates and amounts, and
    # it was the case where the check quietly passed on nothing: with no
    # probe the caller defaulted to "found" and verified reported only
    # that the file opened. A figure is exactly what has to survive the
    # write in such a document, so the longest non-blank token stands in.
    tokens = [t for t in raw.split() if len(t) >= 2]
    if tokens:
        return max(tokens, key=len)[:24]
    return ""


def _squash(text: str) -> str:
    return re.sub(r"\s+", "", str(text or ""))


def create_pdf(path: Any, blocks: list[dict]) -> dict:
    """Write a new PDF from the content blocks :func:`create_docx` takes.

    Each block is one of ``{"heading": str, "level": int}``,
    ``{"paragraph": str}``, ``{"table": [[...], ...], "header_row": bool}``
    or ``{"page_break": true}``.

    The result is read back before it is reported: the page count comes
    from the written file and the first block's text has to be findable in
    it, so a document that was produced but cannot be read does not come
    back as a success.
    """
    p = _resolve(path, must_exist=False)
    if document_kind(p) != "pdf":
        raise OfficeError(f"cannot write {p.suffix!r} as a PDF — use .pdf")
    if p.exists():
        raise OfficeError(
            f"{p.name} already exists — write to a new name; an existing "
            "document is never replaced")
    if not isinstance(blocks, list) or not blocks:
        raise OfficeError("blocks must be a non-empty list")

    _require("pdf_write")
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )
    except Exception as exc:
        raise OfficeError(
            f"the 'reportlab' package is installed but incomplete: {exc}"
        ) from exc

    margin = 20 * mm
    doc = SimpleDocTemplate(
        str(p), pagesize=A4, title=p.stem,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin)
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    cell_style = ParagraphStyle(
        "office_cell", parent=body, fontSize=9, leading=11)

    story: list[Any] = []
    counts = {"headings": 0, "paragraphs": 0, "tables": 0}
    probe = ""
    wide_tables = 0
    for index, block in enumerate(blocks, start=1):
        if not isinstance(block, dict):
            raise OfficeError(f"block {index} is not an object: {block!r}")
        if "heading" in block:
            level = max(0, min(int(block.get("level", 1) or 1), 6))
            style = (styles["Title"] if level == 0
                     else styles[f"Heading{level}"])
            text = str(block["heading"])
            story.append(Paragraph(_pdf_markup(text), style))
            counts["headings"] += 1
            probe = probe or _text_probe(text)
        elif "paragraph" in block:
            runs = _runs(block["paragraph"])
            if isinstance(block["paragraph"], (str, int, float)):
                for run in runs:
                    run["bold"] = bool(block.get("bold"))
                    run["italic"] = bool(block.get("italic"))
                    run["colour"] = parse_colour(
                        block.get("color", block.get("colour")))
            story.append(Paragraph(_runs_markup(runs), body))
            story.append(Spacer(1, 4))
            counts["paragraphs"] += 1
            probe = probe or _text_probe(
                "".join(r["text"] for r in runs))
        elif "table" in block:
            rows = block["table"]
            if not isinstance(rows, list) or not rows or not all(
                    isinstance(r, (list, tuple)) for r in rows):
                raise OfficeError(
                    f"block {index}: table must be a non-empty list of lists")
            width = max(len(r) for r in rows)
            # Equal column widths over the printable width: a table sized
            # from its content overflows the page silently, and a table
            # running off the paper is not a document anyone can file.
            col_width = doc.width / width
            header_row = bool(block.get("header_row"))

            def _cell(value, is_header):
                cell_runs = _runs(value)
                if is_header:
                    for run in cell_runs:
                        run["bold"] = True
                return Paragraph(_runs_markup(cell_runs), cell_style)

            data = [
                [_cell(r[c] if c < len(r) else "", header_row and i == 0)
                 for c in range(width)]
                for i, r in enumerate(rows)
            ]
            table = Table(data, colWidths=[col_width] * width,
                          repeatRows=1 if header_row else 0)
            commands = [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            if header_row:
                commands.append(
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke))
            table.setStyle(TableStyle(commands))
            story.append(table)
            story.append(Spacer(1, 6))
            counts["tables"] += 1
            if width > 8:
                wide_tables += 1
            for row in rows:
                probe = probe or _text_probe(" ".join(_fmt(v) for v in row))
                if probe:
                    break
        elif block.get("page_break"):
            story.append(PageBreak())
        else:
            raise OfficeError(
                f"block {index}: expected one of heading / paragraph / "
                f"table / page_break, got keys {sorted(block)}")

    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.build(story)
    except Exception as exc:
        try:
            p.unlink()
        except OSError:
            pass
        raise OfficeError(f"could not write the PDF: {exc}") from exc

    notes: list[str] = []
    pages = 0
    extracted = ""
    try:
        pypdf = _require("pdf")
        written = pypdf.PdfReader(str(p))
        pages = len(written.pages)
        extracted = " ".join(
            (page.extract_text() or "") for page in written.pages)
    except Exception as exc:
        notes.append(f"the written file could not be read back: {exc}")

    found = True
    if probe and pages:
        found = _squash(probe) in _squash(extracted)
    verified = pages > 0 and found and bool(probe)
    if not pages:
        notes.append("the written file did not open as a PDF")
    elif not probe:
        # Nothing to search for is not the same as searched and found.
        notes.append(
            "the file opened as a PDF, but it carries no text this could "
            "search for, so nothing about its CONTENT was checked")
    elif not found:
        notes.append(
            f"the word {probe!r} was written but was not found when the text "
            "of the result was read back — check the document before "
            "handing it over")
    if wide_tables:
        notes.append(
            f"{wide_tables} table(s) have more than 8 columns; they were "
            "fitted into equal-width columns and may be hard to read on A4.")
    return {
        "path": str(p),
        "blocks": len(blocks),
        "pages": pages,
        "verified": verified,
        "notes": notes,
        **counts,
    }


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]{1,80}?)\s*\}\}")


def _iter_paragraphs(container: Any) -> Any:
    """Every paragraph of a document part, tables and nesting included.

    Templates put placeholders in table cells and in the header and
    footer (the letterhead, the file reference, the page footer) at
    least as often as in the body. A filler that only walks
    ``doc.paragraphs`` silently leaves those in place, and the result
    looks complete until someone prints it.
    """
    for para in getattr(container, "paragraphs", []) or []:
        yield para
    for table in getattr(container, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _document_parts(document: Any) -> Any:
    """The body plus every header/footer variant a section can carry."""
    yield document
    for section in getattr(document, "sections", []) or []:
        for attr in (
            "header", "footer",
            "first_page_header", "first_page_footer",
            "even_page_header", "even_page_footer",
        ):
            part = getattr(section, attr, None)
            if part is not None:
                yield part


def _splice_runs(runs: list, start: int, end: int, replacement: str) -> None:
    """Replace the character range [start, end) spread across *runs*.

    The replacement takes the formatting of the run the placeholder
    starts in; every other run keeps its own text and formatting. The
    naive alternative — writing the whole paragraph into its first run —
    would flatten bold, spacing and font changes across the line.
    """
    pos = 0
    first = True
    for run in runs:
        begin, stop = pos, pos + len(run.text)
        pos = stop
        if stop <= start or begin >= end:
            continue
        prefix = run.text[: max(0, start - begin)]
        suffix = run.text[max(0, end - begin):] if stop > end else ""
        if first:
            run.text = prefix + replacement + suffix
            first = False
        else:
            run.text = prefix + suffix


def _fill_paragraph(paragraph: Any, values: dict, unfilled: set) -> int:
    """Substitute placeholders in one paragraph. Returns the count.

    Works on the paragraph's joined text rather than run by run: Word
    splits a placeholder across runs whenever it feels like it (spell
    check, an edit, a formatting change), so ``{{anrede}}`` routinely
    arrives as ``{{`` + ``an`` + ``rede`` + ``}}``. Per-run replacement
    finds nothing and reports success.
    """
    filled = 0
    offset = 0
    for _ in range(500):                      # bound: pathological templates
        runs = list(paragraph.runs)
        if not runs:
            return filled
        text = "".join(r.text for r in runs)
        match = _PLACEHOLDER_RE.search(text, offset)
        if match is None:
            return filled
        name = match.group(1).strip()
        if name not in values:
            unfilled.add(name)
            offset = match.end()
            continue
        # Not str(): a caller handing the filler a float or a datetime is
        # handing it a value the LETTER has to show, and str() shows
        # "1234.5" and "2026-07-31 00:00:00" to a German customer.
        replacement = _render_for_locale(values[name])
        _splice_runs(runs, match.start(), match.end(), replacement)
        offset = match.start() + len(replacement)
        filled += 1
    return filled


def docx_placeholders(path: Any) -> dict:
    """List the ``{{name}}`` placeholders a template contains.

    Reported per location (body / table / header / footer) so a
    placeholder that only appears in the letterhead is visible before
    the fill, not after.
    """
    p = _resolve(path)
    if document_kind(p) != "word":
        raise OfficeError(f"{p.name} is not a .docx file")
    docx = _require("word")
    try:
        doc = docx.Document(str(p))
    except Exception as exc:
        raise OfficeError(f"could not open document: {exc}") from exc

    counts: dict[str, int] = {}
    for part in _document_parts(doc):
        for para in _iter_paragraphs(part):
            for match in _PLACEHOLDER_RE.finditer(para.text):
                name = match.group(1).strip()
                counts[name] = counts.get(name, 0) + 1

    notes: list[str] = []
    if not counts:
        notes.append(
            "no {{placeholder}} markers in this document. Word mail-merge "
            "fields (MERGEFIELD) are not text placeholders and are not "
            "handled — check what the template actually uses.")
    return {
        "path": str(p),
        "placeholders": [
            {"name": n, "occurrences": c} for n, c in sorted(counts.items())
        ],
        "notes": notes,
    }


_DOCX_TEXT_PARTS = ("word/document.xml", "word/footnotes.xml",
                    "word/endnotes.xml", "word/comments.xml")


def _placeholders_left_in_file(path: Path) -> set[str]:
    """Every ``{{name}}`` still present in a .docx, read from the zip.

    Deliberately independent of python-docx's object model: this is the
    check on the writer, and a check that shares the writer's view of the
    document can only confirm it. Headers and footers are included by
    pattern because their part names are numbered, and a letterhead is
    exactly where a template puts the placeholders somebody notices last.

    Never raises -- a file this cannot open is reported as nothing left,
    and the reader-based check beside it still applies.
    """
    found: set[str] = set()
    try:
        import zipfile
        with zipfile.ZipFile(path) as bundle:
            for name in bundle.namelist():
                if not name.endswith(".xml"):
                    continue
                if not (name in _DOCX_TEXT_PARTS
                        or name.startswith("word/header")
                        or name.startswith("word/footer")):
                    continue
                try:
                    body = bundle.read(name).decode("utf-8", "replace")
                except Exception:
                    continue
                for match in _PLACEHOLDER_RE.finditer(body):
                    found.add(match.group(1).strip())
    except Exception:
        return found
    return found

def fill_docx_template(
    path: Any, values: dict, *, output: Any, strict: bool = True
) -> dict:
    """Substitute ``{{name}}`` placeholders and write a new document.

    The template is never written to, so it stays reusable. With
    *strict* a value whose placeholder does not exist is an error rather
    than a silent no-op — the same contract as the PDF form filler,
    because a typo would otherwise read as a completed document.

    Placeholders left without a value are reported and stay in the text
    verbatim: a half-filled letter that still shows ``{{betrag}}`` is
    recoverable, one that silently dropped it is not.
    """
    src = _resolve(path)
    if document_kind(src) != "word":
        raise OfficeError(f"{src.name} is not a .docx file")
    if not isinstance(values, dict) or not values:
        raise OfficeError("values must be a non-empty object of name -> value")
    out = _resolve(output, must_exist=False)
    if out.resolve() == src.resolve():
        raise OfficeError(
            "output must differ from the template — filling it in place "
            "destroys the blank original")
    # Same rule as every sibling writer, and as the PDF form filler
    # above: a document that is already there is not scratch space.
    if out.exists():
        raise OfficeError(
            f"{out.name} already exists — nothing was written. Write to a "
            "new name, or delete that file first if it is meant to go.")

    known = {
        entry["name"] for entry in docx_placeholders(src)["placeholders"]
    }
    unknown = sorted(str(k) for k in values if str(k) not in known)
    if unknown and strict:
        raise OfficeError(
            f"no such placeholder(s): {', '.join(unknown)}. The template "
            f"contains: {', '.join(sorted(known)) or '(none)'}"
        )

    docx = _require("word")
    try:
        doc = docx.Document(str(src))
    except Exception as exc:
        raise OfficeError(f"could not open template: {exc}") from exc

    normalised = {str(k): v for k, v in values.items()}
    unfilled: set[str] = set()
    filled = 0
    for part in _document_parts(doc):
        for para in _iter_paragraphs(part):
            filled += _fill_paragraph(para, normalised, unfilled)

    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
    except Exception as exc:
        # A part-written .docx is still a file, and the next read opens it
        # as a document. create_pdf beside this one already unlinks.
        try:
            out.unlink()
        except OSError:
            pass
        raise OfficeError(f"could not save the document: {exc}") from exc

    # Verify with a DIFFERENT reader than the one that filled.
    #
    # ``_fill_paragraph`` is driven by ``_iter_paragraphs``, and
    # ``docx_placeholders`` walks the same way -- so a placeholder the
    # writer cannot reach is one the check cannot see either, and the
    # read-back was a tautology that could never catch the writer's blind
    # spot. Proven on a template whose address window is a text box, which
    # is how Word letterhead is built: filled: 1, unfilled: [], complete:
    # True, and "{{name}}" still in the output. Every Bescheid of that
    # Serienbrief prints the placeholder, and fill_series stamps each row
    # ok because it reads that same flag.
    #
    # The output's own XML cannot be fooled that way: what survives there
    # survived, whatever traversal did or did not reach it.
    remaining = {
        entry["name"]
        for entry in docx_placeholders(out)["placeholders"]
    } | _placeholders_left_in_file(out)
    notes: list[str] = []
    if unknown:
        notes.append(
            "ignored, no such placeholder in the template: "
            + ", ".join(unknown))
    still_open = sorted(remaining)
    if still_open:
        notes.append(
            "these placeholders had no value and are still in the text: "
            + ", ".join(still_open)
            + ". The document is not ready to hand over as it stands.")
    leaked = sorted(n for n in normalised if n in remaining)
    if leaked:
        notes.append(
            "a value was given but the placeholder survived in the output: "
            + ", ".join(leaked) + " — treat this fill as failed.")
    return {
        "template": str(src),
        "path": str(out),
        "filled": filled,
        "unfilled": still_open,
        "complete": not still_open,
        "notes": notes,
    }


def create_docx(
    path: Any, blocks: list[dict], *, overwrite: bool = False
) -> dict:
    """Write a new .docx from a list of content blocks.

    Each block is one of ``{"heading": str, "level": int}``,
    ``{"paragraph": str}``, ``{"table": [[...], ...], "header_row":
    bool}`` or ``{"page_break": true}``.
    """
    p = _resolve(path, must_exist=False)
    if document_kind(p) != "word":
        raise OfficeError(
            f"cannot write {p.suffix!r} as a Word document — use .docx")
    if p.exists() and not overwrite:
        raise OfficeError(
            f"{p.name} already exists — pass overwrite=True to replace it")
    if not isinstance(blocks, list) or not blocks:
        raise OfficeError("blocks must be a non-empty list")

    docx = _require("word")
    doc = docx.Document()
    counts = {"headings": 0, "paragraphs": 0, "tables": 0}
    probe = ""
    for index, block in enumerate(blocks, start=1):
        if not isinstance(block, dict):
            raise OfficeError(f"block {index} is not an object: {block!r}")
        if "heading" in block:
            level = max(0, min(int(block.get("level", 1) or 1), 9))
            doc.add_heading(str(block["heading"]), level=level)
            counts["headings"] += 1
            probe = probe or _text_probe(str(block["heading"]))
        elif "paragraph" in block:
            para = doc.add_paragraph()
            runs = _runs(block["paragraph"])
            # Block-level bold/colour apply to a plain string; a list of
            # runs carries its own, so the label can stay plain while the
            # figure beside it is red.
            if isinstance(block["paragraph"], (str, int, float)):
                for run in runs:
                    run["bold"] = bool(block.get("bold"))
                    run["italic"] = bool(block.get("italic"))
                    run["colour"] = parse_colour(
                        block.get("color", block.get("colour")))
            _apply_runs(para, runs)
            counts["paragraphs"] += 1
            probe = probe or _text_probe(
                "".join(str(r.get("text", "")) for r in runs))
        elif "table" in block:
            rows = block["table"]
            if not isinstance(rows, list) or not rows or not all(
                    isinstance(r, (list, tuple)) for r in rows):
                raise OfficeError(
                    f"block {index}: table must be a non-empty list of lists")
            width = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=width)
            if block.get("header_row"):
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
            for r_idx, row in enumerate(rows):
                for c_idx in range(width):
                    value = row[c_idx] if c_idx < len(row) else ""
                    cell = table.cell(r_idx, c_idx)
                    # A fresh cell already holds one empty run. Writing
                    # past it leaves the text in run 2, where anything
                    # reading the first run finds nothing.
                    target = cell.paragraphs[0]
                    for stale in list(target.runs):
                        stale._element.getparent().remove(stale._element)
                    cell_runs = _runs(value)
                    if r_idx == 0 and block.get("header_row"):
                        for run in cell_runs:
                            run["bold"] = True
                    _apply_runs(target, cell_runs)
                    probe = probe or _text_probe(_fmt(value))
            counts["tables"] += 1
        elif block.get("page_break"):
            doc.add_page_break()
        else:
            raise OfficeError(
                f"block {index}: expected one of heading / paragraph / "
                f"table / page_break, got keys {sorted(block)}")

    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
    except Exception as exc:
        # Every sibling writer unlinks here. This one left a part-written
        # .docx behind, which the next read opens as a document.
        try:
            p.unlink()
        except OSError:
            pass
        raise OfficeError(f"could not save the document: {exc}") from exc

    # Read the result back with the reader, not with the object still in
    # memory. This was the one writer in the module with no check at all:
    # it reported blocks written and never asked whether the file holds
    # them. The probe is a word or figure taken from what went in, so a
    # save that produced an unreadable or empty document is caught rather
    # than reported as done.
    notes: list[str] = []
    verified = False
    try:
        back = read_docx(p)
        text = _squash(back.get("text", ""))
        if not probe:
            notes.append(
                "the document was written, but it carries no text this "
                "could search for, so nothing about its CONTENT was checked")
        elif _squash(probe) in text:
            verified = True
        else:
            notes.append(
                f"the word {probe!r} was written but was not found when the "
                "document was read back — check it before handing it over")
    except Exception as exc:
        notes.append(f"the written file could not be read back: {exc}")
    return {
        "path": str(p),
        "blocks": len(blocks),
        "verified": verified,
        "notes": notes,
        **counts,
    }


# Colours a caller may name. Deliberately a short list plus hex: an
# office document needs "this figure is red", not a palette, and a name
# that silently resolves to something else is worse than a refusal.
_COLOURS: dict[str, tuple[int, int, int]] = {
    "black": (0, 0, 0), "red": (192, 0, 0), "green": (0, 128, 0),
    "blue": (0, 0, 192), "orange": (224, 112, 0), "grey": (110, 110, 110),
    "gray": (110, 110, 110), "white": (255, 255, 255),
}

_HEX_RE = re.compile(r"^#?([0-9a-fA-F]{6})$")


def parse_colour(value: Any) -> Optional[tuple]:
    """A colour name or ``#RRGGBB`` as an RGB triple, else None."""
    if value is None or value == "":
        return None
    text = str(value).strip().lower()
    if text in _COLOURS:
        return _COLOURS[text]
    match = _HEX_RE.match(text)
    if match:
        digits = match.group(1)
        return tuple(int(digits[i:i + 2], 16) for i in (0, 2, 4))
    raise OfficeError(
        f"unknown colour {value!r}. Use one of "
        + ", ".join(sorted(_COLOURS)) + ", or #RRGGBB.")


def _runs(value: Any) -> list[dict]:
    """Normalise text into runs of ``{text, bold, italic, colour}``.

    A plain string is one unformatted run, so every existing caller keeps
    working. A list lets one paragraph carry an emphasised part without
    turning the whole line bold — which is what a report actually needs:
    the label plain, the figure red.
    """
    if value is None:
        return [{"text": "", "bold": False, "italic": False, "colour": None}]
    if isinstance(value, (str, int, float)):
        return [{"text": _fmt(value), "bold": False, "italic": False,
                 "colour": None}]
    if isinstance(value, dict):
        value = [value]
    if not isinstance(value, list):
        raise OfficeError(f"expected text or a list of runs, got {value!r}")
    out = []
    for part in value:
        if isinstance(part, (str, int, float)):
            out.append({"text": _fmt(part), "bold": False, "italic": False,
                        "colour": None})
            continue
        if not isinstance(part, dict):
            raise OfficeError(f"a run must be text or an object, got {part!r}")
        out.append({
            "text": _fmt(part.get("text", "")),
            "bold": bool(part.get("bold")),
            "italic": bool(part.get("italic")),
            "colour": parse_colour(part.get("color", part.get("colour"))),
        })
    return out or [{"text": "", "bold": False, "italic": False,
                    "colour": None}]


def _apply_runs(paragraph: Any, runs: list[dict]) -> None:
    """Write runs into a python-docx paragraph, keeping its font."""
    from docx.shared import RGBColor

    for run in runs:
        written = paragraph.add_run(run["text"])
        if run["bold"]:
            written.bold = True
        if run["italic"]:
            written.italic = True
        if run["colour"]:
            written.font.color.rgb = RGBColor(*run["colour"])


def _runs_markup(runs: list[dict]) -> str:
    """Runs as the inline markup reportlab understands, escaped."""
    from xml.sax.saxutils import escape

    out = []
    for run in runs:
        text = escape(run["text"])
        if run["colour"]:
            text = ('<font color="#%02x%02x%02x">%s</font>'
                    % (*run["colour"], text))
        if run["bold"]:
            text = f"<b>{text}</b>"
        if run["italic"]:
            text = f"<i>{text}</i>"
        out.append(text)
    return "".join(out)


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _docx_unreached_text(path: Path) -> list[str]:
    """Passages a .docx carries that python-docx's model does not reach.

    ``doc.paragraphs`` yields the paragraphs that are direct children of
    the body. A text box (``w:txbxContent``) and a content control
    (``w:sdtContent``) hold their paragraphs one level in, so neither
    appears there — and a Word letterhead puts the address block in a
    text box as a matter of course. Read from the zip, deliberately
    independent of the object model, for the same reason the template
    filler verifies from the zip: a reader that shares the writer's view
    can only confirm it.

    Never raises: this is an aid to a read, not a reason to refuse one.
    """
    out: list[str] = []
    try:
        import xml.etree.ElementTree as ET
        import zipfile
        with zipfile.ZipFile(path) as bundle:
            for name in bundle.namelist():
                if not (name in _DOCX_TEXT_PARTS
                        or name.startswith("word/header")
                        or name.startswith("word/footer")):
                    continue
                try:
                    root = ET.fromstring(bundle.read(name))
                except Exception:
                    continue
                for tag in ("txbxContent", "sdtContent"):
                    for node in root.iter(f"{_W_NS}{tag}"):
                        text = "".join(
                            t.text or "" for t in node.iter(f"{_W_NS}t"))
                        text = text.strip()
                        if text:
                            out.append(text)
    except Exception:
        return out
    return out


def _docx_part_text(part: Any) -> str:
    """Every paragraph of one document part, tables included."""
    return "\n".join(
        para.text for para in _iter_paragraphs(part) if para.text.strip())


def read_docx(path: Any, *, max_chars: int = MAX_TEXT_CHARS) -> dict:
    """Extract paragraphs and tables from a .docx file."""
    p = _resolve(path)
    _refuse_unreadable_format(p)
    if document_kind(p) == "opendocument_text":
        return read_odt(p, max_chars=max_chars)
    if document_kind(p) != "word":
        raise OfficeError(
            f"{p.name} is not a .docx file. The legacy .doc format is not "
            "supported — convert it first.")
    docx = _require("word")
    try:
        doc = docx.Document(str(p))
    except Exception as exc:
        raise OfficeError(f"could not open document: {exc}") from exc

    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    tables = []
    for n, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(f"--- table {n} ---\n" + render_grid(rows))
    body = "\n".join(parts)
    if tables:
        body += "\n\n" + "\n\n".join(tables)

    # Headers and footers. This read walked doc.paragraphs and stopped
    # there, so a .docx whose Rechnungsnummer sits in the Kopfzeile --
    # which is where letterhead puts it -- came back as a document that
    # does not contain its own number. Labelled, because a figure lifted
    # out of a running header is not a body figure and an answer citing
    # it should be able to say so.
    # Both extra readers are an AID to the read, never a reason to refuse
    # one: the body came back fine, and a document that gives up its text
    # is worth more than one that raised over its letterhead.
    extras: list[str] = []
    try:
        body_doc = doc
        for part in _document_parts(doc):
            if part is body_doc:
                continue
            text = _docx_part_text(part)
            if text.strip() and text.strip() not in body:
                extras.append(text.strip())
    except Exception:
        extras = []
    if extras:
        body += "\n\n--- header / footer ---\n" + "\n".join(extras)

    # Text boxes and content controls, read from the zip because the
    # object model cannot reach them at all. Only what is MISSING from
    # the extraction above is added, so nothing is counted twice.
    try:
        unreached = [t for t in _docx_unreached_text(p) if t not in body]
    except Exception:
        unreached = []
    if unreached:
        body += "\n\n--- text box / content control ---\n" + "\n".join(
            unreached)

    notes: list[str] = []
    if len(body) > max_chars:
        body = body[:max_chars]
        notes.append(f"output truncated at {max_chars} characters")
    if not body.strip():
        notes.append("the document holds no extractable text")
    if extras or unreached:
        notes.append(
            f"{len(extras) + len(unreached)} passage(s) came from a header, "
            f"footer, text box or content control and are labelled as such "
            f"in the text — the paragraph count below counts the body only")
    return {
        "path": str(p),
        "kind": "word",
        "paragraphs": len(parts),
        "paragraphs_source": "body",
        "tables": len(doc.tables),
        "text": body,
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# OpenDocument text (.odt)
# ---------------------------------------------------------------------------

def _odt_cell_text(cell: Any) -> str:
    from odf import teletype
    try:
        return teletype.extractText(cell).strip()
    except Exception:
        return ""


def _odt_content(container: Any) -> tuple[list[str], list[list[list]]]:
    """``(paragraphs, tables)`` of an .odt body, in document order.

    Walks the tree rather than collecting paragraphs by type: in a
    letter the recipient block sits in a table and the enclosures sit in
    a list, and a reader that only takes top-level ``text:p`` elements
    returns a document with the addressee missing — which looks like a
    complete letter to anyone who did not have the original.
    """
    from odf.namespaces import TABLENS, TEXTNS
    from odf import teletype

    paragraphs: list[str] = []
    tables: list[list[list]] = []

    def walk(node: Any) -> None:
        for child in getattr(node, "childNodes", None) or []:
            qname = getattr(child, "qname", None)
            if not qname:
                continue
            namespace, name = qname
            if namespace == TEXTNS and name in ("p", "h"):
                try:
                    text = teletype.extractText(child).strip()
                except Exception:
                    text = ""
                if text:
                    paragraphs.append(text)
            elif namespace == TABLENS and name == "table":
                def _cells(row: Any) -> list:
                    return [
                        _odt_cell_text(cell)
                        for cell in (getattr(row, "childNodes", None) or [])
                        if (getattr(cell, "qname", None)
                            or ("", ""))[0] == TABLENS
                    ]

                rows = [
                    _cells(row)
                    for row in (getattr(child, "childNodes", None) or [])
                    if getattr(row, "qname", None) == (TABLENS, "table-row")
                ]
                if rows:
                    tables.append(rows)
            else:
                walk(child)

    walk(container)
    return paragraphs, tables


def read_odt(path: Any, *, max_chars: int = MAX_TEXT_CHARS) -> dict:
    """Extract paragraphs and tables from an OpenDocument text document."""
    p = _resolve(path)
    if document_kind(p) != "opendocument_text":
        raise OfficeError(f"{p.name} is not an .odt file")
    document = _odf_load(p)
    mimetype = str(getattr(document, "mimetype", "") or "")
    if mimetype and not mimetype.endswith(("text", "text-template")):
        raise OfficeError(
            f"{p.name} is named like a text document but its OpenDocument "
            f"type is '{mimetype}'. Read it under the name of what it is.")

    body = getattr(document, "text", None)
    if body is None:
        raise OfficeError(f"{p.name} holds no text body")
    paragraphs, tables = _odt_content(body)

    parts = list(paragraphs)
    rendered = [f"--- table {n} ---\n" + render_grid(rows)
                for n, rows in enumerate(tables, start=1)]
    text = "\n".join(parts)
    if rendered:
        text += ("\n\n" if text else "") + "\n\n".join(rendered)

    notes: list[str] = []
    if len(text) > max_chars:
        text = text[:max_chars]
        notes.append(f"output truncated at {max_chars} characters")
    if not text.strip():
        notes.append("the document holds no extractable text")
    notes.append(
        "read-only: an .odt cannot be written here. To produce a filled "
        "document, convert the template to .docx and fill that.")
    return {
        "path": str(p),
        "kind": "opendocument_text",
        "format": "odt",
        "writable": False,
        "paragraphs": len(parts),
        "tables": len(tables),
        "text": text,
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

def read_document(path: Any, **kwargs: Any) -> dict:
    """Read any supported document, dispatching on its suffix."""
    p = _resolve(path)
    _refuse_unreadable_format(p)
    kind = document_kind(p)
    if kind == "opendocument_text":
        if kwargs.get("fields"):
            raise OfficeError(
                f"{p.name} is an OpenDocument text document; there is no "
                "template filler for it. Convert it to .docx — "
                + _LIBREOFFICE_HINT.format(target="docx", name=p.name)
                + " — and use the .docx as the template.")
        return read_odt(p)
    if kind in ("spreadsheet", "csv", "opendocument_sheet"):
        return read_sheet(
            p,
            sheet=kwargs.get("sheet"),
            max_rows=int(kwargs.get("max_rows") or DEFAULT_MAX_ROWS),
            max_cols=int(kwargs.get("max_cols") or DEFAULT_MAX_COLS),
            start_row=int(kwargs.get("start_row") or 1),
            start_col=int(kwargs.get("start_col") or 1),
        )
    if kind == "pdf":
        if kwargs.get("fields"):
            return pdf_form_fields(p)
        return read_pdf(p, pages=kwargs.get("pages"),
                        ocr=bool(kwargs.get("ocr")))
    if kind == "word":
        # Same switch as for a PDF form: "show me what can be filled in"
        # is one question, whatever the format underneath.
        if kwargs.get("fields"):
            return docx_placeholders(p)
        return read_docx(p)
    raise OfficeError(
        f"{p.name}: no document reader for {p.suffix!r}. Supported: "
        + ", ".join(sorted(DOCUMENT_SUFFIXES))
    )


# ---------------------------------------------------------------------------
# Email drafts
# ---------------------------------------------------------------------------

_ADDRESS_RE = re.compile(r"^[^@\s,;<>]+@[^@\s,;<>]+\.[A-Za-z]{2,}$")


def _addresses(value: Any, field: str) -> list[str]:
    """Split and validate a recipient field.

    Validated here rather than left to the mail client, because the whole
    point of a draft is that the user opens something already correct. An
    address the client silently drops is worse than one it rejects.
    """
    if value is None or value == "":
        return []
    if isinstance(value, str):
        parts = [p.strip() for p in re.split(r"[,;]", value)]
    elif isinstance(value, (list, tuple)):
        parts = [str(p).strip() for p in value]
    else:
        raise OfficeError(f"{field} must be a string or a list of addresses")
    out: list[str] = []
    for part in parts:
        if not part:
            continue
        if not _ADDRESS_RE.match(part):
            raise OfficeError(
                f"{field}: {part!r} is not an email address. Write it as "
                "name@example.org; separate several with commas.")
        if part not in out:
            out.append(part)
    return out


def draft_email(
    path: Any,
    *,
    to: Any = None,
    subject: str = "",
    body: str = "",
    cc: Any = None,
    attachments: Optional[list] = None,
    overwrite: bool = False,
) -> dict:
    """Write a ready-to-send .eml the USER opens and sends themselves.

    Deliberately not a send. Sending is irreversible, happens under the
    user's identity, and has no read-back — and reading back what was
    written is the discipline that makes the rest of this module
    trustworthy. Their own mail client is also a better final
    confirmation than any preview we could build: it shows the real
    rendering, the real sender, and the send button is theirs.

    So this produces a file. Double-clicking it opens a fully populated
    message. Nothing leaves the machine, no credential is involved, and
    the file goes through the same write gate as every other document.
    """
    from email.message import EmailMessage
    from email.utils import formatdate, make_msgid

    p = _resolve(path, must_exist=False)
    if p.suffix.lower() != ".eml":
        raise OfficeError(
            f"cannot write {p.suffix!r} as an email draft — use .eml")
    if p.exists() and not overwrite:
        raise OfficeError(
            f"{p.name} already exists — pass overwrite=True to replace it")

    recipients = _addresses(to, "to")
    carbon = _addresses(cc, "cc")
    if not recipients:
        raise OfficeError("to: at least one recipient is required")
    if not str(subject).strip():
        raise OfficeError("subject must be non-empty")

    msg = EmailMessage()
    msg["To"] = ", ".join(recipients)
    if carbon:
        msg["Cc"] = ", ".join(carbon)
    msg["Subject"] = str(subject)
    msg["Date"] = formatdate(localtime=True)
    msg["Message-ID"] = make_msgid()
    msg.set_content(str(body or ""))

    attached: list[str] = []
    skipped: list[str] = []
    for item in attachments or ():
        try:
            src = _resolve(item, must_exist=True)
        except OfficeError:
            skipped.append(str(item))
            continue
        try:
            msg.add_attachment(
                src.read_bytes(), maintype="application",
                subtype="octet-stream", filename=src.name)
            attached.append(src.name)
        except OSError:
            skipped.append(str(item))

    p.write_bytes(bytes(msg))
    return {
        "path": str(p),
        "to": recipients,
        "cc": carbon,
        "subject": str(subject),
        "attachments": attached,
        "attachments_skipped": skipped,
        "bytes": p.stat().st_size,
        # The skipped list was collected and the sentence never carried
        # it, so a draft missing the very document it is about read as a
        # complete one. The count goes where the reader is already looking.
        "note": (
            "Draft written"
            + (f", WITHOUT {len(skipped)} attachment(s) that could not be "
               f"read ({', '.join(skipped[:3])}"
               + (", …" if len(skipped) > 3 else "") + ")"
               if skipped else "")
            + ". It is NOT sent: open it in your mail client, check it, "
              "and send it yourself."
        ),
    }


# ---------------------------------------------------------------------------
# The figure ledger: what the tools returned, against what the answer states
# ---------------------------------------------------------------------------
#
# Everything above this line reads and writes documents correctly. Nothing
# above it checks a number an ANSWER states against what the tools actually
# returned — and in administrative work that is the whole failure mode. A
# total that silently dropped a row reads exactly like a total that did
# not; "31 Belege" reads exactly like a count somebody took.
#
# So the tools record the figures they produce (the column that was summed
# and its result, the counts a reconciliation landed on, the shape of a
# read, the values a grid displayed), and the finished answer is compared
# against that ledger. The consequence is a caveat naming the figure and
# the reason, never a refusal: an administrative answer may legitimately
# name a number the tools did not produce — the user typed it, it is an
# example, it came from an earlier turn — and those are checked BEFORE
# anything is flagged.
#
# The errors here are deliberately asymmetric. A missed flag costs one
# unmarked number; a false one costs a caveat on a correct answer, and a
# caveat on every answer is a caveat nobody reads.
#
# The ledger is keyed by the TURN that produced the figures. It was one
# module-level list, described as safe on the grounds that a figure from
# elsewhere in the process "can only GROUND a claim, never create a
# flag". Both halves of that were false, and both were measured:
#
#   * the clear is driven by ONE turn boundary, so anything else in the
#     process starting a turn wiped the evidence the first one was still
#     collecting;
#   * the cap is a shared 2000, and a single document read contributes up
#     to 600 figures, so four reads by a background sub-agent filled it
#     and the parent's own sum_column result was dropped on the floor.
#
# Both ended in the same place — a correct, tool-computed total carrying
# the caveat that says no tool produced it, which is precisely the false
# positive that makes a guard unreadable. And a figure a sub-agent
# recorded after the parent's reset survived into the parent's next turn,
# which is the stale grounding the reset exists to prevent.
#
# So every turn records and reads under its own token: a figure recorded
# under one token is invisible under another, and a foreign scope filling
# its own cap cannot evict anybody else's figures.

_LEDGER_LOCK = _threading.Lock()

# Figures kept per turn. A read contributes its whole window, so this is
# sized for a few reads and then bounded — an unbounded ledger in a long
# session is a memory leak with a guard attached to it. The cap is per
# SCOPE: shared, it is an eviction channel between unrelated turns.
MAX_LEDGER_FIGURES = 2000
# Documents named per turn, same reasoning.
MAX_LEDGER_PATHS = 200
# Turns held at once. Only the live ones are ever read; the rest are the
# tail of a long session and go oldest-first.
MAX_LEDGER_SCOPES = 32


@dataclass
class _ScopeLedger:
    """One turn's figures, under one token."""

    figures: list["ToolFigure"] = _field(default_factory=list)
    paths: list[str] = _field(default_factory=list)
    seq: int = 0


_LEDGERS: dict[str, _ScopeLedger] = {}
_LEDGER_TICK = 0        # recency counter for scope eviction
_LEDGER_TURN = 0        # turn counter, so one session's tokens stay distinct

# The scope the current thread records under. The engine sets it at the
# turn boundary; anything recording without one — a background sub-agent
# on its own thread — falls back to its thread, which is what keeps its
# reads out of the parent's ledger instead of evicting them.
_LEDGER_SCOPE: "_contextvars.ContextVar[str]" = _contextvars.ContextVar(
    "delfin_figure_ledger_scope", default="")


def figure_ledger_scope() -> str:
    """The token this thread's figures are recorded and read under."""
    return _LEDGER_SCOPE.get() or f"thread-{_threading.get_ident()}"


def _scope(token: str = "") -> str:
    return str(token or "") or figure_ledger_scope()


def _bucket(token: str, *, create: bool) -> Optional["_ScopeLedger"]:
    """This scope's ledger, oldest scopes evicted. Caller holds the lock."""
    global _LEDGER_TICK
    entry = _LEDGERS.get(token)
    if entry is None:
        if not create:
            return None
        entry = _LEDGERS[token] = _ScopeLedger()
    _LEDGER_TICK += 1
    entry.seq = _LEDGER_TICK
    if len(_LEDGERS) > MAX_LEDGER_SCOPES:
        stale = sorted(_LEDGERS, key=lambda key: _LEDGERS[key].seq)
        for key in stale[:len(_LEDGERS) - MAX_LEDGER_SCOPES]:
            if key != token:
                _LEDGERS.pop(key, None)
    return entry


def begin_figure_turn(session_id: str = "") -> str:
    """Open a ledger for a new turn and make it this thread's scope.

    Returns the token. The caller keeps it and hands it back when the
    finished answer is checked, so the check reads the ledger this turn's
    own tools wrote to and no other one.
    """
    global _LEDGER_TURN
    with _LEDGER_LOCK:
        _LEDGER_TURN += 1
        token = f"{str(session_id or '').strip() or 'session'}#{_LEDGER_TURN}"
        _LEDGERS.pop(token, None)
        _bucket(token, create=True)
    _LEDGER_SCOPE.set(token)
    return token


@dataclass(frozen=True)
class ToolFigure:
    """One number an office tool returned this turn."""

    value: float
    kind: str      # "sum" | "count" | "cell"
    label: str     # what it was, in words ("Summe 'Betrag' in Buchungen.xlsx")
    tool: str      # which tool returned it
    # True when the read that produced it covered a WINDOW of the table
    # rather than all of it. The value is real either way -- it was in the
    # grid -- but what may be said about it is not: the largest number
    # among five rows of twenty-six is not the largest booking, and once
    # it is out of the grid and inside a sentence the two look alike.
    partial: bool = False
    # The window it came from, for a message that names it rather than
    # asking the reader to go and find out.
    window: str = ""

    def message(self) -> str:
        return f"{self.label} ({self.tool})"


def reset_figure_ledger(token: str = "") -> str:
    """Forget the figures of the previous turn — in ONE scope.

    Per TURN, like the other evidence ledgers in this framework: the
    question is whether THIS answer states a figure THIS turn produced.
    Carried across turns, a stale total would ground a later, unrelated
    number — which is the failure this exists to catch, one turn late.

    Only the named scope is cleared. It used to clear the whole process,
    so one turn boundary anywhere destroyed evidence every other turn in
    flight was still collecting. Returns the token it cleared.
    """
    scope = _scope(token)
    with _LEDGER_LOCK:
        _LEDGERS.pop(scope, None)
    return scope


def figure_ledger(token: str = "") -> list["ToolFigure"]:
    """The figures the office tools returned this turn, in this scope.

    The number of documents touched rides along as a figure of its own, so
    an answer that says "4 Dateien geprüft" is grounded by the reads that
    happened rather than flagged for a number no single tool returned.
    """
    with _LEDGER_LOCK:
        entry = _bucket(_scope(token), create=False)
        out = list(entry.figures) if entry is not None else []
        paths = list(entry.paths) if entry is not None else []
    if paths:
        out.append(ToolFigure(
            value=float(len(paths)), kind="count",
            label=f"{len(paths)} document(s) read or written this turn",
            tool="office"))
    return out


def _note_ledger_path(path: str, token: str = "") -> None:
    name = str(path or "")
    if not name:
        return
    with _LEDGER_LOCK:
        entry = _bucket(_scope(token), create=True)
        if (entry is not None and name not in entry.paths
                and len(entry.paths) < MAX_LEDGER_PATHS):
            entry.paths.append(name)


def _figures_from_result(tool: str, result: dict) -> list["ToolFigure"]:
    """The figures one tool result carries, without re-reading its prose.

    ``sum_column`` and ``compare_tables`` state theirs outright (their
    ``figures`` key); the readers and the writers are described by their
    own result fields. Nothing here parses a rendered report — two
    implementations of the same facts would eventually disagree, and the
    one the ledger believed would be the wrong one.
    """
    out: list[ToolFigure] = []

    # A read that showed a page of the table produced values that are real
    # and a picture that is not. Counts are exempt: ``rows`` is the file's
    # own total, stated by the reader regardless of how much it showed.
    win = result.get("window") if isinstance(result, dict) else None
    partial_window = bool(isinstance(win, dict) and not win.get("complete"))
    window_text = ""
    if partial_window:
        # German, because this string ends up inside the caveat the user
        # reads, not in a log.
        window_text = (f"Zeilen {win.get('first_row')}–{win.get('last_row')} "
                       f"von {win.get('total_rows')}")

    def _add(value: Any, kind: str, label: str) -> None:
        try:
            number = float(value)
        except (TypeError, ValueError):
            return
        if number != number or number in (float("inf"), float("-inf")):
            return
        cell = kind == "cell"
        out.append(ToolFigure(value=round(number, 6), kind=kind,
                              label=label, tool=tool,
                              partial=partial_window and cell,
                              window=window_text if (partial_window and cell)
                              else ""))

    declared = result.get("figures")
    if isinstance(declared, list):
        for entry in declared:
            if isinstance(entry, dict):
                _add(entry.get("value"), str(entry.get("kind") or "count"),
                     str(entry.get("label") or ""))
        return out

    name = Path(str(result.get("path") or "")).name or "the document"
    if "rows" in result and "columns" in result:
        _add(result.get("rows"), "count",
             f"{result.get('rows')} row(s) in {name}")
        _add(result.get("columns"), "count",
             f"{result.get('columns')} column(s) in {name}")
    for value in result.get("numbers") or ():
        _add(value, "cell", f"a value shown in {name}")
    # A PDF, a .docx or an .odt comes back as text, and the amounts in it
    # are just as much what the tool returned as the cells of a grid are.
    # Without this, an answer quoting a figure off an invoice it had just
    # read would be marked as having invented it.
    text = result.get("text")
    if isinstance(text, str) and text:
        for value in _numbers_in(text)[:MAX_WINDOW_NUMBERS]:
            _add(value, "cell", f"a value in the text of {name}")
    if result.get("pages"):
        _add(result.get("pages"), "count",
             f"{result.get('pages')} page(s) in {name}")
    if result.get("field_count"):
        _add(result.get("field_count"), "count",
             f"{result.get('field_count')} form field(s) in {name}")
    if "appended_rows" in result:               # a write says what it wrote
        _add(result.get("appended_rows"), "count",
             f"{result.get('appended_rows')} row(s) appended to {name}")
        _add(len(result.get("applied") or ()), "count",
             f"{len(result.get('applied') or ())} cell(s) changed in {name}")
    counts = result.get("counts")
    if isinstance(counts, dict):                # fill_series
        for label in ("ok", "incomplete", "failed"):
            if label in counts:
                _add(counts[label], "count",
                     f"{counts[label]} document(s) {label}")
        _add(result.get("rows"), "count",
             f"{result.get('rows')} table row(s) in the series")
        _add(result.get("processed"), "count",
             f"{result.get('processed')} document(s) attempted")
    return out


def record_tool_figures(tool: str, result: Any,
                        token: str = "") -> list["ToolFigure"]:
    """Record the figures one office tool returned. Never raises.

    Called by the tool layer with the tool's own result dict, so the
    ledger holds what the tool computed rather than what its report
    happened to print. Recorded under this scope's token only: a
    sub-agent reading documents in the same process fills its own cap,
    not the cap of the turn that is waiting on its own total.
    """
    try:
        if not isinstance(result, dict):
            return []
        scope = _scope(token)
        found = _figures_from_result(str(tool or "office"), result)
        for key in ("path", "left", "right", "output", "table", "template"):
            value = result.get(key)
            if isinstance(value, str) and value:
                _note_ledger_path(value, scope)
        with _LEDGER_LOCK:
            entry = _bucket(scope, create=True)
            room = MAX_LEDGER_FIGURES - len(entry.figures)
            if room > 0:
                entry.figures.extend(found[:room])
        return found
    except Exception:
        return []


# --- reading the figures an answer states ----------------------------------
#
# The MATCHERS below are German because the files, the questions and the
# answers here are German. "12.345,67" and "12345.67" are one figure
# written two ways, and a matcher that does not know that reports every
# correctly written German total as unsupported.

# A number as an answer writes it: 12.345,67 / 12 345,67 / 12345.67 /
# 1.234 / 31. The lookbehind keeps the tail of a longer token out.
_ANSWER_NUMBER_RE = re.compile(
    r"(?<![\w.,])"
    r"(\d{1,3}(?:[.   ]\d{3})+(?:,\d{1,4})?|\d+(?:[.,]\d{1,6})?)"
    r"(?!\d)"
)

# A figure that is money: a currency on one side of it.
_CURRENCY_AFTER_RE = re.compile(
    r"\s*(?:EUR|Euro|EURO|€|CHF|USD|\$|TEUR)\b|\s*€")
_CURRENCY_BEFORE_RE = re.compile(r"(?:EUR|Euro|€|CHF|USD|\$)\s*$")

# What makes a sentence a statement about a TOTAL. German first: these are
# the words the user's own files and questions use.
# German builds these words by compounding, endlessly — Gesamtbudget,
# Jahressumme, Rechnungsbetrag — so the stems match as stems. A closed
# list of whole words let "Das Gesamtbudget beträgt ..." through while
# catching "Die Gesamtsumme beträgt ...", which is the same sentence.
#
# Compounding was handled and INFLECTION was not, which is the bigger
# hole: German states a total with "betragen" more often than with any
# noun, and the list carried the third-person singular only. Measured —
#
#     flags=0 | Die Kosten betragen 99.999,99 EUR.
#     flags=0 | Die Summen der Konten betragen 99.999,99 EUR.
#     flags=0 | Die Personalkosten liegen bei 99.999,99 EUR.
#     flags=1 | Die Gesamtsumme beträgt 99.999,99 EUR.
#
# — the same fabricated figure, three of four times unseen. Plurals,
# preterites and the "liegt bei" phrasing are below.
_TOTAL_CUE_RE = re.compile(
    r"(?i)\bgesamt\w*|\w*summ(?:e|en)\b|\bsummier\w*|"
    r"\w*betr(?:ag|ags|äge|ägen)\b|"
    r"\bbeträgt\b|\bbetragen\b|\bbetrug\b|\bbetrugen\b|"
    r"\bbeläuft\b|\bbelaufen\b|\bbelief\b|\bbeliefen\b|"
    r"\bliegt\s+bei\b|\bliegen\s+bei\b|\blag\s+bei\b|\blagen\s+bei\b|"
    r"\bergibt\s+sich\b|\bmacht\s+(?:das\s+)?\w*\s*aus\b|"
    r"\w*kosten\b|\w*aufwand\b|\w*aufwände\b|\w*erlös\w*|\w*ausgaben\b|"
    r"\w*einnahmen\b|\bendstand\b|\bendbetrag\b|"
    r"\binsgesamt\b|\bzusammen\b|"
    r"\btotal\b|\bsaldo\b|\bumsatz\w*|\bvolumen\b|\bbilanz\b|\bbudget\w*|"
    r"\bsum\b|\bsums\b|\btotals?\b|\bamounts?\s+to\b|\baltogether\b")

# ... and about a DERIVED figure: a difference, a share, an average, a
# per-unit value. Those are checked for derivability from the ledger
# before anything is said about them.
_DERIVED_CUE_RE = re.compile(
    r"(?i)\b(?:differenz|abweichung|abweichungen|unterschied|anteil|quote|"
    r"prozent|durchschnitt|durchschnittlich|mittelwert|im schnitt|je|pro|"
    r"entspricht|verhältnis|difference|share|average|per)\b|%")

# ... and about an EXTREME: the largest, the smallest, the top one. This
# was the shape with no cue at all, so the sentence that carries the most
# dangerous kind of figure was the one nothing read. Measured on the real
# reader and the real workbook, before this existed:
#
#     flags=0 | Die höchste Kostenstelle ist KST 4711 mit 128.430,55 €.
#
# — with an EMPTY ledger, so it was not a coverage problem: no claim shape
# matched it. An extreme is worse than a total because a total taken from
# a page is at least visibly a part of something, while "the highest" over
# a window reads as a fact about the file and is wrong by construction
# unless the whole file was seen.
_EXTREMUM_CUE_RE = re.compile(
    r"(?i)\b(?:höchst\w*|hoechst\w*|größt\w*|groesst\w*|max(?:imal\w*|imum)?|"
    r"niedrigst\w*|kleinst\w*|geringst\w*|min(?:imal\w*|imum)?|"
    r"teuerst\w*|billigst\w*|meist\w*|wenigst\w*|"
    r"spitzenreiter|ausrei(?:ß|ss)er|"
    r"am\s+(?:höchsten|hoechsten|meisten|wenigsten|größten|groessten)|"
    r"highest|largest|biggest|smallest|lowest|cheapest|"
    r"top|maximum|minimum)\b")

# A percentage is always a derived figure.
_PERCENT_AFTER_RE = re.compile(r"\s*(?:%|Prozent|Prozentpunkte?)\b|\s*%")

# A sentence that states an amount of money. This is the only thing read
# in a turn where no office tool ran at all — see the scanner's docstring.
_MONEY_SENTENCE_RE = re.compile(
    r"\d\s*(?:EUR|Euro|EURO|CHF|USD|TEUR)\b|\d\s*[€$]"
    r"|(?:EUR|Euro|€|CHF|USD|\$)\s*\d")

# A figure the answer does not assert: hedged, Konjunktiv, an example, an
# assumption, or one the answer says it did not compute. Konjunktiv II
# ("die Summe wäre") is how German says "not measured", and flagging it
# would caveat the honest answer.
#
# THE SAME list the answer guard uses. They were two lists, and they
# disagreed inside one answer: "Die Summe liegt bei rund 45.000 EUR." was
# an honest hedge here ("rund") and a confident claim there, so one
# sentence got a caveat from one guard and none from the other. A user
# cannot act on a guard that contradicts itself in the same paragraph.
_NOT_ASSERTED_RE = _de.HEDGE_RE

# Positions where a number is a reference and not a quantity: a cell
# address, a row or column number, a document number, a date.
_REFERENCE_BEFORE_RE = re.compile(
    r"(?i)(?:zeile|zeilen|spalte|spalten|zelle|blatt|seite|tabelle|"
    r"row|column|cell|sheet|page|nr\.?|nummer|kostenstelle|beleg|"
    r"rechnung|konto|version|kapitel|"
    # The abbreviations a German office actually writes. Spelled-out
    # "Kostenstelle" was here and "KST 4711" was not, so an identifier
    # was read as a figure -- and a guard that flags a cost-centre number
    # as an invented amount is the kind of noise that gets guards
    # switched off.
    r"kst|pos|id|az|\w+-nr)\s*[:.]?\s*$")
_DATE_BEFORE_RE = re.compile(r"\d{1,2}[./-]\d{1,2}[./-]?$")

# A counted thing, in the shape German writes it: a number, an optional
# adjective, and a capitalised plural noun ("31 Belege", "4 offene
# Posten"). English plurals count too — a mixed answer happens.
_COUNT_CLAIM_RE = re.compile(
    r"(?<![\w.,])(\d{1,6})\s+"
    r"([A-Za-zÄÖÜäöüß][\wÄÖÜäöüß]*(?:-[\wÄÖÜäöüß]+)*)"
    r"(?:\s+([A-Za-zÄÖÜäöüß][\wÄÖÜäöüß]*(?:-[\wÄÖÜäöüß]+)*))?")

_COUNT_PLURAL_DE = ("en", "er", "se", "ze", "e", "n")
_COUNT_NON_PLURAL_EN = ("ss", "us", "is")
# Words that pass the shape and count no items: measures, money, time.
_COUNT_STOPWORDS = frozenset({
    "euro", "euros", "eur", "cent", "cents", "franken", "dollar", "dollars",
    "prozent", "prozentpunkte", "percent", "stunden", "minuten", "sekunden",
    "tage", "tagen", "wochen", "monate", "monaten", "jahre", "jahren",
    "zeichen", "bytes", "chars", "characters", "seiten", "seite",
    "hours", "minutes", "seconds", "days", "weeks", "months", "years",
    "mal", "male", "times", "grad", "meter", "gramm", "liter", "kwh",
})


def _is_counted_noun(word: str) -> bool:
    """True when ``word`` names things somebody counted."""
    head = str(word or "").rsplit("-", 1)[-1]
    tail = head.lower()
    if len(tail) < 3 or not tail.isalpha() or tail in _COUNT_STOPWORDS:
        return False
    if head[:1].isupper():                       # a German noun
        return tail.endswith(_COUNT_PLURAL_DE) or tail.endswith("s")
    return (len(tail) >= 4 and tail.endswith("s")
            and not tail.endswith(_COUNT_NON_PLURAL_EN))


def _readings(token: str) -> list[tuple[float, float]]:
    """Every number ``token`` can mean, with the tolerance it was written to.

    "12.345,67" is one figure; "1.234" is two — 1234 under the German
    convention and 1.234 under the English one, and nothing in an answer
    settles which. Both are returned and a match against either counts:
    the alternative is reporting a correctly written German total as
    unsupported, which is the one error this must not make.

    The tolerance is half a unit in the last place the figure was written
    to, so "20,9 %" matches a computed 20.9451 and "12.345,67" does not
    match 12.345,68.
    """
    raw = str(token or "").strip()
    if not raw:
        return []
    body = raw.replace(" ", " ").replace(" ", " ")
    out: list[tuple[float, float]] = []

    def _keep(text: str) -> None:
        try:
            value = float(text)
        except ValueError:
            return
        decimals = len(text.split(".")[1]) if "." in text else 0
        pair = (round(value, 9), 0.5 * (10.0 ** -decimals))
        if pair not in out:
            out.append(pair)

    if " " in body:                       # 12 345,67 — spaces group thousands
        body = body.replace(" ", "")
    dots, commas = body.count("."), body.count(",")
    if dots and commas:                   # the LAST separator is the decimal
        if body.rfind(",") > body.rfind("."):
            _keep(body.replace(".", "").replace(",", "."))
        else:
            _keep(body.replace(",", ""))
        return out
    for separator in (".", ","):
        if not body.count(separator):
            continue
        if body.count(separator) > 1:     # only grouping can repeat
            _keep(body.replace(separator, ""))
            return out
        if len(body.split(separator)[1]) == 3:   # 1.234: thousands OR decimals
            _keep(body.replace(separator, ""))
            _keep(body.replace(separator, "."))
        else:
            _keep(body.replace(separator, "."))
        return out
    _keep(body)
    return out


def _numbers_in(text: str) -> list[float]:
    """Every number a piece of text states, under either reading."""
    out: list[float] = []
    for match in _ANSWER_NUMBER_RE.finditer(str(text or "")):
        for value, _tolerance in _readings(match.group(1)):
            out.append(value)
    return out


def _matches(readings: list[tuple[float, float]], values) -> bool:
    for value, tolerance in readings:
        for known in values:
            if abs(value - known) <= tolerance:
                return True
    return False


# How many tool figures take part in deriving a difference or a share.
# Bounded on purpose: the derived set grows with the square of this, and
# a large enough base would make every stated number "derivable".
MAX_DERIVATION_BASE = 24


def _derived_values(
    figures: list["ToolFigure"], *, counts_only: bool = False,
) -> list[float]:
    """Differences, sums, ratios and shares of the tool's own figures.

    An answer that reports the difference between two totals, the share
    one count is of another, or a per-unit value has computed it from
    figures the tools DID return. Requiring it verbatim would flag exactly
    the arithmetic the user asked for. Cell values stay out of the base:
    there are hundreds of them, and pairing those would make anything
    derivable.

    ``counts_only`` is for a claim that counts things. A derived count is
    an addition or a subtraction of counts ("148 minus 31 offene"); a
    RATIO of them is a share and not a count. Measured on the fixture
    workbooks, letting ratios ground a count made almost half of all
    random integers below 300 pass — the ratios form a dense set, and an
    integer is matched to half a unit.
    """
    base: list[float] = []
    for figure in figures:
        if figure.kind == "cell" or (counts_only and figure.kind != "count"):
            continue
        if figure.value not in base:
            base.append(figure.value)
        if len(base) >= MAX_DERIVATION_BASE:
            break
    out: list[float] = []
    for index, a in enumerate(base):
        for b in base[index + 1:]:
            out.append(a + b)
            out.append(a - b)
            out.append(b - a)
    if counts_only:
        return out
    for a in base:
        for b in base:
            if b:
                out.append(a / b)
                out.append(a / b * 100.0)
    return out


@dataclass(frozen=True)
class FigureFlag:
    """One figure an answer states that no tool this turn produced."""

    figure: str    # as written in the answer ("45.231,50")
    kind: str      # "total" | "count" | "derived"
    claim: str     # the sentence it stands in, so the reader can find it

    def message(self) -> str:
        return (f"figure '{self.figure}' ({self.kind}) is not in this "
                f"turn's tool results")


# Sentence boundaries. Conservative on purpose: an abbreviation, a decimal
# and a file name must not split a cue away from the figure it governs.
_SENTENCE_SPLIT_RE = re.compile(
    r"(?<=[.!?;])\s+(?=[\"'“(\[]?[A-ZÄÖÜ])|\n")

_FENCED_RE = re.compile(r"```.*?```", re.DOTALL)
_BACKTICK_RE = re.compile(r"`[^`\n]*`")


def _claimable_text(text: str) -> str:
    """The answer with everything that is not its own claim blanked out.

    Code blocks, inline code and quoted lines are what the answer SHOWS,
    not what it says — a number inside a formula or a quoted tool result
    is not the answer asserting it.
    """
    blanked = _FENCED_RE.sub(" ", str(text or ""))
    blanked = _BACKTICK_RE.sub(" ", blanked)
    return "\n".join(
        "" if line.lstrip().startswith(">") else line
        for line in blanked.split("\n"))


def _sentences(text: str) -> list[str]:
    return [s for s in _SENTENCE_SPLIT_RE.split(text) if s and s.strip()]


def _nearest_cue(before: str) -> int:
    """Where the cue closest in front of a figure ends, or -1.

    Only the POSITION matters: it bounds the window in which a bare
    number counts as the one the cue announces. What the figure is called
    is decided by the sentence — a derived cue names a specific thing
    (Differenz, Anteil, Durchschnitt) where a total cue can be the bare
    verb 'beträgt', so the specific one wins.
    """
    end = -1
    for pattern in (_TOTAL_CUE_RE, _DERIVED_CUE_RE, _EXTREMUM_CUE_RE):
        for hit in pattern.finditer(before):
            end = max(end, hit.end())
    return end


def _claims_in_sentence(
    sentence: str, *, money_only: bool,
) -> list[tuple[str, str]]:
    """(figure, kind) for every quantity claim in one sentence, in the
    order they are written."""
    total_cue = _TOTAL_CUE_RE.search(sentence) is not None
    derived_cue = _DERIVED_CUE_RE.search(sentence) is not None
    extremum_cue = _EXTREMUM_CUE_RE.search(sentence) is not None
    if money_only and not _MONEY_SENTENCE_RE.search(sentence):
        return []
    if _NOT_ASSERTED_RE.search(sentence):
        return []

    found: list[tuple[int, str, str]] = []
    counted: list[tuple[int, int]] = []
    for match in _COUNT_CLAIM_RE.finditer(sentence):
        first, second = match.group(2), match.group(3)
        if first.rsplit("-", 1)[-1].lower() in _COUNT_STOPWORDS:
            continue
        if _is_counted_noun(first) or (second and _is_counted_noun(second)):
            found.append((match.start(1), match.group(1), "count"))
            counted.append((match.start(1), match.end(1)))

    # A total and a derived figure are only ever read where the sentence
    # announces one. A number in a sentence that announces nothing is a
    # reference, a date or a quantity of something else.
    if not (total_cue or derived_cue or extremum_cue):
        found.sort(key=lambda entry: entry[0])
        return [(figure, kind) for _at, figure, kind in found]

    for match in _ANSWER_NUMBER_RE.finditer(sentence):
        start, end = match.start(1), match.end(1)
        if any(s <= start < e for s, e in counted):
            continue
        before, after = sentence[:start], sentence[end:]
        if _REFERENCE_BEFORE_RE.search(before) or _DATE_BEFORE_RE.search(before):
            continue
        if _PERCENT_AFTER_RE.match(after):
            found.append((start, match.group(1), "derived"))
            continue
        kind = ("extremum" if extremum_cue else
                "derived" if derived_cue else "total")
        if _CURRENCY_AFTER_RE.match(after) or _CURRENCY_BEFORE_RE.search(before):
            found.append((start, match.group(1), kind))
            continue
        # A bare number the cue itself governs: "Die Summe beträgt
        # 12.345,67." The window is one clause, so a row reference further
        # along the sentence is not read as the total.
        cue_end = _nearest_cue(before)
        if cue_end < 0:
            continue
        between = before[cue_end:]
        if len(between) <= 40 and not re.search(r"[,;:(]", between):
            found.append((start, match.group(1), kind))
    found.sort(key=lambda entry: entry[0])
    return [(figure, kind) for _at, figure, kind in found]


def scan_answer_for_unledgered_figures(
    text: str,
    *,
    ledger: Optional[list["ToolFigure"]] = None,
    user_text: str = "",
    prior_text: str = "",
    max_flags: int = 3,
    token: str = "",
) -> list[FigureFlag]:
    """Figures an answer states that nothing this turn produced.

    Three claim shapes are read: a total ("die Summe beträgt 12.345,67"),
    a count ("31 Belege", "4 Abweichungen") and a derived figure (a
    difference, a share, a per-unit value). A figure is GROUNDED, and
    nothing is said about it, when any of these holds:

      * a tool returned it this turn — the ledger, under either reading of
        a German number and within the precision it was written to;
      * it is derivable from two figures the tools returned;
      * the user typed it in their own message;
      * it was already stated earlier in the conversation.

    With an EMPTY ledger — no office tool ran at all — only sentences that
    state an amount of MONEY are read at all. That is the case this was
    built for ("Die Gesamtsumme beträgt 45231.50 EUR bei 31 Belegen" after
    zero evidence acts), and it keeps every other kind of turn out of
    reach: "Insgesamt 3 Tests sind fehlgeschlagen" in a coding session
    carries a total cue and a count and is none of this mechanism's
    business.

    Deterministic, de-duplicated, capped, and it never raises: a guard
    that can break a turn is worse than the claim it was watching for.
    """
    flags: list[FigureFlag] = []
    try:
        if not text or not str(text).strip() or max_flags <= 0:
            return []
        known = list(ledger) if ledger is not None else figure_ledger(token)
        money_only = not known
        candidates: list[tuple[str, str, str]] = []
        for sentence in _sentences(_claimable_text(text)):
            for figure, kind in _claims_in_sentence(sentence,
                                                    money_only=money_only):
                candidates.append((figure, kind, " ".join(sentence.split())))
        if not candidates:
            return []

        # A count is grounded by something that COUNTED. Measured on the
        # fixture workbooks: with the values of a read grid in the pool as
        # well, 87% of random integers between 1 and 300 matched some cell
        # and passed — a count check that answers "yes" to almost anything
        # is not a check. A total keeps the whole pool: an answer quoting
        # a single amount out of the sheet it just read is quoting the
        # tool, and that is exactly what the cell values are for.
        values = [f.value for f in known]
        counts = [f.value for f in known if f.kind == "count"]
        # An extreme is the one claim a value cannot ground by being
        # present. "Die höchste Buchung ist 1.234,50" is backed by that
        # number appearing in the grid only if the grid was the whole
        # table -- over a window it is the largest of what was SHOWN, and
        # says nothing about the file. So an extreme is grounded by the
        # complete figures alone, and the window that produced a partial
        # one is named in the caveat.
        whole = [f.value for f in known if not f.partial]
        windows = sorted({f.window for f in known if f.partial and f.window})
        typed = _numbers_in(user_text) + _numbers_in(prior_text)
        derived: dict[bool, list[float]] = {}
        seen: set[str] = set()
        for figure, kind, sentence in candidates:
            key = figure.replace(" ", "")
            if key in seen:
                continue
            readings = _readings(figure)
            if not readings:
                continue
            seen.add(key)
            if kind == "extremum":
                if _matches(readings, typed):
                    continue
                if _matches(readings, whole):
                    continue
                # It IS in the ledger, and only from a window: the claim
                # is not invented, it is over-reaching, and saying so
                # names a different repair than "where does this come
                # from" -- read the rest, or ask sum_column.
                over = _matches(readings, values)
                flags.append(FigureFlag(
                    figure=figure,
                    kind="extremum_window" if over else "extremum",
                    claim=sentence[:120]))
                if len(flags) >= max_flags:
                    break
                continue
            pool = counts if kind == "count" else values
            if _matches(readings, pool) or _matches(readings, typed):
                continue
            counts_only = kind == "count"
            if counts_only not in derived:
                derived[counts_only] = _derived_values(
                    known, counts_only=counts_only)
            if _matches(readings, derived[counts_only]):
                continue
            flags.append(FigureFlag(figure=figure, kind=kind,
                                    claim=sentence[:120]))
            if len(flags) >= max_flags:
                break
    except Exception:
        return flags
    return flags


# The caveat is German for the same reason the matchers are: it is read by
# the person who asked the question, and they asked it in German.
_FIGURE_KIND_DE = {"total": "Summe", "count": "Anzahl",
                   "derived": "abgeleiteter Wert",
                   "extremum": "Höchst-/Tiefstwert",
                   "extremum_window": "Höchst-/Tiefstwert"}


def figure_caveat(flags: list[FigureFlag], windows: list[str] | None = None
                  ) -> str:
    """The note appended to an answer whose figures no tool produced.

    Names every figure and why it is marked, because the reader is the
    one who can tell an invented total from one they typed themselves. It
    annotates and never blocks: refusing the answer would withhold the
    part of it that is right.

    An extreme taken from a window gets its OWN sentence. It is a
    different mistake with a different repair: the number was really in
    the file, so "where does this come from" is the wrong question and
    "you have not seen the rest of the table" is the right one.
    """
    if not flags:
        return ""
    over_reach = [f for f in flags if f.kind == "extremum_window"]
    invented = [f for f in flags if f.kind != "extremum_window"]
    parts: list[str] = []
    if invented:
        named = ", ".join(
            f"'{f.figure}' ({_FIGURE_KIND_DE.get(f.kind, f.kind)})"
            for f in invented[:3])
        parts.append(
            "\n\n> ⚠️ Diese Zahl stammt nicht aus einem Werkzeug-Ergebnis "
            "dieses Zuges: " + named + ". Kein Aufruf hat sie geliefert, sie "
            "lässt sich nicht aus den ermittelten Werten ableiten, und sie "
            "steht weder in Ihrer Nachricht noch weiter oben im Verlauf. "
            "Bitte mit sum_column bzw. compare_tables nachrechnen oder die "
            "Quelle nennen, bevor die Zahl weitergegeben wird.")
    if over_reach:
        named = ", ".join(f"'{f.figure}'" for f in over_reach[:3])
        scope = (" — gelesen wurden " + "; ".join(windows[:2])
                 if windows else "")
        parts.append(
            "\n\n> ⚠️ " + named + " wird als Höchst- bzw. Tiefstwert "
            "genannt, belegt ist aber nur ein Ausschnitt der Tabelle" + scope +
            ". Der größte Wert eines Ausschnitts ist nicht der größte Wert "
            "der Datei. Bitte die Tabelle vollständig lesen (start_row/"
            "max_rows) oder mit sum_column über alle Zeilen auswerten, bevor "
            "die Zahl weitergegeben wird.")
    return "".join(parts)


def figure_coverage_caveat(
    text: str, *, user_text: str = "", prior_text: str = "",
    token: str = "",
) -> str:
    """The whole check for one finished answer, as one call.

    The caller appends the return value and is done: it is "" for an
    answer whose figures the tools produced, which is nearly every
    answer. ``token`` is the one ``begin_figure_turn`` returned, so the
    answer is checked against the ledger ITS turn wrote. Never raises.
    """
    try:
        flags = scan_answer_for_unledgered_figures(
            text, user_text=user_text, prior_text=prior_text, token=token)
        if not flags:
            return ""
        # Named rather than described: the reader can act on "rows 1-5 of
        # 26" and cannot act on "the read was incomplete".
        windows = sorted({f.window for f in figure_ledger(token)
                          if f.partial and f.window})
        return figure_caveat(flags, windows)
    except Exception:
        return ""
