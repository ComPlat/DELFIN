"""Combined DELFIN Word report assembled from DELFIN_Data.json and generated plots."""

from __future__ import annotations

from dataclasses import dataclass
import re
from pathlib import Path
from typing import Any, Dict, Iterable, Optional
import json

from delfin.common.logging import get_logger
from delfin.ir_spectrum import parse_ir_spectrum

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    DOCX_AVAILABLE = False

HARTREE_TO_EV = 27.211386245988


def _translate_state(orca_state: str) -> str:
    """Translate ORCA notation (0-1A/1-3A) to S0/S1/T1 for readability."""
    match = re.match(r'(\d+)-([13])A', str(orca_state))
    if not match:
        return str(orca_state)

    root_number = int(match.group(1))
    multiplicity = match.group(2)
    if multiplicity == "1":
        return f"S{root_number}"
    if multiplicity == "3":
        return f"T{root_number}"
    return str(orca_state)


def _wavelength_to_rgb(wavelength: float) -> tuple[int, int, int]:
    """Map wavelength (nm) to approximate RGB (copied from AFP plotting logic)."""
    if wavelength < 380:
        r, g, b = 0.5, 0.0, 0.5
    elif wavelength < 440:
        r = -(wavelength - 440) / (440 - 380)
        g = 0.0
        b = 1.0
    elif wavelength < 490:
        r = 0.0
        g = (wavelength - 440) / (490 - 440)
        b = 1.0
    elif wavelength < 510:
        r = 0.0
        g = 1.0
        b = -(wavelength - 510) / (510 - 490)
    elif wavelength < 580:
        r = (wavelength - 510) / (580 - 510)
        g = 1.0
        b = 0.0
    elif wavelength < 645:
        r = 1.0
        g = -(wavelength - 645) / (645 - 580)
        b = 0.0
    elif wavelength < 781:
        r, g, b = 1.0, 0.0, 0.0
    else:
        r, g, b = 0.5, 0.0, 0.0

    if wavelength < 420:
        factor = 0.3 + 0.7 * (wavelength - 380) / (420 - 380)
    elif wavelength > 700:
        factor = 0.3 + 0.7 * (780 - wavelength) / (780 - 700)
    else:
        factor = 1.0

    r = int(max(0, min(1, r * factor)) * 255)
    g = int(max(0, min(1, g * factor)) * 255)
    b = int(max(0, min(1, b * factor)) * 255)
    return r, g, b


@dataclass
class ReportAssets:
    """Container for optional plot assets referenced in the final report."""

    afp_png: Optional[Path] = None
    smiles_png: Optional[Path] = None
    uv_vis_pngs: Dict[str, Path] | None = None  # keyed by state name, e.g., "S0", "S1", "T1"
    ir_png: Optional[Path] = None
    energy_level_png: Optional[Path] = None
    vertical_excitation_png: Optional[Path] = None
    correlation_png: Optional[Path] = None
    dipole_moment_png: Optional[Path] = None
    mo_pngs: Dict[str, Path] | None = None  # keyed by orbital name, e.g., "HOMO", "LUMO+1"
    esp_png: Optional[Path] = None


def _prevent_row_splits(table) -> None:
    """Prevent individual table rows from splitting across pages."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Set cantSplit property on each row to prevent splitting
        for i, row in enumerate(table.rows):
            tr = row._element
            tr_pr = tr.xpath('./w:trPr')
            if tr_pr:
                tr_pr = tr_pr[0]
            else:
                tr_pr = OxmlElement('w:trPr')
                tr.insert(0, tr_pr)

            # Set cantSplit to prevent this row from splitting across pages
            cant_split = OxmlElement('w:cantSplit')
            tr_pr.append(cant_split)

            # Keep first 5 rows together with heading using keepNext
            if i < 5:
                keep_next = OxmlElement('w:keepNext')
                tr_pr.append(keep_next)
    except Exception as exc:
        logger.warning(f"Failed to set table row properties: {exc}")


def _keep_heading_with_table(heading_paragraph) -> None:
    """Keep heading with the following table."""
    try:
        heading_paragraph.paragraph_format.keep_with_next = True
    except Exception as exc:
        logger.warning(f"Failed to set heading keep_with_next: {exc}")


def _add_key_value_table(doc: Document, title: str, rows: Iterable[tuple[str, str]]) -> None:
    """Render a simple two-column table."""
    heading = doc.add_heading(title, level=2)
    _keep_heading_with_table(heading)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    _prevent_row_splits(table)


def _add_plot_if_exists(doc: Document, title: str, image_path: Optional[Path]) -> None:
    if not image_path:
        return
    if not image_path.exists():
        logger.warning("Plot '%s' not found at %s", title, image_path)
        return
    doc.add_heading(title, level=2)
    doc.add_picture(str(image_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _xyz_body_to_mol(xyz_lines: list[str]):
    """Convert XYZ body (no header lines) to RDKit Mol via bond perception."""
    try:
        from rdkit import Chem, RDLogger
        from rdkit.Chem import rdDetermineBonds
        from rdkit.Geometry import Point3D

        RDLogger.DisableLog("rdApp.*")
    except Exception:
        return None

    atoms = []
    coords = []
    for line in xyz_lines:
        if not line.strip():
            continue
        parts = line.split()
        if len(parts) < 4:
            return None
        symbol = parts[0]
        try:
            x, y, z = map(float, parts[1:4])
        except Exception:
            return None
        atoms.append(symbol)
        coords.append(Point3D(x, y, z))

    if not atoms:
        return None

    try:
        mol = Chem.RWMol()
        for sym in atoms:
            mol.AddAtom(Chem.Atom(sym))
        conf = Chem.Conformer(len(atoms))
        for idx, pt in enumerate(coords):
            conf.SetAtomPosition(idx, pt)
        mol.AddConformer(conf)
        rdDetermineBonds.DetermineBonds(mol)
        Chem.SanitizeMol(mol)
        return mol
    except Exception:
        return None


def _mol_to_smiles(mol):
    try:
        from rdkit import Chem
        smi = Chem.MolToSmiles(mol)
        return smi
    except Exception:
        return None


def _generate_smiles_image(project_dir: Path) -> Optional[Path]:
    """Render SMILES from input.txt (SMILES or XYZ body) to a PNG using RDKit."""
    smiles_file = project_dir / "input.txt"
    if not smiles_file.exists():
        logger.warning("SMILES source 'input.txt' not found for RDKit rendering")
        return None

    try:
        lines = [line.strip() for line in smiles_file.read_text(encoding="utf-8", errors="ignore").splitlines()]
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to read input.txt: %s", exc)
        return None

    smiles = ""
    mol_from_xyz = None

    # Detect if content looks like XYZ body (no header lines)
    xyz_candidate = [ln for ln in lines if ln and not ln.startswith("#")]
    if xyz_candidate and all(len(ln.split()) >= 4 and ln.split()[0][0].isalpha() for ln in xyz_candidate):
        # Try XYZ -> Mol
        mol_from_xyz = _xyz_body_to_mol(xyz_candidate)

    if mol_from_xyz:
        smiles = _mol_to_smiles(mol_from_xyz) or ""
    else:
        # Fallback: treat first non-empty line as SMILES
        smiles = next((ln for ln in lines if ln), "")

    if not smiles:
        logger.warning("No SMILES or usable XYZ detected in input.txt")
        return None

    try:
        from rdkit import Chem, RDLogger
        from rdkit.Chem import AllChem, Draw

        RDLogger.DisableLog("rdApp.*")
    except Exception as exc:  # noqa: BLE001
        logger.warning("RDKit not available for SMILES rendering: %s", exc)
        return None

    mol = Chem.MolFromSmiles(smiles) if not mol_from_xyz else mol_from_xyz
    if mol is None:
        logger.warning("Invalid SMILES in input.txt: %s", smiles)
        return None

    try:
        # Hide hydrogens in the depiction for a cleaner schematic
        mol_no_h = Chem.RemoveHs(mol, updateExplicitCount=True)

        # Normalize the molecule (standardize functional groups)
        try:
            from rdkit.Chem.MolStandardize import rdMolStandardize
            normalizer = rdMolStandardize.Normalizer()
            mol_no_h = normalizer.normalize(mol_no_h)
        except Exception:
            pass

        # Try to use CoordGen for better 2D coordinates (if available)
        # CoordGen produces publication-quality layouts
        coordgen_success = False
        try:
            from rdkit.Chem import rdCoordGen
            params = rdCoordGen.CoordGenParams()
            params.SetCoordgenScaling(1.5)  # Slightly larger structure
            params.SetMinimizerPrecision(1e-4)  # Higher precision
            coordgen_success = rdCoordGen.AddCoords(mol_no_h, params)
        except Exception:
            pass

        # If CoordGen failed or not available, use optimized Compute2DCoords
        if not coordgen_success:
            # Use AllChem with optimized parameters and ring templates
            AllChem.Compute2DCoords(mol_no_h,
                                   canonOrient=True,
                                   clearConfs=True,
                                   coordMap={},
                                   nFlipsPerSample=3,
                                   nSample=100,
                                   sampleSeed=42,
                                   permuteDeg4Nodes=True,
                                   bondLength=1.5)

            # Try to straighten the depiction for cleaner appearance
            try:
                from rdkit.Chem import rdDepictor
                rdDepictor.StraightenDepiction(mol_no_h)
            except Exception:
                pass

        # Use high-resolution rendering with better options
        drawer_options = Draw.MolDrawOptions()
        drawer_options.clearBackground = True
        drawer_options.bondLineWidth = 3
        drawer_options.multipleBondOffset = 0.15
        drawer_options.addStereoAnnotation = True
        drawer_options.addAtomIndices = False

        # Use larger canvas with more detail
        img = Draw.MolToImage(mol_no_h, size=(500, 400), options=drawer_options, kekulize=True)
        out_path = project_dir / "SMILES.png"
        img.save(out_path)
        return out_path
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to render SMILES to image: %s", exc)
        return None


def _render_color_chip(rgb: tuple[int, int, int], project_dir: Path, label: str) -> Optional[Path]:
    """Create a small PNG color chip for Word compatibility."""
    try:
        from PIL import Image
    except Exception:
        return None

    try:
        chip = Image.new("RGB", (60, 40), rgb)
        safe_label = (
            label.replace("→", "-")
            .replace("₀", "0")
            .replace("₁", "1")
            .replace("₂", "2")
            .replace("₃", "3")
            .replace("₄", "4")
            .replace("₅", "5")
            .replace("₆", "6")
        )
        out_path = project_dir / f"color_chip_{safe_label}.png"
        chip.save(out_path)
        return out_path
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to create color chip: %s", exc)
        return None


def _format_energy_block(optimization: Dict[str, Any]) -> str:
    def _fmt(value, fmt: str) -> Optional[str]:
        try:
            return fmt.format(float(value))
        except Exception:
            return None

    parts: list[str] = []
    hartree = _fmt(optimization.get("hartree"), "{:.6f}")
    if hartree:
        parts.append(f"{hartree} Eh")
    ev = _fmt(optimization.get("eV"), "{:.3f}")
    if ev:
        parts.append(f"{ev} eV")
    kj = _fmt(optimization.get("kJ_mol"), "{:.2f}")
    if kj:
        parts.append(f"{kj} kJ/mol")
    return ", ".join(parts) if parts else "n/a"


def _energy_ev(optimization: Dict[str, Any]) -> Optional[float]:
    if not optimization:
        return None
    if optimization.get("eV") is not None:
        try:
            return float(optimization["eV"])
        except Exception:
            return None
    if optimization.get("hartree") is not None:
        try:
            return float(optimization["hartree"]) * HARTREE_TO_EV
        except Exception:
            return None
    return None


def _parse_goat_conformer_count(project_dir: Path) -> Optional[int]:
    """Parse 'Conformers below 3 kcal/mol: X' from XTB_GOAT output."""
    goat_candidates = [
        project_dir / "XTB_GOAT" / "output_XTB_GOAT.out",
        project_dir / "XTB2_GOAT" / "output_XTB_GOAT.out",
        project_dir / "output_XTB_GOAT.out",
    ]

    for goat_file in goat_candidates:
        if not goat_file.exists():
            continue
        try:
            content = goat_file.read_text(encoding="utf-8", errors="ignore")
            match = re.search(r'Conformers below 3 kcal/mol:\s*(\d+)', content)
            if match:
                return int(match.group(1))
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to parse GOAT conformer count from %s: %s", goat_file, exc)
    return None


def _build_summary_text(data: Dict[str, Any], project_dir: Path) -> tuple[Optional[str], list[tuple[str, float, tuple[int, int, int]]]]:
    meta = data.get("metadata", {}) or {}
    name = meta.get("NAME") or meta.get("name") or project_dir.name
    functional = meta.get("functional") or "unknown functional"
    basis = meta.get("basis_set") or "unknown basis"
    aux_basis = meta.get("auxiliary_basis")
    ri_method = meta.get("ri_method")
    solvent_model = meta.get("implicit_solvation")
    solvent = meta.get("solvent")
    dispersion = meta.get("dispersion_correction")

    gs = data.get("ground_state_S0", {}) or {}
    gs_opt = gs.get("optimization", {}) or {}
    gs_orb = gs.get("orbitals", {}) or {}
    gs_dip = gs.get("dipole_moment", {}) or {}
    s0_e = _energy_ev(gs_opt)

    homo = gs_orb.get("homo_eV")
    lumo = gs_orb.get("lumo_eV")
    gap = gs_orb.get("gap_eV")

    # Vibrational highlights from IR spectrum if available
    vib_frequencies: list[float] = []
    negative_freqs = 0
    ir_file = None
    for candidate in ["ESD/S0.out", "S0.out", "initial.out"]:
        cand_path = project_dir / candidate
        if cand_path.exists():
            ir_file = cand_path
            break
    if ir_file:
        try:
            modes = parse_ir_spectrum(ir_file)
            negative_freqs = sum(1 for m in modes if m.frequency_cm1 < 0)
            vib_frequencies = sorted(
                [m.frequency_cm1 for m in modes],
                key=lambda x: x,
            )
            # Keep top 5 by intensity for summary
            modes_sorted = sorted(modes, key=lambda m: m.intensity_km_mol, reverse=True)
            vib_frequencies = [m.frequency_cm1 for m in modes_sorted[:5]]
        except Exception as exc:  # noqa: BLE001
            logger.warning("IR parsing failed for summary: %s", exc)

    excited = data.get("excited_states", {}) or {}
    excited_count = len(excited)
    s1_e = _energy_ev((excited.get("S1") or {}).get("optimization", {}) or {})
    t1_e = _energy_ev((excited.get("T1") or {}).get("optimization", {}) or {})

    # Extract dipole moments from excited states
    s1_dip = ((excited.get("S1") or {}).get("dipole_moment", {}) or {})
    t1_dip = ((excited.get("T1") or {}).get("dipole_moment", {}) or {})

    def rel_energy(ref, val):
        if ref is None or val is None:
            return None
        return val - ref

    s1_gap = rel_energy(s0_e, s1_e)
    t1_gap = rel_energy(s0_e, t1_e)
    est = None
    if s1_gap is not None and t1_gap is not None:
        est = abs(s1_gap - t1_gap)

    # Absorption highlights from S0 transitions
    s0_abs = (gs.get("tddft_absorption") or {}).get("transitions", []) or []
    if s0_abs:
        top_abs = sorted(
            s0_abs,
            key=lambda t: t.get("oscillator_strength", 0.0) or 0.0,
            reverse=True,
        )[:2]
        abs_peaks = [t.get("wavelength_nm") for t in top_abs if t.get("wavelength_nm")]
    else:
        abs_peaks = []

    def _pick_transition(transitions: list[Dict[str, Any]], state_a: str, state_b: str) -> Optional[tuple[float, float]]:
        """Pick strongest transition between two states (order-agnostic). Returns (fosc, wavelength_nm)."""
        best: tuple[float, float] | None = None
        target_set = {state_a, state_b}
        for t in transitions or []:
            from_state = _translate_state(t.get("from_state", ""))
            to_state = _translate_state(t.get("to_state", ""))
            if {from_state, to_state} != target_set:
                continue
            try:
                fosc = float(t.get("oscillator_strength", 0.0) or 0.0)
            except Exception:
                fosc = 0.0
            wl = t.get("wavelength_nm")
            if wl is None:
                continue
            if best is None or fosc > best[0]:
                best = (fosc, wl)
        return best

    # Key transitions for AFP plot: S0↔S1, S1↔S0, T1↔S0 (order-agnostic)
    s0_to_s1 = _pick_transition(s0_abs, "S0", "S1")

    s1_transitions = (excited.get("S1", {}) or {}).get("tddft_from_geometry", {}).get("transitions", [])
    s1_to_s0 = _pick_transition(s1_transitions, "S0", "S1")

    t1_transitions = (excited.get("T1", {}) or {}).get("tddft_from_geometry", {}).get("transitions", [])
    t1_to_s0 = _pick_transition(t1_transitions, "S0", "T1")

    s0_to_s1_nm = s0_to_s1[1] if s0_to_s1 else None
    s1_to_s0_nm = s1_to_s0[1] if s1_to_s0 else None
    t1_to_s0_nm = t1_to_s0[1] if t1_to_s0 else None

    # Redox potentials from DELFIN.txt and OCCUPIER data
    delfin_summary = data.get("delfin_summary", {}) or {}
    redox_potentials_fc = delfin_summary.get("redox_potentials_vs_fc", {})
    oxidized_states = data.get("oxidized_states", {}) or {}
    reduced_states = data.get("reduced_states", {}) or {}
    occupier_data = data.get("occupier", {}) or {}
    control_flags = data.get("control_flags", {}) or {}

    def fmt(val, suffix=""):
        if val is None:
            return "n/a"
        try:
            return f"{float(val):.2f}{suffix}"
        except Exception:
            return str(val) + suffix

    # Parse GOAT conformer count if available
    goat_conformer_count = _parse_goat_conformer_count(project_dir)

    # Build summary
    parts = []

    # Opening statement about workflow
    opening = "The SMILES of the target system were converted into 3D coordinates using RDKit. "
    if goat_conformer_count is not None:
        opening += (
            f"A global geometry optimization with the GOAT algorithm and GFN-xTB in ORCA "
            f"identified {goat_conformer_count} conformer{'s' if goat_conformer_count != 1 else ''} "
            f"below 3 kcal/mol. The energetically lowest conformer provided the starting structure for subsequent high-level DFT calculations using ORCA. "
        )
    else:
        opening += (
            "A global geometry optimization with the GOAT algorithm and GFN-xTB in ORCA "
            "provided the starting structure for subsequent high-level DFT calculations using ORCA. "
        )
    opening += "The entire workflow, from initial structure generation to high-level DFT calculations, was automated by DELFIN."
    parts.append(opening)

    # Theory level with all details
    theory_parts = [f"{functional}/{basis}"]
    if ri_method:
        theory_parts.append(ri_method)
    if aux_basis:
        theory_parts.append(aux_basis)
    if dispersion:
        theory_parts.append(dispersion)
    if solvent_model and solvent:
        theory_parts.append(f"{solvent_model}({solvent})")
    elif solvent_model:
        theory_parts.append(solvent_model)

    parts.append(f"The calculations were performed at the {' '.join(theory_parts)} level of theory.")

    # Ground state properties (NO SCF energy)
    parts.append(f"HOMO/LUMO energies are {fmt(homo, ' eV')} and {fmt(lumo, ' eV')}, yielding a gap of {fmt(gap, ' eV')}.")

    # Dipole moments for S0, S1, T1
    dipole_parts = []
    s0_dipole = gs_dip.get('magnitude_debye')
    s1_dipole = s1_dip.get('magnitude_debye')
    t1_dipole = t1_dip.get('magnitude_debye')

    if s0_dipole is not None:
        dipole_parts.append(f"S₀: {fmt(s0_dipole, ' D')}")
    if s1_dipole is not None:
        dipole_parts.append(f"S₁: {fmt(s1_dipole, ' D')}")
    if t1_dipole is not None:
        dipole_parts.append(f"T₁: {fmt(t1_dipole, ' D')}")

    if dipole_parts:
        parts.append(f"The permanent dipole moments μ are {', '.join(dipole_parts)}.")

    # Polarizability (if available)
    polarizability = gs.get("polarizability", {}) or {}
    if polarizability:
        alpha_iso_au = polarizability.get("isotropic_au")
        alpha_iso_a3 = polarizability.get("isotropic_angstrom3")
        if alpha_iso_au is not None and alpha_iso_a3 is not None:
            parts.append(f"The isotropic polarizability α is {alpha_iso_au:.2f} a.u. ({alpha_iso_a3:.2f} Å³).")

    # Hyperpolarizability (if available)
    hyperpol = gs.get("hyperpolarizability", {}) or {}
    if hyperpol:
        beta_tot_esu = hyperpol.get("beta_tot_esu")
        beta_mu_esu = hyperpol.get("beta_mu_esu")
        if beta_tot_esu is not None:
            # Format in scientific notation with proper subscript notation
            beta_tot_str = f"{beta_tot_esu:.2e}"
            parts.append(f"The static hyperpolarizability β\u209C\u2092\u209C is {beta_tot_str} esu.")
        if beta_mu_esu is not None:
            beta_mu_str = f"{beta_mu_esu:.2e}"
            parts.append(f"The dipole-projected hyperpolarizability β\u03BC is {beta_mu_str} esu.")

    if vib_frequencies:
        vib_list = ", ".join(f"{freq:.0f}" for freq in vib_frequencies)
        parts.append(f"The most intense vibrational frequencies are approximately {vib_list} cm⁻¹.")
    if ir_file:
        parts.append(f"Negative frequencies detected: {negative_freqs}.")

    # Excited states + TDDFT settings from S1
    tddft_settings = (excited.get("S1", {}) or {}).get("tddft_settings", {})
    if tddft_settings:
        tokens = []
        for key in ["nroots", "maxdim", "followiroot"]:
            if key in tddft_settings:
                tokens.append(f"{key.capitalize()}={tddft_settings[key]}")
        settings_str = f" using TDDFT ({', '.join(tokens)})" if tokens else " using TDDFT"
    else:
        settings_str = ""
    if excited:
        state_list = ", ".join(sorted(excited.keys()))
        parts.append(f"In total, {excited_count} excited states were calculated{settings_str}: {state_list}.")
    else:
        parts.append(f"In total, {excited_count} excited states were calculated{settings_str}.")
    if est is not None:
        parts.append(f"The singlet–triplet energy gap ΔEₛₜ(S₁–T₁) is {fmt(est, ' eV')} (absolute difference).")

    # Absorption peaks
    if abs_peaks:
        peak_str = " and ".join(f"{w:.0f} nm" for w in abs_peaks if isinstance(w, (int, float)))
        if peak_str:
            parts.append(f"The most intense absorption peaks (S₀→Sₙ) are at {peak_str}.")

    # AFP spectrum wavelengths (absorption/fluorescence/phosphorescence)
    color_boxes: list[tuple[str, float, tuple[int, int, int]]] = []
    if any(v is not None for v in (s0_to_s1_nm, s1_to_s0_nm, t1_to_s0_nm)):
        def fmt_nm(val):
            try:
                return f"{float(val):.0f} nm"
            except Exception:
                return "n/a"

        def fmt_fosc(val):
            try:
                return f"{float(val):.4f}"
            except Exception:
                return "n/a"

        s0_to_s1_fosc = s0_to_s1[0] if s0_to_s1 else None
        s1_to_s0_fosc = s1_to_s0[0] if s1_to_s0 else None
        t1_to_s0_fosc = t1_to_s0[0] if t1_to_s0 else None

        afp_parts = [
            f"S₀→S₁ (absorption): {fmt_nm(s0_to_s1_nm)} (fosc = {fmt_fosc(s0_to_s1_fosc)})",
            f"S₁→S₀ (fluorescence): {fmt_nm(s1_to_s0_nm)} (fosc = {fmt_fosc(s1_to_s0_fosc)})",
            f"T₁→S₀ (phosphorescence): {fmt_nm(t1_to_s0_nm)} (fosc = {fmt_fosc(t1_to_s0_fosc)})",
        ]
        parts.append(f"Transitions (fosc = oscillator strength, transition intensity, dimensionless): {', '.join(afp_parts)}.")

        # Collect color chips for these wavelengths
        for label, wl in [
            ("S₀→S₁", s0_to_s1_nm),
            ("S₁→S₀", s1_to_s0_nm),
            ("T₁→S₀", t1_to_s0_nm),
        ]:
            if wl:
                color_boxes.append((label, wl, _wavelength_to_rgb(float(wl))))

    # Redox potentials / oxidation-reduction context
    if oxidized_states or reduced_states:
        redox_parts = []
        def _step_label(name: str) -> str:
            match = re.search(r"_step_(\d+)", name)
            return f" step {match.group(1)}" if match else ""

        for state_name, state_data in oxidized_states.items():
            opt = state_data.get("optimization", {}) or {}
            charge = opt.get("charge")
            mult = opt.get("multiplicity")
            energy = opt.get("hartree")
            if charge is not None and energy is not None:
                step_info = _step_label(state_name)
                # Try to get OCCUPIER info for this state
                occ_key = state_name.replace("_step_", "")
                occ_info = occupier_data.get(occ_key, {})
                preferred_idx = occ_info.get("preferred_index")
                if preferred_idx is not None and occ_info.get("entries"):
                    preferred_entry = next((e for e in occ_info["entries"] if e.get("index") == preferred_idx), None)
                    if preferred_entry:
                        mult = preferred_entry.get("multiplicity", mult)
                        unpaired_alpha = preferred_entry.get("unpaired_alpha")
                        unpaired_beta = preferred_entry.get("unpaired_beta")
                        spin_cont = preferred_entry.get("spin_contamination")
                        info = f"oxidation{step_info} (charge {charge:+d}, multiplicity {mult}"
                        if unpaired_alpha is not None and unpaired_beta is not None:
                            info += f", α/β electrons: {unpaired_alpha}/{unpaired_beta}"
                        if spin_cont is not None:
                            info += f", spin contamination: {spin_cont:.3f}"
                        info += ")"
                        redox_parts.append(info)
                    else:
                        redox_parts.append(f"oxidation{step_info} (charge {charge:+d}, multiplicity {mult})")
                else:
                    redox_parts.append(f"oxidation{step_info} (charge {charge:+d}, multiplicity {mult})")

        for state_name, state_data in reduced_states.items():
            opt = state_data.get("optimization", {}) or {}
            charge = opt.get("charge")
            mult = opt.get("multiplicity")
            energy = opt.get("hartree")
            if charge is not None and energy is not None:
                step_info = _step_label(state_name)
                occ_key = state_name.replace("_step_", "")
                occ_info = occupier_data.get(occ_key, {})
                preferred_idx = occ_info.get("preferred_index")
                if preferred_idx is not None and occ_info.get("entries"):
                    preferred_entry = next((e for e in occ_info["entries"] if e.get("index") == preferred_idx), None)
                    if preferred_entry:
                        mult = preferred_entry.get("multiplicity", mult)
                        unpaired_alpha = preferred_entry.get("unpaired_alpha")
                        unpaired_beta = preferred_entry.get("unpaired_beta")
                        spin_cont = preferred_entry.get("spin_contamination")
                        info = f"reduction{step_info} (charge {charge:+d}, multiplicity {mult}"
                        if unpaired_alpha is not None and unpaired_beta is not None:
                            info += f", α/β electrons: {unpaired_alpha}/{unpaired_beta}"
                        if spin_cont is not None:
                            info += f", spin contamination: {spin_cont:.3f}"
                        info += ")"
                        redox_parts.append(info)
                    else:
                        redox_parts.append(f"reduction{step_info} (charge {charge:+d}, multiplicity {mult})")
                else:
                    redox_parts.append(f"reduction{step_info} (charge {charge:+d}, multiplicity {mult})")

        if redox_parts:
            parts.append(f"Redox calculations were performed for {', '.join(redox_parts)}.")

    if not redox_potentials_fc or (len(redox_potentials_fc) == 1 and 'E_ref' in redox_potentials_fc):
        parts.append("Redox potentials were not calculated.")

    if not occupier_data:
        parts.append("No spin state analysis were done using OCCUPIER.")

    if control_flags.get("imag"):
        parts.append("IMAG was used to eliminate imaginary frequencies.")

    # Add redox potentials vs Fc+/Fc from DELFIN.txt
    if redox_potentials_fc and len(redox_potentials_fc) > 1:
        redox_strs = []
        label_map = {
            "E_red": "E_red1",
            "E_red_2": "E_red2",
            "E_red_3": "E_red3",
            "E_ox": "E_ox1",
            "E_ox_2": "E_ox2",
            "E_ox_3": "E_ox3",
        }
        # Reductions
        for key in ['E_red', 'E_red_2', 'E_red_3']:
            if key in redox_potentials_fc:
                val = redox_potentials_fc[key]
                label = label_map.get(key, key)
                redox_strs.append(f"{label}: {val:.3f} V")
        # Oxidations
        for key in ['E_ox', 'E_ox_2', 'E_ox_3']:
            if key in redox_potentials_fc:
                val = redox_potentials_fc[key]
                label = label_map.get(key, key)
                redox_strs.append(f"{label}: {val:.3f} V")
        # Reference
        if 'E_ref' in redox_potentials_fc:
            redox_strs.append(f"E_ref: {redox_potentials_fc['E_ref']:.3f} V")

        if redox_strs:
            parts.append(f"Redox potentials (vs. Fc⁺/Fc): {', '.join(redox_strs)}.")

    return " ".join(parts), color_boxes


def _style_header_row(row) -> None:
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True


def _format_state_with_subscript(paragraph, state_text: str) -> None:
    """Format state text with subscript numbers (e.g., S₀, S₁, T₁)."""
    import re
    # Match patterns like S0, S1, T1, T2, etc.
    match = re.match(r'([ST])(\d+)', state_text)
    if match:
        letter = match.group(1)
        number = match.group(2)
        # Add letter without bold
        run1 = paragraph.add_run(letter)
        run1.font.bold = False
        # Add number as subscript
        run2 = paragraph.add_run(number)
        run2.font.subscript = True
        run2.font.bold = False
    else:
        # Fallback for non-standard formats
        run = paragraph.add_run(state_text)
        run.font.bold = False


def _add_heading_with_subscript(doc: Document, title: str, level: int = 2) -> None:
    """Add a heading with subscript formatting for state labels (e.g., S₀, T₁)."""
    import re

    heading = doc.add_heading(level=level)

    # Split title by state patterns (S0, S1, T1, etc.)
    parts = re.split(r'([ST]\d+)', title)

    for part in parts:
        if re.match(r'[ST]\d+', part):
            # This is a state label - format with subscript
            match = re.match(r'([ST])(\d+)', part)
            if match:
                letter = match.group(1)
                number = match.group(2)
                run1 = heading.add_run(letter)
                run2 = heading.add_run(number)
                run2.font.subscript = True
        else:
            # Regular text
            heading.add_run(part)


def _format_significant_figures(value, sig_figs: int = 3) -> str:
    """Format a number to specified significant figures."""
    try:
        val = float(value)
        if val == 0:
            return "0.00"
        from math import log10, floor
        return f"{val:.{sig_figs - 1 - int(floor(log10(abs(val))))}f}"
    except (ValueError, TypeError):
        return str(value)


def _add_state_table(doc: Document, title: str, states: Dict[str, Any]) -> None:
    if not states:
        return
    heading = doc.add_heading(title, level=2)
    _keep_heading_with_table(heading)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["State", "Type", "Charge", "Multiplicity", "Energy (Eh)", "ZPE (Eh)"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
    _style_header_row(table.rows[0])

    for state_name, entry in sorted(states.items()):
        opt = entry.get("optimization", {}) or {}
        thermo = entry.get("thermochemistry", {}) or {}
        row = table.add_row().cells

        # Format State with subscript
        row[0].text = ""
        _format_state_with_subscript(row[0].paragraphs[0], state_name)

        row[1].text = str(entry.get("_type", ""))
        row[2].text = str(opt.get("charge", ""))
        row[3].text = str(opt.get("multiplicity", ""))

        # Only show Hartree energy with 3 significant figures
        hartree = opt.get("hartree")
        row[4].text = _format_significant_figures(hartree, sig_figs=3) if hartree is not None else ""

        row[5].text = (
            _format_significant_figures(thermo.get('zero_point_energy_hartree'), sig_figs=3)
            if "zero_point_energy_hartree" in thermo
            else ""
        )

    _prevent_row_splits(table)


def _add_transition_table(
    doc: Document,
    title: str,
    transitions: list[Dict[str, Any]],
    limit: int | None = None,
) -> None:
    if not transitions:
        return
    _add_heading_with_subscript(doc, title, level=3)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["From", "To", "Energy (eV)", "Wavelength (nm)", "fosc"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
    _style_header_row(table.rows[0])

    def _wl_nm(entry: Dict[str, Any]) -> float:
        try:
            return float(entry.get("wavelength_nm", float("inf")))
        except Exception:
            return float("inf")

    # Sort by wavelength (nm) ascending
    sorted_transitions = sorted(transitions, key=_wl_nm)
    max_items = len(sorted_transitions) if limit is None else limit
    for trans in sorted_transitions[:max_items]:
        row = table.add_row().cells

        # Format "From" state with subscript
        from_state = _translate_state(trans.get("from_state", ""))
        row[0].text = ""  # Clear default text
        _format_state_with_subscript(row[0].paragraphs[0], from_state)

        # Format "To" state with subscript
        to_state = _translate_state(trans.get("to_state", ""))
        row[1].text = ""  # Clear default text
        _format_state_with_subscript(row[1].paragraphs[0], to_state)

        # Format numeric values with 3 significant figures
        energy = trans.get('energy_eV', '')
        row[2].text = _format_significant_figures(energy) if energy != '' else ''

        wavelength = trans.get('wavelength_nm', '')
        row[3].text = _format_significant_figures(wavelength) if wavelength != '' else ''

        fosc = trans.get('oscillator_strength', '')
        row[4].text = _format_significant_figures(fosc) if fosc != '' else ''

    _prevent_row_splits(table)


def _add_rate_table(doc: Document, title: str, entries: Dict[str, Any]) -> None:
    if not entries:
        return
    heading = doc.add_heading(title, level=2)
    _keep_heading_with_table(heading)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    headers = ["Transition", "Rate (s⁻¹)", "Temperature (K)", "Δ0-0 (cm⁻¹)", "SOC (cm⁻¹)", "FC (%)", "HT (%)"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
    _style_header_row(table.rows[0])

    def _fmt_sci(val) -> str:
        try:
            return f"{float(val):.3e}"
        except Exception:
            return str(val) if val is not None else ""

    def _format_soc(rec: Dict[str, Any]) -> str:
        re_part = rec.get("soc_re_cm1")
        im_part = rec.get("soc_im_cm1")
        if re_part is None and im_part is None:
            return ""
        if re_part is not None and im_part is not None:
            return f"{_fmt_sci(re_part)}+i{_fmt_sci(im_part)}"
        if re_part is not None:
            return _fmt_sci(re_part)
        return f"i{_fmt_sci(im_part)}"

    def _add_row(label: str, rec: Dict[str, Any]) -> None:
        row = table.add_row().cells
        row[0].text = label
        rate_val = rec.get("rate_s1") or rec.get("total_rate_s1") or rec.get("rate")
        row[1].text = _fmt_sci(rate_val) if rate_val is not None else ""
        row[2].text = str(rec.get("temperature_K", ""))
        row[3].text = str(rec.get("delta_E_cm1", ""))
        row[4].text = _format_soc(rec)
        row[5].text = str(rec.get("fc_percent", ""))
        row[6].text = str(rec.get("ht_percent", ""))

    for name, record in sorted(entries.items()):
        ms_comps = record.get("ms_components", {})
        if ms_comps:
            ms_rec_list = []
            for ms_key, ms_rec in sorted(ms_comps.items()):
                pretty = {
                    "ms_0": "Ms=0",
                    "ms_p1": "Ms=+1",
                    "ms_m1": "Ms=-1",
                }.get(ms_key, ms_key.replace("_", "="))
                _add_row(f"{name} ({pretty})", ms_rec)
                ms_rec_list.append(ms_rec)
            # Add total if available
            total_rate = record.get("total_rate_s1") or record.get("rate_s1")
            if total_rate is not None:
                total_rec = {"total_rate_s1": total_rate}
                if ms_rec_list:
                    sample = ms_rec_list[0]
                    for key in ["temperature_K", "delta_E_cm1"]:
                        if sample.get(key) is not None:
                            total_rec[key] = sample.get(key)
                _add_row(f"{name} (total)", total_rec)
        else:
            _add_row(name, record)

    _prevent_row_splits(table)


def _extract_frontier_orbitals(orbital_data: Optional[Dict[str, Any]]) -> list[tuple[str, str, str]]:
    """Extract LUMO+3 to HOMO-3 orbital energies from orbital data.

    Returns:
        List of tuples (Index, Energy, Orbital) where Orbital is empty string.
    """
    if not orbital_data or "orbital_list" not in orbital_data:
        return []

    orbital_list = orbital_data["orbital_list"]

    # Find HOMO index (last occupied orbital)
    homo_idx = None
    for orbital in orbital_list:
        if orbital.get("occupancy", 0) > 1e-3:
            homo_idx = orbital.get("index")

    if homo_idx is None:
        return []

    # Define target orbitals relative to HOMO (LUMO is HOMO+1)
    target_orbitals = [
        ("LUMO+3", homo_idx + 4),
        ("LUMO+2", homo_idx + 3),
        ("LUMO+1", homo_idx + 2),
        ("LUMO", homo_idx + 1),
        ("HOMO", homo_idx),
        ("HOMO-1", homo_idx - 1),
        ("HOMO-2", homo_idx - 2),
        ("HOMO-3", homo_idx - 3),
    ]

    # Build index to orbital mapping
    idx_to_orbital = {orb["index"]: orb for orb in orbital_list}

    # Extract energies
    rows = []
    for label, idx in target_orbitals:
        if idx in idx_to_orbital:
            energy_ev = idx_to_orbital[idx].get("energy_eV")
            if energy_ev is not None:
                energy_str = f"{energy_ev:.4f}"
            else:
                energy_str = "n/a"
        else:
            energy_str = "n/a"
        rows.append((label, energy_str, ""))

    return rows


def _add_frontier_orbital_table(doc: Document, orbital_data: Optional[Dict[str, Any]], mo_pngs: Dict[str, Path] | None = None) -> None:
    """Add table showing frontier orbital energies (LUMO+3 to HOMO-3) with MO visualizations."""
    rows = _extract_frontier_orbitals(orbital_data)
    if not rows:
        return

    heading = doc.add_heading("Frontier Orbital Energies", level=2)
    _keep_heading_with_table(heading)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    # Header row
    headers = ["Index", "Energy (eV)", "Orbital"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
    _style_header_row(table.rows[0])

    # Data rows
    for label, energy, orbital in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = energy

        # Add MO image if available
        if mo_pngs and label in mo_pngs:
            mo_png = mo_pngs[label]
            if mo_png.exists():
                paragraph = row_cells[2].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(str(mo_png), width=Inches(2.0))
            else:
                row_cells[2].text = orbital
        else:
            row_cells[2].text = orbital

    _prevent_row_splits(table)


def _create_energy_level_plot(data: Dict[str, Any], output_path: Path) -> Optional[Path]:
    """Create energy level diagram for S, T, oxidized and reduced states.

    Args:
        data: DELFIN data dictionary
        output_path: Path to save the PNG plot

    Returns:
        Path to saved plot or None if creation failed
    """
    try:
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
    except ImportError:
        logger.warning("matplotlib not available; skipping energy level plot")
        return None

    # Extract S and T states with their energies
    s_states = []
    t_states = []

    # Ground state S0
    gs = data.get("ground_state_S0", {}) or {}
    if gs:
        opt = gs.get("optimization", {}) or {}
        hartree = opt.get("hartree")
        if hartree is not None:
            s_states.append(("S0", float(hartree)))

    # Excited states
    excited = data.get("excited_states", {}) or {}
    for state_name, entry in excited.items():
        if not (state_name.startswith("S") or state_name.startswith("T")):
            continue
        opt = entry.get("optimization", {}) or {}
        hartree = opt.get("hartree")
        if hartree is not None:
            if state_name.startswith("S"):
                s_states.append((state_name, float(hartree)))
            elif state_name.startswith("T"):
                t_states.append((state_name, float(hartree)))

    if not s_states and not t_states:
        logger.warning("No S or T states found for energy level plot")
        return None

    # Sort by energy
    s_states.sort(key=lambda x: x[1])
    t_states.sort(key=lambda x: x[1])

    # Create plot with variable width based on number of lanes
    num_lanes = sum([1 if s_states else 0, 1 if t_states else 0])
    fig_width = 4 + num_lanes * 2
    fig, ax = plt.subplots(figsize=(fig_width, 6))

    # Define x positions for the lanes
    lane_spacing = 1.2
    current_x = 1.0
    lane_positions = {}
    lane_labels = {}

    if s_states:
        lane_positions['S'] = current_x
        lane_labels['S'] = "Singlet States"
        current_x += lane_spacing
    if t_states:
        lane_positions['T'] = current_x
        lane_labels['T'] = "Triplet States"

    line_width = 0.3

    def _avoid_label_overlap(states_list, x_pos):
        """Adjust label positions to avoid overlap by placing them side-by-side."""
        # Group states that are very close in energy (within 0.005 Eh ~0.14 eV)
        overlap_threshold = 0.005
        label_offsets = []
        overlap_groups = []
        current_group = []

        # Group overlapping states
        for i, (name, energy) in enumerate(states_list):
            if not current_group:
                current_group = [i]
            else:
                # Check if this state overlaps with the last one in current group
                prev_idx = current_group[-1]
                prev_energy = states_list[prev_idx][1]
                if abs(energy - prev_energy) < overlap_threshold:
                    current_group.append(i)
                else:
                    # Start a new group
                    overlap_groups.append(current_group)
                    current_group = [i]

        if current_group:
            overlap_groups.append(current_group)

        # Assign offsets: for groups with multiple states, spread them horizontally
        for group in overlap_groups:
            if len(group) == 1:
                label_offsets.append(0)
            else:
                # Spread labels horizontally for overlapping states
                for j, idx in enumerate(group):
                    offset = j * 0.15  # Horizontal spacing between labels
                    label_offsets.append(offset)

        return label_offsets

    # Plot S states
    if s_states:
        s_x = lane_positions['S']
        label_offsets = _avoid_label_overlap(s_states, s_x)
        for idx, (state_name, energy) in enumerate(s_states):
            ax.plot([s_x - line_width, s_x + line_width], [energy, energy], 'b-', linewidth=2)
            # Format label with subscript
            label = state_name.replace("S", "$S_").replace("0", "{0}") \
                              .replace("1", "{1}").replace("2", "{2}") \
                              .replace("3", "{3}").replace("4", "{4}") \
                              .replace("5", "{5}").replace("6", "{6}") + "$"
            # Apply offset to avoid overlap
            x_offset = s_x + line_width + 0.05 + label_offsets[idx]
            ax.text(x_offset, energy, label, va='center', fontsize=10, ha='left')

    # Plot T states
    if t_states:
        t_x = lane_positions['T']
        label_offsets = _avoid_label_overlap(t_states, t_x)
        for idx, (state_name, energy) in enumerate(t_states):
            ax.plot([t_x - line_width, t_x + line_width], [energy, energy], 'r-', linewidth=2)
            # Format label with subscript
            label = state_name.replace("T", "$T_").replace("1", "{1}") \
                              .replace("2", "{2}").replace("3", "{3}") \
                              .replace("4", "{4}").replace("5", "{5}") \
                              .replace("6", "{6}") + "$"
            # Apply offset to avoid overlap
            x_offset = t_x + line_width + 0.05 + label_offsets[idx]
            ax.text(x_offset, energy, label, va='center', fontsize=10, ha='left')

    # Styling
    ax.set_xlim(0.3, current_x + 0.8)
    ax.set_ylabel("Energy (Eh)", fontsize=12)

    # Set x-ticks for all lanes
    xtick_positions = [lane_positions[k] for k in sorted(lane_positions.keys(),
                                                          key=lambda x: lane_positions[x])]
    xtick_labels = [lane_labels[k] for k in sorted(lane_positions.keys(),
                                                     key=lambda x: lane_positions[x])]
    ax.set_xticks(xtick_positions)
    ax.set_xticklabels(xtick_labels, fontsize=11)
    ax.tick_params(axis='both', labelsize=11)

    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.title("Energy Level Diagram (Optimized State Energies)", fontsize=14, fontweight='bold')
    plt.tight_layout()

    # Save plot
    try:
        plt.savefig(str(output_path), dpi=300, bbox_inches='tight')
        plt.close(fig)
        logger.info(f"Energy level plot saved to {output_path}")
        return output_path
    except Exception as exc:
        logger.error(f"Failed to save energy level plot: {exc}")
        plt.close(fig)
        return None


def _create_vertical_excitation_plot(data: Dict[str, Any], output_path: Path) -> Optional[Path]:
    """Create vertical excitation energy diagram with separate lanes for S0, S1, T1.

    This shows all vertical excitation energies (from TDDFT) for each state on its own axis.

    Args:
        data: DELFIN data dictionary
        output_path: Path to save the PNG plot

    Returns:
        Path to saved plot or None if creation failed
    """
    try:
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
    except ImportError:
        logger.warning("matplotlib not available; skipping vertical excitation plot")
        return None

    # Collect transitions for S0, S1, T1
    # Structure: state_name -> list of (from_state, to_state, energy_eV)
    transitions_data = {}

    # S0 absorption transitions
    gs = data.get("ground_state_S0", {}) or {}
    s0_abs = (gs.get("tddft_absorption") or {}).get("transitions", []) or []
    s0_transitions = []
    for t in s0_abs:
        from_state = _translate_state(t.get("from_state", ""))
        to_state = _translate_state(t.get("to_state", ""))
        energy_ev = t.get("energy_eV")
        if energy_ev is not None:
            s0_transitions.append((from_state, to_state, float(energy_ev)))
    if s0_transitions:
        transitions_data["S0"] = s0_transitions

    # S1 and T1 transitions
    excited = data.get("excited_states", {}) or {}
    for state_name in ["S1", "T1"]:
        state_data = excited.get(state_name, {}) or {}
        transitions = (state_data.get("tddft_from_geometry") or {}).get("transitions", []) or []
        state_transitions = []
        for t in transitions:
            from_state = _translate_state(t.get("from_state", ""))
            to_state = _translate_state(t.get("to_state", ""))
            energy_ev = t.get("energy_eV")
            if energy_ev is not None:
                state_transitions.append((from_state, to_state, float(energy_ev)))
        if state_transitions:
            transitions_data[state_name] = state_transitions

    if not transitions_data:
        logger.warning("No vertical excitation data found for S0, S1, or T1")
        return None

    # Create plot with lanes for each state
    num_lanes = len(transitions_data)
    fig_width = 4 + num_lanes * 2.5
    fig, ax = plt.subplots(figsize=(fig_width, 6))

    # Define x positions for lanes
    lane_spacing = 1.5
    current_x = 1.0
    lane_positions = {}
    lane_labels = {}
    lane_colors = {
        "S0": "blue",
        "S1": "green",
        "T1": "red"
    }

    for state in ["S0", "S1", "T1"]:
        if state in transitions_data:
            lane_positions[state] = current_x
            lane_labels[state] = f"{state} Geometry"
            current_x += lane_spacing

    line_width = 0.3

    # Plot transitions for each state
    for state, x_pos in lane_positions.items():
        transitions = transitions_data[state]
        # Sort by energy
        transitions.sort(key=lambda x: x[2])
        color = lane_colors.get(state, "gray")

        for idx, (from_state, to_state, energy) in enumerate(transitions):
            # Plot horizontal line for this transition
            ax.plot([x_pos - line_width, x_pos + line_width], [energy, energy],
                    color=color, linewidth=2, solid_capstyle='butt')
            # Add label only for first 3 transitions
            if idx < 3:
                label = f"{from_state}→{to_state}"
                ax.text(x_pos + line_width + 0.05, energy,
                       f"{label} ({energy:.2f} eV)",
                       va='center', fontsize=8, ha='left')

    # Styling
    ax.set_xlim(0.3, current_x + 0.8)
    ax.set_ylabel("Vertical Excitation Energy (eV)", fontsize=12)

    # Set x-ticks for all lanes
    xtick_positions = [lane_positions[k] for k in sorted(lane_positions.keys(),
                                                          key=lambda x: ["S0", "S1", "T1"].index(x))]
    xtick_labels = [lane_labels[k] for k in sorted(lane_positions.keys(),
                                                     key=lambda x: ["S0", "S1", "T1"].index(x))]
    ax.set_xticks(xtick_positions)
    ax.set_xticklabels(xtick_labels, fontsize=11)
    ax.tick_params(axis='both', labelsize=11)

    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.title("Vertical Excitation Energies (TDDFT)", fontsize=14, fontweight='bold')
    plt.tight_layout()

    # Save plot
    try:
        plt.savefig(str(output_path), dpi=300, bbox_inches='tight')
        plt.close(fig)
        logger.info(f"Vertical excitation plot saved to {output_path}")
        return output_path
    except Exception as exc:
        logger.error(f"Failed to save vertical excitation plot: {exc}")
        plt.close(fig)
        return None


def _format_state_label(state_name: str) -> str:
    """Format state name with subscript notation for matplotlib.

    Examples: S0 -> $S_0$, S1 -> $S_1$, T1 -> $T_1$
    """
    import re
    match = re.match(r'([ST])(\d+)', state_name)
    if match:
        letter = match.group(1)
        number = match.group(2)
        return f"${letter}_{{{number}}}$"
    return state_name


def _create_dipole_moment_plot(project_dir: Path, data: Dict[str, Any], output_path: Path) -> Optional[Path]:
    """Create 3D visualization of S0 geometry with dipole moment vector using PyMOL.

    Args:
        project_dir: Project directory containing S0.xyz
        data: DELFIN data dictionary
        output_path: Path to save the PNG plot

    Returns:
        Path to saved plot or None if creation failed
    """
    def _autocrop_white(im, threshold: int = 10, pad: int = 8):
        try:
            from PIL import Image, ImageChops
        except Exception:
            return im

        bg = Image.new("RGB", im.size, "white")
        diff = ImageChops.difference(im.convert("RGB"), bg).convert("L")
        diff = diff.point(lambda p: 255 if p > threshold else 0)
        bbox = diff.getbbox()
        if not bbox:
            return im
        left, upper, right, lower = bbox
        left = max(0, left - pad)
        upper = max(0, upper - pad)
        right = min(im.width, right + pad)
        lower = min(im.height, lower + pad)
        return im.crop((left, upper, right, lower))

    def _montage(png_paths: list[Path], out_png: Path, cols: int = 3, pad_px: int = 10) -> Optional[Path]:
        try:
            from PIL import Image
        except Exception:
            return None

        imgs = []
        for p in png_paths:
            if not p.exists():
                continue
            im = Image.open(p).convert("RGB")
            im = _autocrop_white(im, threshold=10, pad=8)
            imgs.append(im)
        if not imgs:
            return None

        cols = max(1, int(cols))
        rows = (len(imgs) + cols - 1) // cols
        pad_px = max(0, int(pad_px))

        tile_w = min(im.width for im in imgs)
        tile_h = min(im.height for im in imgs)
        imgs = [im.resize((tile_w, tile_h), Image.Resampling.LANCZOS) for im in imgs]

        out = Image.new("RGB", (cols * tile_w + (cols - 1) * pad_px, rows * tile_h + (rows - 1) * pad_px), "white")
        for idx, im in enumerate(imgs):
            r = idx // cols
            c = idx % cols
            out.paste(im, (c * (tile_w + pad_px), r * (tile_h + pad_px)))

        out_png.parent.mkdir(parents=True, exist_ok=True)
        out.save(out_png)
        return out_png if out_png.exists() else None

    try:
        import pymol
        from pymol import cmd
        import numpy as np
        from ase.io import read, write
    except ImportError as e:
        logger.warning(f"PyMOL or required dependencies not available; skipping dipole moment plot: {e}")
        return None

    # Get dipole moment from data
    gs = data.get("ground_state_S0", {}) or {}
    dipole = gs.get("dipole_moment", {}) or {}
    dipole_x = dipole.get("x_au")
    dipole_y = dipole.get("y_au")
    dipole_z = dipole.get("z_au")
    dipole_mag = dipole.get("magnitude_debye")

    if dipole_x is None or dipole_y is None or dipole_z is None:
        logger.warning("Dipole moment components not found in data")
        return None

    # Read S0.xyz file
    xyz_file = project_dir / "ESD" / "S0.xyz"
    if not xyz_file.exists():
        logger.warning(f"S0.xyz not found at {xyz_file}")
        return None

    try:
        # Read structure using ASE
        atoms = read(str(xyz_file))

        # Dipole moment vector (in a.u.)
        dipole_vec = np.array([dipole_x, dipole_y, dipole_z])

        # Calculate rotation matrix to align dipole moment with X-axis
        dipole_norm = dipole_vec / np.linalg.norm(dipole_vec)
        x_axis = np.array([1.0, 0.0, 0.0])

        # Rotation axis (cross product)
        rotation_axis = np.cross(dipole_norm, x_axis)
        rotation_axis_norm = np.linalg.norm(rotation_axis)

        # Rotation angle
        cos_angle = np.dot(dipole_norm, x_axis)
        angle = np.arccos(np.clip(cos_angle, -1.0, 1.0))

        # Apply rotation to atoms if dipole is not already aligned with X
        if rotation_axis_norm > 1e-6:  # Not already aligned
            rotation_axis = rotation_axis / rotation_axis_norm

            # Rodrigues' rotation formula
            K = np.array([
                [0, -rotation_axis[2], rotation_axis[1]],
                [rotation_axis[2], 0, -rotation_axis[0]],
                [-rotation_axis[1], rotation_axis[0], 0]
            ])
            R = np.eye(3) + np.sin(angle) * K + (1 - np.cos(angle)) * (K @ K)

            # Rotate all atoms
            positions = atoms.get_positions()
            center = atoms.get_center_of_mass()
            positions_centered = positions - center
            positions_rotated = (R @ positions_centered.T).T + center
            atoms.set_positions(positions_rotated)

            # Rotate dipole vector
            dipole_vec = R @ dipole_vec

        # Save rotated structure to temporary file
        temp_xyz = project_dir / "temp_rotated_dipole.xyz"
        write(str(temp_xyz), atoms)

        # Get center of mass for dipole arrow placement
        center = atoms.get_center_of_mass()

        # Initialize PyMOL in quiet mode (no GUI)
        pymol.finish_launching(['pymol', '-c'])
        cmd.reinitialize()

        # Load the rotated structure
        cmd.load(str(temp_xyz), 'molecule')

        # Set up visualization style
        cmd.hide('everything', 'molecule')
        cmd.show('sticks', 'molecule')
        cmd.show('spheres', 'molecule')
        cmd.set('sphere_scale', 0.25)
        cmd.set('stick_radius', 0.15)
        cmd.set('stick_quality', 15)
        cmd.set('sphere_quality', 3)

        # Use ball-and-stick representation
        cmd.set('valence', 1)

        # Set element colors - carbon should be gray, not green
        cmd.color('gray50', 'elem C')
        cmd.color('white', 'elem H')
        cmd.color('slate', 'elem N')
        cmd.color('red', 'elem O')
        cmd.color('yellow', 'elem S')

        # Set background to white
        cmd.bg_color('white')

        # Enable ray tracing settings for high quality
        cmd.set('ray_trace_mode', 1)
        cmd.set('ray_shadows', 1)
        cmd.set('ray_opaque_background', 1)
        cmd.set('antialias', 2)
        cmd.set('ambient', 0.4)
        cmd.set('direct', 0.6)
        cmd.set('specular', 0.5)
        cmd.set('shininess', 10)
        cmd.set('depth_cue', 0)

        # Draw dipole moment vector as CGO (Compiled Graphics Object)
        from pymol.cgo import CYLINDER, CONE, COLOR

        # Use dipole vector directly for visualization (no artificial scaling)
        dipole_vec_scaled = dipole_vec
        arrow_start = center
        arrow_end = center + dipole_vec_scaled

        # Arrow shaft radius and cone dimensions
        shaft_radius = 0.15
        cone_radius = 0.3
        cone_length = 0.8

        # Calculate cone base position
        arrow_direction = dipole_vec_scaled / np.linalg.norm(dipole_vec_scaled)
        cone_base = arrow_end - arrow_direction * cone_length

        # Create CGO arrow (red color: RGB = 1.0, 0.0, 0.0)
        arrow_obj = [
            COLOR, 1.0, 0.0, 0.0,  # Red color
            CYLINDER, float(arrow_start[0]), float(arrow_start[1]), float(arrow_start[2]),
                     float(cone_base[0]), float(cone_base[1]), float(cone_base[2]),
                     shaft_radius,
                     1.0, 0.0, 0.0,  # Start color (red)
                     1.0, 0.0, 0.0,  # End color (red)
            CONE, float(cone_base[0]), float(cone_base[1]), float(cone_base[2]),
                 float(arrow_end[0]), float(arrow_end[1]), float(arrow_end[2]),
                 cone_radius, 0.0,
                 1.0, 0.0, 0.0,  # Base color (red)
                 1.0, 0.0, 0.0,  # Tip color (red)
                 1.0, 1.0,  # Caps
        ]
        cmd.load_cgo(arrow_obj, 'dipole_arrow')

        # Set camera view to emphasize X-axis alignment
        cmd.orient('molecule')
        cmd.zoom('all', buffer=2)

        # Render 6 views (3x2 montage) for a more informative figure
        view_files = [
            output_path.with_name(output_path.stem + "_v1.png"),
            output_path.with_name(output_path.stem + "_v2.png"),
            output_path.with_name(output_path.stem + "_v3.png"),
            output_path.with_name(output_path.stem + "_v4.png"),
            output_path.with_name(output_path.stem + "_v5.png"),
            output_path.with_name(output_path.stem + "_v6.png"),
        ]
        # Azimuth (y) x Elevation (x)
        y_angles = [-60, 0, 60]
        x_angles = [15, -30]
        k = 0
        for xa in x_angles:
            for ya in y_angles:
                cmd.orient('molecule')
                cmd.turn('y', ya)
                cmd.turn('x', xa)
                cmd.zoom('all', buffer=2)
                cmd.ray(1200, 1000)
                cmd.png(str(view_files[k]), dpi=300)
                k += 1

        montage_path = _montage(view_files, output_path, cols=3, pad_px=10)
        for vf in view_files:
            try:
                if vf.exists():
                    vf.unlink()
            except Exception:
                pass
        if montage_path is None:
            # Fallback: single view (legacy)
            cmd.orient('molecule')
            cmd.turn('y', -60)
            cmd.turn('x', 15)
            cmd.zoom('all', buffer=2)
            cmd.ray(2400, 2000)
            cmd.png(str(output_path), dpi=300)

        # Clean up
        cmd.delete('all')
        cmd.reinitialize()
        if temp_xyz.exists():
            temp_xyz.unlink()

        logger.info(f"Dipole moment plot saved to {output_path}")
        return output_path

    except Exception as exc:
        logger.error(f"Failed to create dipole moment plot: {exc}")
        import traceback
        logger.error(traceback.format_exc())
        return None


def _create_mo_visualizations(project_dir: Path, orbital_indices: list[int]) -> Dict[str, Path]:
    """Generate molecular orbital visualizations using ORCA and PyMOL.

    Args:
        project_dir: Project directory containing ESD/S0.gbw
        orbital_indices: List of orbital indices relative to HOMO (e.g., [-3, -2, -1, 0, 1, 2, 3, 4] for HOMO-3 to LUMO+3)

    Returns:
        Dictionary mapping orbital names (e.g., 'HOMO', 'LUMO+1') to PNG paths
    """
    import subprocess
    import shutil

    try:
        import pymol
        from pymol import cmd
    except ImportError:
        logger.warning("PyMOL not available; skipping MO visualizations")
        return {}

    gbw_file = project_dir / "ESD" / "S0.gbw"
    if not gbw_file.exists():
        logger.warning(f"S0.gbw not found at {gbw_file}")
        return {}

    # Find HOMO orbital number from S0.out
    s0_out = project_dir / "ESD" / "S0.out"
    if not s0_out.exists():
        logger.warning(f"S0.out not found at {s0_out}")
        return {}

    homo_number = None
    try:
        with open(s0_out, 'r') as f:
            lines = f.readlines()
            for i, line in enumerate(lines):
                if "ORBITAL ENERGIES" in line:
                    # Find the last occupied orbital
                    for j in range(i, min(i + 300, len(lines))):
                        if "NO   OCC" in lines[j]:
                            # Parse orbital lines
                            for k in range(j + 1, min(j + 250, len(lines))):
                                parts = lines[k].split()
                                if len(parts) >= 2:
                                    try:
                                        orb_num = int(parts[0])
                                        occ = float(parts[1])
                                        # HOMO is the last orbital with occupation 2.0
                                        if occ > 1.9:
                                            homo_number = orb_num
                                        elif occ < 0.1 and homo_number is not None:
                                            # Found LUMO, stop searching
                                            break
                                    except (ValueError, IndexError):
                                        continue
                            break
                    if homo_number is not None:
                        break
    except Exception as exc:
        logger.error(f"Failed to find HOMO orbital number: {exc}")
        return {}

    if homo_number is None:
        logger.warning("Could not determine HOMO orbital number")
        return {}

    logger.info(f"HOMO orbital number: {homo_number}")

    mo_images = {}

    # Work directly in ESD directory where .gbw file is located
    esd_dir = gbw_file.parent

    for rel_idx in orbital_indices:
        # Calculate absolute MO number
        abs_mo_number = homo_number + rel_idx

        # Determine orbital name
        if rel_idx == 0:
            mo_name = "HOMO"
        elif rel_idx < 0:
            mo_name = f"HOMO{rel_idx}"  # e.g., HOMO-1
        elif rel_idx == 1:
            mo_name = "LUMO"
        else:
            mo_name = f"LUMO+{rel_idx-1}"  # e.g., LUMO+1 for index 2

        logger.info(f"Generating MO visualization for {mo_name} (orbital {abs_mo_number})")

        try:
            # Prepare orca_plot input for interactive mode
            # Menu sequence:
            # 1 - Enter type of plot -> 1 (MO-PLOT)
            # 2 - Enter no of orbital -> abs_mo_number
            # 4 - Enter number of grid intervals -> 100
            # 11 - Generate the plot
            # 12 - exit
            orca_input = f"1\n1\n2\n{abs_mo_number}\n4\n100\n11\n12\n"

            # Run orca_plot in interactive mode from ESD directory
            result = subprocess.run(
                ["/opt/orca/orca_plot", "S0.gbw", "-i"],
                input=orca_input.encode(),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                cwd=esd_dir,
                timeout=120
            )

            if result.returncode != 0:
                logger.warning(f"orca_plot failed for {mo_name}: {result.stderr.decode()}")
                continue

            # Find generated cube file (ORCA creates S0.mo<X>a.cube or similar)
            cube_files = sorted(esd_dir.glob("*.cube"), key=lambda p: p.stat().st_mtime)
            if not cube_files:
                logger.warning(f"No cube file generated for {mo_name}")
                continue

            cube_file = cube_files[-1]  # Use the most recently created

            # Visualize with PyMOL
            pymol.finish_launching(['pymol', '-c'])
            cmd.reinitialize()

            # Load molecular structure from S0.xyz
            xyz_file = project_dir / "ESD" / "S0.xyz"
            if xyz_file.exists():
                cmd.load(str(xyz_file), 'molecule')

                # Display molecule as sticks
                cmd.hide('everything', 'molecule')
                cmd.show('sticks', 'molecule')
                cmd.set('stick_radius', 0.15)
                cmd.set('stick_quality', 15)

                # Set element colors
                cmd.color('gray50', 'elem C')
                cmd.color('white', 'elem H')
                cmd.color('slate', 'elem N')
                cmd.color('red', 'elem O')
                cmd.color('yellow', 'elem S')

            # Load cube file for MO isosurface
            cmd.load(str(cube_file), 'orbital')

            # Create isosurface
            cmd.isosurface('positive', 'orbital', 0.03)
            cmd.isosurface('negative', 'orbital', -0.03)

            # Color isosurfaces (blue for positive, red for negative)
            cmd.color('blue', 'positive')
            cmd.color('red', 'negative')

            # Set transparency for isosurfaces
            cmd.set('transparency', 0.3, 'positive')
            cmd.set('transparency', 0.3, 'negative')

            # Set background to white
            cmd.bg_color('white')

            # Enable ray tracing
            cmd.set('ray_trace_mode', 1)
            cmd.set('ray_shadows', 0)
            cmd.set('ray_opaque_background', 1)
            cmd.set('antialias', 2)

            # Orient and zoom
            cmd.orient()
            cmd.zoom('all', buffer=2.5)

            # Render
            output_png = project_dir / f"MO_{mo_name}.png"
            cmd.ray(1200, 1200)
            cmd.png(str(output_png), dpi=150)

            # Clean up cube file
            if cube_file.exists():
                cube_file.unlink()

            # Clean up PyMOL
            cmd.delete('all')
            cmd.reinitialize()

            mo_images[mo_name] = output_png
            logger.info(f"MO visualization saved: {output_png}")

        except Exception as exc:
            logger.error(f"Failed to create MO visualization for {mo_name}: {exc}")
            continue

    return mo_images


def _create_correlation_plot(data: Dict[str, Any], output_path: Path) -> Optional[Path]:
    """Create correlation diagram between S0 vertical excitations and optimized state energies.

    Left axis: S0 TDDFT vertical excitations (eV)
    Right axis: Final single point energies of optimized states, normalized to S0=0 eV
    Dashed lines connect corresponding states.

    Args:
        data: DELFIN data dictionary
        output_path: Path to save the PNG plot

    Returns:
        Path to saved plot or None if creation failed
    """
    try:
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')  # Non-interactive backend
    except ImportError:
        logger.warning("matplotlib not available; skipping correlation plot")
        return None

    # Get S0 energy for normalization
    gs = data.get("ground_state_S0", {}) or {}
    gs_opt = gs.get("optimization", {}) or {}
    s0_hartree = gs_opt.get("hartree")
    if s0_hartree is None:
        logger.warning("S0 energy not found for correlation plot")
        return None
    s0_energy_ev = float(s0_hartree) * HARTREE_TO_EV

    # Collect optimized state energies (right side), normalized to S0=0
    excited = data.get("excited_states", {}) or {}
    right_states_dict = {"S0": 0.0}  # S0 at 0 eV reference
    for state_name in sorted(excited.keys()):
        state_data = excited[state_name] or {}
        opt = state_data.get("optimization", {}) or {}
        hartree = opt.get("hartree")
        if hartree is not None:
            # Convert to eV and normalize to S0=0
            energy_ev = float(hartree) * HARTREE_TO_EV - s0_energy_ev
            right_states_dict[state_name] = energy_ev

    # Collect S0 vertical excitations (left side) - only for states that have FSPE
    s0_abs = (gs.get("tddft_absorption") or {}).get("transitions", []) or []
    left_states_dict = {"S0": 0.0}  # S0 ground state at 0 eV
    for t in s0_abs:
        from_state = _translate_state(t.get("from_state", ""))
        to_state = _translate_state(t.get("to_state", ""))
        energy_ev = t.get("energy_eV")
        # Only include if there is a corresponding optimized state
        if energy_ev is not None and to_state in right_states_dict:
            left_states_dict[to_state] = float(energy_ev)

    # Find common states
    common_states = set(left_states_dict.keys()) & set(right_states_dict.keys())
    if len(common_states) <= 1:
        logger.warning("Insufficient common states for correlation plot")
        return None

    # Convert to lists for plotting - keep state ordering separate for left and right
    # Sort by energy for visual clarity on each side
    left_states = [(state, left_states_dict[state]) for state in sorted(common_states, key=lambda x: left_states_dict[x])]
    right_states = [(state, right_states_dict[state]) for state in sorted(common_states, key=lambda x: right_states_dict[x])]

    # Create mapping for connecting corresponding states (not by sorted order)
    state_connections = {state: (left_states_dict[state], right_states_dict[state]) for state in common_states}

    # Create plot with layout matching Energy Level Diagram
    num_lanes = 2  # left and right sides
    fig_width = 4 + num_lanes * 2
    fig, ax = plt.subplots(figsize=(fig_width, 6))

    # Define x positions (matching Energy Level Diagram style)
    lane_spacing = 1.2
    left_x = 1.0
    right_x = left_x + lane_spacing
    line_width = 0.3

    # Helper function to avoid label overlap
    def _get_label_offsets(states_list):
        """Calculate vertical offsets for labels to avoid overlap."""
        overlap_threshold = 0.15  # eV
        offsets = []
        for i, (_, energy) in enumerate(states_list):
            offset = 0
            for j in range(i):
                _, prev_energy = states_list[j]
                if abs(energy - prev_energy) < overlap_threshold:
                    offset += 0.08
            offsets.append(offset)
        return offsets

    # Plot left side (vertical excitations)
    left_offsets = _get_label_offsets(left_states)
    for idx, (state_name, energy) in enumerate(left_states):
        color = "blue" if state_name.startswith("S") else "red" if state_name.startswith("T") else "gray"
        ax.plot([left_x - line_width, left_x + line_width], [energy, energy],
                color=color, linewidth=2, solid_capstyle='butt')
        label = _format_state_label(state_name)
        ax.text(left_x + line_width + 0.05, energy + left_offsets[idx], label,
                ha='left', va='center', fontsize=10)

    # Plot right side (optimized states)
    right_offsets = _get_label_offsets(right_states)
    for idx, (state_name, energy) in enumerate(right_states):
        color = "blue" if state_name.startswith("S") else "red" if state_name.startswith("T") else "gray"
        ax.plot([right_x - line_width, right_x + line_width], [energy, energy],
                color=color, linewidth=2, solid_capstyle='butt')
        label = _format_state_label(state_name)
        ax.text(right_x + line_width + 0.05, energy + right_offsets[idx], label,
                ha='left', va='center', fontsize=10)

    # Connect corresponding states with dashed lines
    # Use state_connections to match T1↔T1, S1↔S1, etc., not by sorted order
    for state_name, (left_energy, right_energy) in state_connections.items():
        color = "blue" if state_name.startswith("S") else "red" if state_name.startswith("T") else "gray"
        ax.plot([left_x + line_width, right_x - line_width], [left_energy, right_energy],
                color=color, linestyle='--', linewidth=1, alpha=0.6)

    # Styling (matching Energy Level Diagram)
    ax.set_xlim(0.3, right_x + 0.8)
    ax.set_ylabel("Energy (eV)", fontsize=12)
    ax.set_xticks([left_x, right_x])
    ax.set_xticklabels(["S0 Vertical Excitations\n(TDDFT)", "Optimized States\n(FSPE, S0=0 eV)"], fontsize=11)
    ax.tick_params(axis='both', labelsize=11)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.title("Correlation: Vertical Excitations vs. Optimized State Energies", fontsize=14, fontweight='bold')
    plt.tight_layout()

    # Save plot
    try:
        plt.savefig(str(output_path), dpi=300, bbox_inches='tight')
        plt.close(fig)
        logger.info(f"Correlation plot saved to {output_path}")
        return output_path
    except Exception as exc:
        logger.error(f"Failed to save correlation plot: {exc}")
        plt.close(fig)
        return None


def generate_combined_docx_report(
    project_dir: Path,
    json_path: Path,
    output_docx: Path,
    assets: ReportAssets | None = None,
) -> Optional[Path]:
    """
    Build DELFIN.docx using the collected JSON data and available plots.

    Args:
        project_dir: Workspace directory (used for relative paths)
        json_path: Path to DELFIN_Data.json
        output_docx: Destination DOCX path
        assets: Optional plot paths to embed

    Returns:
        Path to generated DOCX or None if generation failed.
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed; cannot build DELFIN.docx")
        return None

    if assets is None:
        assets = ReportAssets()

    if assets.smiles_png is None:
        assets.smiles_png = _generate_smiles_image(project_dir)

    if not json_path.exists():
        logger.error("JSON data file not found: %s", json_path)
        return None

    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to read %s: %s", json_path, exc)
        return None

    doc = Document()

    # Title and metadata
    meta = data.get("metadata", {}) or {}
    molecule_name = meta.get("name") or meta.get("NAME") or project_dir.name
    heading = doc.add_heading(f"DELFIN Report – {molecule_name}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary_text, color_boxes = _build_summary_text(data, project_dir)
    if summary_text:
        summary_para = doc.add_paragraph(summary_text)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Color chips for key transitions (render before structure)
    if color_boxes:
        chip_para = doc.add_paragraph()
        chip_para.add_run("Transition colors: ")
        for label, wl, rgb in color_boxes:
            chip_path = _render_color_chip(rgb, project_dir, label)
            if chip_path and chip_path.exists():
                run = chip_para.add_run()
                run.add_picture(str(chip_path), width=Inches(0.25))
            else:
                run_box = chip_para.add_run("[]")
                run_box.font.color.rgb = RGBColor(*rgb)
                run_box.font.bold = True
            chip_para.add_run(f" {label} ({wl:.0f} nm)  ")

    # SMILES picture (if available) near the top
    if assets.smiles_png and assets.smiles_png.exists():
        doc.add_heading("Structure from SMILES (input.txt)", level=2)
        doc.add_picture(str(assets.smiles_png), width=Inches(2.25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = [
        ("Functional", meta.get("functional", "")),
        ("Basis set", meta.get("basis_set", "")),
        ("Auxiliary basis", meta.get("auxiliary_basis", "")),
        ("RI method", meta.get("ri_method", "")),
        ("Dispersion correction", meta.get("dispersion_correction", "")),
        ("Solvent model", meta.get("implicit_solvation", "")),
        ("Solvent", meta.get("solvent", "")),
        ("Charge", meta.get("charge", "")),
        ("Multiplicity", meta.get("multiplicity", "")),
    ]
    _add_key_value_table(doc, "Calculation Setup", meta_rows)

    # Add frontier orbital energies table for S0
    gs = data.get("ground_state_S0", {}) or {}
    gs_orbitals = gs.get("orbitals") if gs else None
    _add_frontier_orbital_table(doc, gs_orbitals, mo_pngs=assets.mo_pngs)

    # Collect state summaries for one consolidated table
    state_rows: Dict[str, Dict[str, Any]] = {}

    if gs:
        gs_entry = dict(gs)
        gs_entry["_type"] = "Ground"
        state_rows["S0"] = gs_entry
        # Absorption transitions from S0_TDDFT - start on new page
        doc.add_page_break()
        s0_abs = (gs.get("tddft_absorption") or {}).get("transitions", [])
        _add_transition_table(doc, "Absorption transitions (S0)", s0_abs)

    excited = data.get("excited_states", {}) or {}
    for name, entry in excited.items():
        entry = dict(entry)
        entry["_type"] = "Excited"
        state_rows[name] = entry
        transitions = (entry.get("tddft_from_geometry") or {}).get("transitions", [])
        if transitions:
            _add_transition_table(doc, f"Vertical transitions from {name}", transitions)

    for name, entry in (data.get("oxidized_states", {}) or {}).items():
        row = dict(entry)
        row["_type"] = "Oxidized"
        state_rows[name] = row

    for name, entry in (data.get("reduced_states", {}) or {}).items():
        row = dict(entry)
        row["_type"] = "Reduced"
        state_rows[name] = row

    # Add vertical excitation diagram
    _add_plot_if_exists(doc, "Vertical Excitation Energies", assets.vertical_excitation_png)

    # Add energetics overview table
    _add_state_table(doc, "Energetics overview", state_rows)

    # Add energy level diagram
    _add_plot_if_exists(doc, "Energy Level Diagram (Optimized State Energies)", assets.energy_level_png)

    # Add correlation diagram
    _add_plot_if_exists(doc, "Correlation: Vertical Excitations vs. Optimized State Energies", assets.correlation_png)

    # Rates
    _add_rate_table(doc, "Intersystem crossing", data.get("intersystem_crossing", {}) or {})
    _add_rate_table(doc, "Internal conversion", data.get("internal_conversion", {}) or {})

    # Plots
    _add_plot_if_exists(doc, "AFP spectrum", assets.afp_png)
    if assets.uv_vis_pngs:
        for state, png_path in sorted(assets.uv_vis_pngs.items()):
            label = {
                "S0": "Absorption spectrum (S0)",
                "S1": "Fluorescence spectrum (S1)",
                "T1": "Phosphorescence spectrum (T1)",
            }.get(state, f"Spectrum ({state})")
            # Use heading with subscript for spectrum labels
            if label:
                _add_heading_with_subscript(doc, label, level=2)
                doc.add_picture(str(png_path), width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # IR spectrum
    if assets.ir_png:
        _add_heading_with_subscript(doc, "IR spectrum (S0)", level=2)
        doc.add_picture(str(assets.ir_png), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dipole moment visualization
    if assets.dipole_moment_png:
        _add_heading_with_subscript(doc, "Dipole Moment Visualization (S0)", level=2)
        doc.add_picture(str(assets.dipole_moment_png), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption with dipole magnitude
        gs = data.get("ground_state_S0", {}) or {}
        dipole = gs.get("dipole_moment", {}) or {}
        dipole_mag = dipole.get("magnitude_debye")
        if dipole_mag is not None:
            caption = doc.add_paragraph(f"μ = {dipole_mag:.2f} D")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.runs[0]
            caption_run.font.size = Pt(11)
            caption_run.font.italic = True

    # Electrostatic potential (ESP) visualization (add at the very end)
    if assets.esp_png and assets.esp_png.exists():
        _add_heading_with_subscript(doc, "Electrostatic Potential Plot (S0)", level=2)
        doc.add_picture(str(assets.esp_png), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif assets.esp_png:
        logger.warning("Plot 'Electrostatic Potential Plot (S0)' not found at %s", assets.esp_png)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    logger.info("DELFIN.docx written to %s", output_docx)
    return output_docx
