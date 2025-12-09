"""Generate UV-Vis spectrum reports in Word format."""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-GUI backend
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

from delfin.common.logging import get_logger
from delfin.uv_vis_spectrum import (
    Transition,
    parse_absorption_spectrum,
    gaussian_broadening,
    filter_significant_transitions
)

logger = get_logger(__name__)


def create_spectrum_plot(
    transitions: List[Transition],
    output_png: Path,
    wavelength_range: tuple[float, float] = (200, 800),
    fwhm: float = 20.0,
    dpi: int = 300
) -> None:
    """Create UV-Vis spectrum plot with stick spectrum and broadened curve.

    Args:
        transitions: List of transitions
        output_png: Path to save PNG image
        wavelength_range: Wavelength range (min, max) in nm
        fwhm: Full width at half maximum for Gaussian broadening (nm)
        dpi: Resolution of output image
    """
    if not MATPLOTLIB_AVAILABLE:
        logger.warning("matplotlib not available, skipping spectrum plot")
        return

    fig, ax = plt.subplots(figsize=(10, 6))

    # Filter for significant transitions
    significant = filter_significant_transitions(transitions, min_fosc=0.001)

    if not significant:
        logger.warning("No significant transitions found for plotting")
        plt.close(fig)
        return

    # Create stick spectrum
    for trans in significant:
        ax.stem([trans.wavelength_nm], [trans.fosc],
                linefmt='gray', markerfmt='o', basefmt=' ',
                label='_nolegend_')

    # Add broadened spectrum
    wavelengths, intensities = gaussian_broadening(
        transitions,
        wavelength_range=wavelength_range,
        num_points=2000,
        fwhm=fwhm
    )
    ax.plot(wavelengths, intensities, 'b-', linewidth=2, label='Broadened spectrum')

    # Formatting
    ax.set_xlabel('Wavelength (nm)', fontsize=12)
    ax.set_ylabel('Oscillator strength / Intensity (arb. units)', fontsize=12)
    ax.set_title('UV-Vis Absorption Spectrum', fontsize=14, fontweight='bold')
    ax.set_xlim(wavelength_range)
    ax.set_ylim(bottom=0)
    ax.grid(True, alpha=0.3)
    ax.legend()

    plt.tight_layout()
    plt.savefig(output_png, dpi=dpi, bbox_inches='tight')
    plt.close(fig)
    logger.info(f"Spectrum plot saved to {output_png}")


def generate_uv_vis_word_report(
    esd_output_file: Path,
    output_docx: Path,
    min_fosc: float = 0.001,
    fwhm: float = 20.0
) -> None:
    """Generate a Word document with UV-Vis spectrum data and plot.

    Args:
        esd_output_file: Path to S0_TDDFT.out or S0.out file
        output_docx: Path to output Word document
        min_fosc: Minimum oscillator strength to include in table
        fwhm: Full width at half maximum for Gaussian broadening (nm)
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return

    # Parse spectrum
    transitions = parse_absorption_spectrum(esd_output_file)
    if not transitions:
        logger.warning(f"No transitions found in {esd_output_file}")
        return

    # Filter significant transitions for table
    significant = filter_significant_transitions(transitions, min_fosc=min_fosc)

    # Create Word document
    doc = Document()

    # Title
    title = doc.add_heading('UV-Vis Absorption Spectrum', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Source file info
    doc.add_paragraph(f"Source: {esd_output_file.name}")
    doc.add_paragraph(f"Total transitions: {len(transitions)}")
    doc.add_paragraph(f"Significant transitions (fosc ≥ {min_fosc}): {len(significant)}")
    doc.add_paragraph()

    # Table heading
    doc.add_heading('Transition Data', level=2)

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    headers = ['Transition', 'Energy (eV)', 'Wavelength (nm)', 'fosc', 'Type', 'Intensity']
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for trans in significant:
        row_cells = table.add_row().cells
        row_cells[0].text = trans.readable_transition
        row_cells[1].text = f"{trans.energy_ev:.4f}"
        row_cells[2].text = f"{trans.wavelength_nm:.1f}"
        row_cells[3].text = f"{trans.fosc:.6f}"

        # Determine transition type
        if 'T' in trans.to_state and 'S' in trans.from_state:
            trans_type = "Forbidden (S→T)"
        elif 'S' in trans.to_state and 'S' in trans.from_state:
            trans_type = "Allowed (S→S)"
        elif 'T' in trans.to_state and 'T' in trans.from_state:
            trans_type = "Allowed (T→T)"
        else:
            trans_type = "Unknown"
        row_cells[4].text = trans_type

        # Qualitative intensity
        if trans.fosc < 0.01:
            intensity = "Very weak"
        elif trans.fosc < 0.1:
            intensity = "Weak"
        elif trans.fosc < 0.5:
            intensity = "Medium"
        else:
            intensity = "Strong"
        row_cells[5].text = intensity

    doc.add_paragraph()

    # Create spectrum plot
    doc.add_heading('Simulated Spectrum', level=2)
    doc.add_paragraph(f"Gaussian broadening with FWHM = {fwhm} nm")
    doc.add_paragraph()

    # Generate plot
    temp_png = output_docx.parent / f"{output_docx.stem}_spectrum.png"
    create_spectrum_plot(
        transitions,
        temp_png,
        wavelength_range=(200, 800),
        fwhm=fwhm,
        dpi=300
    )

    # Insert plot into document
    if temp_png.exists():
        doc.add_picture(str(temp_png), width=Inches(6))
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add notes section
    doc.add_page_break()
    doc.add_heading('Notes', level=2)
    notes = [
        "• fosc: Oscillator strength (dimensionless)",
        "• Allowed transitions (S→S): High oscillator strength",
        "• Forbidden transitions (S→T): Zero or very low oscillator strength",
        "• The spectrum is simulated using Gaussian broadening",
        f"• Only transitions with fosc ≥ {min_fosc} are shown in the table"
    ]
    for note in notes:
        doc.add_paragraph(note)

    # Save document
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    logger.info(f"UV-Vis report saved to {output_docx}")

    # Clean up temporary PNG (optional - keep it for reference)
    # if temp_png.exists():
    #     temp_png.unlink()


def generate_all_esd_uv_vis_reports(esd_dir: Path, output_dir: Optional[Path] = None) -> None:
    """Generate UV-Vis reports for all relevant ESD output files.

    Args:
        esd_dir: ESD working directory
        output_dir: Output directory for reports (default: esd_dir)
    """
    if output_dir is None:
        output_dir = esd_dir

    output_dir.mkdir(parents=True, exist_ok=True)

    # Look for S0_TDDFT.out first, fallback to S0.out
    s0_output = esd_dir / "S0_TDDFT.out"
    if not s0_output.exists():
        s0_output = esd_dir / "S0.out"

    if not s0_output.exists():
        logger.warning(f"No S0 output file found in {esd_dir}")
        return

    # Generate report
    output_docx = output_dir / "UV_Vis_Spectrum.docx"
    generate_uv_vis_word_report(
        s0_output,
        output_docx,
        min_fosc=0.001,
        fwhm=20.0
    )
